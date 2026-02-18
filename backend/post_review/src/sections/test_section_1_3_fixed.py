#!/usr/bin/env python3
"""
Section 1_3 Individual Testing - FIXED IMPLEMENTATION
Tests portfolio selection (conservative/balanced/growth) and action taken deletions
"""

# This file requires pre-extracted images from pdf_section_splitter.py
# Run: splitter = PDFSectionSplitter(pdf_path); splitter.extract_all_sections()
# Then this file will work with the pre-extracted section image

import os
import json
import base64
import requests
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt
from datetime import datetime



# This file is designed to work only through the Post-Review UI system
# Run: python post_review_modern_ui.py
# Do not run this file directly

class Section13FixedTester:
    def __init__(self, pdf_path=None):
        # This test section works with pre-extracted PNG images only
        # PDF sectioning is handled by PDFSectionSplitter
        # Word processing is handled by unified_section_implementations.py
        
        # Make test directory relative to the project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.test_dir = os.path.join(project_root, "data", "test_sections", "section_1_3_test")
        
        # Get API key
        from dotenv import load_dotenv
        # Load environment variables from project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        env_path = os.path.join(project_root, '.env')
        load_dotenv(env_path)
        self.api_key = os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("OpenAI API key required. Set OPENAI_API_KEY in .env file.")
        
        self.api_url = "https://api.openai.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    

    
    def load_extracted_image(self) -> Image.Image:
        """Load pre-extracted section image (extracted by pdf_section_splitter.py)"""
        try:
            # Ensure test directory exists
            os.makedirs(self.test_dir, exist_ok=True)
            
            # Check for pre-extracted image
            image_filename = "section_1_3_extracted.png"
            image_path = os.path.join(self.test_dir, image_filename)
            
            if os.path.exists(image_path):
                print(f"✅ Found pre-extracted image: {image_path}")
                return image_path
            else:
                print(f"❌ Pre-extracted image not found: {image_path}")
                print(f"   💡 Make sure to run pdf_section_splitter.py first!")
                return None
                
        except Exception as e:
            print(f"Error loading pre-extracted image: {e}")
            return None
    
    def encode_image(self, image: Image.Image) -> str:
        """Encode PIL image to base64"""
        import io
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    def analyze_section_1_3_enhanced(self, section_image: Image.Image) -> dict:
        """Enhanced analysis for Section 1_3"""
        
        prompt = """
        Analyze this Section 1_3 for TWO CRITICAL TASKS:
        
        TASK 1: PORTFOLIO SELECTION (Left Box):
        - Find the text "conservative / balanced / growth" (or similar variations)
        - Look for CIRCLES around ONE word (this word STAYS)
        - Look for LINES/CROSSES through the OTHER TWO words (these get DELETED)
        - Be very specific about which word is circled and which are crossed out
        
        TASK 2: DOT POINT DELETIONS (Right Box - ACTION TAKEN):
        - Examine EACH individual dot point/bullet point carefully (should be 4 total dot points)
        - Look for ANY handwritten marks that INTERRUPT the text flow:
          * Diagonal lines crossing through text (even if they continue from one dot point to another)
          * X marks over text
          * Crosses through sentences
          * Continuous diagonal lines that may pass through multiple dot points
          * Any handwritten marks that disrupt the sentence
        - CRITICAL: A single diagonal line may start at dot point 2, pass through dot point 3, and end at dot point 4
        - If a continuous diagonal line passes through multiple dot points, ALL affected dot points should be deleted
        - If ANY handwritten mark interrupts a dot point's text, that ENTIRE dot point should be deleted
        - Expected result based on visual inspection: dot points 2, 3, and 4 should have diagonal interruptions
        - Count the dot points and identify which specific ones are interrupted
        - Be very careful to detect subtle diagonal lines or crosses
        
        CRITICAL DETECTION RULES:
        1. For portfolio: ONLY the circled word stays, crossed-out words are deleted
        2. For dot points: ANY interruption by handwritten marks = delete that entire dot point
        3. Look for continuous diagonal lines that may cross multiple dot points
        4. Don't miss diagonal lines that might be subtle
        
        Return detailed JSON:
        {
            "portfolio_selection": {
                "conservative_status": "normal/circled/crossed",
                "balanced_status": "normal/circled/crossed", 
                "growth_status": "normal/circled/crossed",
                "selected_word": "conservative/balanced/growth",
                "deleted_words": ["word1", "word2"],
                "confidence": 0.0-1.0
            },
            "dot_point_analysis": {
                "total_dot_points_found": 0,
                "dot_points_with_interruptions": [
                    {
                        "dot_point_number": 1,
                        "dot_point_text": "beginning of text...",
                        "interruption_type": "diagonal_line/x_mark/cross",
                        "interruption_description": "describe the mark",
                        "should_delete": true/false
                    }
                ],
                "total_deletions_needed": 0,
                "deletion_confidence": 0.0-1.0
            },
            "visual_description": "comprehensive description of all marks and selections seen"
        }
        """
        
        base64_image = self.encode_image(section_image)
        
        payload = {
            "model": "gpt-4.1",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 1500,
            "temperature": 0.1
        }
        
        try:
            response = requests.post(self.api_url, headers=self.headers, json=payload, timeout=90)
            response.raise_for_status()
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            return {
                "raw_analysis": content,
                "success": True
            }
            
        except Exception as e:
            return {
                "raw_analysis": f"Error analyzing Section 1_3: {str(e)}",
                "success": False
            }
    
    def extract_json_from_analysis(self, raw_analysis: str) -> dict:
        """Extract JSON from GPT-4o analysis response"""
        try:
            # Look for JSON in code blocks
            if "```json" in raw_analysis:
                start = raw_analysis.find("```json") + 7
                end = raw_analysis.find("```", start)
                json_str = raw_analysis[start:end].strip()
                return json.loads(json_str)
            
            # Look for JSON without code blocks
            import re
            json_match = re.search(r'\{.*\}', raw_analysis, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
                
            return {}
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            return {}
    
    def find_and_modify_portfolio_text(self, doc: Document, selected_word: str) -> bool:
        """Find and modify the portfolio selection text in the Word document"""
        print(f"   🔍 Searching for portfolio text to replace with: '{selected_word}'")
        
        # Search patterns to look for
        search_patterns = [
            "conservative / balanced / growth",
            "Conservative / Balanced / Growth", 
            "conservative/balanced/growth",
            "Conservative/Balanced/Growth",
            "conservative, balanced, growth",
            "Conservative, Balanced, Growth"
        ]
        
        # First, search in all paragraphs
        for para in doc.paragraphs:
            para_text = para.text
            for pattern in search_patterns:
                if pattern in para_text:
                    print(f"      📍 Found in paragraph: '{para_text[:80]}...'")
                    old_text = para_text
                    new_text = para_text.replace(pattern, selected_word)
                    
                    # Clear and replace the paragraph
                    para.clear()
                    run = para.add_run(new_text)
                    run.font.name = 'Verdana'
                    run.font.size = Pt(9)
                    
                    print(f"      ✅ Replaced '{pattern}' with '{selected_word}'")
                    return True
        
        # If not found in paragraphs, search in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    for pattern in search_patterns:
                        if pattern in cell_text:
                            print(f"      📍 Found in table {table_idx}, row {row_idx}, cell {cell_idx}: '{cell_text[:80]}...'")
                            
                            # Modify each paragraph in the cell
                            for para in cell.paragraphs:
                                if pattern in para.text:
                                    old_text = para.text
                                    new_text = para.text.replace(pattern, selected_word)
                                    
                                    para.clear()
                                    run = para.add_run(new_text)
                                    run.font.name = 'Verdana'
                                    run.font.size = Pt(9)
                                    
                                    print(f"      ✅ Replaced '{pattern}' with '{selected_word}' in table")
                                    return True
        
        print(f"      ❌ Portfolio text pattern not found in document")
        return False
    
    def find_and_delete_action_taken_items(self, doc: Document, deletions_list: list) -> int:
        """Find and delete ENTIRE DOT POINTS (not just text content) based on crossed-out items"""
        print(f"   🗑️ Searching for {len(deletions_list)} items to delete")
        
        deleted_count = 0
        
        # Get the text content of items to delete
        items_to_delete = []
        for deletion_info in deletions_list:
            if deletion_info.get("should_delete", False):
                dot_point_text = deletion_info.get("dot_point_text", "")
                if dot_point_text:
                    items_to_delete.append(dot_point_text.strip())
        
        print(f"      📋 Items to delete:")
        for i, item in enumerate(items_to_delete):
            print(f"         {i+1}. '{item[:60]}...'")
        
        # Search through all paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                # Check if this paragraph contains any of the items to delete
                for item_to_delete in items_to_delete:
                    # Use partial matching to handle formatting differences
                    if self.text_similarity(para_text, item_to_delete) > 0.7:
                        print(f"      ❌ Deleting entire dot point: '{para_text[:60]}...'")
                        self.delete_paragraph(para)
                        deleted_count += 1
                        break
        
        # Search through all table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        para_text = para.text.strip()
                        if para_text:
                            # Check if this paragraph contains any of the items to delete
                            for item_to_delete in items_to_delete:
                                if self.text_similarity(para_text, item_to_delete) > 0.7:
                                    print(f"      ❌ Deleting entire table dot point: '{para_text[:60]}...'")
                                    self.delete_paragraph(para)
                                    deleted_count += 1
                                    break
        
        return deleted_count
    
    def delete_paragraph(self, paragraph):
        """Delete an entire paragraph (dot point) from the document"""
        try:
            # Get the paragraph's parent element
            p = paragraph._element
            # Remove the paragraph element from its parent
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None
            print(f"         ✅ Successfully deleted entire dot point paragraph")
        except Exception as e:
            # Fallback: clear the paragraph content if deletion fails
            print(f"         ⚠️ Fallback to clearing paragraph content: {e}")
            paragraph.clear()
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two text strings"""
        # Simple word-based similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def apply_section_1_3_changes(self, doc: Document, analysis_data: dict) -> list:
        """Apply Section 1_3 changes to the Word document"""
        changes_applied = []
        
        print(f"🔧 Applying Section 1_3 Changes...")
        
        # Apply portfolio selection
        portfolio_selection = analysis_data.get("portfolio_selection", {})
        selected_word = portfolio_selection.get("selected_word")
        
        if selected_word:
            print(f"   🎯 Portfolio Selection: '{selected_word}' selected")
            
            if self.find_and_modify_portfolio_text(doc, selected_word):
                changes_applied.append({
                    "type": "portfolio_selection",
                    "selected_portfolio": selected_word,
                    "old_pattern": "conservative / balanced / growth",
                    "new_text": selected_word
                })
                print(f"   ✅ Portfolio selection applied successfully")
            else:
                print(f"   ❌ Portfolio selection failed - text not found")
        
        # Apply dot point deletions
        dot_point_analysis = analysis_data.get("dot_point_analysis", {})
        dot_points_to_delete = dot_point_analysis.get("dot_points_with_interruptions", [])
        
        if dot_points_to_delete:
            print(f"   🗑️ Processing {len(dot_points_to_delete)} dot point deletions")
            
            deleted_count = self.find_and_delete_action_taken_items(doc, dot_points_to_delete)
            
            if deleted_count > 0:
                changes_applied.append({
                    "type": "dot_point_deletions",
                    "deleted_count": deleted_count,
                    "total_requested": len(dot_points_to_delete)
                })
                print(f"   ✅ Successfully deleted {deleted_count} items")
            else:
                print(f"   ❌ No items were deleted - text not found")
        
        return changes_applied
    
    def run_section_1_3_test(self):
        """Run complete Section 1_3 test"""
        print("🧪 Section 1_3 Individual Test - FIXED IMPLEMENTATION")
        print("=" * 60)
        
        # Step 1: Extract section image
        print("📷 Step 1: Extracting Section 1_3 image...")
        section_image = self.load_extracted_image()
        if not section_image:
            print("❌ Failed to extract Section 1_3 image")
            return
        
        # Save section image for reference
        os.makedirs(self.test_dir, exist_ok=True)
        image_path = os.path.join(self.test_dir, "section_1_3_extracted.png")
        section_image.save(image_path)
        print(f"   💾 Section image saved: {image_path}")
        
        # Step 2: Analyze with GPT-4o
        print("\n🔍 Step 2: Analyzing with enhanced GPT-4o...")
        analysis_result = self.analyze_section_1_3_enhanced(section_image)
        
        if not analysis_result["success"]:
            print("❌ Analysis failed")
            return
        
        # Step 3: Parse analysis results
        print("\n📋 Step 3: Parsing analysis results...")
        analysis_data = self.extract_json_from_analysis(analysis_result["raw_analysis"])
        
        # Save analysis results
        analysis_file = os.path.join(self.test_dir, "section_1_3_analysis.json")
        with open(analysis_file, 'w', encoding='utf-8') as f:
            json.dump({
                "raw_analysis": analysis_result["raw_analysis"],
                "parsed_data": analysis_data
            }, f, indent=2, ensure_ascii=False)
        
        print(f"   💾 Analysis saved: {analysis_file}")
        
        # Step 4: Save analysis results for unified processing
        print("💾 Step 4: Saving analysis results...")
        self.save_analysis_results(analysis_data)
        
        print("📝 Analysis-only mode: Word document processing handled by unified_section_implementations.py")

    def save_analysis_results(self, analysis_data: dict):
        """Save analysis results to JSON file for unified processing"""
        try:
            analysis_file = os.path.join(self.test_dir, "section_1_3_analysis.json")
            analysis_result = {
                "raw_analysis": json.dumps(analysis_data),
                "parsed_data": analysis_data
            }
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_result, f, indent=2, ensure_ascii=False)
            print(f"Analysis saved: {analysis_file}")
        except Exception as e:
            print(f"Error saving analysis: {e}")

    def run_analysis_only(self, progress_callback=None):
        """Run analysis only - called by master processor"""
        try:
            if progress_callback:
                progress_callback("Extracting Section 1_3 image...")
            
            # Load pre-extracted section image  
            section_image_path = self.load_extracted_image()
            if not section_image_path:
                print("Failed to load Section 1_3 pre-extracted image")
                return False
            print(f"Section image loaded: {section_image_path}")
            
            if progress_callback:
                progress_callback("Analyzing Section 1_3 with GPT-4o...")
            
            # Load Image object and analyze with GPT-4o
            section_image = Image.open(section_image_path)
            analysis_result = self.analyze_section_1_3_enhanced(section_image)
            if not analysis_result or "raw_analysis" not in analysis_result:
                print("Failed to get analysis from GPT-4o")
                return False
            
            # Parse analysis
            analysis_data = self.extract_json_from_analysis(analysis_result["raw_analysis"])
            
            # Save analysis results
            analysis_result_full = {
                "raw_analysis": analysis_result["raw_analysis"],
                "parsed_data": analysis_data
            }
            
            analysis_file = os.path.join(self.test_dir, "section_1_3_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_result_full, f, indent=2, ensure_ascii=False)
            
            print("Analysis completed and saved successfully")
            return True
            
        except Exception as e:
            print(f"Error in run_analysis_only: {e}")
            return False

def main():
    """Run Section 1_3 test"""
    import argparse

    parser = argparse.ArgumentParser(description="Section 1_3 Tester")
    parser.add_argument("--pdf-path", help="Path to PDF file for analysis")
    parser.add_argument("--analysis-only", action="store_true",
                       help="Run only analysis (no Word processing)")
    args = parser.parse_args()

    # Create tester with dynamic PDF path
    tester = Section13FixedTester(pdf_path=args.pdf_path)

    if args.analysis_only:
        success = tester.run_analysis_only()
        print(f"Analysis {'completed successfully' if success else 'failed'}")
    else:
        tester.run_test()

if __name__ == "__main__":
    main()

