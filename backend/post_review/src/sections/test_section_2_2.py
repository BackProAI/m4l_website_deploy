#!/usr/bin/env python3
"""

Section 2_2 Individual Testing

# This file requires pre-extracted images from pdf_section_splitter.py
# Run: splitter = PDFSectionSplitter(pdf_path); splitter.extract_all_sections()
# Then this file will work with the pre-extracted section image

Test Section 2_2 Implementation
Two-part section for portfolio selection and sell/purchase additions
"""

import os
import json
import base64
import requests
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt
from datetime import datetime
import re

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from src.utils.section_implementations_reference import SectionImplementationsReference

# This file is designed to work only through the Post-Review UI system
# Run: python post_review_modern_ui.py
# Do not run this file directly

class Section2_2Tester:
    def __init__(self, pdf_path=None):
        self.section_name = "Section_2_2"
        
        # Two-part section - no main test directory needed, uses part directories only
        project_root = os.path.dirname(os.path.abspath(__file__))
        # self.test_dir no longer used - parts save to their own directories
        
        # This test section works with pre-extracted PNG images only
        # PDF sectioning is handled by PDFSectionSplitter
        # Word processing is handled by unified_section_implementations.py
        
        # Get API key
        from dotenv import load_dotenv
        # Load environment variables from project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        env_path = os.path.join(project_root, '.env')
        load_dotenv(env_path)
        self.api_key = os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("OpenAI API key required. Set OPENAI_API_KEY in .env file.")
        
        self.api_url = "https://api.openai.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    
    def load_part1_image(self) -> str:
        """Load part1 specific image"""
        try:
            # Get project root (go up from src/sections/ to project root)
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            image_path = os.path.join(project_root, "data", "test_sections", "section_2_2_part1_test", "section_2_2_part1_extracted.png")
            
            if os.path.exists(image_path):
                print(f"✅ Found Part 1 image: {image_path}")
                # Verify it's a valid image
                try:
                    with Image.open(image_path) as img:
                        print(f"   📏 Part 1 image size: {img.size}")
                        return image_path
                except Exception as e:
                    print(f"❌ Invalid Part 1 image file: {e}")
                    return None
            else:
                print(f"❌ Part 1 image not found: {image_path}")
                return None
                
        except Exception as e:
            print(f"Error loading Part 1 image: {e}")
            return None
    
    def load_part2_image(self) -> str:
        """Load part2 specific image"""
        try:
            # Get project root (go up from src/sections/ to project root)
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            image_path = os.path.join(project_root, "data", "test_sections", "section_2_2_part2_test", "section_2_2_part2_extracted.png")
            
            if os.path.exists(image_path):
                print(f"✅ Found Part 2 image: {image_path}")
                # Verify it's a valid image
                try:
                    with Image.open(image_path) as img:
                        print(f"   📏 Part 2 image size: {img.size}")
                        return image_path
                except Exception as e:
                    print(f"❌ Invalid Part 2 image file: {e}")
                    return None
            else:
                print(f"❌ Part 2 image not found: {image_path}")
                return None
                
        except Exception as e:
            print(f"Error loading Part 2 image: {e}")
            return None

    def load_extracted_image(self) -> str:
        """Load pre-extracted section image (extracted by pdf_section_splitter.py) - LEGACY METHOD"""
        try:
            # Get project root (go up from src/sections/ to project root)
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            
            # Check multiple possible locations for the extracted images
            possible_locations = [
                # New format: separate part directories
                os.path.join(project_root, "data", "test_sections", "section_2_2_part1_test", "section_2_2_part1_extracted.png"),
                os.path.join(project_root, "data", "test_sections", "section_2_2_part2_test", "section_2_2_part2_extracted.png"),
                # No legacy format needed - removed main directory
            ]
            
            for image_path in possible_locations:
                if os.path.exists(image_path):
                    print(f"✅ Found pre-extracted image: {image_path}")
                    
                    # Verify it's a valid image
                    try:
                        with Image.open(image_path) as img:
                            print(f"   📏 Image size: {img.size}")
                            return image_path
                    except Exception as e:
                        print(f"❌ Invalid image file: {e}")
                        continue  # Try next location
            
            print(f"❌ Pre-extracted image not found in any of these locations:")
            for location in possible_locations:
                print(f"   - {location}")
            print(f"   💡 Make sure to run pdf_section_splitter.py first!")
            return None
                
        except Exception as e:
            print(f"Error loading pre-extracted image: {e}")
            return None
    
    def get_section_2_2_part1_analysis_prompt(self) -> str:
        """
        PROMPT for Section 2_2 PART 1 - Portfolio selection + sell/purchase additions only
        """
        return """
        Analyze this Section 2_2 PART 1 for PORTFOLIO SELECTION and SELL/PURCHASE ADDITIONS:
        
        SECTION LAYOUT: This section has TWO BOXES side by side (LEFT BOX and RIGHT BOX)
        
        TASK 1: LEFT BOX ANALYSIS - PORTFOLIO SELECTION:
        - Look for the text "Invest in a conservative / balanced / growth style of asset allocation"
        - Examine the three words: "conservative", "balanced", "growth"
        - CRITICAL DETECTION RULES:
          * TWO words will have lines through them (crossed out)
          * ONE word will be circled (selected)
          * The CIRCLED word should be kept, the CROSSED OUT words should be deleted
        - Identify which word is circled (this one stays)
        - Identify which words are crossed out (these get deleted)
        
        TASK 2: RIGHT BOX ANALYSIS - SELL/PURCHASE ADDITIONS ONLY:
        
        SUB-TASK 2A: SELL DOT POINTS:
        - Look for a section with "Sell…" followed by dot points (should be 3 sub-dot points)
        - Structure should be:
          Sell...
          •
          •  
          •
        - Check if there is handwritten text near or under each "Sell" dot point
        - Extract handwritten text to populate each dot point
        
        SUB-TASK 2B: PURCHASE DOT POINTS:
        - Look for a section with "Purchase…" followed by dot points (should be 2 sub-dot points)
        - Structure should be:
          Purchase...
          •
          •
        - Check if there is handwritten text near or under each "Purchase" dot point
        - Extract handwritten text to populate each dot point
        
        Return detailed JSON:
        {
            "left_box_portfolio_selection": {
                "portfolio_text_found": true/false,
                "conservative_status": "circled/crossed_out/unmarked",
                "balanced_status": "circled/crossed_out/unmarked", 
                "growth_status": "circled/crossed_out/unmarked",
                "selected_word": "conservative/balanced/growth",
                "words_to_delete": ["conservative", "balanced", "growth"]
            },
            "right_box_sell_additions": {
                "sell_section_found": true/false,
                "has_handwritten_text": true/false,
                "handwritten_sell_items": [
                    "handwritten item 1",
                    "handwritten item 2",
                    "handwritten item 3"
                ]
            },
            "right_box_purchase_additions": {
                "purchase_section_found": true/false,
                "has_handwritten_text": true/false,
                "handwritten_purchase_items": [
                    "handwritten item 1",
                    "handwritten item 2"
                ]
            },
            "visual_description": "comprehensive description of all marks and handwriting detected"
        }
        """
    
    def get_section_2_2_part2_analysis_prompt(self) -> str:
        """
        PROMPT for Section 2_2 PART 2 - Time unit selection only
        """
        return """
        Analyze this Section 2_2 PART 2 for TIME UNIT SELECTION:
        
        SECTION LAYOUT: This section has TWO BOXES side by side (LEFT BOX and RIGHT BOX)
        
        TASK 1: LEFT BOX ANALYSIS:
        - The left box should be empty or contain no relevant content for this part
        - Mark as empty/irrelevant
        
        TASK 2: RIGHT BOX ANALYSIS - TIME UNIT SELECTION ONLY:
        
        SUB-TASK 2A: TIME UNIT SELECTION:
        - Look for text "Trades will take approximately ____ days / months to complete."
        - Check for handwritten number above or near the "____" (blank space)
        - Check the words "days" and "months" for markings:
          * CASE 1: One circled + one lined out → Keep circled, delete lined out
          * CASE 2: Only one lined out → Delete lined out, keep the other
          * CASE 3: Only one circled → Keep circled, delete the other
          * CASE 4: Both unmarked → Keep both (no changes)
        - Extract the handwritten number to replace "____"
        
        Return detailed JSON:
        {
            "left_box_content": {
                "has_content": false,
                "description": "Left box is empty for Part 2"
            },
            "right_box_time_selection": {
                "time_text_found": true/false,
                "handwritten_number": "number to replace ____",
                "days_status": "circled/lined_out/crossed_out/unmarked",
                "months_status": "circled/lined_out/crossed_out/unmarked",
                "selection_case": "circled_and_lined/only_lined/only_circled/both_unmarked",
                "selected_time_unit": "days/months/both",
                "time_unit_to_delete": "days/months/none"
            },
            "visual_description": "comprehensive description of all marks and handwriting detected"
        }
        """
    
    def analyze_section_2_2_part1(self, image_path: str) -> dict:
        """Analyze Section 2_2 Part 1 with GPT-4o (Portfolio + Sell/Purchase)"""
        try:
            # Load and encode image
            with open(image_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')
            
            prompt = self.get_section_2_2_part1_analysis_prompt()
        
            payload = {
                "model": "gpt-4.1",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 1500,
                "temperature": 0.1
            }
        
            response = requests.post(self.api_url, headers=self.headers, json=payload, timeout=90)
            response.raise_for_status()
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            return {
                "raw_analysis": content,
                "success": True
            }
            
        except Exception as e:
            return {
                "raw_analysis": f"Error analyzing Section 2_2 Part 1: {str(e)}",
                "success": False
            }
    
    def analyze_section_2_2_part2(self, image_path: str) -> dict:
        """Analyze Section 2_2 Part 2 with GPT-4o (Time selection only)"""
        try:
            # Load and encode image
            with open(image_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')
            
            prompt = self.get_section_2_2_part2_analysis_prompt()
        
            payload = {
                "model": "gpt-4.1",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 1500,
                "temperature": 0.1
            }
        
            response = requests.post(self.api_url, headers=self.headers, json=payload, timeout=90)
            response.raise_for_status()
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            return {
                "raw_analysis": content,
                "success": True
            }
            
        except Exception as e:
            return {
                "raw_analysis": f"Error analyzing Section 2_2 Part 2: {str(e)}",
                "success": False
            }
    
    def extract_json_from_analysis(self, raw_analysis: str) -> dict:
        """Extract JSON from GPT-4o analysis response"""
        try:
            # Look for JSON in code blocks
            if "```json" in raw_analysis:
                start = raw_analysis.find("```json") + 7
                end = raw_analysis.find("```", start)
                json_str = raw_analysis[start:end].strip()
                return json.loads(json_str)
            
            # Look for JSON without code blocks
            import re
            json_match = re.search(r'\{.*\}', raw_analysis, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
                
            return {}
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            return {}
    
    def debug_document_structure(self, doc: Document):
        """Debug the document structure to understand where Section 2_2 content is"""
        print("🔍 DEBUG: Page 2 Document Structure Analysis for Section 2_2")
        print("=" * 50)
        
        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")
        print(f"📋 Total tables: {len(doc.tables)}")
        
        # Show all table content
        print("\n📊 All table content:")
        for table_idx, table in enumerate(doc.tables):
            print(f"   Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"      Row {row_idx}, Cell {cell_idx}: '{cell_text[:80]}...'")
        
        print("=" * 50)
    
    def find_section_2_2_table_row(self, doc: Document) -> tuple:
        """Find the table and row index for Section 2_2 content"""
        # Section 2_2 should contain portfolio selection and sell/purchase content
        # Look for text patterns that match the GPT-4o analysis
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) >= 1 and len(table.columns) >= 2:
                # Check each row for Section 2_2 content
                for row_idx, row in enumerate(table.rows):
                    if len(row.cells) >= 2:
                        left_cell = row.cells[0].text.strip().lower()
                        right_cell = row.cells[1].text.strip().lower()
                        
                        # Look for Section 2_2 indicators
                        section_2_2_keywords = [
                            "invest", "conservative", "balanced", "growth", 
                            "sell", "purchase", "trades", "days", "months", 
                            "asset allocation", "approximately"
                        ]
                        
                        # Check if this row contains Section 2_2 content
                        combined_text = left_cell + " " + right_cell
                        keyword_matches = sum(1 for keyword in section_2_2_keywords if keyword in combined_text)
                        
                        if keyword_matches >= 3:  # At least 3 keywords match
                            print(f"🎯 Found Section 2_2 in Table {table_idx}, Row {row_idx}")
                            print(f"   Left: '{left_cell[:60]}...'")
                            print(f"   Right: '{right_cell[:60]}...'")
                            print(f"   Keyword matches: {keyword_matches}")
                            return table_idx, row_idx
        
        # Fallback: assume first table, second row (Row 1) for Section 2_2
        if len(doc.tables) > 0 and len(doc.tables[0].rows) > 1:
            print(f"🎯 Fallback: Using Table 0, Row 1 for Section 2_2")
            return 0, 1
        
        return None, None
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two text strings"""
        # Simple word-based similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def apply_portfolio_selection(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, portfolio_data: dict) -> bool:
        """Apply portfolio selection changes - same logic as Section 1_3"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                
                selected_word = portfolio_data.get("selected_word", "")
                words_to_delete = portfolio_data.get("words_to_delete", [])
                
                print(f"   📊 Portfolio Selection:")
                print(f"      • Selected (keep): {selected_word}")
                print(f"      • Delete: {words_to_delete}")
                
                if not selected_word:
                    print("      ❌ No selected word found")
                    return False
                
                # Find and update the portfolio text
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if "conservative" in para_text.lower() and "balanced" in para_text.lower() and "growth" in para_text.lower():
                        print(f"      🎯 Found portfolio text: '{para_text[:60]}...'")
                        
                        # Create the new text with only the selected word
                        new_text = para_text
                        
                        # Replace the "conservative / balanced / growth" part with just the selected word
                        portfolio_pattern = r"conservative\s*/\s*balanced\s*/\s*growth"
                        if re.search(portfolio_pattern, new_text, re.IGNORECASE):
                            new_text = re.sub(portfolio_pattern, selected_word, new_text, flags=re.IGNORECASE)
                            
                            # Update the paragraph
                            para.clear()
                            para.add_run(new_text)
                            
                            print(f"      ✅ Updated to: '{new_text[:60]}...'")
                            return True
                
                print("      ❌ Portfolio text not found")
                return False
            return False
        except Exception as e:
            print(f"Error applying portfolio selection: {e}")
            return False
    
    def apply_sell_purchase_additions(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, sell_data: dict, purchase_data: dict) -> int:
        """Apply handwritten additions to sell and purchase dot points"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                changes_applied = 0
                
                # Get handwritten items
                sell_items = sell_data.get("handwritten_sell_items", [])
                purchase_items = purchase_data.get("handwritten_purchase_items", [])
                
                print(f"   📝 Sell/Purchase Additions:")
                print(f"      • Sell items to add: {len(sell_items)}")
                print(f"      • Purchase items to add: {len(purchase_items)}")
                
                paragraphs = list(cell.paragraphs)
                sell_mode = False
                purchase_mode = False
                sell_dot_count = 0
                purchase_dot_count = 0
                
                for para_idx, para in enumerate(paragraphs):
                    para_text = para.text.strip()
                    print(f"      🔍 Para {para_idx}: '{para_text[:50]}...' (sell_mode: {sell_mode}, purchase_mode: {purchase_mode})")
                    
                    if "sell" in para_text.lower() and len(para_text) < 20:  # "Sell" header
                        sell_mode = True
                        purchase_mode = False
                        sell_dot_count = 0
                        print(f"      🎯 Found Sell section at para {para_idx}")
                    elif "purchase" in para_text.lower() and len(para_text) < 20:  # "Purchase" header
                        sell_mode = False
                        purchase_mode = True
                        purchase_dot_count = 0
                        print(f"      🎯 Found Purchase section at para {para_idx}")
                    elif sell_mode and sell_dot_count < len(sell_items):
                        # This is a sell dot point - populate with handwritten text (even if empty)
                        if not ("sell" in para_text.lower() and len(para_text) < 20):  # Skip the header itself
                            sell_text = sell_items[sell_dot_count]
                            para.clear()
                            para.add_run(sell_text)
                            print(f"      ✅ Sell dot {sell_dot_count + 1}: '{sell_text}'")
                            sell_dot_count += 1
                            changes_applied += 1
                    elif purchase_mode and purchase_dot_count < len(purchase_items):
                        # This is a purchase dot point - populate with handwritten text (even if empty)
                        if not ("purchase" in para_text.lower() and len(para_text) < 20):  # Skip the header itself
                            purchase_text = purchase_items[purchase_dot_count]
                            para.clear()
                            para.add_run(purchase_text)
                            print(f"      ✅ Purchase dot {purchase_dot_count + 1}: '{purchase_text}'")
                            purchase_dot_count += 1
                            changes_applied += 1
                
                return changes_applied
            return 0
        except Exception as e:
            print(f"Error applying sell/purchase additions: {e}")
            return 0
    
    def apply_time_unit_selection(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, time_data: dict) -> bool:
        """Apply time unit selection and number replacement with enhanced logic"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                
                handwritten_number = time_data.get("handwritten_number", "")
                days_status = time_data.get("days_status", "unmarked")
                months_status = time_data.get("months_status", "unmarked")
                selection_case = time_data.get("selection_case", "both_unmarked")
                selected_time_unit = time_data.get("selected_time_unit", "both")
                time_unit_to_delete = time_data.get("time_unit_to_delete", "none")
                
                print(f"   ⏰ Time Unit Selection:")
                print(f"      • Number to replace ____: '{handwritten_number}'")
                print(f"      • Days status: {days_status}")
                print(f"      • Months status: {months_status}")
                print(f"      • Selection case: {selection_case}")
                print(f"      • Selected time unit (keep): '{selected_time_unit}'")
                print(f"      • Time unit to delete: '{time_unit_to_delete}'")
                
                if not handwritten_number:
                    print("      ❌ Missing handwritten number")
                    return False
                
                # Find and update the time text
                found_time_text = False
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    print(f"      🔍 Checking paragraph: '{para_text}'")
                    
                    # Look for time text with flexible matching
                    if ("trade" in para_text.lower() and ("approximate" in para_text.lower() or "approx" in para_text.lower())) or \
                       ("____" in para_text and ("days" in para_text.lower() or "months" in para_text.lower())) or \
                       ("days" in para_text.lower() and "months" in para_text.lower() and ("complete" in para_text.lower() or "take" in para_text.lower())):
                        print(f"      🎯 Found time text: '{para_text}'")
                        
                        # Replace ____ with the handwritten number
                        new_text = para_text.replace("____", handwritten_number)
                        
                        # Apply time unit selection based on the case
                        if selection_case == "circled_and_lined":
                            # Case 1: One circled + one lined out → Keep circled, delete lined out
                            if time_unit_to_delete and time_unit_to_delete != "none":
                                time_pattern = rf"\s*/\s*{re.escape(time_unit_to_delete)}|\s*{re.escape(time_unit_to_delete)}\s*/?"
                                new_text = re.sub(time_pattern, "", new_text, flags=re.IGNORECASE)
                                print(f"      🔄 CASE 1: Circled + Lined - Keeping '{selected_time_unit}', Deleting '{time_unit_to_delete}'")
                        
                        elif selection_case == "only_lined":
                            # Case 2: Only one lined out → Delete lined out, keep the other
                            if time_unit_to_delete and time_unit_to_delete != "none":
                                time_pattern = rf"\s*/\s*{re.escape(time_unit_to_delete)}|\s*{re.escape(time_unit_to_delete)}\s*/?"
                                new_text = re.sub(time_pattern, "", new_text, flags=re.IGNORECASE)
                                print(f"      🔄 CASE 2: Only Lined - Deleting '{time_unit_to_delete}', Keeping '{selected_time_unit}'")
                        
                        elif selection_case == "only_circled":
                            # Case 3: Only one circled → Keep circled, delete the other
                            if time_unit_to_delete and time_unit_to_delete != "none":
                                time_pattern = rf"\s*/\s*{re.escape(time_unit_to_delete)}|\s*{re.escape(time_unit_to_delete)}\s*/?"
                                new_text = re.sub(time_pattern, "", new_text, flags=re.IGNORECASE)
                                print(f"      🔄 CASE 3: Only Circled - Keeping '{selected_time_unit}', Deleting '{time_unit_to_delete}'")
                        
                        elif selection_case == "both_unmarked":
                            # Case 4: Both unmarked → Keep both (no changes)
                            print(f"      🔄 CASE 4: Both Unmarked - Keeping both 'days' and 'months'")
                        
                        # Update the paragraph
                        para.clear()
                        para.add_run(new_text)
                        
                        print(f"      ✅ Updated to: '{new_text}'")
                        found_time_text = True
                        return True
                
                if not found_time_text:
                    # Try to find it in the full cell text
                    full_cell_text = cell.text
                    print(f"      🔍 Searching full cell text: '{full_cell_text[:200]}...'")
                    
                    if "____" in full_cell_text or ("days" in full_cell_text.lower() and "months" in full_cell_text.lower()):
                        # Find the paragraph containing this text
                        for para in cell.paragraphs:
                            if "____" in para.text or ("days" in para.text.lower() and "months" in para.text.lower()):
                                para_text = para.text.strip()
                                print(f"      🎯 Found time text in full search: '{para_text}'")
                                
                                # Replace ____ with the handwritten number
                                new_text = para_text.replace("____", handwritten_number)
                                
                                # Apply time unit selection based on the case (same logic as above)
                                if selection_case == "circled_and_lined":
                                    if time_unit_to_delete and time_unit_to_delete != "none":
                                        time_pattern = rf"\s*/\s*{re.escape(time_unit_to_delete)}|\s*{re.escape(time_unit_to_delete)}\s*/?"
                                        new_text = re.sub(time_pattern, "", new_text, flags=re.IGNORECASE)
                                        print(f"      🔄 FALLBACK CASE 1: Circled + Lined - Keeping '{selected_time_unit}', Deleting '{time_unit_to_delete}'")
                                
                                elif selection_case == "only_lined":
                                    if time_unit_to_delete and time_unit_to_delete != "none":
                                        time_pattern = rf"\s*/\s*{re.escape(time_unit_to_delete)}|\s*{re.escape(time_unit_to_delete)}\s*/?"
                                        new_text = re.sub(time_pattern, "", new_text, flags=re.IGNORECASE)
                                        print(f"      🔄 FALLBACK CASE 2: Only Lined - Deleting '{time_unit_to_delete}', Keeping '{selected_time_unit}'")
                                
                                elif selection_case == "only_circled":
                                    if time_unit_to_delete and time_unit_to_delete != "none":
                                        time_pattern = rf"\s*/\s*{re.escape(time_unit_to_delete)}|\s*{re.escape(time_unit_to_delete)}\s*/?"
                                        new_text = re.sub(time_pattern, "", new_text, flags=re.IGNORECASE)
                                        print(f"      🔄 FALLBACK CASE 3: Only Circled - Keeping '{selected_time_unit}', Deleting '{time_unit_to_delete}'")
                                
                                elif selection_case == "both_unmarked":
                                    print(f"      🔄 FALLBACK CASE 4: Both Unmarked - Keeping both 'days' and 'months'")
                                
                                # Update the paragraph
                                para.clear()
                                para.add_run(new_text)
                                
                                print(f"      ✅ Updated to: '{new_text}'")
                                return True
                
                print("      ❌ Time text not found")
                return False
            return False
        except Exception as e:
            print(f"Error applying time unit selection: {e}")
            return False
    
    def apply_section_2_2_changes(self, doc: Document, analysis_data: dict) -> list:
        """Apply Section 2_2 changes to the Word document"""
        changes_applied = []
        
        print(f"🔧 Applying Section 2_2 Changes...")
        
        # Debug document structure first
        self.debug_document_structure(doc)
        
        # Find Section 2_2 table and row
        table_idx, row_idx = self.find_section_2_2_table_row(doc)
        
        if table_idx is None or row_idx is None:
            print("❌ Could not find Section 2_2 table row")
            return changes_applied
        
        print(f"🎯 Found Section 2_2 in Table {table_idx}, Row {row_idx}")
        
        # Get analysis data
        portfolio_data = analysis_data.get("left_box_portfolio_selection", {})
        sell_data = analysis_data.get("right_box_sell_additions", {})
        purchase_data = analysis_data.get("right_box_purchase_additions", {})
        time_data = analysis_data.get("right_box_time_selection", {})
        
        # Apply left box portfolio selection
        if portfolio_data.get("portfolio_text_found", False):
            print(f"\n📦 LEFT BOX: Portfolio Selection")
            success = self.apply_portfolio_selection(doc, table_idx, row_idx, 0, portfolio_data)
            
            if success:
                changes_applied.append({
                    "type": "portfolio_selection",
                    "section": "Section_2_2",
                    "selected_word": portfolio_data.get("selected_word", ""),
                    "deleted_words": portfolio_data.get("words_to_delete", [])
                })
                print(f"      ✅ Portfolio selection applied")
        
        # Apply right box sell/purchase additions
        if sell_data.get("has_handwritten_text", False) or purchase_data.get("has_handwritten_text", False):
            print(f"\n📦 RIGHT BOX: Sell/Purchase Additions")
            additions_count = self.apply_sell_purchase_additions(doc, table_idx, row_idx, 1, sell_data, purchase_data)
            
            if additions_count > 0:
                changes_applied.append({
                    "type": "sell_purchase_additions",
                    "section": "Section_2_2",
                    "additions_applied": additions_count,
                    "sell_items": len(sell_data.get("handwritten_sell_items", [])),
                    "purchase_items": len(purchase_data.get("handwritten_purchase_items", []))
                })
                print(f"      ✅ {additions_count} sell/purchase additions applied")
        
        # Apply right box time unit selection
        if time_data.get("time_text_found", False):
            print(f"\n📦 RIGHT BOX: Time Unit Selection")
            success = self.apply_time_unit_selection(doc, table_idx, row_idx, 1, time_data)
            
            if success:
                changes_applied.append({
                    "type": "time_unit_selection",
                    "section": "Section_2_2",
                    "number_replacement": time_data.get("handwritten_number", ""),
                    "selected_unit": time_data.get("selected_time_unit", ""),
                    "deleted_unit": time_data.get("time_unit_to_delete", ""),
                    "selection_case": time_data.get("selection_case", "unknown"),
                    "days_status": time_data.get("days_status", "unmarked"),
                    "months_status": time_data.get("months_status", "unmarked")
                })
                print(f"      ✅ Time unit selection applied")
        
        return changes_applied
    
    def save_analysis_results(self, analysis_data: dict):
        """Save analysis results to JSON file for unified processing - LEGACY METHOD"""
        try:
            # Note: This method is legacy - two-part sections save to part directories
            print("⚠️ save_analysis_results is legacy for two-part sections")
            return  # Skip saving to main directory
        except Exception as e:
            print(f"❌ Error saving analysis: {e}")
    
    def run_analysis_only(self, progress_callback=None):
        """Run only image extraction and analysis (no Word processing) - now handles two parts"""
        print(f"Starting Section_2_2 Analysis - TWO-PART FORMAT")
        print("Analysis only - Word processing handled by unified processor")
        print("=" * 80)
        
        # Step 1A: Load Part 1 image and analyze
        if progress_callback:
            progress_callback(1, 4, self.section_name, "Loading Part 1 image...")
        
        print(f"📊 Step 1A: Loading Part 1 image...")
        part1_image_path = self.load_part1_image()
        if not part1_image_path:
            print(f"❌ Section_2_2 Part 1 Analysis Failed!")
            return False
        print(f"✅ Part 1 image loaded: {part1_image_path}")
        
        # Step 2A: Run Part 1 analysis
        if progress_callback:
            progress_callback(2, 4, self.section_name, "Analyzing Part 1 with GPT-4o...")
        
        print(f"🤖 Step 2A: Analyzing Section_2_2 Part 1 with GPT-4o...")
        part1_result = self.analyze_section_2_2_part1(part1_image_path)
        
        if not part1_result["success"]:
            print(f"❌ Part 1 GPT-4o analysis failed: {part1_result.get('raw_analysis', 'Unknown error')}")
            return False
        
        part1_data = self.extract_json_from_analysis(part1_result["raw_analysis"])
        
        # Save Part 1 results to part1 directory
        part1_analysis = {
            "section_name": "Section_2_2_Part1",
            "raw_analysis": part1_result["raw_analysis"],
            "parsed_data": part1_data,
            "timestamp": datetime.now().isoformat()
        }
        
        # Get project root (go up from src/sections/ to project root)
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        part1_dir = os.path.join(project_root, "data", "test_sections", "section_2_2_part1_test")
        os.makedirs(part1_dir, exist_ok=True)
        part1_file = os.path.join(part1_dir, "section_2_2_part1_analysis.json")
        with open(part1_file, 'w', encoding='utf-8') as f:
            json.dump(part1_analysis, f, indent=2, ensure_ascii=False)
        print(f"💾 Part 1 analysis saved: {part1_file}")
        
        # Step 1B: Load Part 2 image
        print(f"📊 Step 1B: Loading Part 2 image...")
        part2_image_path = self.load_part2_image()
        if not part2_image_path:
            print(f"❌ Section_2_2 Part 2 Analysis Failed!")
            return False
        print(f"✅ Part 2 image loaded: {part2_image_path}")
        
        # Step 2B: Run Part 2 analysis
        if progress_callback:
            progress_callback(3, 4, self.section_name, "Analyzing Part 2 with GPT-4o...")
        
        print(f"🤖 Step 2B: Analyzing Section_2_2 Part 2 with GPT-4o...")
        part2_result = self.analyze_section_2_2_part2(part2_image_path)
        
        if not part2_result["success"]:
            print(f"❌ Part 2 GPT-4o analysis failed: {part2_result.get('raw_analysis', 'Unknown error')}")
            return False
        
        part2_data = self.extract_json_from_analysis(part2_result["raw_analysis"])
        
        # Save Part 2 results to part2 directory
        part2_analysis = {
            "section_name": "Section_2_2_Part2",
            "raw_analysis": part2_result["raw_analysis"],
            "parsed_data": part2_data,
            "timestamp": datetime.now().isoformat()
        }
        
        part2_dir = os.path.join(project_root, "data", "test_sections", "section_2_2_part2_test")
        os.makedirs(part2_dir, exist_ok=True)
        part2_file = os.path.join(part2_dir, "section_2_2_part2_analysis.json")
        with open(part2_file, 'w', encoding='utf-8') as f:
            json.dump(part2_analysis, f, indent=2, ensure_ascii=False)
        print(f"💾 Part 2 analysis saved: {part2_file}")
        
        if progress_callback:
            progress_callback(4, 4, self.section_name, "Analysis complete...")
        
        print(f"✅ Section_2_2 Two-Part Analysis Complete!")
        print(f"📄 Saved: section_2_2_part1_analysis.json")
        print(f"📄 Saved: section_2_2_part2_analysis.json")
        print("💼 Ready for unified Word processing (parts will be combined)")
        
        return True
    
    def run_section_2_2_test(self):
        """Run complete Section 2_2 test - now split into Part 1 and Part 2"""
        print("🧪 Section 2_2 Individual Test - Page 2, Section 2 (SPLIT INTO 2 PARTS)")
        print("=" * 70)
        
        # Step 1: Extract section image
        print("📷 Step 1: Extracting Section 2_2 image from Page 2...")
        section_image_path = self.load_extracted_image()
        if not section_image_path:
            print("❌ Failed to extract Section 2_2 image")
            return
        
        print(f"   💾 Using part directories for image storage")
        
        # Step 2A: Analyze PART 1 with GPT-4o (Portfolio + Sell/Purchase)
        print("\n🔍 Step 2A: Analyzing PART 1 with GPT-4o (Portfolio + Sell/Purchase)...")
        part1_result = self.analyze_section_2_2_part1(section_image_path)
        
        if not part1_result["success"]:
            print("❌ Part 1 Analysis failed")
            return
        
        part1_data = self.extract_json_from_analysis(part1_result["raw_analysis"])
        
        # Save Part 1 analysis results to part1 directory
        # Get project root (go up from src/sections/ to project root)
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        part1_dir = os.path.join(project_root, "data", "test_sections", "section_2_2_part1_test")
        os.makedirs(part1_dir, exist_ok=True)
        part1_file = os.path.join(part1_dir, "section_2_2_part1_analysis.json")
        with open(part1_file, 'w', encoding='utf-8') as f:
            json.dump({
                "raw_analysis": part1_result["raw_analysis"],
                "parsed_data": part1_data
            }, f, indent=2, ensure_ascii=False)
        print(f"   💾 Part 1 analysis saved: {part1_file}")
        
        # Step 2B: Analyze PART 2 with GPT-4o (Time selection only)
        print("\n🔍 Step 2B: Analyzing PART 2 with GPT-4o (Time selection only)...")
        part2_result = self.analyze_section_2_2_part2(section_image_path)
        
        if not part2_result["success"]:
            print("❌ Part 2 Analysis failed")
            return
        
        part2_data = self.extract_json_from_analysis(part2_result["raw_analysis"])
        
        # Save Part 2 analysis results to part2 directory
        part2_dir = os.path.join(project_root, "data", "test_sections", "section_2_2_part2_test")
        os.makedirs(part2_dir, exist_ok=True)
        part2_file = os.path.join(part2_dir, "section_2_2_part2_analysis.json")
        with open(part2_file, 'w', encoding='utf-8') as f:
            json.dump({
                "raw_analysis": part2_result["raw_analysis"],
                "parsed_data": part2_data
            }, f, indent=2, ensure_ascii=False)
        print(f"   💾 Part 2 analysis saved: {part2_file}")
        
        # Step 3: Display analysis results for both parts
        print("\n📊 Step 3: Analysis Results:")
        
        print(f"\n   🎯 PART 1 RESULTS:")
        portfolio = part1_data.get("left_box_portfolio_selection", {})
        sell = part1_data.get("right_box_sell_additions", {})
        purchase = part1_data.get("right_box_purchase_additions", {})
        
        print(f"      📦 Left Box - Portfolio Selection:")
        print(f"         • Portfolio text found: {portfolio.get('portfolio_text_found', False)}")
        print(f"         • Selected word (keep): {portfolio.get('selected_word', 'None')}")
        print(f"         • Words to delete: {portfolio.get('words_to_delete', [])}")
        
        print(f"      📦 Right Box - Sell Additions:")
        print(f"         • Sell section found: {sell.get('sell_section_found', False)}")
        print(f"         • Has handwritten text: {sell.get('has_handwritten_text', False)}")
        if sell.get('handwritten_sell_items'):
            for i, item in enumerate(sell.get('handwritten_sell_items', [])):
                print(f"            Sell {i+1}: '{item}'")
        
        print(f"      📦 Right Box - Purchase Additions:")
        print(f"         • Purchase section found: {purchase.get('purchase_section_found', False)}")
        print(f"         • Has handwritten text: {purchase.get('has_handwritten_text', False)}")
        if purchase.get('handwritten_purchase_items'):
            for i, item in enumerate(purchase.get('handwritten_purchase_items', [])):
                print(f"            Purchase {i+1}: '{item}'")
        
        print(f"\n   🎯 PART 2 RESULTS:")
        left_box = part2_data.get("left_box_content", {})
        time = part2_data.get("right_box_time_selection", {})
        
        print(f"      📦 Left Box - Content:")
        print(f"         • Has content: {left_box.get('has_content', False)}")
        print(f"         • Description: {left_box.get('description', 'None')}")
        
        print(f"      📦 Right Box - Time Unit Selection:")
        print(f"         • Time text found: {time.get('time_text_found', False)}")
        print(f"         • Handwritten number: '{time.get('handwritten_number', 'None')}'")
        print(f"         • Days status: {time.get('days_status', 'unmarked')}")
        print(f"         • Months status: {time.get('months_status', 'unmarked')}")
        print(f"         • Selection case: {time.get('selection_case', 'both_unmarked')}")
        print(f"         • Selected time unit (keep): '{time.get('selected_time_unit', 'None')}'")
        print(f"         • Time unit to delete: '{time.get('time_unit_to_delete', 'None')}'")
        
        print("\n📝 Analysis-only mode: Word document processing handled by unified_section_implementations.py")
        print("✅ Both Part 1 and Part 2 analysis completed successfully!")
    
    def run_test(self):
        """Run the complete Section 2_2 test"""
        print("🚀 Starting Section 2_2 Test (Portfolio selection and additions)")
        print("📄 Test section: Section 2_2 - PNG image analysis only")
        print("🔍 Section: Section_2_2")
        print("=" * 80)
        
        # Run the two-part test
        self.run_section_2_2_test()

def main():
    """Run Section 2_2 test"""
    import argparse

    parser = argparse.ArgumentParser(description="Section 2_2 Tester")
    parser.add_argument("--pdf-path", help="Path to PDF file for analysis")
    parser.add_argument("--analysis-only", action="store_true",
                       help="Run only analysis (no Word processing)")
    args = parser.parse_args()

    # Create tester with dynamic PDF path
    tester = Section2_2Tester(pdf_path=args.pdf_path)

    if args.analysis_only:
        success = tester.run_analysis_only()
        print(f"Analysis {'completed successfully' if success else 'failed'}")
    else:
        tester.run_section_2_2_test()

if __name__ == "__main__":
    main()
