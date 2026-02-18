#!/usr/bin/env python3
"""
Section 1_4 Individual Testing - IMPROVED IMPLEMENTATION
Tests bottom boxes with diagonal/cross detection and row deletion rule
ACTUALLY DELETES THE CONTENT
"""

# This file requires pre-extracted images from pdf_section_splitter.py
# Run: splitter = PDFSectionSplitter(pdf_path); splitter.extract_all_sections()
# Then this file will work with the pre-extracted section image

import os
import json
import base64
import requests
from pathlib import Path
from PIL import Image

from docx import Document
from docx.shared import Pt
from datetime import datetime



# This file is designed to work only through the Post-Review UI system
# Run: python post_review_modern_ui.py
# Do not run this file directly

class Section14ImprovedTester:
    def __init__(self, pdf_path=None):
        # This test section works with pre-extracted PNG images only
        # PDF sectioning is handled by PDFSectionSplitter
        # Word processing is handled by unified_section_implementations.py
        self.section_name = "Section_1_4"  # Add missing section name
        
        # Make test directory relative to the project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.test_dir = os.path.join(project_root, "data", "test_sections", "section_1_4_test")
        
        # Get API key
        from dotenv import load_dotenv
        # Load environment variables from project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        env_path = os.path.join(project_root, '.env')
        load_dotenv(env_path)
        self.api_key = os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("OpenAI API key required. Set OPENAI_API_KEY in .env file.")
        
        self.api_url = "https://api.openai.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    

    
    def load_extracted_image(self) -> str:
        """Load pre-extracted section image (extracted by pdf_section_splitter.py)"""
        try:
            # Ensure test directory exists
            os.makedirs(self.test_dir, exist_ok=True)
            
            # Check for pre-extracted image
            image_filename = "section_1_4_extracted.png"
            image_path = os.path.join(self.test_dir, image_filename)
            
            if os.path.exists(image_path):
                print(f"✅ Found pre-extracted image: {image_path}")
                return image_path
            else:
                print(f"❌ Pre-extracted image not found: {image_path}")
                print(f"   💡 Make sure to run pdf_section_splitter.py first!")
                return None
            
        except Exception as e:
            print(f"Error loading pre-extracted image: {e}")
            return None
    

    
    def encode_image(self, image: Image.Image) -> str:
        """Encode PIL image to base64"""
        import io
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    def analyze_section_1_4_enhanced(self, section_image: Image.Image) -> dict:
        """Enhanced analysis for Section 1_4 with left/right box detection and row deletion rule"""
        
        prompt = """
        Analyze this Section 1_4 for CRITICAL BOX PROCESSING WITH ALL MARK TYPES:
        
        SECTION LAYOUT: This section has TWO BOXES side by side (LEFT BOX and RIGHT BOX)
        
        CRITICAL MARK DETECTION RULES:
        1. **DIAGONAL LINES/CROSSES**: Delete affected sentence(s) completely
        2. **HORIZONTAL STRIKETHROUGHS**: Delete affected sentence(s) completely  
        3. **HORIZONTAL ARROWS**: Replace text (find original text and replacement text)
        4. **ROW DELETION RULE**: If BOTH boxes have marks through ALL sentences → delete entire row
        
        TASK 1: LEFT BOX ANALYSIS:
        - Examine the LEFT box for ALL types of handwritten marks:
          * Diagonal lines crossing through text → DELETE sentence
          * X marks over text → DELETE sentence
          * Crosses through sentences → DELETE sentence
          * Horizontal strikethrough lines → DELETE sentence
          * Horizontal arrows pointing to text → REPLACE text
          * Any other handwritten marks that disrupt sentences → DELETE sentence
        - For each mark, identify the specific sentence affected
        - For arrows: identify BOTH the original text and the replacement text
        
        TASK 2: RIGHT BOX ANALYSIS:
        - Examine the RIGHT box for ALL types of handwritten marks:
          * Diagonal lines crossing through text → DELETE sentence
          * X marks over text → DELETE sentence
          * Crosses through sentences → DELETE sentence
          * Horizontal strikethrough lines → DELETE sentence
          * Horizontal arrows pointing to text → REPLACE text
          * Any other handwritten marks that disrupt sentences → DELETE sentence
        - For each mark, identify the specific sentence affected
        - For arrows: identify BOTH the original text and the replacement text
        
        TASK 3: ROW DELETION RULE:
        - CRITICAL: If BOTH the left box AND right box have deletion marks (diagonal/cross/strikethrough) through ALL their sentences
        - Then the ENTIRE ROW should be deleted (both boxes completely removed)
        - This rule applies to deletion marks only, NOT replacement arrows
        
        DETECTION RULES:
        1. Diagonal lines/crosses/strikethroughs = DELETE affected sentence
        2. Horizontal arrows = REPLACE text (original → replacement)
        3. If both boxes have deletion marks through ALL sentences = DELETE entire row
        4. Look for subtle marks that may span multiple sentences
        5. Distinguish between deletion marks (diagonal/cross/strikethrough) and replacement arrows (horizontal)
        
        Return detailed JSON:
        {
            "left_box_analysis": {
                "has_interruptions": true/false,
                "sentences_to_delete": [
                    {
                        "sentence_number": 1,
                        "sentence_text": "sentence content...",
                        "mark_type": "diagonal_line/x_mark/cross/strikethrough",
                        "mark_description": "describe the deletion mark",
                        "action": "delete"
                    }
                ],
                "sentences_to_replace": [
                    {
                        "sentence_number": 2,
                        "original_text": "original sentence text...",
                        "replacement_text": "handwritten replacement...",
                        "mark_type": "horizontal_arrow",
                        "mark_description": "describe the arrow and replacement",
                        "action": "replace"
                    }
                ],
                "total_sentences": 0,
                "deletion_count": 0,
                "replacement_count": 0,
                "all_sentences_have_deletion_marks": true/false
            },
            "right_box_analysis": {
                "has_interruptions": true/false,
                "sentences_to_delete": [
                    {
                        "sentence_number": 1,
                        "sentence_text": "sentence content...",
                        "mark_type": "diagonal_line/x_mark/cross/strikethrough",
                        "mark_description": "describe the deletion mark",
                        "action": "delete"
                    }
                ],
                "sentences_to_replace": [
                    {
                        "sentence_number": 2,
                        "original_text": "original sentence text...",
                        "replacement_text": "handwritten replacement...",
                        "mark_type": "horizontal_arrow",
                        "mark_description": "describe the arrow and replacement",
                        "action": "replace"
                    }
                ],
                "total_sentences": 0,
                "deletion_count": 0,
                "replacement_count": 0,
                "all_sentences_have_deletion_marks": true/false
            },
            "row_deletion_rule": {
                "left_box_all_deletion_marks": true/false,
                "right_box_all_deletion_marks": true/false,
                "delete_entire_row": true/false,
                "deletion_reason": "both boxes have deletion marks (diagonal/cross/strikethrough) through all sentences"
            },
            "visual_description": "comprehensive description of all marks detected in both boxes, specifying mark types and locations"
        }
        """
        
        base64_image = self.encode_image(section_image)
        
        payload = {
            "model": "gpt-4.1",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 1500,
            "temperature": 0.1
        }
        
        try:
            response = requests.post(self.api_url, headers=self.headers, json=payload, timeout=90)
            response.raise_for_status()
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            return {
                "raw_analysis": content,
                "success": True
            }
            
        except Exception as e:
            return {
                "raw_analysis": f"Error analyzing Section 1_4: {str(e)}",
                "success": False
            }
    
    def analyze_with_gpt4o(self, image_path: str) -> dict:
        """Analyze Section 1_4 with GPT-4o for table row modifications"""
        try:
            # Read and encode the image
            with open(image_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            # Section 1_4 specific prompt for table analysis
            prompt = """
You are analyzing Section 1_4 of a document that contains a table with rows that can be modified.

Look for:
1. Diagonal lines or crosses through table rows (indicating row deletion)
2. Strikethrough text (line through text, indicating text deletion)
3. Handwritten text or arrows (indicating replacements or additions)
4. Any other handwritten marks or modifications

Return your analysis as JSON with this structure:
{
    "has_deletion_marks": boolean,
    "has_replacement_marks": boolean,
    "has_additions": boolean,
    "has_any_handwriting": boolean,
    "row_modifications": [
        {
            "row_number": int,
            "modification_type": "deletion|replacement|addition",
            "description": "description of what to do",
            "should_delete_row": boolean,
            "replacement_text": "text if replacement"
        }
    ],
    "explanation": "detailed explanation of what you see and what actions to take"
}
"""
            
            # Make API call using requests
            payload = {
                "model": "gpt-4.1",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 1000
            }
            
            response = requests.post(self.api_url, headers=self.headers, json=payload)
            response.raise_for_status()
            result = response.json()
            
            return {
                "success": True,
                "raw_analysis": result["choices"][0]["message"]["content"]
            }
            
        except Exception as e:
            print(f"GPT-4o analysis error: {e}")
            return {"success": False, "error": str(e)}

    def extract_json_from_analysis(self, raw_analysis: str) -> dict:
        """Extract JSON from GPT-4o analysis response"""
        try:
            # Look for JSON in code blocks
            if "```json" in raw_analysis:
                start = raw_analysis.find("```json") + 7
                end = raw_analysis.find("```", start)
                json_str = raw_analysis[start:end].strip()
                return json.loads(json_str)
            
            # Look for JSON without code blocks
            import re
            json_match = re.search(r'\{.*\}', raw_analysis, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
                
            return {}
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            return {}
    
    def debug_document_structure(self, doc: Document):
        """Debug the document structure to understand where Section 1_4 content is"""
        print("🔍 DEBUG: Document Structure Analysis")
        print("=" * 50)
        
        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")
        print(f"📋 Total tables: {len(doc.tables)}")
        
        # Show all paragraphs with content
        print("\n📝 All paragraphs with content:")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"   Para {i}: '{para.text[:80]}...'")
        
        # Show all table content
        print("\n📊 All table content:")
        for table_idx, table in enumerate(doc.tables):
            print(f"   Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"      Row {row_idx}, Cell {cell_idx}: '{cell_text[:60]}...'")
        
        print("=" * 50)
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two text strings"""
        # Simple word-based similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def find_text_with_keywords(self, text: str, keywords: list) -> bool:
        """Check if text contains any of the keywords"""
        text_lower = text.lower()
        for keyword in keywords:
            if keyword.lower() in text_lower:
                return True
        return False
    
    def comprehensive_text_search_and_delete(self, doc: Document, sentences_to_delete: list) -> int:
        """Comprehensive search and delete using multiple strategies"""
        deleted_count = 0
        
        print(f"🔍 COMPREHENSIVE SEARCH for {len(sentences_to_delete)} sentences:")
        for sentence in sentences_to_delete:
            print(f"   Target: '{sentence[:60]}...'")
        
        # Strategy 1: Exact text matching
        print("\n📝 Strategy 1: Searching paragraphs...")
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if para_text:
                for sentence in sentences_to_delete:
                    # Try exact match first
                    if sentence.lower() in para_text.lower():
                        print(f"   ✅ FOUND EXACT in Para {i}: '{para_text[:60]}...'")
                        para.clear()
                        deleted_count += 1
                        break
                    # Try similarity match
                    elif self.text_similarity(para_text, sentence) > 0.6:
                        print(f"   ✅ FOUND SIMILAR in Para {i}: '{para_text[:60]}...' (similarity: {self.text_similarity(para_text, sentence):.2f})")
                        para.clear()
                        deleted_count += 1
                        break
        
        # Strategy 2: Table cell search
        print("\n📊 Strategy 2: Searching table cells...")
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        para_text = para.text.strip()
                        if para_text:
                            for sentence in sentences_to_delete:
                                # Try exact match first
                                if sentence.lower() in para_text.lower():
                                    print(f"   ✅ FOUND EXACT in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: '{para_text[:60]}...'")
                                    para.clear()
                                    deleted_count += 1
                                    break
                                # Try similarity match
                                elif self.text_similarity(para_text, sentence) > 0.6:
                                    print(f"   ✅ FOUND SIMILAR in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: '{para_text[:60]}...' (similarity: {self.text_similarity(para_text, sentence):.2f})")
                                    para.clear()
                                    deleted_count += 1
                                    break
        
        # Strategy 3: Keyword-based search (fallback)
        if deleted_count == 0:
            print("\n🔍 Strategy 3: Keyword-based search...")
            
            # Extract keywords from sentences to delete
            all_keywords = []
            for sentence in sentences_to_delete:
                words = sentence.split()
                # Get important words (longer than 3 characters)
                keywords = [word.strip('.,!?') for word in words if len(word.strip('.,!?')) > 3]
                all_keywords.extend(keywords)
            
            print(f"   Keywords: {all_keywords}")
            
            # Search for content with these keywords
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text and self.find_text_with_keywords(para_text, all_keywords):
                    print(f"   ✅ FOUND KEYWORDS in Para {i}: '{para_text[:60]}...'")
                    para.clear()
                    deleted_count += 1
            
            # Search table cells for keywords
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            para_text = para.text.strip()
                            if para_text and self.find_text_with_keywords(para_text, all_keywords):
                                print(f"   ✅ FOUND KEYWORDS in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: '{para_text[:60]}...'")
                                para.clear()
                                deleted_count += 1
        
        return deleted_count
    
    def find_section_1_4_table_row(self, doc: Document) -> tuple:
        """Find the table and row index for Section 1_4 content"""
        # Section 1_4 content should be in the bottom row of the main content table
        # Look for the table that contains "ITEMS DISCUSSED" and "ACTION TAKEN"
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) >= 3 and len(table.columns) >= 2:  # Should have at least 3 rows, 2 columns
                # Check if this is the main content table
                header_row = table.rows[0]
                if len(header_row.cells) >= 2:
                    left_header = header_row.cells[0].text.strip().upper()
                    right_header = header_row.cells[1].text.strip().upper()
                    
                    if "ITEMS DISCUSSED" in left_header and "ACTION TAKEN" in right_header:
                        # This is the main content table, Section 1_4 should be in the last row (row 2)
                        if len(table.rows) >= 3:
                            return table_idx, 2  # Row 2 is Section 1_4
        
        return None, None
    
    def delete_table_row(self, doc: Document, table_idx: int, row_idx: int) -> bool:
        """Delete an entire table row"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows):
                # Get the row element
                row = table.rows[row_idx]
                # Remove the row from the table
                table._tbl.remove(row._tr)
                return True
            return False
        except Exception as e:
            print(f"Error deleting table row: {e}")
            return False
    
    def delete_cell_content_only(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int) -> bool:
        """Delete content from a specific table cell only"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                # Clear all paragraphs in the cell
                for para in cell.paragraphs:
                    para.clear()
                return True
            return False
        except Exception as e:
            print(f"Error deleting cell content: {e}")
            return False

    def apply_section_1_4_changes(self, doc: Document, analysis_data: dict) -> list:
        """Apply Section 1_4 changes: deletions, replacements, and row deletion rule"""
        changes_applied = []
        
        print(f"🔧 Applying Section 1_4 Changes - ALL MARK TYPES (DELETE, REPLACE, ROW DELETION)...")
        
        # Debug document structure first
        self.debug_document_structure(doc)
        
        # Find Section 1_4 table and row
        table_idx, row_idx = self.find_section_1_4_table_row(doc)
        
        if table_idx is None or row_idx is None:
            print("❌ Could not find Section 1_4 table row")
            return changes_applied
        
        print(f"🎯 Found Section 1_4 in Table {table_idx}, Row {row_idx}")
        
        # Get analysis data
        left_box_analysis = analysis_data.get("left_box_analysis", {})
        right_box_analysis = analysis_data.get("right_box_analysis", {})
        row_deletion_rule = analysis_data.get("row_deletion_rule", {})
        
        left_box_all_deletion_marks = row_deletion_rule.get("left_box_all_deletion_marks", False)
        right_box_all_deletion_marks = row_deletion_rule.get("right_box_all_deletion_marks", False)
        delete_entire_row = row_deletion_rule.get("delete_entire_row", False)
        
        print(f"📊 Analysis Results:")
        print(f"   • Left box all deletion marks: {left_box_all_deletion_marks}")
        print(f"   • Right box all deletion marks: {right_box_all_deletion_marks}")
        print(f"   • Delete entire row: {delete_entire_row}")
        
        if delete_entire_row and left_box_all_deletion_marks and right_box_all_deletion_marks:
            # CASE 1: Both boxes have deletion marks through ALL sentences -> DELETE ENTIRE TABLE ROW
            print(f"\n🚨 CASE 1: COMPLETE ROW DELETION")
            print(f"   Both left and right boxes have deletion marks through all sentences")
            print(f"   -> DELETING ENTIRE TABLE ROW {row_idx}")
            
            success = self.delete_table_row(doc, table_idx, row_idx)
            
            if success:
                changes_applied.append({
                    "type": "complete_table_row_deletion",
                    "section": "Section_1_4",
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "deletion_reason": "Both boxes have deletion marks through all sentences",
                    "left_box_all_deletion_marks": left_box_all_deletion_marks,
                    "right_box_all_deletion_marks": right_box_all_deletion_marks
                })
                print(f"✅ COMPLETE TABLE ROW DELETED SUCCESSFULLY")
            else:
                print(f"❌ FAILED TO DELETE TABLE ROW")
        
        else:
            # CASE 2: Individual sentence processing (deletions and replacements)
            print(f"\n📦 CASE 2: INDIVIDUAL SENTENCE PROCESSING")
            print(f"   Processing individual deletions and replacements within boxes")
            
            # Process left box (Cell 0)
            print(f"\n🔍 LEFT BOX (Cell 0) PROCESSING:")
            left_deletions = left_box_analysis.get("sentences_to_delete", [])
            left_replacements = left_box_analysis.get("sentences_to_replace", [])
            
            if left_deletions:
                print(f"   🗑️ Found {len(left_deletions)} sentences to delete in left box")
                for deletion in left_deletions:
                    sentence_text = deletion.get("sentence_text", "")
                    mark_type = deletion.get("mark_type", "unknown")
                    print(f"      - DELETE: '{sentence_text[:50]}...' (mark: {mark_type})")
                
                # Apply deletions
                deleted_count = self.apply_sentence_deletions(doc, table_idx, row_idx, 0, left_deletions)
                if deleted_count > 0:
                    changes_applied.append({
                        "type": "left_box_sentence_deletions",
                        "deleted_count": deleted_count,
                        "total_requested": len(left_deletions)
                    })
            
            if left_replacements:
                print(f"   🔄 Found {len(left_replacements)} sentences to replace in left box")
                for replacement in left_replacements:
                    original_text = replacement.get("original_text", "")
                    replacement_text = replacement.get("replacement_text", "")
                    print(f"      - REPLACE: '{original_text[:30]}...' → '{replacement_text}'")
                
                # Apply replacements
                replaced_count = self.apply_sentence_replacements(doc, table_idx, row_idx, 0, left_replacements)
                if replaced_count > 0:
                    changes_applied.append({
                        "type": "left_box_sentence_replacements",
                        "replaced_count": replaced_count,
                        "total_requested": len(left_replacements)
                    })
            
            # Process right box (Cell 1)
            print(f"\n🔍 RIGHT BOX (Cell 1) PROCESSING:")
            right_deletions = right_box_analysis.get("sentences_to_delete", [])
            right_replacements = right_box_analysis.get("sentences_to_replace", [])
            
            if right_deletions:
                print(f"   🗑️ Found {len(right_deletions)} sentences to delete in right box")
                for deletion in right_deletions:
                    sentence_text = deletion.get("sentence_text", "")
                    mark_type = deletion.get("mark_type", "unknown")
                    print(f"      - DELETE: '{sentence_text[:50]}...' (mark: {mark_type})")
                
                # Apply deletions
                deleted_count = self.apply_sentence_deletions(doc, table_idx, row_idx, 1, right_deletions)
                if deleted_count > 0:
                    changes_applied.append({
                        "type": "right_box_sentence_deletions",
                        "deleted_count": deleted_count,
                        "total_requested": len(right_deletions)
                    })
            
            if right_replacements:
                print(f"   🔄 Found {len(right_replacements)} sentences to replace in right box")
                for replacement in right_replacements:
                    original_text = replacement.get("original_text", "")
                    replacement_text = replacement.get("replacement_text", "")
                    print(f"      - REPLACE: '{original_text[:30]}...' → '{replacement_text}'")
                
                # Apply replacements
                replaced_count = self.apply_sentence_replacements(doc, table_idx, row_idx, 1, right_replacements)
                if replaced_count > 0:
                    changes_applied.append({
                        "type": "right_box_sentence_replacements",
                        "replaced_count": replaced_count,
                        "total_requested": len(right_replacements)
                    })
        
        return changes_applied
    
    def apply_sentence_deletions(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, deletions: list) -> int:
        """Apply sentence deletions to a specific table cell"""
        deleted_count = 0
        
        try:
            table = doc.tables[table_idx]
            cell = table.rows[row_idx].cells[cell_idx]
            
            # Get all sentences to delete
            sentences_to_delete = []
            for deletion in deletions:
                sentence_text = deletion.get("sentence_text", "").strip()
                if sentence_text:
                    sentences_to_delete.append(sentence_text)
            
            print(f"      🔍 Searching cell for {len(sentences_to_delete)} sentences to delete:")
            
            # Search through all paragraphs in the cell
            for para in cell.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    for sentence in sentences_to_delete:
                        # Try exact match first
                        if sentence.lower() in para_text.lower():
                            print(f"         ✅ FOUND & DELETED: '{para_text[:50]}...'")
                            para.clear()
                            deleted_count += 1
                            break
                        # Try similarity match
                        elif self.text_similarity(para_text, sentence) > 0.6:
                            print(f"         ✅ FOUND SIMILAR & DELETED: '{para_text[:50]}...' (sim: {self.text_similarity(para_text, sentence):.2f})")
                            para.clear()
                            deleted_count += 1
                            break
            
        except Exception as e:
            print(f"         ❌ Error applying sentence deletions: {e}")
        
        return deleted_count
    
    def apply_sentence_replacements(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, replacements: list) -> int:
        """Apply sentence replacements to a specific table cell"""
        replaced_count = 0
        
        try:
            table = doc.tables[table_idx]
            cell = table.rows[row_idx].cells[cell_idx]
            
            print(f"      🔍 Searching cell for {len(replacements)} sentences to replace:")
            
            # Apply each replacement
            for replacement in replacements:
                original_text = replacement.get("original_text", "").strip()
                replacement_text = replacement.get("replacement_text", "").strip()
                
                if not original_text or not replacement_text:
                    continue
                
                # Search through all paragraphs in the cell
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        # Try exact match first
                        if original_text.lower() in para_text.lower():
                            # Replace the text
                            new_text = para_text.replace(original_text, replacement_text)
                            para.clear()
                            para.add_run(new_text)
                            print(f"         ✅ REPLACED: '{original_text}' → '{replacement_text}'")
                            replaced_count += 1
                            break
                        # Try similarity match
                        elif self.text_similarity(para_text, original_text) > 0.7:
                            # Replace entire paragraph with replacement text
                            para.clear()
                            para.add_run(replacement_text)
                            print(f"         ✅ SIMILAR MATCH REPLACED: '{para_text[:30]}...' → '{replacement_text}'")
                            replaced_count += 1
                            break
                        
        except Exception as e:
            print(f"         ❌ Error applying sentence replacements: {e}")
        
        return replaced_count
    
    def save_analysis_results(self, analysis_data: dict):
        """Save analysis results to JSON file for unified processing"""
        try:
            # Ensure directory exists
            os.makedirs(self.test_dir, exist_ok=True)
            analysis_file = os.path.join(self.test_dir, "section_1_4_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_data, f, indent=2, ensure_ascii=False)
            print(f"Analysis saved: {analysis_file}")
        except Exception as e:
            print(f"Error saving analysis: {e}")
    
    def run_analysis_only(self, progress_callback=None):
        """Run only image extraction and analysis (no Word processing)"""
        print("Starting Section_1_4 Analysis")
        print("Analysis only - Word processing handled by unified processor")
        print("=" * 80)
        
        if progress_callback:
            progress_callback(1, 3, self.section_name, "Extracting section image...")
        
        # Step 1: Load pre-extracted section image
        print("Step 1: Loading pre-extracted Section_1_4 image...")
        section_image_path = self.load_extracted_image()
        if not section_image_path:
            print("Section_1_4 Analysis Failed!")
            return False
        print(f"Section image loaded: {section_image_path}")
        
        if progress_callback:
            progress_callback(2, 3, self.section_name, "Analyzing with GPT-4o...")
        
        # Step 2: Analyze with GPT-4o
        print("Step 2: Analyzing Section_1_4 with GPT-4o...")
        analysis_result = self.analyze_with_gpt4o(section_image_path)
        
        if not analysis_result.get("success"):
            print("GPT-4o analysis failed")
            return False
        
        print("GPT-4o analysis completed")
        
        # Step 3: Parse and save results
        analysis_data = self.extract_json_from_analysis(analysis_result.get("raw_analysis", ""))
        if not analysis_data:
            print("Could not parse GPT-4o analysis")
            return False
        
        if progress_callback:
            progress_callback(3, 3, self.section_name, "Saving analysis results...")
        
        # Save analysis results for unified processing (include raw_analysis for master processor)
        analysis_result_full = {
            "raw_analysis": analysis_result.get("raw_analysis", ""),
            "parsed_data": analysis_data
        }
        self.save_analysis_results(analysis_result_full)
        
        print("Section_1_4 Analysis Complete!")
        return True

    def run_section_1_4_test(self):
        """Run complete Section 1_4 test with improved implementation"""
        print("🧪 Section 1_4 Individual Test - IMPROVED IMPLEMENTATION")
        print("=" * 70)
        
        # Step 1: Load pre-extracted section image
        print("📷 Step 1: Loading pre-extracted Section 1_4 image...")
        section_image_path = self.load_extracted_image()
        if not section_image_path:
            print("❌ Failed to load pre-extracted Section 1_4 image")
            return
        
        # Load the image
        section_image = Image.open(section_image_path)
        print(f"   💾 Section image loaded: {section_image_path}")
        
        # Step 2: Analyze with GPT-4o (reuse existing analysis if available)
        analysis_file = os.path.join(self.test_dir, "section_1_4_analysis.json")
        if os.path.exists(analysis_file):
            print(f"\n🔍 Step 2: Loading existing analysis from {analysis_file}...")
            with open(analysis_file, 'r', encoding='utf-8') as f:
                analysis_result = json.load(f)
                analysis_data = analysis_result.get('parsed_data', {})
        else:
            print("\n🔍 Step 2: Analyzing with enhanced GPT-4o...")
            analysis_result = self.analyze_section_1_4_enhanced(section_image)
            
            if not analysis_result["success"]:
                print("❌ Analysis failed")
                return
            
            analysis_data = self.extract_json_from_analysis(analysis_result["raw_analysis"])
            
            # Save analysis results
            analysis_file = os.path.join(self.test_dir, "section_1_4_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "raw_analysis": analysis_result["raw_analysis"],
                    "parsed_data": analysis_data
                }, f, indent=2, ensure_ascii=False)
        
        # Step 3: Apply changes to test document
        print("\n🔧 Step 3: Applying changes to test document with IMPROVED IMPLEMENTATION...")
        
        if os.path.exists(self.test_word_path):
            # Create a copy for processing
            doc = Document(self.test_word_path)
            changes_applied = self.apply_section_1_4_changes(doc, analysis_data)
            
            # Save processed document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(self.test_dir, f"section_1_4_processed_{timestamp}.docx")
            doc.save(output_path)
            
            print(f"\n💾 Processed document saved: {output_path}")
            print(f"🔧 Changes applied: {len(changes_applied)}")
            
            for change in changes_applied:
                change_type = change.get('type', 'unknown')
                if change_type == 'row_deletion':
                    sentences_deleted = change.get('sentences_deleted', 0)
                    total_sentences = change.get('total_sentences', 0)
                    print(f"   🚨 Row Deletion: {sentences_deleted}/{total_sentences} sentences deleted")
                elif 'deletions' in change_type:
                    print(f"   ✅ {change_type}: {change.get('deleted_count', 0)} sentences deleted")
        
        print("\n" + "=" * 70)
        print("🎉 Section 1_4 IMPROVED Test Complete!")
        print(f"Check the {self.test_dir} folder for results")

def main():
    """Run Section 1_4 test"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Section 1_4 Tester")
    parser.add_argument("--pdf-path", help="Path to PDF file for analysis")
    parser.add_argument("--analysis-only", action="store_true", 
                       help="Run only analysis (no Word processing)")
    args = parser.parse_args()
    
    # Create tester with dynamic PDF path
    tester = Section14ImprovedTester(pdf_path=args.pdf_path)
    
    if args.analysis_only:
        success = tester.run_analysis_only()
        print(f"Analysis {'completed successfully' if success else 'failed'}")
    else:
        tester.run_test()

if __name__ == "__main__":
    main()

