#!/usr/bin/env python3
"""
Section 2_1 Individual Testing
Tests Page 2, Section 1 with same rules as Section 1_4 but with more dot points
"""

# This file requires pre-extracted images from pdf_section_splitter.py
# Run: splitter = PDFSectionSplitter(pdf_path); splitter.extract_all_sections()
# Then this file will work with the pre-extracted section image

import os
import json
import base64
import requests
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt
from datetime import datetime



# This file is designed to work only through the Post-Review UI system
# Run: python post_review_modern_ui.py
# Do not run this file directly

class Section21Tester:
    def __init__(self, pdf_path=None):
        # This test section works with pre-extracted PNG images only
        # PDF sectioning is handled by PDFSectionSplitter
        # Word processing is handled by unified_section_implementations.py
        
        # Make test directory relative to the project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.test_dir = os.path.join(project_root, "data", "test_sections", "section_2_1_test")
        
        # Get API key
        from dotenv import load_dotenv
        # Load environment variables from project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        env_path = os.path.join(project_root, '.env')
        load_dotenv(env_path)
        self.api_key = os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("OpenAI API key required. Set OPENAI_API_KEY in .env file.")
        
        self.api_url = "https://api.openai.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    

    
    def load_extracted_image(self) -> Image.Image:
        """Load pre-extracted section image (extracted by pdf_section_splitter.py)"""
        try:
            # Ensure test directory exists
            os.makedirs(self.test_dir, exist_ok=True)
            
            # Check for pre-extracted image
            image_filename = "section_2_1_extracted.png"
            image_path = os.path.join(self.test_dir, image_filename)
            
            if os.path.exists(image_path):
                print(f"✅ Found pre-extracted image: {image_path}")
                
                # Load and return the image
                section_image = Image.open(image_path)
                print(f"   📏 Image size: {section_image.size}")
                return section_image
            else:
                print(f"❌ Pre-extracted image not found: {image_path}")
                print(f"   💡 Make sure to run pdf_section_splitter.py first!")
                return None
                
        except Exception as e:
            print(f"Error loading pre-extracted image: {e}")
            return None
    
    def encode_image(self, image: Image.Image) -> str:
        """Encode PIL image to base64"""
        import io
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    def get_section_2_1_analysis_prompt(self) -> str:
        """
        WORKING PROMPT for Section 2_1 - Same rules as Section 1_4 but for Page 2
        """
        return """
        Analyze this Section 2_1 for HANDWRITTEN MARKS - Both DELETIONS and ADDITIONS:
        
        SECTION LAYOUT: This section has TWO BOXES side by side (LEFT BOX and RIGHT BOX)
        
        CRITICAL DISTINCTION:
        1. Handwritten marks THROUGH/OVER text = DELETE the sentence/dot point
        2. Handwritten text AFTER/AROUND text (not crossing through) = APPEND after full stop
        
        TASK 1: LEFT BOX ANALYSIS:
        - Examine the LEFT box for handwritten marks
        - DELETION MARKS (should_delete: true):
          * Diagonal lines crossing THROUGH text
          * X marks OVER text  
          * Crosses THROUGH sentences/dot points
          * Any handwritten marks that physically cross/interrupt the printed text
        - ADDITION MARKS (should_delete: false):
          * Handwritten text AFTER the sentence (after the full stop)
          * Handwritten notes AROUND text but not crossing through it
          * Handwritten annotations that supplement rather than replace text
        - Count how many sentences/dot points are affected in the LEFT box
        
        TASK 2: RIGHT BOX ANALYSIS - MULTIPLE DOT POINTS:
        - Examine the RIGHT box for handwritten marks
        - CRITICAL: This box has MULTIPLE DOT POINTS - analyze each one separately  
        - DELETION MARKS (should_delete: true):
          * Diagonal lines crossing THROUGH text
          * X marks OVER text
          * Crosses THROUGH sentences/dot points
          * Any handwritten marks that physically cross/interrupt the printed text
        - ADDITION MARKS (should_delete: false):
          * Handwritten text AFTER the sentence (after the full stop)
          * Handwritten notes AROUND text but not crossing through it
          * Handwritten annotations that supplement rather than replace text
        - CONTINUOUS DIAGONAL LINE DETECTION (CRITICAL):
          * TRACE THE FULL PATH: A single diagonal line may START at one dot point, PASS THROUGH multiple others, and END at another
          * FOLLOW THE ENTIRE LINE: Look carefully at where the line begins and where it ends - don't miss any dot points it touches
          * ALL TOUCHED = ALL DELETED: If a continuous diagonal line touches/passes through multiple dot points, ALL affected dot points should be deleted
          * SPECIFIC EXAMPLE: Line starts at dot point 2, passes through dot point 3, continues through dot point 4, ends at dot point 5 → DELETE dot points 2, 3, 4, AND 5
          * VERIFICATION STEP: After identifying a diagonal line, trace its COMPLETE path from start to finish to ensure no dot points are missed
          * COMMON MISTAKE: Don't stop analyzing after finding the first few items - continue tracing the line to its endpoint
        - Count each affected dot point separately and verify the total count
        
        TASK 3: ROW DELETION RULE:
        - CRITICAL: If BOTH the left box AND right box have handwritten diagonal lines or crosses through ALL their sentences/dot points
        - Then the ENTIRE ROW should be deleted (both boxes completely removed)
        - This is different from individual sentence/dot point deletions
        - Look for complete box coverage by handwritten marks
        
        DETECTION RULES:
        1. DELETION CASE: Lines/crosses THROUGH text = delete that specific sentence/dot point
        2. ADDITION CASE: Handwriting AFTER text (not crossing through) = keep sentence + append handwriting
        3. If both boxes have marks through ALL sentences/dot points = delete entire row
        4. Look for subtle diagonal lines, X marks, crosses
        5. Don't miss continuous lines that may pass through multiple dot points
        6. Analyze each dot point in the right box individually
        
        EXAMPLES:
        - "Cross through 'financial year'" → should_delete: true, interruption_type: "cross"
        - "Handwritten text after sentence: 'We can of approx 8-5k top to us'" → should_delete: false, interruption_type: "handwritten_notes"
        - "Diagonal line crossing through dot point" → should_delete: true, interruption_type: "diagonal_line"
        - "Handwritten notes around text but not crossing through" → should_delete: false, interruption_type: "handwritten_notes"
        
        Return detailed JSON:
        {
            "left_box_analysis": {
                "has_interruptions": true/false,
                "interrupted_items": [
                    {
                        "item_number": 1,
                        "item_text": "sentence/dot point content...",
                        "interruption_type": "cross/diagonal_line/x_mark/handwritten_notes",
                        "interruption_description": "SPECIFY LOCATION: 'Cross through text' OR 'Handwritten text after sentence'",
                        "should_delete": true/false
                    }
                ],
                "total_items": 0,
                "interrupted_count": 0,
                "all_items_interrupted": true/false
            },
            "right_box_analysis": {
                "has_interruptions": true/false,
                "interrupted_items": [
                    {
                        "item_number": 1,
                        "item_text": "dot point content...",
                        "interruption_type": "cross/diagonal_line/x_mark/handwritten_notes", 
                        "interruption_description": "SPECIFY LOCATION: 'Cross through text' OR 'Handwritten text after sentence'",
                        "should_delete": true/false,
                        "continuous_line_affected": true/false
                    }
                ],
                "total_items": 0,
                "interrupted_count": 0,
                "all_items_interrupted": true/false,
                "continuous_line_detected": true/false,
                "continuous_line_description": "describe the continuous line and which dot points it affects"
            },
            "row_deletion_rule": {
                "left_box_completely_marked": true/false,
                "right_box_completely_marked": true/false,
                "delete_entire_row": true/false,
                "deletion_reason": "both boxes have marks through all sentences/dot points"
            },
            "visual_description": "comprehensive description of all marks detected in both boxes"
        }
        """
    
    def analyze_section_2_1(self, section_image: Image.Image) -> dict:
        """Analyze Section 2_1 with GPT-4o"""
        
        prompt = self.get_section_2_1_analysis_prompt()
        base64_image = self.encode_image(section_image)
        
        payload = {
            "model": "gpt-4o",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            "max_tokens": 1500,
            "temperature": 0.1
        }
        
        try:
            response = requests.post(self.api_url, headers=self.headers, json=payload, timeout=90)
            response.raise_for_status()
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            return {
                "raw_analysis": content,
                "success": True
            }
            
        except Exception as e:
            return {
                "raw_analysis": f"Error analyzing Section 2_1: {str(e)}",
                "success": False
            }
    
    def extract_json_from_analysis(self, raw_analysis: str) -> dict:
        """Extract JSON from GPT-4o analysis response"""
        try:
            # Look for JSON in code blocks
            if "```json" in raw_analysis:
                start = raw_analysis.find("```json") + 7
                end = raw_analysis.find("```", start)
                json_str = raw_analysis[start:end].strip()
                return json.loads(json_str)
            
            # Look for JSON without code blocks
            import re
            json_match = re.search(r'\{.*\}', raw_analysis, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
                
            return {}
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            return {}
    
    def debug_document_structure(self, doc: Document):
        """Debug the document structure to understand where Section 2_1 content is"""
        print("🔍 DEBUG: Page 2 Document Structure Analysis")
        print("=" * 50)
        
        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")
        print(f"📋 Total tables: {len(doc.tables)}")
        
        # Show all paragraphs with content
        print("\n📝 All paragraphs with content:")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"   Para {i}: '{para.text[:80]}...'")
        
        # Show all table content
        print("\n📊 All table content:")
        for table_idx, table in enumerate(doc.tables):
            print(f"   Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"      Row {row_idx}, Cell {cell_idx}: '{cell_text[:60]}...'")
        
        print("=" * 50)
    
    def find_section_2_1_table_row(self, doc: Document) -> tuple:
        """Find the table and row index for Section 2_1 content"""
        # Section 2_1 should contain superannuation contribution content
        # Look for text patterns that match the GPT-4o analysis
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) >= 1 and len(table.columns) >= 2:
                # Check each row for superannuation/contribution content
                for row_idx, row in enumerate(table.rows):
                    if len(row.cells) >= 2:
                        left_cell = row.cells[0].text.strip().lower()
                        right_cell = row.cells[1].text.strip().lower()
                        
                        # Look for Section 2_1 indicators - superannuation contributions
                        section_2_1_keywords = [
                            "maximise", "superannuation", "contribution", "$30,000", 
                            "concessional", "$120,000", "$360,000", "$300,000", 
                            "downsizer", "non-concessional"
                        ]
                        
                        # Check if this row contains Section 2_1 content
                        combined_text = left_cell + " " + right_cell
                        keyword_matches = sum(1 for keyword in section_2_1_keywords if keyword in combined_text)
                        
                        if keyword_matches >= 2:  # At least 2 keywords match
                            print(f"🎯 Found Section 2_1 in Table {table_idx}, Row {row_idx}")
                            print(f"   Left: '{left_cell[:60]}...'")
                            print(f"   Right: '{right_cell[:60]}...'")
                            print(f"   Keyword matches: {keyword_matches}")
                            return table_idx, row_idx
        
        # Fallback: assume first table, first row (Row 0)
        if len(doc.tables) > 0 and len(doc.tables[0].rows) > 0:
            print(f"🎯 Fallback: Using Table 0, Row 0 for Section 2_1")
            return 0, 0
        
        return None, None
    
    def delete_table_row(self, doc: Document, table_idx: int, row_idx: int) -> bool:
        """Delete an entire table row"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows):
                # Get the row element
                row = table.rows[row_idx]
                # Remove the row from the table
                table._tbl.remove(row._tr)
                return True
            return False
        except Exception as e:
            print(f"Error deleting table row: {e}")
            return False
    
    def delete_specific_dot_points(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, dot_points_to_delete: list) -> int:
        """Delete specific dot points from a table cell - both content AND bullet point structure"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                deleted_count = 0
                
                print(f"   🔍 Analyzing cell paragraphs for dot point deletion...")
                print(f"   🎯 Target dot points to delete (only these should be deleted):")
                
                # Create list of ONLY the dot points that should be deleted
                dots_to_delete = []
                for dot_point in dot_points_to_delete:
                    if dot_point.get("should_delete", False):
                        dot_number = dot_point.get("item_number", "?")
                        dot_text = dot_point.get("item_text", "")
                        dots_to_delete.append({
                            "number": dot_number,
                            "text": dot_text,
                            "matched": False  # Track if we've already matched this dot point
                        })
                        print(f"      Dot {dot_number}: '{dot_text[:60]}...'")
                
                print(f"   📊 Should delete {len(dots_to_delete)} out of {len(dot_points_to_delete)} total dot points")
                
                # Get all paragraphs in the cell
                paragraphs = list(cell.paragraphs)
                paragraphs_to_remove = []  # Track paragraphs to remove
                
                for para_idx, para in enumerate(paragraphs):
                    para_text = para.text.strip()
                    if para_text:
                        print(f"      Para {para_idx}: '{para_text[:60]}...'")
                        
                        # Check if this paragraph matches any UNMATCHED dot point to delete
                        best_match = None
                        best_similarity = 0
                        best_match_type = ""
                        
                        for i, dot_to_delete in enumerate(dots_to_delete):
                            if dot_to_delete["matched"]:  # Skip already matched dot points
                                continue
                                
                            dot_text = dot_to_delete["text"]
                            dot_number = dot_to_delete["number"]
                            
                            # Strategy 1: Direct text similarity (high threshold for accuracy)
                            similarity = self.text_similarity(para_text, dot_text)
                            if similarity > 0.7:  # Higher threshold to avoid false matches
                                if similarity > best_similarity:
                                    best_match = i
                                    best_similarity = similarity
                                    best_match_type = f"SIMILARITY ({similarity:.2f})"
                            
                            # Strategy 2: Financial amount matching (very specific)
                            financial_amounts = ["120,000", "360,000", "300,000"]
                            for amount in financial_amounts:
                                if amount in dot_text and amount in para_text:
                                    # This is a very specific match
                                    best_match = i
                                    best_similarity = 1.0
                                    best_match_type = f"FINANCIAL AMOUNT (${amount})"
                                    break
                            
                            # Strategy 3: Specific superannuation term matching
                            if "non-concessional" in dot_text.lower() and "non-concessional" in para_text.lower():
                                if similarity > 0.4:  # Additional similarity check
                                    best_match = i
                                    best_similarity = similarity
                                    best_match_type = f"NON-CONCESSIONAL ({similarity:.2f})"
                            elif "downsizer" in dot_text.lower() and "downsizer" in para_text.lower():
                                if similarity > 0.4:  # Additional similarity check
                                    best_match = i
                                    best_similarity = similarity
                                    best_match_type = f"DOWNSIZER ({similarity:.2f})"
                        
                        # If we found a match, mark it for deletion
                        if best_match is not None:
                            dots_to_delete[best_match]["matched"] = True
                            dot_number = dots_to_delete[best_match]["number"]
                            print(f"         ✅ MATCH FOUND - Dot {dot_number}: {best_match_type}")
                            print(f"            Will delete: '{para_text[:50]}...'")
                            paragraphs_to_remove.append(para)
                            deleted_count += 1
                        else:
                            print(f"         ⚪ NO MATCH - Keeping: '{para_text[:50]}...'")
                
                # Remove the matched paragraphs (this removes both content AND bullet structure)
                for para in paragraphs_to_remove:
                    # Remove the paragraph element from the cell
                    para._element.getparent().remove(para._element)
                
                print(f"   📊 Final Deletion Summary:")
                print(f"      • Requested deletions: {len(dots_to_delete)}")
                print(f"      • Actual deletions: {deleted_count}")
                print(f"      • Remaining paragraphs: {len(paragraphs) - deleted_count}")
                
                # Show which dot points were matched
                for dot in dots_to_delete:
                    status = "✅ DELETED" if dot["matched"] else "❌ NOT FOUND"
                    print(f"      • Dot {dot['number']}: {status}")
                
                return deleted_count
            return 0
        except Exception as e:
            print(f"Error deleting specific dot points: {e}")
            return 0
    
    def delete_cell_content_only(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int) -> bool:
        """Delete all content from a specific table cell"""
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                # Clear all paragraphs in the cell
                for para in cell.paragraphs:
                    para.clear()
                return True
            return False
        except Exception as e:
            print(f"Error deleting cell content: {e}")
            return False
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two text strings"""
        # Simple word-based similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def apply_section_2_1_changes(self, doc: Document, analysis_data: dict) -> list:
        """Apply Section 2_1 changes to the Word document"""
        changes_applied = []
        
        print(f"🔧 Applying Section 2_1 Changes...")
        
        # Debug document structure first
        self.debug_document_structure(doc)
        
        # Find Section 2_1 table and row
        table_idx, row_idx = self.find_section_2_1_table_row(doc)
        
        if table_idx is None or row_idx is None:
            print("❌ Could not find Section 2_1 table row")
            return changes_applied
        
        print(f"🎯 Found Section 2_1 in Table {table_idx}, Row {row_idx}")
        
        # Get analysis data
        left_box_analysis = analysis_data.get("left_box_analysis", {})
        right_box_analysis = analysis_data.get("right_box_analysis", {})
        row_deletion_rule = analysis_data.get("row_deletion_rule", {})
        
        left_box_marked = row_deletion_rule.get("left_box_completely_marked", False)
        right_box_marked = row_deletion_rule.get("right_box_completely_marked", False)
        delete_entire_row = row_deletion_rule.get("delete_entire_row", False)
        
        print(f"📊 Analysis Results:")
        print(f"   • Left box marked: {left_box_marked}")
        print(f"   • Right box marked: {right_box_marked}")
        print(f"   • Delete entire row: {delete_entire_row}")
        
        if delete_entire_row and left_box_marked and right_box_marked:
            # CASE 1: Both boxes have diagonal/crosses -> DELETE ENTIRE TABLE ROW
            print(f"\n🚨 CASE 1: COMPLETE ROW DELETION")
            print(f"   Both left and right boxes have diagonal lines/crosses through all items")
            print(f"   -> DELETING ENTIRE TABLE ROW {row_idx}")
            
            success = self.delete_table_row(doc, table_idx, row_idx)
            
            if success:
                changes_applied.append({
                    "type": "complete_table_row_deletion",
                    "section": "Section_2_1",
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "deletion_reason": "Both boxes have diagonal/crosses through all items",
                    "left_box_marked": left_box_marked,
                    "right_box_marked": right_box_marked
                })
                print(f"✅ COMPLETE TABLE ROW DELETED SUCCESSFULLY")
            else:
                print(f"❌ FAILED TO DELETE TABLE ROW")
        
        else:
            # CASE 2: Individual box/dot point deletions
            print(f"\n📦 CASE 2: INDIVIDUAL ITEM PROCESSING")
            print(f"   Only some boxes/items have diagonal/crosses -> DELETE SPECIFIC ITEMS")
            
            # Process left box
            if left_box_analysis.get("has_interruptions", False):
                left_items = left_box_analysis.get("interrupted_items", [])
                items_to_delete = [item for item in left_items if item.get("should_delete", False)]
                
                if items_to_delete:
                    print(f"\n🔍 LEFT BOX: Deleting {len(items_to_delete)} interrupted items")
                    deleted_count = self.delete_specific_dot_points(doc, table_idx, row_idx, 0, items_to_delete)
                    
                    if deleted_count > 0:
                        changes_applied.append({
                            "type": "left_box_item_deletions",
                            "section": "Section_2_1",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "cell_index": 0,
                            "deleted_count": deleted_count,
                            "total_requested": len(items_to_delete)
                        })
                        print(f"      ✅ Left box: Successfully deleted {deleted_count} items")
            
            # Process right box (with multiple dot points)
            if right_box_analysis.get("has_interruptions", False):
                right_items = right_box_analysis.get("interrupted_items", [])
                items_to_delete = [item for item in right_items if item.get("should_delete", False)]
                
                if items_to_delete:
                    print(f"\n🔍 RIGHT BOX: Deleting {len(items_to_delete)} interrupted dot points")
                    
                    # Show continuous line detection
                    if right_box_analysis.get("continuous_line_detected", False):
                        continuous_desc = right_box_analysis.get("continuous_line_description", "")
                        print(f"      🔗 CONTINUOUS LINE DETECTED: {continuous_desc}")
                    
                    deleted_count = self.delete_specific_dot_points(doc, table_idx, row_idx, 1, items_to_delete)
                    
                    if deleted_count > 0:
                        changes_applied.append({
                            "type": "right_box_dot_point_deletions",
                            "section": "Section_2_1",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "cell_index": 1,
                            "deleted_count": deleted_count,
                            "total_requested": len(items_to_delete),
                            "continuous_line": right_box_analysis.get("continuous_line_detected", False)
                        })
                        print(f"      ✅ Right box: Successfully deleted {deleted_count} dot points")
        
        return changes_applied
    
    def run_section_2_1_test(self):
        """Run complete Section 2_1 test"""
        print("🧪 Section 2_1 Individual Test - Page 2, Section 1")
        print("=" * 70)
        
        # Step 1: Extract section image
        print("📷 Step 1: Extracting Section 2_1 image from Page 2...")
        section_image = self.load_extracted_image()
        if not section_image:
            print("❌ Failed to extract Section 2_1 image")
            return
        
        # Save section image for reference
        os.makedirs(self.test_dir, exist_ok=True)
        image_path = os.path.join(self.test_dir, "section_2_1_extracted.png")
        section_image.save(image_path)
        print(f"   💾 Section image saved: {image_path}")
        
        # Step 2: Analyze with GPT-4o
        print("\n🔍 Step 2: Analyzing with GPT-4o...")
        analysis_result = self.analyze_section_2_1(section_image)
        
        if not analysis_result["success"]:
            print("❌ Analysis failed")
            return
        
        analysis_data = self.extract_json_from_analysis(analysis_result["raw_analysis"])
        
        # Save analysis results
        analysis_file = os.path.join(self.test_dir, "section_2_1_analysis.json")
        with open(analysis_file, 'w', encoding='utf-8') as f:
            json.dump({
                "raw_analysis": analysis_result["raw_analysis"],
                "parsed_data": analysis_data
            }, f, indent=2, ensure_ascii=False)
        
        # Step 3: Display analysis results
        print("\n📊 Step 3: Analysis Results:")
        
        left_box = analysis_data.get("left_box_analysis", {})
        right_box = analysis_data.get("right_box_analysis", {})
        row_rule = analysis_data.get("row_deletion_rule", {})
        
        print(f"   📦 Left Box Analysis:")
        print(f"      • Has interruptions: {left_box.get('has_interruptions', False)}")
        print(f"      • Total items: {left_box.get('total_items', 0)}")
        print(f"      • Interrupted count: {left_box.get('interrupted_count', 0)}")
        print(f"      • All items interrupted: {left_box.get('all_items_interrupted', False)}")
        
        left_items = left_box.get("interrupted_items", [])
        for item in left_items:
            if item.get("should_delete", False):
                print(f"         Item {item.get('item_number', '?')}: DELETE - {item.get('interruption_type', 'unknown')}")
        
        print(f"   📦 Right Box Analysis:")
        print(f"      • Has interruptions: {right_box.get('has_interruptions', False)}")
        print(f"      • Total items: {right_box.get('total_items', 0)}")
        print(f"      • Interrupted count: {right_box.get('interrupted_count', 0)}")
        print(f"      • All items interrupted: {right_box.get('all_items_interrupted', False)}")
        print(f"      • Continuous line detected: {right_box.get('continuous_line_detected', False)}")
        
        if right_box.get('continuous_line_detected', False):
            print(f"      • Continuous line description: {right_box.get('continuous_line_description', '')}")
        
        right_items = right_box.get("interrupted_items", [])
        for item in right_items:
            if item.get("should_delete", False):
                continuous = " (continuous line)" if item.get("continuous_line_affected", False) else ""
                print(f"         Dot Point {item.get('item_number', '?')}: DELETE - {item.get('interruption_type', 'unknown')}{continuous}")
        
        print(f"   🚨 Row Deletion Rule:")
        print(f"      • Left box completely marked: {row_rule.get('left_box_completely_marked', False)}")
        print(f"      • Right box completely marked: {row_rule.get('right_box_completely_marked', False)}")
        print(f"      • Delete entire row: {row_rule.get('delete_entire_row', False)}")
        
        if row_rule.get('delete_entire_row', False):
            print(f"      🚨 ENTIRE ROW WILL BE DELETED!")
        
        # Step 4: Save analysis results for unified processing
        print("💾 Step 4: Saving analysis results...")
        self.save_analysis_results(analysis_data)
        
        print("📝 Analysis-only mode: Word document processing handled by unified_section_implementations.py")

    def save_analysis_results(self, analysis_data: dict):
        """Save analysis results to JSON file for unified processing"""
        try:
            analysis_file = os.path.join(self.test_dir, "section_2_1_analysis.json")
            analysis_result = {
                "raw_analysis": json.dumps(analysis_data),
                "parsed_data": analysis_data
            }
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_result, f, indent=2, ensure_ascii=False)
            print(f"Analysis saved: {analysis_file}")
        except Exception as e:
            print(f"Error saving analysis: {e}")

    def run_analysis_only(self, progress_callback=None):
        """Run analysis only - called by master processor"""
        try:
            if progress_callback:
                progress_callback("Extracting Section 2_1 image...")
            
            # Extract section image  
            section_image = self.load_extracted_image()
            if not section_image:
                print("Failed to extract Section 2_1 image")
                return False

            # Save section image
            os.makedirs(self.test_dir, exist_ok=True)
            image_path = os.path.join(self.test_dir, "section_2_1_extracted.png")
            section_image.save(image_path)
            
            if progress_callback:
                progress_callback("Analyzing Section 2_1 with GPT-4o...")
            
            # Analyze with GPT-4o
            analysis_result = self.analyze_section_2_1(section_image)
            if not analysis_result or "raw_analysis" not in analysis_result:
                print("Failed to get analysis from GPT-4o")
                return False
            
            # Parse analysis
            analysis_data = self.extract_json_from_analysis(analysis_result["raw_analysis"])
            
            # Save analysis results
            analysis_result_full = {
                "raw_analysis": analysis_result["raw_analysis"],
                "parsed_data": analysis_data
            }
            
            analysis_file = os.path.join(self.test_dir, "section_2_1_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_result_full, f, indent=2, ensure_ascii=False)
            
            print("Analysis completed and saved successfully")
            return True
            
        except Exception as e:
            print(f"Error in run_analysis_only: {e}")
            return False

def main():
    """Run 2 1 test"""
    import argparse

    parser = argparse.ArgumentParser(description="2 1 Tester")
    parser.add_argument("--pdf-path", help="Path to PDF file for analysis")
    parser.add_argument("--analysis-only", action="store_true",
                       help="Run only analysis (no Word processing)")
    args = parser.parse_args()

    # Create tester with dynamic PDF path
    tester = Section21Tester(pdf_path=args.pdf_path)

    if args.analysis_only:
        success = tester.run_analysis_only()
        print(f"Analysis {'completed successfully' if success else 'failed'}")
    else:
        tester.run_test()

if __name__ == "__main__":
    main()

