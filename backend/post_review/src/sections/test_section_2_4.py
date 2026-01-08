#!/usr/bin/env python3
"""

# This file requires pre-extracted images from pdf_section_splitter.py
# Run: splitter = PDFSectionSplitter(pdf_path); splitter.extract_all_sections()
# Then this file will work with the pre-extracted section image

Test script for Section 2_4 - Same rules as Section 2_1
- Diagonal/cross detection in LEFT and RIGHT boxes
- Row deletion if BOTH boxes have ALL items interrupted  
- Individual dot point/sentence deletion for partial interruptions
- Continuous diagonal line detection across multiple dot points
"""

import os
import sys
import json
import base64
import requests

from datetime import datetime
from docx import Document
from PIL import Image

# Add the project root to Python path
project_root = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, project_root)

from dotenv import load_dotenv

# This file is designed to work only through the Post-Review UI system
# Run: python post_review_modern_ui.py
# Do not run this file directly

class Section2_4Tester:
    def __init__(self, pdf_path=None):
        self.section_name = "Section_2_4"
        
        # Make test directory relative to the project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.test_dir = os.path.join(project_root, "data", "test_sections", "section_2_4_test")
        
        # This test section works with pre-extracted PNG images only
        # PDF sectioning is handled by PDFSectionSplitter
        # Word processing is handled by unified_section_implementations.py
        
        # Load environment variables
        # Load environment variables from project root
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        env_path = os.path.join(project_root, '.env')
        load_dotenv(env_path)
        self.api_key = os.getenv('OPENAI_API_KEY')
        if not self.api_key:
            raise ValueError("OpenAI API key required. Set OPENAI_API_KEY in .env file.")
        
        self.api_url = "https://api.openai.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

    def load_extracted_image(self) -> str:
        """Load pre-extracted section image (extracted by pdf_section_splitter.py)"""
        try:
            # Ensure test directory exists
            os.makedirs(self.test_dir, exist_ok=True)
            
            # Check for pre-extracted image
            image_filename = f"section_2_4_extracted.png"
            image_path = os.path.join(self.test_dir, image_filename)
            
            if os.path.exists(image_path):
                print(f"✅ Found pre-extracted image: {image_path}")
                
                # Verify it's a valid image
                try:
                    with Image.open(image_path) as img:
                        print(f"   📏 Image size: {img.size}")
                        return image_path
                except Exception as e:
                    print(f"❌ Invalid image file: {e}")
                    return None
            else:
                print(f"❌ Pre-extracted image not found: {image_path}")
                print(f"   💡 Make sure to run pdf_section_splitter.py first!")
                return None
                
        except Exception as e:
            print(f"Error loading pre-extracted image: {e}")
            return None
    
    def analyze_with_gpt4o(self, image_path: str) -> dict:
        """Analyze section image with GPT-4o"""
        try:
            # Load and encode image
            with open(image_path, 'rb') as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            # Get analysis prompt
            prompt = self.get_section_2_4_analysis_prompt()
            
            payload = {
                "model": "gpt-4o",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 2000,
                "temperature": 0.1
            }
            
            response = requests.post(self.api_url, headers=self.headers, json=payload, timeout=90)
            response.raise_for_status()
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            return {
                "raw_analysis": content,
                "success": True
            }
            
        except Exception as e:
            return {
                "raw_analysis": f"Error analyzing section: {str(e)}",
                "success": False
            }
    
    def apply_changes_to_document(self, analysis_result: dict) -> bool:
        """Apply changes to the test document"""
        try:
            if not os.path.exists(self.test_document):
                print(f"❌ Test document not found: {self.test_document}")
                return False
            
            # Load document
            doc = Document(self.test_document)
            
            # Apply Section 2_4 changes
            changes_applied = self.apply_section_2_4_changes(doc, analysis_result)
            
            # Save processed document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"{self.test_dir}/section_2_4_processed_{timestamp}.docx"
            doc.save(output_path)
            
            print(f"💾 Processed document saved: {output_path}")
            print(f"🔧 Changes applied: {len(changes_applied)}")
            
            # Display changes summary
            if changes_applied:
                print(f"\n📋 Changes Summary:")
                for i, change in enumerate(changes_applied, 1):
                    change_type = change.get("type", "unknown")
                    section = change.get("section", "unknown")
                    print(f"   {i}. {change_type} in {section}")
                    
                    if change_type == "row_deletion":
                        print(f"      🚨 Complete row deleted: {change.get('explanation', 'N/A')}")
                    elif change_type in ["left_box_deletions", "right_box_deletions"]:
                        deleted = change.get("deleted_count", 0)
                        requested = change.get("total_requested", 0)
                        print(f"      🗑️ Deleted {deleted}/{requested} items")
                        if change.get("continuous_line", False):
                            print(f"      🔄 Continuous line detected")
            else:
                print(f"\n📋 No changes were applied")
            
            return True
            
        except Exception as e:
            print(f"❌ Error applying changes: {e}")
            return False
    
    def get_section_2_4_analysis_prompt(self) -> str:
        """
        Analysis prompt for Section 2_4 - SAME RULES AS SECTION 2_1
        """
        return """
        Analyze this Section 2_4 for DIAGONAL/CROSS DELETION MARKS:
        
        SECTION LAYOUT: This section has TWO BOXES side by side (LEFT BOX and RIGHT BOX)
        
        CRITICAL DELETION RULES (Same as Section 2_1):
        
        RULE 1: INDIVIDUAL DELETIONS
        - Look for handwritten diagonal lines or X marks (crosses) that interrupt ANY text
        - If diagonal/cross interrupts a dot point, sentence, or text → DELETE that specific item
        - CONTINUOUS DIAGONAL LINES: A single diagonal line may start at one dot point, pass through others, and end at another
        - ALL dot points that the continuous line passes through should be marked for deletion
        - Example: Line starts at dot 2, goes through dot 3, ends at dot 4 → DELETE dots 2, 3, AND 4
        
        RULE 2: ROW DELETION (Critical)
        - If BOTH left box AND right box have diagonal/cross marks interrupting ALL their content
        - Then the ENTIRE ROW should be deleted (not just individual content)
        
        DETECTION REQUIREMENTS:
        - Distinguish handwritten marks from printed text, table borders, and blank spaces
        - Look for diagonal lines, X marks, crosses, or similar deletion indicators
        - Focus on marks that clearly interrupt or cross through text content
        - Ignore decorative elements, table structure, or background patterns
        
        Return detailed JSON:
        {
            "left_box_analysis": {
                "has_deletion_marks": true/false,
                "total_items": number,
                "interrupted_items": number,
                "all_items_interrupted": true/false,
                "deletion_details": [
                    {
                        "item_number": 1,
                        "item_text": "text content",
                        "has_diagonal_cross": true/false,
                        "should_delete": true/false
                    }
                ]
            },
            "right_box_analysis": {
                "has_deletion_marks": true/false,
                "total_items": number,
                "interrupted_items": number,
                "all_items_interrupted": true/false,
                "continuous_line_detected": true/false,
                "continuous_line_description": "description of the continuous line path",
                "deletion_details": [
                    {
                        "item_number": 1,
                        "item_text": "text content", 
                        "has_diagonal_cross": true/false,
                        "should_delete": true/false
                    }
                ]
            },
            "row_deletion_rule": {
                "both_boxes_all_interrupted": true/false,
                "should_delete_entire_row": true/false,
                "explanation": "reasoning for row deletion decision"
            },
            "visual_description": "comprehensive description of all deletion marks detected"
        }
        """
    
    def find_section_2_4_table_row(self, doc: Document) -> tuple:
        """Find the table and row index for Section 2_4 content"""
        # Section 2_4 should be different from Section 2_1, 2_2, and 2_3 content
        # Look for Section 2_4 specific keywords or patterns
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) >= 4 and len(table.columns) >= 2:  # At least 4 rows for multiple sections
                # Check the 4th row (index 3) for Section 2_4 content
                for row_idx in range(3, len(table.rows)):  # Start from row 3 (4th row)
                    row = table.rows[row_idx]
                    if len(row.cells) >= 2:
                        left_cell = row.cells[0].text.strip().lower()
                        right_cell = row.cells[1].text.strip().lower()
                        
                        # Look for Section 2_4 indicators (different from previous sections)
                        section_2_4_keywords = [
                            "pension", "retirement", "aged", "pension", "centrelink",
                            "social", "security", "benefits", "allowance", "payment",
                            "government", "support", "welfare", "assistance"
                        ]
                        
                        # Check if this row contains Section 2_4 content
                        combined_text = left_cell + " " + right_cell
                        keyword_matches = sum(1 for keyword in section_2_4_keywords if keyword in combined_text)
                        
                        if keyword_matches >= 2:  # At least 2 keywords match
                            return table_idx, row_idx
        
        # Fallback: assume first table, fourth row (Row 3) for Section 2_4
        if len(doc.tables) > 0 and len(doc.tables[0].rows) > 3:
            return 0, 3
        
        return None, None
    
    def delete_specific_dot_points(self, doc: Document, table_idx: int, row_idx: int, cell_idx: int, items_to_delete: list) -> int:
        """
        Delete specific dot points from a table cell - SAME LOGIC AS SECTION 2_1
        Returns the number of deletions performed
        """
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[cell_idx]
                
                # Get all paragraphs that could be dot points
                paragraphs = list(cell.paragraphs)
                deleted_count = 0
                matched_items = set()  # Track which items we've successfully matched
                
                print(f"      🔍 Found {len(paragraphs)} paragraphs in cell")
                print(f"      🎯 Looking to delete {len(items_to_delete)} specific items")
                
                # Create a list of paragraphs to delete (don't modify during iteration)
                paragraphs_to_delete = []
                
                for item in items_to_delete:
                    if item.get("should_delete", False):
                        target_text = item.get("item_text", "").strip()
                        item_number = item.get("item_number", 0)
                        
                        if not target_text:
                            continue
                        
                        # Find matching paragraph using multiple strategies
                        best_match = None
                        best_similarity = 0
                        
                        for para_idx, para in enumerate(paragraphs):
                            if para in paragraphs_to_delete:  # Skip already marked for deletion
                                continue
                                
                            para_text = para.text.strip()
                            if not para_text:  # Skip empty paragraphs
                                continue
                            
                            # Strategy 1: Text similarity
                            similarity = self.text_similarity(target_text.lower(), para_text.lower())
                            
                            # Strategy 2: Financial amount matching
                            if '$' in target_text and '$' in para_text:
                                if any(amount in para_text for amount in ['$120,000', '$120000', '120,000', '120000']):
                                    similarity = max(similarity, 0.8)
                            
                            # Strategy 3: Pension/retirement terms matching
                            pension_terms = ['pension', 'retirement', 'aged', 'centrelink', 'social', 'security', 'benefits']
                            if any(term in target_text.lower() for term in pension_terms) and \
                               any(term in para_text.lower() for term in pension_terms):
                                similarity = max(similarity, 0.7)
                            
                            if similarity > best_similarity and similarity >= 0.6:  # Minimum threshold
                                best_similarity = similarity
                                best_match = para
                        
                        if best_match and f"{item_number}_{target_text[:20]}" not in matched_items:
                            paragraphs_to_delete.append(best_match)
                            matched_items.add(f"{item_number}_{target_text[:20]}")
                            print(f"      ✅ Matched item {item_number}: '{target_text[:50]}...' (similarity: {best_similarity:.2f})")
                
                # Now delete the matched paragraphs
                for para in paragraphs_to_delete:
                    try:
                        # Remove the entire paragraph element (including bullet point structure)
                        para._element.getparent().remove(para._element)
                        deleted_count += 1
                        print(f"      🗑️ Deleted paragraph: '{para.text[:50]}...'")
                    except Exception as e:
                        print(f"      ❌ Error deleting paragraph: {e}")
                
                print(f"      📊 Successfully deleted {deleted_count} out of {len(items_to_delete)} requested items")
                return deleted_count
            
            return 0
        except Exception as e:
            print(f"      ❌ Error in delete_specific_dot_points: {e}")
            return 0
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two texts using word-based Jaccard similarity"""
        if not text1 or not text2:
            return 0.0
        
        # Simple word-based similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def delete_table_row(self, doc: Document, table_idx: int, row_idx: int) -> bool:
        """
        Delete an entire table row - SAME LOGIC AS SECTION 2_1
        """
        try:
            table = doc.tables[table_idx]
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                # Remove the row from the table
                table._tbl.remove(row._tr)
                print(f"      🗑️ Deleted entire table row {row_idx}")
                return True
            return False
        except Exception as e:
            print(f"      ❌ Error deleting table row: {e}")
            return False
    
    def apply_section_2_4_changes(self, doc: Document, analysis_data: dict) -> list:
        """
        Apply Section 2_4 changes - SAME LOGIC AS SECTION 2_1
        """
        changes_applied = []
        
        json_data = self.extract_json_from_analysis(analysis_data.get("raw_analysis", ""))
        
        print(f"🔧 Applying Section 2_4 Changes...")
        
        # Find Section 2_4 table and row
        table_idx, row_idx = self.find_section_2_4_table_row(doc)
        
        if table_idx is None or row_idx is None:
            print("❌ Could not find Section 2_4 table row")
            return changes_applied
        
        print(f"🎯 Found Section 2_4 in Table {table_idx}, Row {row_idx}")
        
        # Get analysis data
        left_box_analysis = json_data.get("left_box_analysis", {})
        right_box_analysis = json_data.get("right_box_analysis", {})
        row_deletion_rule = json_data.get("row_deletion_rule", {})
        
        # Check for row deletion first (highest priority)
        if row_deletion_rule.get("should_delete_entire_row", False):
            print(f"   🚨 ROW DELETION RULE TRIGGERED")
            print(f"      📋 Both boxes have ALL items interrupted")
            
            success = self.delete_table_row(doc, table_idx, row_idx)
            if success:
                changes_applied.append({
                    "type": "row_deletion",
                    "section": "Section_2_4",
                    "explanation": row_deletion_rule.get("explanation", "Both boxes completely interrupted")
                })
                print(f"      ✅ Complete row deletion applied")
                return changes_applied  # Row deleted, no individual processing needed
        
        # Individual deletions for left box
        if left_box_analysis.get("has_deletion_marks", False):
            print(f"   📦 LEFT BOX: Individual Deletions")
            left_items = left_box_analysis.get("deletion_details", [])
            items_to_delete = [item for item in left_items if item.get("should_delete", False)]
            
            if items_to_delete:
                deleted_count = self.delete_specific_dot_points(doc, table_idx, row_idx, 0, items_to_delete)
                if deleted_count > 0:
                    changes_applied.append({
                        "type": "left_box_deletions",
                        "section": "Section_2_4",
                        "deleted_count": deleted_count,
                        "total_requested": len(items_to_delete)
                    })
                    print(f"      ✅ Left box: Successfully deleted {deleted_count} items")
        
        # Individual deletions for right box
        if right_box_analysis.get("has_deletion_marks", False):
            print(f"   📦 RIGHT BOX: Individual Deletions")
            right_items = right_box_analysis.get("deletion_details", [])
            items_to_delete = [item for item in right_items if item.get("should_delete", False)]
            
            if items_to_delete:
                deleted_count = self.delete_specific_dot_points(doc, table_idx, row_idx, 1, items_to_delete)
                if deleted_count > 0:
                    changes_applied.append({
                        "type": "right_box_deletions", 
                        "section": "Section_2_4",
                        "deleted_count": deleted_count,
                        "total_requested": len(items_to_delete),
                        "continuous_line": right_box_analysis.get("continuous_line_detected", False)
                    })
                    print(f"      ✅ Right box: Successfully deleted {deleted_count} items")
        
        return changes_applied
    
    def extract_json_from_analysis(self, raw_analysis: str) -> dict:
        """Extract JSON data from GPT-4o analysis response"""
        try:
            # Try to find JSON in the response
            if '{' in raw_analysis and '}' in raw_analysis:
                start = raw_analysis.find('{')
                end = raw_analysis.rfind('}') + 1
                json_str = raw_analysis[start:end]
                return json.loads(json_str)
        except:
            pass
        
        return {}
    
        # Word processing methods removed - now handled by unified_section_implementations.py

    def save_analysis_results(self, analysis_data: dict):
        """Save analysis results to JSON file for unified processing"""
        try:
            analysis_file = os.path.join(self.test_dir, "section_2_4_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_data, f, indent=2, ensure_ascii=False)
            print(f"💾 Analysis saved: {analysis_file}")
        except Exception as e:
            print(f"❌ Error saving analysis: {e}")
    
    def run_analysis_only(self, progress_callback=None):
        """Run only image extraction and analysis (no Word processing)"""
        print(f"Starting Section_2_4 Analysis")
        print("📄 Analysis only - Word processing handled by unified processor")
        print("=" * 80)
        
        if progress_callback:
            progress_callback(1, 3, self.section_name, "Extracting section image...")
        
        # Step 1: Load PDF and cut section
        print(f"📊 Step 1: Loading PDF and cutting Section_2_4...")
        section_image_path = self.load_extracted_image()
        if not section_image_path:
            print(f"❌ Section_2_4 Analysis Failed!")
            return False
        print(f"✅ Section image loaded: {section_image_path}")
        
        if progress_callback:
            progress_callback(2, 3, self.section_name, "Analyzing with GPT-4o...")
        
        # Step 2: Analyze with GPT-4o
        print(f"🧠 Step 2: Analyzing Section_2_4 with GPT-4o...")
        analysis_result = self.analyze_with_gpt4o(section_image_path)
        
        if not analysis_result.get("success"):
            print("❌ GPT-4o analysis failed")
            return False
        
        print("✅ GPT-4o analysis completed")
        
        # Step 3: Parse and save results
        analysis_data = self.extract_json_from_analysis(analysis_result.get("raw_analysis", ""))
        if not analysis_data:
            print("⚠️ Could not parse GPT-4o analysis")
            return False
        
        if progress_callback:
            progress_callback(3, 3, self.section_name, "Saving analysis results...")
        
        # Save analysis results for unified processing
        self.save_analysis_results(analysis_data)
        
        print(f"🎉 Section_2_4 Analysis Complete!")
        return True
    
    def run_test(self):
        """Run the complete Section 2_4 test"""
        print(f"🚀 Starting Section 2_4 Test (Same rules as Section 2_1)")
        print(f"📄 Test document: {self.test_document}")
        print(f"🔍 Section: {self.section_name}")
        print("=" * 80)
        
        # Step 1: Load and cut the PDF section
        print(f"📊 Step 1: Loading PDF and cutting {self.section_name}...")
        section_image_path = self.load_extracted_image()
        
        if section_image_path is None:
            print(f"❌ Failed to load section image")
            return False
        
        print(f"✅ Section image loaded: {section_image_path}")
        
        # Step 2: Analyze with GPT-4o
        print(f"🧠 Step 2: Analyzing {self.section_name} with GPT-4o...")
        analysis_result = self.analyze_with_gpt4o(section_image_path)
        
        if not analysis_result.get("success", False):
            print(f"❌ GPT-4o analysis failed")
            return False
        
        print(f"✅ GPT-4o analysis completed")
        
        # Step 3: Display analysis results
        print(f"📊 Step 3: Analysis Results:")
        json_data = self.extract_json_from_analysis(analysis_result.get("raw_analysis", ""))
        
        # Save analysis results
        analysis_file = os.path.join(self.test_dir, f"{self.section_name}_analysis.json")
        with open(analysis_file, 'w', encoding='utf-8') as f:
            json.dump({
                "raw_analysis": analysis_result["raw_analysis"],
                "parsed_data": json_data
            }, f, indent=2, ensure_ascii=False)
        print(f"💾 Analysis saved: {analysis_file}")
        
        if json_data:
            left_analysis = json_data.get("left_box_analysis", {})
            right_analysis = json_data.get("right_box_analysis", {})
            row_rule = json_data.get("row_deletion_rule", {})
            
            print(f"   📦 Left Box Analysis:")
            print(f"      • Has deletion marks: {left_analysis.get('has_deletion_marks', False)}")
            print(f"      • Items interrupted: {left_analysis.get('interrupted_items', 0)}/{left_analysis.get('total_items', 0)}")
            print(f"      • All items interrupted: {left_analysis.get('all_items_interrupted', False)}")
            
            print(f"   📦 Right Box Analysis:")
            print(f"      • Has deletion marks: {right_analysis.get('has_deletion_marks', False)}")
            print(f"      • Items interrupted: {right_analysis.get('interrupted_items', 0)}/{right_analysis.get('total_items', 0)}")
            print(f"      • All items interrupted: {right_analysis.get('all_items_interrupted', False)}")
            print(f"      • Continuous line detected: {right_analysis.get('continuous_line_detected', False)}")
            
            print(f"   🚨 Row Deletion Rule:")
            print(f"      • Should delete entire row: {row_rule.get('should_delete_entire_row', False)}")
            print(f"      • Explanation: {row_rule.get('explanation', 'N/A')}")
        
        # Step 4: Save analysis results for unified processing
        print("💾 Step 4: Saving analysis results...")
        self.save_analysis_results(analysis_data)
        
        print("📝 Analysis-only mode: Word document processing handled by unified_section_implementations.py")

def main():
    """Run 2 4 test"""
    import argparse

    parser = argparse.ArgumentParser(description="2 4 Tester")
    parser.add_argument("--pdf-path", help="Path to PDF file for analysis")
    parser.add_argument("--analysis-only", action="store_true",
                       help="Run only analysis (no Word processing)")
    args = parser.parse_args()

    # Create tester with dynamic PDF path
    tester = Section2_4Tester(pdf_path=args.pdf_path)

    if args.analysis_only:
        success = tester.run_analysis_only()
        print(f"Analysis {'completed successfully' if success else 'failed'}")
    else:
        tester.run_test()

if __name__ == "__main__":
    main()

