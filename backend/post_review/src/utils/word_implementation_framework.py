#!/usr/bin/env python3
"""
Word Document Implementation Framework
A production-ready system for applying AI analysis results to Word documents

Based on the sophisticated implementation patterns from the post_review workspace
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from typing import Dict, List, Tuple, Optional, Union


class WordDocumentProcessor:
    """
    Production-ready Word document modification framework
    Implements multiple search strategies and robust error handling
    """
    
    def __init__(self, document_path: str, similarity_threshold: float = 0.6):
        """
        Initialize the Word document processor
        
        Args:
            document_path: Path to the Word document to modify
            similarity_threshold: Minimum similarity score for fuzzy matching (0.0-1.0)
        """
        self.document_path = document_path
        self.similarity_threshold = similarity_threshold
        self.changes_log = []
        
    def load_document(self) -> Document:
        """Load the Word document"""
        return Document(self.document_path)
    
    # =========================================================================
    # CORE TEXT MATCHING STRATEGIES (Production Ready)
    # =========================================================================
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """
        Calculate Jaccard similarity between two text strings
        Handles OCR variations and formatting differences
        """
        if not text1.strip() or not text2.strip():
            return 0.0
            
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def find_text_with_keywords(self, text: str, keywords: List[str]) -> bool:
        """Check if text contains any of the specified keywords"""
        text_lower = text.lower()
        return any(keyword.lower() in text_lower for keyword in keywords)
    
    def normalize_text(self, text: str) -> str:
        """Normalize text for better matching (remove extra spaces, etc.)"""
        return ' '.join(text.split())
    
    # =========================================================================
    # STRATEGY 1: PARAGRAPH-LEVEL MODIFICATIONS
    # =========================================================================
    
    def modify_paragraphs(self, doc: Document, modifications: List[Dict]) -> int:
        """
        Apply modifications to document paragraphs
        
        Args:
            doc: Word document object
            modifications: List of modification instructions
                [
                    {
                        "type": "delete|replace",
                        "target_text": "text to find",
                        "replacement_text": "new text" (for replacements only),
                        "match_strategy": "exact|similarity|keywords"
                    }
                ]
        
        Returns:
            Number of successful modifications
        """
        modification_count = 0
        
        print(f"📝 STRATEGY 1: Modifying {len(modifications)} paragraph targets...")
        
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
                
            for mod in modifications:
                target_text = mod.get("target_text", "").strip()
                mod_type = mod.get("type", "delete")
                match_strategy = mod.get("match_strategy", "exact")
                
                if self._text_matches(para_text, target_text, match_strategy):
                    if mod_type == "delete":
                        print(f"   ✅ DELETING Para {i}: '{para_text[:60]}...'")
                        para.clear()
                        modification_count += 1
                        self._log_change("paragraph_deletion", i, para_text)
                        
                    elif mod_type == "replace":
                        replacement_text = mod.get("replacement_text", "")
                        print(f"   ✅ REPLACING Para {i}: '{target_text}' → '{replacement_text}'")
                        
                        # Preserve formatting while replacing content
                        new_text = para_text.replace(target_text, replacement_text)
                        runs = list(para.runs)  # Preserve formatting from first run
                        
                        para.clear()
                        new_run = para.add_run(new_text)
                        
                        # Copy formatting from original if available
                        if runs:
                            self._copy_run_formatting(runs[0], new_run)
                        
                        modification_count += 1
                        self._log_change("paragraph_replacement", i, para_text, replacement_text)
                    break
        
        return modification_count
    
    # =========================================================================
    # STRATEGY 2: TABLE-LEVEL MODIFICATIONS (Your Specialty!)
    # =========================================================================
    
    def modify_table_cells(self, doc: Document, table_modifications: List[Dict]) -> int:
        """
        Apply modifications to specific table cells
        
        Args:
            table_modifications: List of table modification instructions
                [
                    {
                        "table_index": 0,  # or "auto_detect": True
                        "row_index": 2,
                        "cell_index": 1,
                        "modifications": [
                            {
                                "type": "delete|replace",
                                "target_text": "text to find",
                                "replacement_text": "new text"
                            }
                        ]
                    }
                ]
        
        Returns:
            Number of successful modifications
        """
        modification_count = 0
        
        print(f"📊 STRATEGY 2: Modifying {len(table_modifications)} table targets...")
        
        for table_mod in table_modifications:
            # Auto-detect table if needed
            if table_mod.get("auto_detect", False):
                table_idx = self._auto_detect_table(doc, table_mod.get("detection_criteria", {}))
                if table_idx is None:
                    print(f"   ❌ Could not auto-detect table")
                    continue
            else:
                table_idx = table_mod.get("table_index", 0)
            
            row_idx = table_mod.get("row_index", 0)
            cell_idx = table_mod.get("cell_index", 0)
            
            if table_idx >= len(doc.tables):
                print(f"   ❌ Table {table_idx} not found")
                continue
                
            table = doc.tables[table_idx]
            if row_idx >= len(table.rows):
                print(f"   ❌ Row {row_idx} not found in table {table_idx}")
                continue
                
            if cell_idx >= len(table.rows[row_idx].cells):
                print(f"   ❌ Cell {cell_idx} not found in row {row_idx}")
                continue
            
            cell = table.rows[row_idx].cells[cell_idx]
            cell_modifications = table_mod.get("modifications", [])
            
            # Apply modifications to this cell
            for mod in cell_modifications:
                count = self._modify_cell_content(cell, mod, table_idx, row_idx, cell_idx)
                modification_count += count
        
        return modification_count
    
    def delete_table_rows(self, doc: Document, row_deletions: List[Dict]) -> int:
        """
        Delete entire table rows
        
        Args:
            row_deletions: List of row deletion instructions
                [
                    {
                        "table_index": 0,
                        "row_index": 2,
                        "deletion_reason": "Both boxes marked for deletion"
                    }
                ]
        
        Returns:
            Number of rows successfully deleted
        """
        deletion_count = 0
        
        print(f"🚨 STRATEGY 3: Deleting {len(row_deletions)} table rows...")
        
        # Sort by row_index in reverse order to avoid index shifting
        sorted_deletions = sorted(row_deletions, key=lambda x: x.get("row_index", 0), reverse=True)
        
        for deletion in sorted_deletions:
            table_idx = deletion.get("table_index", 0)
            row_idx = deletion.get("row_index", 0)
            reason = deletion.get("deletion_reason", "Not specified")
            
            if self._delete_table_row(doc, table_idx, row_idx):
                print(f"   ✅ DELETED Table {table_idx}, Row {row_idx}: {reason}")
                deletion_count += 1
                self._log_change("table_row_deletion", f"T{table_idx}R{row_idx}", reason)
            else:
                print(f"   ❌ FAILED to delete Table {table_idx}, Row {row_idx}")
        
        return deletion_count
    
    # =========================================================================
    # STRATEGY 3: COMPREHENSIVE SEARCH (Fallback Strategy)
    # =========================================================================
    
    def comprehensive_search_and_modify(self, doc: Document, targets: List[Dict]) -> int:
        """
        Comprehensive search using all strategies as fallback
        Used when specific location information is not available
        """
        modification_count = 0
        
        print(f"🔍 STRATEGY 4: Comprehensive search for {len(targets)} targets...")
        
        for target in targets:
            target_text = target.get("target_text", "")
            mod_type = target.get("type", "delete")
            
            # Try paragraphs first
            found = self._search_and_modify_paragraphs(doc, target)
            if found:
                modification_count += 1
                continue
            
            # Try table cells
            found = self._search_and_modify_table_cells(doc, target)
            if found:
                modification_count += 1
                continue
            
            # Try keyword-based search as last resort
            if self._keyword_based_modification(doc, target):
                modification_count += 1
        
        return modification_count
    
    # =========================================================================
    # HELPER METHODS (Internal Implementation)
    # =========================================================================
    
    def _text_matches(self, text: str, target: str, strategy: str) -> bool:
        """Check if text matches target using specified strategy"""
        text = self.normalize_text(text)
        target = self.normalize_text(target)
        
        if strategy == "exact":
            return target.lower() in text.lower()
        elif strategy == "similarity":
            return self.text_similarity(text, target) > self.similarity_threshold
        elif strategy == "keywords":
            keywords = target.split()
            return self.find_text_with_keywords(text, keywords)
        else:
            return False
    
    def _modify_cell_content(self, cell, modification: Dict, table_idx: int, row_idx: int, cell_idx: int) -> int:
        """Apply modification to a specific table cell"""
        mod_count = 0
        target_text = modification.get("target_text", "")
        mod_type = modification.get("type", "delete")
        
        for para in cell.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue
            
            if self._text_matches(para_text, target_text, modification.get("match_strategy", "exact")):
                if mod_type == "delete":
                    print(f"      ✅ DELETED cell content: '{para_text[:40]}...'")
                    para.clear()
                    mod_count += 1
                    self._log_change("cell_deletion", f"T{table_idx}R{row_idx}C{cell_idx}", para_text)
                    
                elif mod_type == "replace":
                    replacement_text = modification.get("replacement_text", "")
                    new_text = para_text.replace(target_text, replacement_text)
                    
                    runs = list(para.runs)
                    para.clear()
                    new_run = para.add_run(new_text)
                    
                    if runs:
                        self._copy_run_formatting(runs[0], new_run)
                    
                    print(f"      ✅ REPLACED cell content: '{target_text}' → '{replacement_text}'")
                    mod_count += 1
                    self._log_change("cell_replacement", f"T{table_idx}R{row_idx}C{cell_idx}", para_text, replacement_text)
                break
        
        return mod_count
    
    def _delete_table_row(self, doc: Document, table_idx: int, row_idx: int) -> bool:
        """Delete an entire table row"""
        try:
            if table_idx >= len(doc.tables):
                return False
            
            table = doc.tables[table_idx]
            if row_idx >= len(table.rows):
                return False
            
            # Remove the row from the table
            row = table.rows[row_idx]
            table._tbl.remove(row._tr)
            return True
            
        except Exception as e:
            print(f"      ❌ Error deleting table row: {e}")
            return False
    
    def _auto_detect_table(self, doc: Document, criteria: Dict) -> Optional[int]:
        """Auto-detect table based on content criteria"""
        header_keywords = criteria.get("header_keywords", [])
        min_rows = criteria.get("min_rows", 2)
        min_cols = criteria.get("min_cols", 2)
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) >= min_rows and len(table.columns) >= min_cols:
                # Check header row for keywords
                if header_keywords:
                    header_row = table.rows[0]
                    header_text = " ".join(cell.text for cell in header_row.cells)
                    if any(keyword.lower() in header_text.lower() for keyword in header_keywords):
                        return table_idx
                else:
                    return table_idx  # Return first matching table if no keywords specified
        
        return None
    
    def _search_and_modify_paragraphs(self, doc: Document, target: Dict) -> bool:
        """Search and modify in paragraphs (fallback method)"""
        target_text = target.get("target_text", "")
        
        for para in doc.paragraphs:
            if self._text_matches(para.text, target_text, "similarity"):
                if target.get("type") == "delete":
                    para.clear()
                return True
        return False
    
    def _search_and_modify_table_cells(self, doc: Document, target: Dict) -> bool:
        """Search and modify in table cells (fallback method)"""
        target_text = target.get("target_text", "")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._text_matches(para.text, target_text, "similarity"):
                            if target.get("type") == "delete":
                                para.clear()
                            return True
        return False
    
    def _keyword_based_modification(self, doc: Document, target: Dict) -> bool:
        """Keyword-based modification as last resort"""
        keywords = target.get("target_text", "").split()
        
        # Try paragraphs
        for para in doc.paragraphs:
            if self.find_text_with_keywords(para.text, keywords):
                if target.get("type") == "delete":
                    para.clear()
                return True
        
        # Try table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self.find_text_with_keywords(para.text, keywords):
                            if target.get("type") == "delete":
                                para.clear()
                            return True
        
        return False
    
    def _copy_run_formatting(self, source_run, target_run):
        """Copy formatting from source run to target run"""
        try:
            if hasattr(source_run, 'font') and hasattr(target_run, 'font'):
                target_run.font.name = source_run.font.name
                target_run.font.size = source_run.font.size
                target_run.font.bold = source_run.font.bold
                target_run.font.italic = source_run.font.italic
        except:
            pass  # Fail silently on formatting errors
    
    def _log_change(self, change_type: str, location: str, old_content: str, new_content: str = None):
        """Log changes for audit trail"""
        change_record = {
            "type": change_type,
            "location": location,
            "old_content": old_content[:100] + "..." if len(old_content) > 100 else old_content,
            "new_content": new_content[:100] + "..." if new_content and len(new_content) > 100 else new_content
        }
        self.changes_log.append(change_record)
    
    # =========================================================================
    # MAIN PROCESSING METHOD
    # =========================================================================
    
    def process_modifications(self, analysis_results: Dict, output_path: str = None) -> Dict:
        """
        Main method to process all modifications based on analysis results
        
        Args:
            analysis_results: Dictionary containing all modification instructions
                {
                    "paragraph_modifications": [...],
                    "table_modifications": [...],
                    "row_deletions": [...],
                    "comprehensive_search": [...]
                }
            output_path: Path to save the modified document
        
        Returns:
            Dictionary with processing results
        """
        print("🔧 STARTING WORD DOCUMENT PROCESSING")
        print("=" * 60)
        
        # Load document
        doc = self.load_document()
        total_modifications = 0
        
        # Strategy 1: Paragraph modifications
        if "paragraph_modifications" in analysis_results:
            count = self.modify_paragraphs(doc, analysis_results["paragraph_modifications"])
            total_modifications += count
        
        # Strategy 2: Table cell modifications
        if "table_modifications" in analysis_results:
            count = self.modify_table_cells(doc, analysis_results["table_modifications"])
            total_modifications += count
        
        # Strategy 3: Row deletions (highest priority)
        if "row_deletions" in analysis_results:
            count = self.delete_table_rows(doc, analysis_results["row_deletions"])
            total_modifications += count
        
        # Strategy 4: Comprehensive search (fallback)
        if "comprehensive_search" in analysis_results:
            count = self.comprehensive_search_and_modify(doc, analysis_results["comprehensive_search"])
            total_modifications += count
        
        # Save document
        if output_path:
            doc.save(output_path)
            print(f"\n💾 Document saved: {output_path}")
        
        results = {
            "total_modifications": total_modifications,
            "changes_log": self.changes_log,
            "success": True,
            "output_path": output_path
        }
        
        print(f"✅ PROCESSING COMPLETE: {total_modifications} modifications applied")
        print("=" * 60)
        
        return results


# =============================================================================
# USAGE EXAMPLES AND TEMPLATES
# =============================================================================

class WordProcessorTemplates:
    """Ready-to-use templates for common modification patterns"""
    
    @staticmethod
    def create_deletion_template(sentences_to_delete: List[str]) -> Dict:
        """Create template for sentence deletions"""
        return {
            "comprehensive_search": [
                {
                    "target_text": sentence,
                    "type": "delete",
                    "match_strategy": "similarity"
                }
                for sentence in sentences_to_delete
            ]
        }
    
    @staticmethod
    def create_table_modification_template(table_idx: int, row_idx: int, 
                                         left_box_changes: List[Dict], 
                                         right_box_changes: List[Dict]) -> Dict:
        """Create template for table modifications (like Section 1.4)"""
        return {
            "table_modifications": [
                {
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "cell_index": 0,  # Left box
                    "modifications": left_box_changes
                },
                {
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "cell_index": 1,  # Right box
                    "modifications": right_box_changes
                }
            ]
        }
    
    @staticmethod
    def create_row_deletion_template(table_idx: int, row_idx: int, reason: str) -> Dict:
        """Create template for complete row deletion"""
        return {
            "row_deletions": [
                {
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "deletion_reason": reason
                }
            ]
        }


# =============================================================================
# EXAMPLE USAGE
# =============================================================================

def example_usage():
    """Example of how to use the Word processor framework"""
    
    # Initialize processor
    processor = WordDocumentProcessor("path/to/document.docx", similarity_threshold=0.6)
    
    # Example: Process analysis results from AI
    analysis_results = {
        "table_modifications": [
            {
                "table_index": 0,
                "row_index": 2,
                "cell_index": 0,
                "modifications": [
                    {
                        "type": "delete",
                        "target_text": "Some sentence to delete",
                        "match_strategy": "similarity"
                    },
                    {
                        "type": "replace",
                        "target_text": "Original text",
                        "replacement_text": "New text",
                        "match_strategy": "exact"
                    }
                ]
            }
        ],
        "row_deletions": [
            {
                "table_index": 0,
                "row_index": 3,
                "deletion_reason": "Both boxes marked for deletion"
            }
        ]
    }
    
    # Process modifications
    results = processor.process_modifications(analysis_results, "output_document.docx")
    
    # Check results
    print(f"Modifications applied: {results['total_modifications']}")
    print(f"Changes log: {len(results['changes_log'])} entries")


if __name__ == "__main__":
    example_usage()
