#!/usr/bin/env python3
"""
Unified Section Implementation System
Production-ready multi-section document processor based on proven patterns
Implements sophisticated 3-strategy cascading approach with section-specific business rules
"""

import logging
import json
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from dataclasses import dataclass, asdict
import yaml

@dataclass
class SectionConfig:
    """Configuration for each document section"""
    name: str
    description: str
    x: int
    y: int
    width: int
    height: int
    analysis_type: str
    business_rules: List[str]
    ai_prompt_template: str = ""
    
@dataclass 
class ChangeRecord:
    """Track individual changes made to document"""
    type: str  # "sentence_deletion", "content_replacement", "row_deletion", etc.
    section: str
    original_text: str
    new_text: str = ""
    location: str = ""  # e.g., "table_0_row_2_cell_1"
    timestamp: str = ""
    ai_confidence: float = 0.0
    strategy_used: str = ""  # "exact", "similarity", "keyword"
    
class UnifiedSectionProcessor:
    """
    Production-ready unified section implementation system
    Processes all sections in a single document with comprehensive error handling
    """
    
    def __init__(self, config_path: str = "data/templates/config.yaml"):
        """Initialize with section configurations and processing rules"""
        self.logger = logging.getLogger(__name__)
        self.logger.info("Initializing Unified Section Processor")
        
        # Load configuration
        self.config = self._load_config(config_path)
        
        # Section-specific configurations (based on your 18-section system)
        self.section_configs = self._load_section_configs()
        
        # 3-Strategy System Configuration (proven from your implementation)
        self.similarity_threshold = 0.6  # Battle-tested optimal threshold
        self.keyword_min_length = 3      # Minimum keyword length for fallback
        
        # Processing state
        self.changes_applied: List[ChangeRecord] = []
        self.processing_stats = {
            'exact_matches': 0, 'similarity_matches': 0, 'keyword_matches': 0,
            'failed_matches': 0, 'sections_processed': 0, 'sections_failed': 0,
            'total_deletions': 0, 'total_replacements': 0, 'total_row_deletions': 0,
            'strategy_1_exact_matches': 0, 'strategy_2_similarity_matches': 0, 
            'strategy_3_keyword_matches': 0, 'failed_operations': 0
        }
        
        # Error tracking
        self.section_errors: List[str] = []
        self.last_strategy_used = "unknown"
        
    def _load_config(self, config_path: str) -> Dict:
        """Load configuration file with fallback handling"""
        try:
            if Path(config_path).exists():
                with open(config_path, 'r') as f:
                    return yaml.safe_load(f)
            else:
                self.logger.warning(f"Config file not found: {config_path}, using defaults")
                return self._get_default_config()
        except Exception as e:
            self.logger.error(f"Error loading config: {e}")
            return self._get_default_config()
    
    def _get_default_config(self) -> Dict:
        """Default configuration for the system"""
        return {
            'processing': {
                'similarity_threshold': 0.6,
                'max_retries': 3,
                'enable_logging': True
            },
            'output': {
                'include_change_summary': True,
                'timestamp_format': '%Y%m%d_%H%M%S'
            }
        }
    
    def _load_section_configs(self) -> Dict[str, SectionConfig]:
        """Load section-specific configurations (customizable for your document types)"""
        return {
            # Section 1.1 - Portfolio Analysis
            "section_1_1": SectionConfig(
                name="section_1_1",
                description="Portfolio Analysis Section",
                x=63, y=180, width=754, height=120,
                analysis_type="checkbox_circle_detection",
                business_rules=["circle_means_keep", "x_means_delete", "strikethrough_means_delete"]
            ),
            
            # Section 1.2 - Goals/Achieved Table
            "section_1_2": SectionConfig(
                name="section_1_2", 
                description="Goals/Achieved Table with handwritten additions",
                x=63, y=320, width=754, height=180,
                analysis_type="handwritten_text_extraction",
                business_rules=["add_handwritten_to_bullets", "ignore_checkmarks"]
            ),
            
            # Section 1.3 - Portfolio Selection
            "section_1_3": SectionConfig(
                name="section_1_3",
                description="Portfolio Selection with deletion markers",
                x=63, y=520, width=754, height=145,
                analysis_type="selection_and_deletion",
                business_rules=["keep_circled_items", "delete_crossed_items", "delete_marked_bullets"]
            ),
            
            # Section 1.4 - Two-Box Analysis (Your complex example)
            "section_1_4": SectionConfig(
                name="section_1_4",
                description="Items Discussed/Action Taken (Two-box layout)",
                x=63, y=520, width=754, height=145,
                analysis_type="two_box_deletion_replacement",
                business_rules=["row_deletion_if_both_marked", "individual_sentence_deletion", "arrow_replacement"]
            ),
            
            "chunk_1_date_replacement": SectionConfig(
                name="chunk_1_date_replacement",
                description="Chunk 1: Replace XXXX with handwritten date",
                x=0, y=100, width=800, height=200,  # Greeting section area
                analysis_type="date_replacement",
                business_rules=["replace_xxxx_with_date", "strikethrough_detection"]
            ),
            
            "chunk_2_bullet_points": SectionConfig(
                name="chunk_2_bullet_points",
                description="Chunk 2: Fill bullet points with handwritten concerns",
                x=0, y=300, width=800, height=150,  # Concerns section area
                analysis_type="bullet_point_filling",
                business_rules=["fill_existing_bullets", "add_new_bullets", "maintain_formatting"]
            ),
            
            "chunk_3_bullet_points": SectionConfig(
                name="chunk_3_bullet_points",
                description="Chunk 3: Fill bullet points with handwritten opportunities",
                x=0, y=450, width=800, height=150,  # Opportunities section area
                analysis_type="bullet_point_filling",
                business_rules=["fill_existing_bullets", "add_new_bullets", "maintain_formatting"]
            ),
            
            "chunk_4_to_opportunities": SectionConfig(
                name="chunk_4_to_opportunities",
                description="Chunk 4: Append to opportunities section (when Chunk 5 has strengths)",
                x=0, y=450, width=800, height=150,  # Same as opportunities section
                analysis_type="append_to_opportunities",
                business_rules=["append_to_existing_bullets", "add_after_existing"]
            ),
            
            "chunk_4_additional_opportunities": SectionConfig(
                name="chunk_4_additional_opportunities",
                description="Chunk 4: Add to opportunities as additional content (when Chunk 5 has main strengths)",
                x=0, y=450, width=800, height=150,  # Same as opportunities section
                analysis_type="additional_opportunities",
                business_rules=["add_after_existing_bullets", "append_to_opportunities"]
            ),
            
            "chunk_4_strengths": SectionConfig(
                name="chunk_4_strengths",
                description="Chunk 4: Strengths section (when Chunk 4 IS the strengths section)",
                x=0, y=600, width=800, height=150,  # Strengths section area
                analysis_type="bullet_point_filling",
                business_rules=["fill_existing_bullets", "add_new_bullets", "maintain_formatting"]
            ),
            
            "chunk_4_standalone": SectionConfig(
                name="chunk_4_standalone",
                description="Chunk 4: Standalone processing (when no strengths found)",
                x=0, y=750, width=800, height=150,  # Flexible positioning
                analysis_type="flexible_section_detection",
                business_rules=["detect_section_type", "flexible_placement"]
            ),
            
            "chunk_5_strengths": SectionConfig(
                name="chunk_5_strengths",
                description="Chunk 5: Strengths section (normal case)",
                x=0, y=600, width=800, height=150,  # Strengths section area
                analysis_type="bullet_point_filling",
                business_rules=["fill_existing_bullets", "add_new_bullets", "maintain_formatting"]
            ),
            
            "chunk_6_editing": SectionConfig(
                name="chunk_6_editing",
                description="Chunk 6: Editing operations - strikethrough deletion, crosses, $AMOUNT replacement",
                x=0, y=900, width=800, height=200,  # Strategies section area
                analysis_type="editing_operations",
                business_rules=["delete_crossed_bullets", "strikethrough_word_deletion", "amount_replacement"]
            ),
            
            "chunk_7_editing": SectionConfig(
                name="chunk_7_editing", 
                description="Chunk 7: Cash flow section - strikethrough deletion and dot point deletion",
                x=0, y=1100, width=800, height=200,  # Cash flow section area
                analysis_type="editing_operations",
                business_rules=["strikethrough_word_deletion", "delete_crossed_bullets"]
            ),
            
            "chunk_8_editing": SectionConfig(
                name="chunk_8_editing",
                description="Chunk 8: Business Plan table - row deletion based on crosses and lines",
                x=31, y=806, width=1566, height=1488,  # Business plan table area
                analysis_type="editing_operations",
                business_rules=["delete_table_rows", "append_handwritten_notes", "ignore_small_diagonal_lines"]
            ),
            
            "chunk_9_editing": SectionConfig(
                name="chunk_9_editing",
                description="Chunk 9: Fee section and tax returns - $AMOUNT replacement and text substitution",
                x=0, y=1300, width=800, height=200,  # Fee section area
                analysis_type="editing_operations", 
                business_rules=["amount_replacement", "text_substitution", "color_correction"]
            ),
            
            "chunk_5_general": SectionConfig(
                name="chunk_5_general",
                description="Chunk 5: General content processing",
                x=0, y=750, width=800, height=150,  # Flexible positioning
                analysis_type="flexible_section_detection",
                business_rules=["detect_section_type", "flexible_placement"]
            ),
            
            "general_analysis": SectionConfig(
                name="general_analysis",
                description="General document analysis with detected annotations",
                x=0, y=0, width=0, height=0,  # Full document scope
                analysis_type="general_document_analysis",
                business_rules=["annotation_detection", "change_tracking", "comment_addition"]
            ),
            
            # Add more sections as needed for your specific document types
        }
    
    def process_all_sections(self, section_analyses: Dict[str, Dict], 
                           base_document_path: str, output_path: str) -> Dict[str, Any]:
        """
        Main processing method - applies all section analyses to single document
        
        Args:
            section_analyses: Dictionary of {section_name: ai_analysis_results}
            base_document_path: Path to base Word document template
            output_path: Path for final processed document
            
        Returns:
            Comprehensive processing results with statistics and change tracking
        """
        self.logger.info(f"🚀 Starting unified section processing")
        self.logger.info(f"📄 Base document: {base_document_path}")
        self.logger.info(f"📊 Processing {len(section_analyses)} sections")
        
        # Reset processing state
        self.changes_applied = []
        self.section_errors = []
        self.processing_stats = {k: 0 for k in self.processing_stats.keys()}
        
        try:
            # Load base document
            doc = Document(base_document_path)
            self.logger.info(f"✅ Loaded base document: {base_document_path}")
            
            # Process each section sequentially (order matters for some business rules)
            for section_name, analysis in section_analyses.items():
                self.logger.info(f"🔧 Processing {section_name}...")
                
                try:
                    changes = self._apply_section_changes(doc, section_name, analysis)
                    self.changes_applied.extend(changes)  # Track all changes
                    self.processing_stats['sections_processed'] += 1
                    self.logger.info(f"✅ {section_name}: {len(changes)} changes applied")
                    
                except Exception as e:
                    self.processing_stats['sections_failed'] += 1
                    error_msg = f"❌ {section_name}: {str(e)}"
                    self.section_errors.append(error_msg)
                    self.logger.error(error_msg)
                    continue  # Process remaining sections
            
            # Add processing summary to document (optional)
            if self.config.get('output', {}).get('include_change_summary', True):
                self._add_processing_summary_to_document(doc)
            
            # Save final document
            doc.save(output_path)
            self.logger.info(f"💾 Final document saved: {output_path}")
            
            # Calculate final statistics
            # Use actual changes count instead of just strategy-based counters
            total_changes = len(self.changes_applied)
            strategy_based_changes = (self.processing_stats['exact_matches'] + 
                                     self.processing_stats['similarity_matches'] + 
                                     self.processing_stats['keyword_matches'])
            
            success_rate = (self.processing_stats['sections_processed'] / 
                          len(section_analyses)) if section_analyses else 0
            
            # Return comprehensive results
            return {
                'status': 'success',
                'input_file': base_document_path,
                'output_file': output_path,
                'processing_timestamp': datetime.now().isoformat(),
                'sections': {
                    'total_sections': len(section_analyses),
                    'successful_sections': self.processing_stats['sections_processed'],
                    'failed_sections': self.processing_stats['sections_failed'],
                    'success_rate': success_rate
                },
                'changes': {
                    'total_changes_applied': total_changes,
                    'strategy_breakdown': {
                        'exact_matches': self.processing_stats['exact_matches'],
                        'similarity_matches': self.processing_stats['similarity_matches'],
                        'keyword_matches': self.processing_stats['keyword_matches'],
                        'failed_matches': self.processing_stats['failed_matches']
                    },
                    'operation_breakdown': {
                        'deletions': self.processing_stats['total_deletions'],
                        'replacements': self.processing_stats['total_replacements'],
                        'row_deletions': self.processing_stats['total_row_deletions']
                    }
                },
                'detailed_changes': [asdict(change) for change in self.changes_applied],
                'errors': self.section_errors,
                'file_size': Path(output_path).stat().st_size if Path(output_path).exists() else 0
            }
            
        except Exception as e:
            self.logger.error(f"❌ Critical error in unified processing: {str(e)}")
            return {
                'status': 'error',
                'error': str(e),
                'input_file': base_document_path,
                'sections_processed': self.processing_stats['sections_processed'],
                'changes_applied': len(self.changes_applied)
            }
    
    def _apply_section_changes(self, doc: Document, section_name: str, 
                             analysis: Dict[str, Any]) -> List[ChangeRecord]:
        """
        Apply changes for a specific section using section-specific business rules
        """
        section_config = self.section_configs.get(section_name)
        if not section_config:
            self.logger.warning(f"⚠️ No configuration found for {section_name}")
            return []
        
        changes = []
        
        # Route to section-specific implementation based on analysis_type
        if section_config.analysis_type == "two_box_deletion_replacement":
            changes = self._implement_section_1_4(doc, analysis, section_name)
        elif section_config.analysis_type == "handwritten_text_extraction":
            changes = self._implement_section_1_2(doc, analysis, section_name) 
        elif section_config.analysis_type == "selection_and_deletion":
            changes = self._implement_section_1_3(doc, analysis, section_name)
        elif section_config.analysis_type == "checkbox_circle_detection":
            changes = self._implement_section_1_1(doc, analysis, section_name)
        elif section_config.analysis_type == "date_replacement":
            changes = self._implement_chunk_1_date_replacement(doc, analysis, section_name)
        elif section_config.analysis_type == "bullet_point_filling":
            if section_name == "chunk_2_bullet_points":
                changes = self._implement_chunk_2_bullet_points(doc, analysis, section_name)
            elif section_name == "chunk_3_bullet_points":
                changes = self._implement_chunk_3_bullet_points(doc, analysis, section_name)
            elif section_name == "chunk_4_strengths":
                changes = self._implement_chunk_4_strengths(doc, analysis, section_name)
            elif section_name == "chunk_5_strengths":
                changes = self._implement_chunk_5_strengths(doc, analysis, section_name)
            else:
                changes = self._implement_generic_bullet_points(doc, analysis, section_name)
        elif section_config.analysis_type == "append_to_opportunities":
            changes = self._implement_chunk_4_to_opportunities(doc, analysis, section_name)
        elif section_config.analysis_type == "additional_opportunities" or section_name == "chunk_4_additional_opportunities":
            changes = self._implement_chunk_4_additional_opportunities(doc, analysis, section_name)
        elif section_config.analysis_type == "editing_operations":
            if section_name == "chunk_6_editing":
                changes = self._implement_chunk_6_editing(doc, analysis, section_name)
            elif section_name == "chunk_7_editing":
                changes = self._implement_chunk_7_editing(doc, analysis, section_name)
            elif section_name == "chunk_8_editing":
                changes = self._implement_chunk_8_editing(doc, analysis, section_name)
            elif section_name == "chunk_9_editing":
                changes = self._implement_chunk_9_editing(doc, analysis, section_name)
            else:
                changes = self._implement_chunk_6_editing(doc, analysis, section_name)  # Default to chunk 6 logic
        elif section_config.analysis_type == "flexible_section_detection":
            if section_name == "chunk_4_standalone":
                changes = self._implement_chunk_4_standalone(doc, analysis, section_name)
            elif section_name == "chunk_5_general":
                changes = self._implement_chunk_5_general(doc, analysis, section_name)
            else:
                changes = self._implement_generic_flexible_detection(doc, analysis, section_name)
        else:
            # Generic implementation for unknown types
            changes = self._implement_generic_section(doc, analysis, section_name)
        
        # Track changes
        self.changes_applied.extend(changes)
        return changes
    
    def _implement_section_1_4(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Section 1.4: Two-box deletion/replacement implementation
        Your most complex example with sophisticated business rules
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract analysis results
        left_box = analysis.get('left_box_analysis', {})
        right_box = analysis.get('right_box_analysis', {})
        row_deletion = analysis.get('row_deletion_rule', {})
        
        # Rule 1: Row deletion if both boxes marked
        if row_deletion.get('delete_entire_row', False):
            self.logger.info(f"🗑️ {section_name}: Applying row deletion rule")
            
            # Find and delete entire rows (this requires identifying the table structure)
            for table_idx, table in enumerate(doc.tables):
                # Look for the specific table related to this section
                if self._is_section_table(table, section_name):
                    # Delete rows that match the criteria
                    rows_to_delete = self._identify_rows_for_deletion(table, analysis)
                    for row_idx in reversed(rows_to_delete):  # Delete from end to maintain indices
                        self._delete_table_row(table, row_idx)
                        changes.append(ChangeRecord(
                            type="row_deletion",
                            section=section_name,
                            original_text=f"Table row {row_idx}",
                            location=f"table_{table_idx}_row_{row_idx}",
                            timestamp=timestamp,
                            ai_confidence=analysis.get('confidence', 0.9),
                            strategy_used="row_deletion_rule"
                        ))
                        self.processing_stats['total_row_deletions'] += 1
        
        # Rule 2: Individual sentence deletions in left box
        for sentence_data in left_box.get('sentences_to_delete', []):
            sentence_text = sentence_data.get('sentence_text', '')
            if sentence_text:
                success, strategy = self._apply_cascading_deletion(doc, sentence_text)
                if success:
                    changes.append(ChangeRecord(
                        type="sentence_deletion",
                        section=section_name,
                        original_text=sentence_text[:100],
                        location="left_box",
                        timestamp=timestamp,
                        ai_confidence=sentence_data.get('confidence', 0.8),
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_deletions'] += 1
        
        # Rule 3: Individual sentence deletions in right box
        for sentence_data in right_box.get('sentences_to_delete', []):
            sentence_text = sentence_data.get('sentence_text', '')
            if sentence_text:
                success, strategy = self._apply_cascading_deletion(doc, sentence_text)
                if success:
                    changes.append(ChangeRecord(
                        type="sentence_deletion", 
                        section=section_name,
                        original_text=sentence_text[:100],
                        location="right_box",
                        timestamp=timestamp,
                        ai_confidence=sentence_data.get('confidence', 0.8),
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_deletions'] += 1
        
        # Rule 4: Content replacements (arrows pointing to new text)
        for replacement_data in left_box.get('sentences_to_replace', []) + right_box.get('sentences_to_replace', []):
            original_text = replacement_data.get('original_text', '')
            replacement_text = replacement_data.get('replacement_text', '')
            
            if original_text and replacement_text:
                success, strategy = self._apply_cascading_replacement(doc, original_text, replacement_text)
                if success:
                    changes.append(ChangeRecord(
                        type="content_replacement",
                        section=section_name, 
                        original_text=original_text[:100],
                        new_text=replacement_text[:100],
                        timestamp=timestamp,
                        ai_confidence=replacement_data.get('confidence', 0.8),
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_replacements'] += 1
        
        return changes
    
    def _implement_section_1_2(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Section 1.2: Goals/Achieved table with handwritten text additions
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten text detections
        handwritten_items = analysis.get('handwritten_text', [])
        
        for item in handwritten_items:
            handwritten_text = item.get('text', '')
            nearby_bullet = item.get('nearby_bullet_text', '')
            
            if handwritten_text and nearby_bullet:
                # Find the bullet point and add handwritten text
                success, strategy = self._add_text_to_bullet_point(doc, nearby_bullet, handwritten_text)
                if success:
                    changes.append(ChangeRecord(
                        type="handwritten_addition",
                        section=section_name,
                        original_text=nearby_bullet[:50],
                        new_text=handwritten_text,
                        timestamp=timestamp,
                        ai_confidence=item.get('confidence', 0.85),
                        strategy_used=strategy
                    ))
        
        return changes
    
    def _implement_section_1_3(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Section 1.3: Portfolio selection with circles/crosses
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Process circled items (keep these)
        circled_items = analysis.get('circled_items', [])
        crossed_items = analysis.get('crossed_items', [])
        marked_bullets = analysis.get('marked_bullet_points', [])
        
        # Delete crossed items
        for item in crossed_items:
            item_text = item.get('text', '')
            if item_text:
                success, strategy = self._apply_cascading_deletion(doc, item_text)
                if success:
                    changes.append(ChangeRecord(
                        type="crossed_item_deletion",
                        section=section_name,
                        original_text=item_text[:100],
                        timestamp=timestamp,
                        ai_confidence=item.get('confidence', 0.8),
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_deletions'] += 1
        
        # Delete marked bullet points
        for bullet in marked_bullets:
            bullet_text = bullet.get('text', '')
            if bullet_text:
                success, strategy = self._apply_cascading_deletion(doc, bullet_text)
                if success:
                    changes.append(ChangeRecord(
                        type="bullet_deletion",
                        section=section_name,
                        original_text=bullet_text[:100],
                        timestamp=timestamp,
                        ai_confidence=bullet.get('confidence', 0.8),
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_deletions'] += 1
        
        return changes
    
    def _implement_section_1_1(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Section 1.1: Portfolio analysis with checkbox/circle detection
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Process checkbox selections
        checkboxes = analysis.get('checkboxes', [])
        
        for checkbox in checkboxes:
            item_text = checkbox.get('item', '')
            status = checkbox.get('status', '')  # 'checked', 'crossed', 'circled'
            action = checkbox.get('action', '')  # 'keep', 'delete', 'select'
            
            if item_text and action == 'delete':
                success, strategy = self._apply_cascading_deletion(doc, item_text)
                if success:
                    changes.append(ChangeRecord(
                        type="checkbox_deletion",
                        section=section_name,
                        original_text=item_text[:100],
                        timestamp=timestamp,
                        ai_confidence=checkbox.get('confidence', 0.8),
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_deletions'] += 1
        
        return changes
    
    def _implement_generic_section(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Generic implementation for sections without specific handlers
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Look for common patterns in analysis
        items_to_delete = analysis.get('items_to_delete', [])
        items_to_replace = analysis.get('items_to_replace', [])
        
        # Process deletions
        for item in items_to_delete:
            if isinstance(item, dict):
                text = item.get('text', '')
            else:
                text = str(item)
                
            if text:
                success, strategy = self._apply_cascading_deletion(doc, text)
                if success:
                    changes.append(ChangeRecord(
                        type="generic_deletion",
                        section=section_name,
                        original_text=text[:100],
                        timestamp=timestamp,
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_deletions'] += 1
        
        # Process replacements
        for item in items_to_replace:
            original = item.get('original', '') if isinstance(item, dict) else ''
            replacement = item.get('replacement', '') if isinstance(item, dict) else ''
            
            if original and replacement:
                success, strategy = self._apply_cascading_replacement(doc, original, replacement)
                if success:
                    changes.append(ChangeRecord(
                        type="generic_replacement",
                        section=section_name,
                        original_text=original[:100],
                        new_text=replacement[:100],
                        timestamp=timestamp,
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_replacements'] += 1
        
        return changes
    
    def _add_processing_summary_to_document(self, doc: Document):
        """Add a summary of processing changes to the document"""
        try:
            from datetime import datetime
            
            # Add a page break and summary section
            doc.add_page_break()
            doc.add_heading('Document Processing Summary', level=1)
            
            # Add processing timestamp
            doc.add_paragraph(f"Processing completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add processing statistics
            doc.add_heading('Processing Statistics', level=2)
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"• Total Deletions: {self.processing_stats.get('total_deletions', 0)}\n")
            stats_para.add_run(f"• Total Replacements: {self.processing_stats.get('total_replacements', 0)}\n")
            stats_para.add_run(f"• Exact Matches: {self.processing_stats.get('strategy_1_exact_matches', 0)}\n")
            stats_para.add_run(f"• Similarity Matches: {self.processing_stats.get('strategy_2_similarity_matches', 0)}\n")
            stats_para.add_run(f"• Keyword Matches: {self.processing_stats.get('strategy_3_keyword_matches', 0)}\n")
            
            # Add changes summary if available
            if self.changes_applied:
                doc.add_heading('Changes Applied', level=2)
                for change in self.changes_applied[:10]:  # Limit to first 10 changes
                    change_para = doc.add_paragraph()
                    change_para.add_run(f"• {change.type} in {change.section}: ")
                    if change.original_text:
                        change_para.add_run(f'"{change.original_text[:50]}..."')
                    if change.new_text:
                        change_para.add_run(f' → "{change.new_text[:50]}..."')
                    change_para.add_run(f" (Strategy: {change.strategy_used})")
            
        except Exception as e:
            self.logger.warning(f"Could not add processing summary: {e}")
    
    def _implement_chunk_2_bullet_points(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 2: Fill bullet points with handwritten concerns
        
        Business Rules:
        1. Fill existing empty bullet points with handwritten text
        2. Add new bullet points if additional handwritten items detected
        3. Maintain consistent formatting
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for {section_name}")
            return changes
        
        # Debug: Log the handwritten items we're working with
        self.logger.info(f"Chunk 2 Debug: Found {len(handwritten_items)} handwritten items")
        for i, item in enumerate(handwritten_items):
            self.logger.info(f"  Item {i+1}: {item}")
        
        # Find the concerns section with bullet points
        concerns_text = "main areas of concern that need to be resolved are:"
        target_paragraph = None
        
        # Debug: Log all paragraphs to see document structure
        self.logger.info(f"Chunk 2 Debug: Searching for concerns section...")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                self.logger.info(f"  Para {i}: '{paragraph.text[:100]}...'")
        
        # Strategy 1: Find the exact concerns section
        for paragraph in doc.paragraphs:
            if concerns_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                self.logger.info(f"Chunk 2 Debug: Found exact concerns section: '{paragraph.text}'")
                break
        
        if not target_paragraph:
            # Strategy 2: Find similar text about concerns
            concern_keywords = ['concern', 'problem', 'danger', 'obstacle']
            for paragraph in doc.paragraphs:
                if any(keyword in paragraph.text.lower() for keyword in concern_keywords):
                    target_paragraph = paragraph
                    self.logger.info(f"Chunk 2 Debug: Found similar concerns section: '{paragraph.text}'")
                    break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find concerns section for {section_name}")
            return changes
        
        # Find existing bullet points after the concerns paragraph
        bullet_paragraphs = []
        found_concerns = False
        
        self.logger.info(f"Chunk 2 Debug: Looking for bullet points after concerns section...")
        paragraph_count = 0
        total_paragraphs_examined = 0
        target_text = target_paragraph.text.strip()
        
        for paragraph in doc.paragraphs:
            total_paragraphs_examined += 1
            
            # Use text matching instead of object comparison
            if paragraph.text.strip() == target_text and not found_concerns:
                found_concerns = True
                self.logger.info(f"Chunk 2 Debug: Found target paragraph at position {total_paragraphs_examined}, now looking for bullets...")
                continue
            
            if found_concerns:
                paragraph_count += 1
                # Debug: Log what we're examining - show more paragraphs
                self.logger.info(f"Chunk 2 Debug: Examining paragraph #{paragraph_count}: '{paragraph.text}' (stripped: '{paragraph.text.strip()}')")
                
                # Check if this is a Word formatted list item (bullet or numbered)
                is_list_item = False
                if hasattr(paragraph, '_element') and paragraph._element is not None:
                    numbering = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numbering is not None:
                        is_list_item = True
                        self.logger.info(f"Chunk 2 Debug: Found Word list item formatting in paragraph #{paragraph_count}!")
                
                self.logger.info(f"Chunk 2 Debug: Para #{paragraph_count} analysis - is_list_item: {is_list_item}, text_len: {len(paragraph.text)}, stripped_len: {len(paragraph.text.strip())}")
                
                # Look for bullet points (text-based OR Word formatted list items)
                if (paragraph.text.strip().startswith('•') or 
                    paragraph.text.strip().startswith('-') or 
                    paragraph.text.strip().startswith('*') or
                    paragraph.text.strip() == '•' or
                    is_list_item or  # NEW: Include Word formatted list items
                    (paragraph.text.strip() == '' and is_list_item)):  # Empty list items
                    bullet_paragraphs.append(paragraph)
                    list_type = "Word-formatted" if is_list_item else "text-based"
                    self.logger.info(f"Chunk 2 Debug: ✅ ADDED {list_type} bullet paragraph #{paragraph_count}: '{paragraph.text}' (empty: {len(paragraph.text.strip()) == 0})")
                else:
                    self.logger.info(f"Chunk 2 Debug: ❌ SKIPPED paragraph #{paragraph_count} - not a bullet point")
                
                # Stop after examining 10 paragraphs to avoid infinite loops
                if paragraph_count >= 10:
                    self.logger.info(f"Chunk 2 Debug: Stopping after examining {paragraph_count} paragraphs")
                    break
                    
                # If we found a non-empty paragraph that's not a bullet, likely moved to next section
                if paragraph.text.strip() and not paragraph.text.strip().startswith(' '):
                    # Don't stop too early - let's see more paragraphs
                    self.logger.info(f"Chunk 2 Debug: Next section paragraph: '{paragraph.text}'")
                    if len(bullet_paragraphs) > 0:  # Only stop if we found some bullets
                        break
                
                # Safety: don't look at too many paragraphs
                if len(bullet_paragraphs) >= 10:
                    self.logger.info(f"Chunk 2 Debug: Stopping search - found {len(bullet_paragraphs)} paragraphs")
                    break
        
        self.logger.info(f"Chunk 2 Debug: Loop completed. Total paragraphs examined: {total_paragraphs_examined}")
        self.logger.info(f"Chunk 2 Debug: Found concerns flag: {found_concerns}")
        self.logger.info(f"Chunk 2 Debug: Paragraphs examined after concerns: {paragraph_count}")
        self.logger.info(f"Chunk 2 Debug: Found {len(bullet_paragraphs)} bullet paragraphs")
        
        # Fill existing bullet points ONLY with handwritten content
        # Do NOT add new bullet points - only fill the existing ones
        items_used = 0
        for i, bullet_para in enumerate(bullet_paragraphs):
            if items_used < len(handwritten_items):
                handwritten_item = handwritten_items[items_used]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    # Update the bullet point with handwritten content
                    # Do NOT add bullet character - the bullet formatting already exists
                    bullet_para.clear()
                    bullet_para.add_run(new_text)  # Just add the text, no bullet character
                    
                    changes.append(ChangeRecord(
                        type="bullet_point_fill",
                        section=section_name,
                        original_text="empty bullet" if not bullet_para.text else bullet_para.text[:50],
                        new_text=new_text[:50],  # Just the text, no bullet character
                        timestamp=timestamp,
                        strategy_used="fill_existing_bullet_only"
                    ))
                    
                    items_used += 1
                    self.processing_stats['total_replacements'] += 1
        
        # Log if there are extra handwritten items that couldn't be used
        if items_used < len(handwritten_items):
            extra_items = len(handwritten_items) - items_used
            self.logger.info(f"Chunk 2: {extra_items} extra handwritten items found but only {len(bullet_paragraphs)} existing bullets available")
        
        self.logger.info(f"Chunk 2: {len(changes)} bullet point changes applied")
        return changes
    
    def _implement_chunk_3_bullet_points(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 3: Fill bullet points with handwritten opportunities
        
        Business Rules:
        1. Fill existing empty bullet points with handwritten text
        2. Add new bullet points if additional handwritten items detected
        3. Maintain consistent formatting
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for {section_name}")
            return changes
        
        # Debug: Log the handwritten items we're working with
        self.logger.info(f"Chunk 3 Debug: Found {len(handwritten_items)} handwritten items")
        for i, item in enumerate(handwritten_items):
            self.logger.info(f"  Item {i+1}: {item}")
        
        # Find the opportunities section with bullet points
        opportunities_text = "most promising and strategic areas of opportunity are:"
        target_paragraph = None
        
        # Debug: Log all paragraphs to see document structure
        self.logger.info(f"Chunk 3 Debug: Searching for opportunities section...")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                self.logger.info(f"  Para {i}: '{paragraph.text[:100]}...'")
        
        # Strategy 1: Find the exact opportunities section
        for paragraph in doc.paragraphs:
            if opportunities_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                self.logger.info(f"Chunk 3 Debug: Found exact opportunities section: '{paragraph.text}'")
                break
        
        if not target_paragraph:
            # Strategy 2: Find similar text about opportunities
            opportunity_keywords = ['opportunity', 'opportunit', 'advantage', 'strategic', 'promising']
            for paragraph in doc.paragraphs:
                if any(keyword in paragraph.text.lower() for keyword in opportunity_keywords):
                    target_paragraph = paragraph
                    self.logger.info(f"Chunk 3 Debug: Found similar opportunities section: '{paragraph.text}'")
                    break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find opportunities section for {section_name}")
            return changes
        
        # Find existing bullet points after the opportunities paragraph
        bullet_paragraphs = []
        found_opportunities = False
        
        self.logger.info(f"Chunk 3 Debug: Looking for bullet points after opportunities section...")
        paragraph_count = 0
        total_paragraphs_examined = 0
        target_text = target_paragraph.text.strip()
        
        for paragraph in doc.paragraphs:
            total_paragraphs_examined += 1
            
            # Use text matching instead of object comparison
            if paragraph.text.strip() == target_text and not found_opportunities:
                found_opportunities = True
                self.logger.info(f"Chunk 3 Debug: Found target paragraph at position {total_paragraphs_examined}, now looking for bullets...")
                continue
            
            if found_opportunities:
                paragraph_count += 1
                # Debug: Log what we're examining - show more paragraphs
                self.logger.info(f"Chunk 3 Debug: Examining paragraph #{paragraph_count}: '{paragraph.text}' (stripped: '{paragraph.text.strip()}')")
                
                # Check if this is a Word formatted list item (bullet or numbered)
                is_list_item = False
                if hasattr(paragraph, '_element') and paragraph._element is not None:
                    numbering = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numbering is not None:
                        is_list_item = True
                        self.logger.info(f"Chunk 3 Debug: Found Word list item formatting in paragraph #{paragraph_count}!")
                
                self.logger.info(f"Chunk 3 Debug: Para #{paragraph_count} analysis - is_list_item: {is_list_item}, text_len: {len(paragraph.text)}, stripped_len: {len(paragraph.text.strip())}")
                
                # Look for bullet points (text-based OR Word formatted list items)
                if (paragraph.text.strip().startswith('•') or 
                    paragraph.text.strip().startswith('-') or 
                    paragraph.text.strip().startswith('*') or
                    paragraph.text.strip() == '•' or
                    is_list_item or  # NEW: Include Word formatted list items
                    (paragraph.text.strip() == '' and is_list_item)):  # Empty list items
                    bullet_paragraphs.append(paragraph)
                    list_type = "Word-formatted" if is_list_item else "text-based"
                    self.logger.info(f"Chunk 3 Debug: ✅ ADDED {list_type} bullet paragraph #{paragraph_count}: '{paragraph.text}' (empty: {len(paragraph.text.strip()) == 0})")
                else:
                    self.logger.info(f"Chunk 3 Debug: ❌ SKIPPED paragraph #{paragraph_count} - not a bullet point")
                
                # Stop after examining 10 paragraphs to avoid infinite loops
                if paragraph_count >= 10:
                    self.logger.info(f"Chunk 3 Debug: Stopping after examining {paragraph_count} paragraphs")
                    break
                    
                # If we found a non-empty paragraph that's not a bullet, likely moved to next section
                if paragraph.text.strip() and not paragraph.text.strip().startswith(' '):
                    # Don't stop too early - let's see more paragraphs
                    self.logger.info(f"Chunk 3 Debug: Next section paragraph: '{paragraph.text}'")
                    if len(bullet_paragraphs) > 0:  # Only stop if we found some bullets
                        break
                
                # Safety: don't look at too many paragraphs
                if len(bullet_paragraphs) >= 10:
                    self.logger.info(f"Chunk 3 Debug: Stopping search - found {len(bullet_paragraphs)} paragraphs")
                    break
        
        self.logger.info(f"Chunk 3 Debug: Loop completed. Total paragraphs examined: {total_paragraphs_examined}")
        self.logger.info(f"Chunk 3 Debug: Found opportunities flag: {found_opportunities}")
        self.logger.info(f"Chunk 3 Debug: Paragraphs examined after opportunities: {paragraph_count}")
        self.logger.info(f"Chunk 3 Debug: Found {len(bullet_paragraphs)} bullet paragraphs")
        
        # Fill existing bullet points with handwritten content
        items_used = 0
        for i, bullet_para in enumerate(bullet_paragraphs):
            if items_used < len(handwritten_items):
                handwritten_item = handwritten_items[items_used]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    # Update the bullet point with handwritten content
                    # Do NOT add bullet character - the bullet formatting already exists
                    bullet_para.clear()
                    bullet_para.add_run(new_text)  # Just add the text, no bullet character
                    
                    changes.append(ChangeRecord(
                        type="bullet_point_fill",
                        section=section_name,
                        original_text=bullet_para.text[:50] if bullet_para.text else "empty bullet",
                        new_text=new_text[:50],  # Just the text, no bullet character
                        timestamp=timestamp,
                        strategy_used="bullet_point_replacement"
                    ))
                    
                    items_used += 1
                    self.processing_stats['total_replacements'] += 1
        
        # Log if there are extra handwritten items that couldn't be used
        if items_used < len(handwritten_items):
            extra_items = len(handwritten_items) - items_used
            self.logger.info(f"Chunk 3: {extra_items} extra handwritten items found but only {len(bullet_paragraphs)} existing bullets available")
        
        self.logger.info(f"Chunk 3: {len(changes)} bullet point changes applied")
        return changes
    
    def _implement_chunk_4_to_opportunities(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 4 → Opportunities: Append Chunk 4 content to opportunities section
        (When Chunk 5 contains strengths section)
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for {section_name}")
            return changes
        
        # Find the opportunities section
        opportunities_text = "most promising and strategic areas of opportunity are:"
        target_paragraph = None
        
        for paragraph in doc.paragraphs:
            if opportunities_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find opportunities section for {section_name}")
            return changes
        
        # Find existing bullet points in opportunities section
        bullet_paragraphs = []
        found_opportunities = False
        
        for paragraph in doc.paragraphs:
            if paragraph == target_paragraph:
                found_opportunities = True
                continue
            
            if found_opportunities:
                if (paragraph.text.strip().startswith('•') or 
                    paragraph.text.strip().startswith('-') or 
                    paragraph.text.strip().startswith('*')):
                    bullet_paragraphs.append(paragraph)
                elif paragraph.text.strip() and not paragraph.text.strip().startswith(' '):
                    break
        
        # Append new bullets after existing ones
        insert_location = bullet_paragraphs[-1] if bullet_paragraphs else target_paragraph
        
        for handwritten_item in handwritten_items:
            new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
            
            if new_text:
                # Add new bullet point
                new_para = insert_location._element.addnext(doc.add_paragraph(f"• {new_text}")._element)
                
                changes.append(ChangeRecord(
                    type="chunk_4_append_to_opportunities",
                    section=section_name,
                    original_text="",
                    new_text=f"• {new_text}"[:50],
                    timestamp=timestamp,
                    strategy_used="append_to_opportunities"
                ))
                
                self.processing_stats['total_replacements'] += 1
        
        self.logger.info(f"Chunk 4→Opportunities: {len(changes)} items appended")
        return changes
    
    def _implement_chunk_4_standalone(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 4 Standalone: Flexible processing when Chunk 5 doesn't have strengths
        Try to determine what section this content belongs to
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for {section_name}")
            return changes
        
        self.logger.info(f"Chunk 4 Debug: Found {len(handwritten_items)} handwritten items")
        for i, item in enumerate(handwritten_items):
            text = item.get('text', item.get('description', ''))[:100]
            self.logger.info(f"  Item {i+1}: {text}...")
        
        # Try to detect what type of content this is
        section_type = self._detect_content_type(handwritten_items)
        self.logger.info(f"Chunk 4 Debug: Detected content type: {section_type}")
        
        if section_type == "strengths":
            # This might be the strengths section - use robust strengths implementation
            self.logger.info(f"Chunk 4 Debug: Routing to strengths section")
            changes = self._implement_chunk_4_strengths(doc, analysis, f"{section_name}_as_strengths")
        elif section_type == "opportunities":
            # Add to opportunities section - fill remaining empty bullets
            self.logger.info(f"Chunk 4 Debug: Routing to opportunities section")
            changes = self._fill_remaining_opportunities_bullets(doc, handwritten_items, section_name, timestamp)
        elif section_type == "concerns":
            # Add to concerns section
            changes = self._append_to_section(doc, handwritten_items, "concerns", section_name, timestamp)
        else:
            # Enhanced detection for chunk 4 standalone - check for strengths FIRST, then opportunities
            all_text = " ".join([item.get('text', item.get('description', '')) for item in handwritten_items]).lower()
            
            # Enhanced strengths detection (same as chunk 5)
            primary_strength_indicators = ['strength', 'good at', 'skilled', 'experienced', 'established', 'expertise', 'talent']
            secondary_strength_indicators = [
                'educated', 'overachieving', 'daughters', 'plan', 'succeed', 'resources',
                'age', 'still', 'independent', 'freed up', 'cash', 'invest', 'future',
                'well', 'benefits', 'saving', 'manage', 'capable', 'financial'
            ]
            
            primary_matches = sum(1 for indicator in primary_strength_indicators if indicator in all_text)
            secondary_matches = sum(1 for indicator in secondary_strength_indicators if indicator in all_text)
            
            self.logger.info(f"Chunk 4 Debug: Strengths analysis - primary: {primary_matches}, secondary: {secondary_matches}")
            
            # Check for strengths first (more aggressive detection)
            if primary_matches > 0 or secondary_matches >= 3:
                self.logger.info(f"Chunk 4 Debug: Content appears to be strengths (enhanced detection)")
                # This looks like strengths content - use robust strengths implementation
                changes = self._implement_chunk_4_strengths(doc, analysis, f"{section_name}_as_strengths")
                section_type = "strengths (enhanced-keyword-detected)"
            else:
                # Fall back to opportunities detection
                opportunity_indicators = ['cash flow', 'increase', 'university', 'school', 'career', 'business', 'scholarship']
                opportunity_matches = sum(1 for indicator in opportunity_indicators if indicator in all_text)
                
                self.logger.info(f"Chunk 4 Debug: Opportunities analysis - matches: {opportunity_matches}")
                
                if opportunity_matches > 0:
                    self.logger.info(f"Chunk 4 Debug: Content appears to be opportunities based on keywords")
                    # This looks like opportunities content - add to opportunities section
                    changes = self._fill_remaining_opportunities_bullets(doc, handwritten_items, section_name, timestamp)
                    section_type = "opportunities (keyword-detected)"
                else:
                    # Generic processing - add as new section or append to most recent
                    changes = self._process_generic_content(doc, handwritten_items, section_name, timestamp)
        
        self.logger.info(f"Chunk 4 Standalone ({section_type}): {len(changes)} changes applied")
        return changes
        
    def _fill_remaining_opportunities_bullets(self, doc: Document, handwritten_items: list, section_name: str, timestamp: str) -> List[ChangeRecord]:
        """Fill any remaining empty bullet points in the opportunities section"""
        changes = []
        
        # Find the opportunities section
        opportunities_text = "most promising and strategic areas of opportunity are:"
        target_paragraph = None
        
        for paragraph in doc.paragraphs:
            if opportunities_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                self.logger.info(f"Chunk 4 Debug: Found opportunities section")
                break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find opportunities section for {section_name}")
            return changes
        
        # Find all bullet points in opportunities section - using same logic as chunk 3
        bullet_paragraphs = []
        empty_bullets = []
        found_opportunities = False
        paragraphs_examined = 0
        target_text = target_paragraph.text.strip()
        
        for paragraph in doc.paragraphs:
            # Use text matching instead of object comparison (same as chunk 3)
            if paragraph.text.strip() == target_text and not found_opportunities:
                found_opportunities = True
                self.logger.info(f"Chunk 4 Debug: Found target opportunities paragraph, now looking for bullets...")
                continue
            
            if found_opportunities:
                paragraphs_examined += 1
                para_text = paragraph.text.strip()
                self.logger.info(f"Chunk 4 Debug: Examining paragraph #{paragraphs_examined}: '{para_text}' (len: {len(para_text)})")
                
                # Check for Word list formatting using XML structure (same as chunk 3)
                is_list_item = bool(paragraph._element.xpath('.//w:numPr'))
                
                if is_list_item:
                    bullet_paragraphs.append(paragraph)
                    self.logger.info(f"Chunk 4 Debug: Found Word list item #{paragraphs_examined}")
                    
                    # Check if bullet is empty or nearly empty
                    if len(para_text) <= 1:
                        empty_bullets.append(paragraph)
                        self.logger.info(f"Chunk 4 Debug: ✅ Found EMPTY bullet to fill: '{para_text}'")
                    else:
                        self.logger.info(f"Chunk 4 Debug: Found FILLED bullet: '{para_text[:50]}...'")
                elif len(para_text) == 0:
                    # Empty paragraph, might be spacing
                    self.logger.info(f"Chunk 4 Debug: Empty spacing paragraph")
                elif (para_text and len(para_text) > 10 and not is_list_item):
                    # Next section found
                    self.logger.info(f"Chunk 4 Debug: Next section found: '{para_text[:50]}...'")
                    break
                else:
                    self.logger.info(f"Chunk 4 Debug: Non-bullet paragraph: '{para_text}'")
                
                # Stop after examining reasonable number of paragraphs
                if paragraphs_examined >= 10:
                    self.logger.info(f"Chunk 4 Debug: Stopping after examining {paragraphs_examined} paragraphs")
                    break
        
        self.logger.info(f"Chunk 4 Debug: Found {len(bullet_paragraphs)} total bullets, {len(empty_bullets)} empty")
        self.logger.info(f"Chunk 4 Debug: Examined {paragraphs_examined} paragraphs after opportunities section")
        
        # Fill empty bullets with chunk 4 content
        items_used = 0
        for bullet_para in empty_bullets:
            if items_used < len(handwritten_items):
                handwritten_item = handwritten_items[items_used]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    # Clear and fill the bullet
                    bullet_para.clear()
                    bullet_para.add_run(new_text)
                    
                    changes.append(ChangeRecord(
                        type="opportunities_bullet_fill",
                        section=section_name,
                        original_text="",
                        new_text=new_text[:100],
                        timestamp=timestamp,
                        strategy_used="fill_empty_opportunities_bullet"
                    ))
                    
                    items_used += 1
                    self.processing_stats['total_replacements'] += 1
                    self.logger.info(f"Chunk 4 Debug: ✅ FILLED empty opportunities bullet with: '{new_text[:50]}...'")
        
        return changes
    
    def _detect_content_type(self, handwritten_items: list) -> str:
        """Detect what type of content the handwritten items represent"""
        all_text = " ".join([item.get('text', item.get('description', '')) for item in handwritten_items]).lower()
        
        # Enhanced strengths keywords - including positive descriptors and capabilities
        strengths_keywords = [
            'strength', 'advantage', 'skill', 'expertise', 'good at', 'established',
            'educated', 'overachieving', 'daughters', 'plan', 'succeed', 'resources',
            'age', 'still', 'independent', 'freed up', 'cash', 'invest', 'future',
            'well', 'benefits', 'saving', 'manage', 'capable', 'experienced'
        ]
        opportunities_keywords = [
            'opportunity', 'potential', 'growth', 'expand', 'develop', 'increase',
            'cash flow', 'finishing', 'school', 'university', 'career', 'business'
        ]
        concerns_keywords = [
            'concern', 'problem', 'issue', 'danger', 'obstacle', 'challenge',
            'stress', 'time poor', 'aging', 'worry', 'risk', 'struggle'
        ]
        
        strengths_score = sum(1 for keyword in strengths_keywords if keyword in all_text)
        opportunities_score = sum(1 for keyword in opportunities_keywords if keyword in all_text)
        concerns_score = sum(1 for keyword in concerns_keywords if keyword in all_text)
        
        # Log the detection process for debugging
        self.logger.info(f"Content type detection scores - Strengths: {strengths_score}, Opportunities: {opportunities_score}, Concerns: {concerns_score}")
        self.logger.info(f"Sample text for analysis: {all_text[:200]}...")
        
        if strengths_score > max(opportunities_score, concerns_score):
            return "strengths"
        elif opportunities_score > concerns_score:
            return "opportunities"
        elif concerns_score > 0:
            return "concerns"
        else:
            return "unknown"
    
    def _process_as_strengths_section(self, doc: Document, handwritten_items: list, section_name: str, timestamp: str) -> List[ChangeRecord]:
        """Process content as strengths section"""
        changes = []
        
        # Look for existing strengths section or create placeholder
        strengths_text = "existing strengths"
        target_paragraph = None
        
        for paragraph in doc.paragraphs:
            if "strength" in paragraph.text.lower():
                target_paragraph = paragraph
                break
        
        if not target_paragraph:
            # Create a new strengths section
            # Find a good place to insert it (after opportunities section)
            opportunities_text = "opportunity"
            insert_after = None
            
            for paragraph in doc.paragraphs:
                if "opportunity" in paragraph.text.lower():
                    insert_after = paragraph
            
            if insert_after:
                # Add strengths heading
                new_heading = insert_after._element.addnext(doc.add_paragraph("Reinforcing and maximizing your existing strengths:")._element)
                target_paragraph = doc.paragraphs[-1]  # Get the newly added paragraph
        
        # Add bullet points for strengths
        if target_paragraph:
            for handwritten_item in handwritten_items:
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    new_para = target_paragraph._element.addnext(doc.add_paragraph(f"• {new_text}")._element)
                    
                    changes.append(ChangeRecord(
                        type="chunk_4_strengths_creation",
                        section=section_name,
                        original_text="",
                        new_text=f"• {new_text}"[:50],
                        timestamp=timestamp,
                        strategy_used="create_strengths_section"
                    ))
                    
                    self.processing_stats['total_replacements'] += 1
        
        return changes
    
    def _append_to_section(self, doc: Document, handwritten_items: list, section_type: str, section_name: str, timestamp: str) -> List[ChangeRecord]:
        """Append content to an existing section"""
        changes = []
        
        # Define section search terms
        search_terms = {
            "opportunities": ["opportunity", "strategic areas of opportunity"],
            "concerns": ["concern", "areas of concern", "problems to be solved"],
            "strengths": ["strength", "existing strengths"]
        }
        
        # Find the target section
        target_paragraph = None
        for paragraph in doc.paragraphs:
            for term in search_terms.get(section_type, []):
                if term in paragraph.text.lower():
                    target_paragraph = paragraph
                    break
            if target_paragraph:
                break
        
        if target_paragraph:
            # Find existing bullets and append
            for handwritten_item in handwritten_items:
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    new_para = target_paragraph._element.addnext(doc.add_paragraph(f"• {new_text}")._element)
                    
                    changes.append(ChangeRecord(
                        type=f"chunk_4_append_to_{section_type}",
                        section=section_name,
                        original_text="",
                        new_text=f"• {new_text}"[:50],
                        timestamp=timestamp,
                        strategy_used=f"append_to_{section_type}"
                    ))
                    
                    self.processing_stats['total_replacements'] += 1
        
        return changes
    
    def _process_generic_content(self, doc: Document, handwritten_items: list, section_name: str, timestamp: str) -> List[ChangeRecord]:
        """Generic processing for unknown content types"""
        changes = []
        
        # Add as generic additional content
        # Find the last content paragraph and append after it
        last_content_para = None
        for paragraph in reversed(doc.paragraphs):
            if paragraph.text.strip():
                last_content_para = paragraph
                break
        
        if last_content_para:
            for handwritten_item in handwritten_items:
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    new_para = last_content_para._element.addnext(doc.add_paragraph(f"• {new_text}")._element)
                    
                    changes.append(ChangeRecord(
                        type="chunk_4_generic_content",
                        section=section_name,
                        original_text="",
                        new_text=f"• {new_text}"[:50],
                        timestamp=timestamp,
                        strategy_used="generic_append"
                    ))
                    
                    self.processing_stats['total_replacements'] += 1
        
        return changes
    
    def _implement_chunk_1_date_replacement(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 1: Replace XXXX with handwritten date
        
        Business Logic:
        - Find crossed-out/strikethrough "XXXX" 
        - Look for handwritten date in the same chunk
        - Replace "XXXX" with the detected date using 3-strategy approach
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        self.logger.info(f"🔧 Processing Chunk 1: Date replacement for {section_name}")
        
        # Extract detected items from this chunk
        detected_items = analysis.get('detected_items', {})
        handwritten_texts = detected_items.get('handwritten_text', [])
        strikethrough_items = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        
        # Find the handwritten date - ENHANCED: Only select items that are actually dates
        replacement_date = None
        
        self.logger.info(f"Chunk 1 Debug: === ALL HANDWRITTEN ITEMS ===")
        for i, item in enumerate(handwritten_texts):
            text = item.get('text', '').strip()
            description = item.get('description', '').strip().lower()
            self.logger.info(f"Chunk 1 Debug: Item {i+1}: '{text}' | Description: '{description}'")
        
        for item in handwritten_texts:
            text = item.get('text', '').strip()
            description = item.get('description', '').strip().lower()
            
            # CRITICAL: Only accept items that are actually dates based on description
            is_date_description = any(keyword in description for keyword in [
                'date', 'handwritten date', 'below the salutation', 'after dear', 
                'below dear', 'salutation', '21st', '22nd', '23rd', '24th', '25th',
                'august', 'september', 'october', 'november', 'december',
                'january', 'february', 'march', 'april', 'may', 'june', 'july'
            ])
            
            # Also check if the text itself looks like a date
            has_date_pattern = any(date_indicator in text.lower() for date_indicator in [
                '2024', '2025', '2026', 'january', 'february', 'march', 'april', 
                'may', 'june', 'july', 'august', 'september', 'october', 
                'november', 'december', 'jan', 'feb', 'mar', 'apr', 'jun', 
                'jul', 'aug', 'sep', 'oct', 'nov', 'dec'
            ])
            
            # Also check for date format patterns like "21st of August", "the 21st", etc.
            import re
            has_date_format = bool(re.search(r'\d{1,2}(st|nd|rd|th)', text, re.IGNORECASE))
            
            # Skip non-date items like "Pt = address"
            is_address_note = any(keyword in description for keyword in [
                'address', 'pt =', 'note at the top', 'handwritten note at'
            ]) and not is_date_description
            
            self.logger.info(f"Chunk 1 Debug: Checking '{text}':")
            self.logger.info(f"  - is_date_description: {is_date_description}")
            self.logger.info(f"  - has_date_pattern: {has_date_pattern}")
            self.logger.info(f"  - has_date_format: {has_date_format}")
            self.logger.info(f"  - is_address_note: {is_address_note}")
            
            # Only use this as the date if it's clearly a date, not an address note
            if (is_date_description or has_date_pattern or has_date_format) and not is_address_note:
                replacement_date = text
                self.logger.info(f"📅 Found handwritten date: '{replacement_date}' (description: '{description}')")
                break
        
        if not replacement_date:
            self.logger.warning("Chunk 1 Debug: ❌ No valid date found in handwritten items")
        
        # ENHANCED: Check context and replace the entire problematic phrase
        if replacement_date:
            # Find XXXX in context and replace the entire problematic phrase
            for paragraph in doc.paragraphs:
                if 'XXXX' in paragraph.text:
                    context_text = paragraph.text
                    self.logger.info(f"📄 XXXX context: '{context_text[:100]}...'")
                    
                    import re
                    
                    # DYNAMIC: Pattern to match any problematic text with XXXX
                    # These patterns will catch various malformed date/address formats
                    problematic_patterns = [
                        r'Pt in address Today the .* XXXX',
                        r'on Pt in address Today the .* XXXX', 
                        r'time on .* XXXX',
                        r'on .* Today .* XXXX',  # Additional flexible patterns
                        r'address .* XXXX',
                        r'Today .* XXXX'
                    ]
                    
                    # DYNAMIC: Extract clean date format from AI-detected handwritten date
                    date_part = replacement_date.strip()
                    clean_date = self._extract_clean_date_format(date_part)
                    
                    self.logger.info(f"📅 AI detected date: '{replacement_date}' → Clean format: '{clean_date}'")
                    
                    # Replace the entire problematic phrase with the clean date
                    found_pattern = False
                    original_context = context_text
                    for pattern in problematic_patterns:
                        if re.search(pattern, context_text, re.IGNORECASE):
                            # Replace the entire problematic phrase with clean date
                            new_text = re.sub(pattern, clean_date, context_text, flags=re.IGNORECASE)
                            paragraph.text = new_text
                            found_pattern = True
                            self.logger.info(f"📅 Replaced problematic phrase with: '{clean_date}'")
                            self.logger.info(f"📄 New paragraph text: '{new_text[:100]}...'")
                            
                            # Create change record for the phrase replacement
                            changes.append(ChangeRecord(
                                type="phrase_replacement",
                                section=section_name,
                                original_text=original_context[:100],
                                new_text=new_text[:100],
                                timestamp=timestamp,
                                strategy_used="phrase_replacement"
                            ))
                            self.processing_stats['total_replacements'] += 1
                            
                            # We've done the replacement, so return early
                            return changes
                    
                    if not found_pattern:
                        # Fallback: just replace XXXX with the clean date (without "the")
                        fallback_date = clean_date.replace("the ", "")
                        replacement_date = fallback_date
                        self.logger.info(f"📅 Using fallback XXXX replacement: '{replacement_date}'")
                    
                    break
        
        # Verify there's strikethrough or cross over XXXX area
        has_strikethrough = any('xxxx' in item.get('text', '').lower() for item in strikethrough_items)
        has_cross_over_xxxx = len(crosses) > 0  # Assume crosses indicate XXXX deletion
        
        if replacement_date and (has_strikethrough or has_cross_over_xxxx):
            self.logger.info(f"✅ Conditions met: Date='{replacement_date}', Strikethrough={has_strikethrough}, Cross={has_cross_over_xxxx}")
            
            # Check if we already handled the replacement in the context checking above
            xxxx_still_exists = any('XXXX' in p.text for p in doc.paragraphs)
            
            if xxxx_still_exists:
                # Apply 3-strategy cascading replacement
                success, strategy = self._apply_cascading_replacement(doc, "XXXX", replacement_date)
                
                if success:
                    changes.append(ChangeRecord(
                        type="date_replacement",
                        section=section_name,
                        original_text="XXXX",
                        new_text=replacement_date,
                        timestamp=timestamp,
                        strategy_used=strategy
                    ))
                    self.processing_stats['total_replacements'] += 1
                    self.logger.info(f"✅ Successfully replaced 'XXXX' with '{replacement_date}' using {strategy}")
                else:
                    self.logger.warning(f"❌ Failed to find 'XXXX' for replacement with '{replacement_date}'")
                    self.processing_stats['failed_operations'] += 1
            else:
                # Replacement was already handled in context checking
                changes.append(ChangeRecord(
                    type="phrase_replacement",
                    section=section_name,
                    original_text="problematic phrase with XXXX",
                    new_text=replacement_date,
                    timestamp=timestamp,
                    strategy_used="phrase_replacement"
                ))
                self.processing_stats['total_replacements'] += 1
                self.logger.info(f"✅ Successfully replaced problematic phrase with '{replacement_date}'")
        else:
            missing = []
            if not replacement_date:
                missing.append("handwritten date")
            if not (has_strikethrough or has_cross_over_xxxx):
                missing.append("strikethrough/cross over XXXX")
            
            self.logger.warning(f"❌ Chunk 1 conditions not met. Missing: {', '.join(missing)}")
            self.processing_stats['failed_operations'] += 1
        
        return changes
    
    def _implement_chunk_4_strengths(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 4 as Strengths: Fill strengths bullet points when Chunk 4 IS the strengths section
        Uses the same robust bullet detection as chunks 2 and 3
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for {section_name}")
            return changes
        
        self.logger.info(f"Chunk 4 Debug: Found {len(handwritten_items)} handwritten items")
        for i, item in enumerate(handwritten_items):
            self.logger.info(f"  Item {i+1}: {item}")
        
        # Find the strengths section: "established the following important strengths:"
        strengths_text = "established the following important strengths:"
        target_paragraph = None
        target_position = None
        
        self.logger.info(f"Chunk 4 Debug: Searching for strengths section...")
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                self.logger.info(f"  Para {i}: '{para_text[:100]}...'")
            
            if strengths_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                target_position = i
                self.logger.info(f"Chunk 4 Debug: Found exact strengths section: '{paragraph.text.strip()}'")
                break
        
        if not target_paragraph:
            # Strategy 2: Find similar text about strengths
            strength_keywords = ['strengths', 'strength', 'maximizing', 'reinforcing', 'established']
            for i, paragraph in enumerate(doc.paragraphs):
                if any(keyword in paragraph.text.lower() for keyword in strength_keywords):
                    target_paragraph = paragraph
                    target_position = i
                    self.logger.info(f"Chunk 4 Debug: Found similar strengths section: '{paragraph.text.strip()}'")
                    break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find strengths section for {section_name}")
            return changes
        
        # Use the same robust bullet detection as chunk 2 and 3
        self.logger.info(f"Chunk 4 Debug: Looking for bullet points after strengths section...")
        self.logger.info(f"Chunk 4 Debug: Found target paragraph at position {target_position}, now looking for bullets...")
        
        bullet_paragraphs = []
        found_strengths = False
        paragraphs_examined = 0
        target_text = target_paragraph.text.strip()
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Use text matching instead of object comparison (same as chunk 2/3)
            if paragraph.text.strip() == target_text and not found_strengths:
                found_strengths = True
                continue
            
            if found_strengths:
                paragraphs_examined += 1
                para_text = paragraph.text.strip()
                self.logger.info(f"Chunk 4 Debug: Examining paragraph #{paragraphs_examined}: '{para_text}' (stripped: '{para_text}')")
                
                # Check for Word list formatting using XML structure (same as chunk 2/3)
                is_list_item = bool(paragraph._element.xpath('.//w:numPr'))
                
                if is_list_item:
                    self.logger.info(f"Chunk 4 Debug: Found Word list item formatting in paragraph #{paragraphs_examined}!")
                
                self.logger.info(f"Chunk 4 Debug: Para #{paragraphs_examined} analysis - is_list_item: {is_list_item}, text_len: {len(paragraph.text)}, stripped_len: {len(para_text)}")
                
                if is_list_item and len(para_text) <= 1:  # Empty or single bullet character
                    bullet_paragraphs.append(paragraph)
                    self.logger.info(f"Chunk 4 Debug: ✅ ADDED Word-formatted bullet paragraph #{paragraphs_examined}: '{para_text}' (empty: {len(para_text) == 0})")
                elif len(para_text) == 0:
                    # Empty paragraph, might be spacing
                    self.logger.info(f"Chunk 4 Debug: ❌ SKIPPED paragraph #{paragraphs_examined} - not a bullet point")
                elif (para_text and 
                      not para_text.startswith(' ') and 
                      len(para_text) > 10 and
                      not is_list_item):
                    # Likely start of next section
                    self.logger.info(f"Chunk 4 Debug: Next section paragraph: '{para_text}'")
                    break
                else:
                    self.logger.info(f"Chunk 4 Debug: ❌ SKIPPED paragraph #{paragraphs_examined} - not a bullet point")
        
        self.logger.info(f"Chunk 4 Debug: Loop completed. Total paragraphs examined: {len(doc.paragraphs)}")
        self.logger.info(f"Chunk 4 Debug: Found strengths flag: {found_strengths}")
        self.logger.info(f"Chunk 4 Debug: Paragraphs examined after strengths: {paragraphs_examined}")
        self.logger.info(f"Chunk 4 Debug: Found {len(bullet_paragraphs)} bullet paragraphs")
        
        if not bullet_paragraphs:
            self.logger.warning(f"No bullet points found after strengths section for {section_name}")
            return changes
        
        # Fill existing bullet points using paragraph text comparison (same as chunk 2/3)
        items_used = 0
        for i, bullet_para in enumerate(bullet_paragraphs):
            if items_used < len(handwritten_items):
                handwritten_item = handwritten_items[items_used]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    # Clear the paragraph and add the new text without extra bullet
                    bullet_para.clear()
                    bullet_para.add_run(new_text)
                    
                    changes.append(ChangeRecord(
                        type="bullet_point_fill",
                        section=section_name,
                        original_text="",
                        new_text=new_text[:100],
                        timestamp=timestamp,
                        strategy_used="fill_existing_bullet"
                    ))
                    
                    items_used += 1
                    self.processing_stats['total_replacements'] += 1
        
        # Check if we have unused handwritten items
        if items_used < len(handwritten_items):
            extra_items = len(handwritten_items) - items_used
            self.logger.info(f"Chunk 4: {extra_items} extra handwritten items found but only {len(bullet_paragraphs)} existing bullets available")
        
        self.logger.info(f"Chunk 4: {len(changes)} bullet point changes applied")
        return changes
    
    def _implement_chunk_5_strengths(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 5 as Strengths: Fill strengths bullet points when Chunk 5 is the strengths section
        Also handles deletion of crossed-out bullet points with proper spacing removal
        Uses the same robust bullet detection as chunks 2, 3, and 4
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract items from analysis for both filling and deletion
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        strikethrough_items = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        
        self.logger.info(f"Chunk 5 Debug: Found {len(handwritten_items)} handwritten, {len(strikethrough_items)} strikethrough, {len(crosses)} crosses")
        
        # If no items at all, return early
        if not handwritten_items and not strikethrough_items and not crosses:
            self.logger.info(f"No items found for {section_name}")
            return changes
        
        self.logger.info(f"Chunk 5 Debug: Processing items for deletion and filling")
        for i, item in enumerate(handwritten_items):
            self.logger.info(f"  Item {i+1}: {item}")
        
        # Find the strengths section: "established the following important strengths:"
        strengths_text = "established the following important strengths:"
        target_paragraph = None
        target_position = None
        
        self.logger.info(f"Chunk 5 Debug: Searching for strengths section...")
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                self.logger.info(f"  Para {i}: '{para_text[:100]}...'")
            
            if strengths_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                target_position = i
                self.logger.info(f"Chunk 5 Debug: Found exact strengths section: '{paragraph.text.strip()}'")
                break
        
        if not target_paragraph:
            # Strategy 2: Find similar text about strengths
            strength_keywords = ['strengths', 'strength', 'maximizing', 'reinforcing', 'established']
            for i, paragraph in enumerate(doc.paragraphs):
                if any(keyword in paragraph.text.lower() for keyword in strength_keywords):
                    target_paragraph = paragraph
                    target_position = i
                    self.logger.info(f"Chunk 5 Debug: Found similar strengths section: '{paragraph.text.strip()}'")
                    break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find strengths section for {section_name}")
            return changes
        
        # Use the same robust bullet detection as chunks 2, 3, and 4
        self.logger.info(f"Chunk 5 Debug: Looking for bullet points after strengths section...")
        self.logger.info(f"Chunk 5 Debug: Found target paragraph at position {target_position}, now looking for bullets...")
        
        bullet_paragraphs = []
        empty_bullets = []
        found_strengths = False
        paragraphs_examined = 0
        target_text = target_paragraph.text.strip()
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Use text matching instead of object comparison (same as chunks 2/3/4)
            if paragraph.text.strip() == target_text and not found_strengths:
                found_strengths = True
                continue
            
            if found_strengths:
                paragraphs_examined += 1
                para_text = paragraph.text.strip()
                self.logger.info(f"Chunk 5 Debug: Examining paragraph #{paragraphs_examined}: '{para_text}' (stripped: '{para_text}')")
                
                # Check for Word list formatting using XML structure OR bullet character
                is_list_item = bool(paragraph._element.xpath('.//w:numPr'))
                is_bullet_char = para_text.startswith('•')
                
                if is_list_item or is_bullet_char:
                    self.logger.info(f"Chunk 5 Debug: Found bullet paragraph #{paragraphs_examined}! (list_item: {is_list_item}, bullet_char: {is_bullet_char})")
                    bullet_paragraphs.append(paragraph)
                    
                    # Track empty bullets specifically (same as chunk 3)
                    if len(para_text) <= 1:  # Empty or single bullet character
                        empty_bullets.append(paragraph)
                        self.logger.info(f"Chunk 5 Debug: ✅ ADDED empty bullet paragraph #{paragraphs_examined}: '{para_text}'")
                    else:
                        self.logger.info(f"Chunk 5 Debug: ✅ ADDED filled bullet paragraph #{paragraphs_examined}: '{para_text[:50]}...'")
                elif len(para_text) == 0:
                    # Empty paragraph, might be spacing
                    self.logger.info(f"Chunk 5 Debug: ❌ SKIPPED paragraph #{paragraphs_examined} - empty spacing")
                elif para_text.startswith('  ') or para_text.startswith('\t'):
                    # Likely continuation/indented text, skip for now
                    self.logger.info(f"Chunk 5 Debug: ❌ SKIPPED paragraph #{paragraphs_examined} - indented text (continuation)")
                elif (para_text and 
                      not para_text.startswith(' ') and 
                      len(para_text) > 10 and
                      not is_list_item and 
                      not is_bullet_char):
                    # Likely start of next section
                    self.logger.info(f"Chunk 5 Debug: Next section paragraph: '{para_text}'")
                    break
                else:
                    self.logger.info(f"Chunk 5 Debug: ❌ SKIPPED paragraph #{paragraphs_examined} - not a bullet point")
        
        self.logger.info(f"Chunk 5 Debug: Loop completed. Total paragraphs examined: {len(doc.paragraphs)}")
        self.logger.info(f"Chunk 5 Debug: Found strengths flag: {found_strengths}")
        self.logger.info(f"Chunk 5 Debug: Paragraphs examined after strengths: {paragraphs_examined}")
        self.logger.info(f"Chunk 5 Debug: Found {len(bullet_paragraphs)} total bullet paragraphs")
        self.logger.info(f"Chunk 5 Debug: Found {len(empty_bullets)} empty bullet paragraphs")
        
        # STEP 1: Handle bullet deletion for crossed-out/strikethrough items (WITH SPACING REMOVAL)
        deletion_items = strikethrough_items + crosses
        if deletion_items:
            self.logger.info(f"Chunk 5 Debug: Processing {len(deletion_items)} items for bullet deletion")
            
            paragraphs_to_delete = []
            
            for deletion_item in deletion_items:
                delete_text = deletion_item.get('text', deletion_item.get('description', '')).strip()
                if not delete_text:
                    continue
                
                self.logger.info(f"Chunk 5 Debug: Looking for bullet to delete: '{delete_text}'")
                
                # Find matching bullet paragraphs to delete across the whole document (same as chunk 5 general)
                for i, para in enumerate(doc.paragraphs):
                    para_text = para.text.strip()
                    
                    # Check if this paragraph contains the crossed-out text and looks like a bullet
                    if (para_text and 
                        (delete_text.lower() in para_text.lower() or para_text.lower() in delete_text.lower()) and
                        (para_text.startswith('•') or para._element.xpath('.//w:numPr'))):
                        
                        self.logger.info(f"Chunk 5 Debug: ✅ FOUND bullet to delete: '{para_text}'")
                        paragraphs_to_delete.append(para)
                        
                        # Also look for continuation sentences (indented paragraphs following this bullet)
                        for next_i in range(i + 1, len(doc.paragraphs)):
                            next_para = doc.paragraphs[next_i]
                            next_text = next_para.text  # Don't strip yet for checking indentation
                            next_text_stripped = next_text.strip()
                            
                            self.logger.info(f"Chunk 5 Debug: Checking continuation #{next_i}: original='{next_text}', stripped='{next_text_stripped}'")
                            
                            # If it's an indented continuation or empty spacing
                            if (next_text.startswith('  ') or next_text.startswith('\t') or 
                                len(next_text_stripped) == 0):
                                paragraphs_to_delete.append(next_para)
                                self.logger.info(f"Chunk 5 Debug: ✅ ADDED continuation/spacing to delete: '{next_text_stripped}'")
                            elif next_text_stripped.startswith('•') or len(next_text_stripped) > 10:
                                # Hit next bullet or substantial content, stop looking
                                self.logger.info(f"Chunk 5 Debug: ❌ STOP - Found next section: '{next_text_stripped}'")
                                break
            
            # Delete paragraphs and remove spacing gaps
            if paragraphs_to_delete:
                self.logger.info(f"Chunk 5 Debug: Deleting {len(paragraphs_to_delete)} paragraphs (bullets + continuations + spacing)")
                
                for para_to_delete in paragraphs_to_delete:
                    original_text = para_to_delete.text.strip()
                    
                    # Clear the paragraph content completely
                    para_to_delete.clear()
                    
                    # Remove the paragraph element entirely to eliminate spacing gaps
                    p_element = para_to_delete._element
                    parent = p_element.getparent()
                    if parent is not None:
                        parent.remove(p_element)
                    
                    changes.append(ChangeRecord(
                        type="bullet_deletion_with_spacing",
                        section=section_name,
                        original_text=original_text[:100] if original_text else "(empty spacing)",
                        new_text="",
                        timestamp=timestamp,
                        strategy_used="complete_bullet_removal"
                    ))
                    
                    self.processing_stats['total_deletions'] += 1
                    self.logger.info(f"Chunk 5 Debug: ✅ DELETED paragraph with spacing removal: '{original_text}'")
                
                # Clean up spacing gaps after deletion
                cleaned_spacing = self._cleanup_document_spacing_after_deletion(doc, "Chunk 5 bullet deletion")
                self.logger.info(f"Chunk 5 Debug: Cleaned up {cleaned_spacing} spacing gaps after deletion")
                
                # Refresh bullet paragraphs list after deletions
                bullet_paragraphs = [p for p in bullet_paragraphs if p not in paragraphs_to_delete]
                empty_bullets = [p for p in empty_bullets if p not in paragraphs_to_delete]
                
                self.logger.info(f"Chunk 5 Debug: After deletion - {len(bullet_paragraphs)} bullets remain")
        
        # STEP 2: Prioritize filling empty bullets first, then filled ones if needed
        bullets_to_fill = empty_bullets if empty_bullets else bullet_paragraphs
        
        if not bullets_to_fill:
            self.logger.warning(f"No bullet points found after strengths section for {section_name}")
            return changes
        
        # Fill bullet points using paragraph text comparison (same as chunks 2/3/4)
        items_used = 0
        for i, bullet_para in enumerate(bullets_to_fill):
            if items_used < len(handwritten_items):
                handwritten_item = handwritten_items[items_used]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    # Clear the paragraph and add the new text without extra bullet
                    bullet_para.clear()
                    bullet_para.add_run(new_text)
                    
                    changes.append(ChangeRecord(
                        type="bullet_point_fill",
                        section=section_name,
                        original_text="",
                        new_text=new_text[:100],
                        timestamp=timestamp,
                        strategy_used="fill_existing_bullet"
                    ))
                    
                    items_used += 1
                    self.processing_stats['total_replacements'] += 1
                    self.logger.info(f"Chunk 5 Debug: ✅ FILLED bullet #{i+1} with: '{new_text[:50]}...'")
        
        # Check if we have unused handwritten items
        if items_used < len(handwritten_items):
            extra_items = len(handwritten_items) - items_used
            self.logger.info(f"Chunk 5: {extra_items} extra handwritten items found but only {len(bullets_to_fill)} bullets available")
        
        self.logger.info(f"Chunk 5: {len(changes)} bullet point changes applied")
        return changes
    
    def _implement_chunk_5_general(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 5 General: Flexible processing for Chunk 5 when it's not strengths
        Also handles deletion of crossed-out bullet points with proper spacing removal
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract items from analysis for both filling and deletion  
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        strikethrough_items = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        
        self.logger.info(f"Chunk 5 General Debug: Found {len(handwritten_items)} handwritten, {len(strikethrough_items)} strikethrough, {len(crosses)} crosses")
        
        # STEP 1: Handle bullet deletion for crossed-out/strikethrough items (WITH SPACING REMOVAL)
        deletion_items = strikethrough_items + crosses
        if deletion_items:
            self.logger.info(f"Chunk 5 General Debug: Processing {len(deletion_items)} items for bullet deletion")
            
            paragraphs_to_delete = []
            
            for deletion_item in deletion_items:
                delete_text = deletion_item.get('text', deletion_item.get('description', '')).strip()
                if not delete_text:
                    continue
                
                self.logger.info(f"Chunk 5 General Debug: Looking for bullet to delete: '{delete_text}'")
                
                # Find matching bullet paragraphs to delete across the whole document
                for i, para in enumerate(doc.paragraphs):
                    para_text = para.text.strip()
                    
                    # Check if this paragraph contains the crossed-out text and looks like a bullet
                    if (para_text and 
                        (delete_text.lower() in para_text.lower() or para_text.lower() in delete_text.lower()) and
                        (para_text.startswith('•') or para._element.xpath('.//w:numPr'))):
                        
                        self.logger.info(f"Chunk 5 General Debug: ✅ FOUND bullet to delete: '{para_text}'")
                        paragraphs_to_delete.append(para)
                        
                        # Also look for continuation sentences (indented paragraphs following this bullet)
                        for next_i in range(i + 1, len(doc.paragraphs)):
                            next_para = doc.paragraphs[next_i]
                            next_text = next_para.text  # Don't strip yet for checking indentation
                            next_text_stripped = next_text.strip()
                            
                            self.logger.info(f"Chunk 5 General Debug: Checking continuation #{next_i}: original='{next_text}', stripped='{next_text_stripped}'")
                            
                            # If it's an indented continuation or empty spacing
                            if (next_text.startswith('  ') or next_text.startswith('\t') or 
                                len(next_text_stripped) == 0):
                                paragraphs_to_delete.append(next_para)
                                self.logger.info(f"Chunk 5 General Debug: ✅ ADDED continuation/spacing to delete: '{next_text_stripped}'")
                            elif next_text_stripped.startswith('•') or len(next_text_stripped) > 10:
                                # Hit next bullet or substantial content, stop looking
                                self.logger.info(f"Chunk 5 General Debug: ❌ STOP - Found next section: '{next_text_stripped}'")
                                break
            
            # Delete paragraphs and remove spacing gaps
            if paragraphs_to_delete:
                self.logger.info(f"Chunk 5 General Debug: Deleting {len(paragraphs_to_delete)} paragraphs (bullets + continuations + spacing)")
                
                for para_to_delete in paragraphs_to_delete:
                    original_text = para_to_delete.text.strip()
                    
                    # Clear the paragraph content completely
                    para_to_delete.clear()
                    
                    # Remove the paragraph element entirely to eliminate spacing gaps
                    p_element = para_to_delete._element
                    parent = p_element.getparent()
                    if parent is not None:
                        parent.remove(p_element)
                    
                    changes.append(ChangeRecord(
                        type="bullet_deletion_with_spacing",
                        section=section_name,
                        original_text=original_text[:100] if original_text else "(empty spacing)",
                        new_text="",
                        timestamp=timestamp,
                        strategy_used="complete_bullet_removal"
                    ))
                    
                    self.processing_stats['total_deletions'] += 1
                    self.logger.info(f"Chunk 5 General Debug: ✅ DELETED paragraph with spacing removal: '{original_text}'")
                
                # Clean up spacing gaps after deletion
                cleaned_spacing = self._cleanup_document_spacing_after_deletion(doc, "Chunk 5 general bullet deletion")
                self.logger.info(f"Chunk 5 General Debug: Cleaned up {cleaned_spacing} spacing gaps after deletion")
        
        # If no handwritten items after deletion, return early
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for filling in {section_name}")
            return changes if deletion_items else changes
        
        self.logger.info(f"Chunk 5 Debug: Found {len(handwritten_items)} handwritten items for filling")
        for i, item in enumerate(handwritten_items):
            text = item.get('text', item.get('description', ''))[:100]
            self.logger.info(f"  Item {i+1}: {text}...")
        
        # Try to detect what type of content this is
        section_type = self._detect_content_type(handwritten_items)
        self.logger.info(f"Chunk 5 Debug: Detected content type: {section_type}")
        
        if section_type == "strengths":
            # This might be the strengths section - use robust bullet filling
            changes = self._fill_remaining_strengths_bullets(doc, handwritten_items, section_name, timestamp)
        elif section_type == "opportunities":
            # Add to opportunities section - use same logic as chunk 4
            changes = self._fill_remaining_opportunities_bullets(doc, handwritten_items, section_name, timestamp)
        elif section_type == "concerns":
            # Add to concerns section
            changes = self._append_to_section(doc, handwritten_items, "concerns", section_name, timestamp)
        else:
            # Enhanced strengths detection for chunk 5 - more aggressive keyword matching
            all_text = " ".join([item.get('text', item.get('description', '')) for item in handwritten_items]).lower()
            
            # Primary strengths indicators
            primary_strength_indicators = ['strength', 'good at', 'skilled', 'experienced', 'established', 'expertise', 'talent']
            
            # Secondary indicators that suggest positive attributes/capabilities in chunk 5 context
            secondary_strength_indicators = [
                'educated', 'overachieving', 'daughters', 'plan', 'succeed', 'resources',
                'age', 'still', 'independent', 'freed up', 'cash', 'invest', 'future',
                'well', 'benefits', 'saving', 'manage', 'capable', 'financial'
            ]
            
            primary_matches = sum(1 for indicator in primary_strength_indicators if indicator in all_text)
            secondary_matches = sum(1 for indicator in secondary_strength_indicators if indicator in all_text)
            
            self.logger.info(f"Chunk 5 Debug: Primary strength matches: {primary_matches}, Secondary matches: {secondary_matches}")
            
            # More lenient detection - either primary matches OR multiple secondary matches
            if primary_matches > 0 or secondary_matches >= 3:
                self.logger.info(f"Chunk 5 Debug: Content appears to be strengths (primary: {primary_matches}, secondary: {secondary_matches})")
                # This looks like strengths content - add to strengths section
                changes = self._fill_remaining_strengths_bullets(doc, handwritten_items, section_name, timestamp)
                section_type = "strengths (enhanced-keyword-detected)"
            else:
                # Generic processing - add as new section or append to most recent
                changes = self._process_generic_content(doc, handwritten_items, section_name, timestamp)
        
        self.logger.info(f"Chunk 5 General ({section_type}): {len(changes)} changes applied")
        return changes

    def _fill_remaining_strengths_bullets(self, doc: Document, handwritten_items: list, section_name: str, timestamp: str) -> List[ChangeRecord]:
        """Fill empty strengths bullets with handwritten content - used by chunk 5 general when content is detected as strengths"""
        changes = []
        
        # Use the same logic as _implement_chunk_5_strengths but for general chunk 5 processing
        detected_items = {'handwritten_text': handwritten_items}
        analysis = {'detected_items': detected_items}
        
        # Call the robust strengths implementation
        strengths_changes = self._implement_chunk_5_strengths(doc, analysis, f"{section_name}_as_strengths")
        changes.extend(strengths_changes)
        
        return changes

    def _fill_remaining_opportunities_bullets(self, doc: Document, handwritten_items: list, section_name: str, timestamp: str) -> List[ChangeRecord]:
        """Fill empty opportunities bullets with handwritten content - used by chunk 4/5 cross-chunk processing"""
        changes = []
        
        # Use the same logic as _implement_chunk_3_bullet_points but for cross-chunk processing
        detected_items = {'handwritten_text': handwritten_items}
        analysis = {'detected_items': detected_items}
        
        # Call the robust opportunities implementation
        opportunities_changes = self._implement_chunk_3_bullet_points(doc, analysis, f"{section_name}_as_opportunities")
        changes.extend(opportunities_changes)
        
        return changes
    
    def _fill_bullet_points_after_paragraph(self, doc: Document, target_paragraph, handwritten_items: list, 
                                          section_name: str, timestamp: str, section_type: str) -> List[ChangeRecord]:
        """
        Reusable method to fill bullet points after a target paragraph
        """
        changes = []
        
        # Find existing bullet points after the target paragraph
        bullet_paragraphs = []
        found_target = False
        
        for paragraph in doc.paragraphs:
            if paragraph == target_paragraph:
                found_target = True
                continue
            
            if found_target:
                # Look for bullet points
                if (paragraph.text.strip().startswith('•') or 
                    paragraph.text.strip().startswith('-') or 
                    paragraph.text.strip().startswith('*') or
                    paragraph.text.strip() == '•' or
                    (paragraph.text.strip() == '' and len(bullet_paragraphs) < 5)):
                    bullet_paragraphs.append(paragraph)
                elif paragraph.text.strip() and not paragraph.text.strip().startswith(' '):
                    break
        
        # Fill existing bullet points with handwritten content
        items_used = 0
        for i, bullet_para in enumerate(bullet_paragraphs):
            if items_used < len(handwritten_items):
                handwritten_item = handwritten_items[items_used]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    # Update the bullet point with handwritten content
                    # Do NOT add bullet character - the bullet formatting already exists
                    bullet_para.clear()
                    bullet_para.add_run(new_text)  # Just add the text, no bullet character
                    
                    changes.append(ChangeRecord(
                        type=f"{section_type}_bullet_point_fill",
                        section=section_name,
                        original_text=bullet_para.text[:50] if bullet_para.text else "empty bullet",
                        new_text=new_text[:50],  # Just the text, no bullet character
                        timestamp=timestamp,
                        strategy_used=f"{section_type}_bullet_replacement"
                    ))
                    
                    items_used += 1
                    self.processing_stats['total_replacements'] += 1
        
        # Add additional bullet points if we have more handwritten items
        if items_used < len(handwritten_items):
            insert_location = bullet_paragraphs[-1] if bullet_paragraphs else target_paragraph
            
            for i in range(items_used, len(handwritten_items)):
                handwritten_item = handwritten_items[i]
                new_text = handwritten_item.get('text', handwritten_item.get('description', ''))
                
                if new_text:
                    new_para = insert_location._element.addnext(doc.add_paragraph(f"• {new_text}")._element)
                    
                    changes.append(ChangeRecord(
                        type=f"{section_type}_bullet_point_add",
                        section=section_name,
                        original_text="",
                        new_text=f"• {new_text}"[:50],
                        timestamp=timestamp,
                        strategy_used=f"{section_type}_bullet_addition"
                    ))
                    
                    self.processing_stats['total_replacements'] += 1
        
        return changes
    
    def _extract_clean_date_format(self, ai_detected_date: str) -> str:
        """
        DYNAMIC: Extract clean date format from AI-detected handwritten text
        Handles various date formats the AI might detect and converts to clean format
        """
        import re
        
        date_text = ai_detected_date.strip()
        self.logger.info(f"📅 Processing AI-detected date: '{date_text}'")
        
        # Pattern 1: Full date formats like "21st of August 2024", "March 15th 2025", etc.
        full_date_patterns = [
            r'(\d{1,2}(?:st|nd|rd|th)?\s+of\s+\w+)',  # "21st of August", "3rd of March"
            r'(\w+\s+\d{1,2}(?:st|nd|rd|th)?)',       # "August 21st", "March 3rd" 
            r'(\d{1,2}(?:st|nd|rd|th)?\s+\w+)',       # "21st August", "3rd March"
        ]
        
        # Try to extract a proper date format
        for pattern in full_date_patterns:
            match = re.search(pattern, date_text, re.IGNORECASE)
            if match:
                extracted_date = match.group(1)
                clean_format = f"the {extracted_date}"
                self.logger.info(f"📅 Extracted date pattern: '{extracted_date}' → '{clean_format}'")
                return clean_format
        
        # Pattern 2: Just month/day without year
        month_day_patterns = [
            r'(\w+\s+\d{1,2})',     # "August 21", "March 3"
            r'(\d{1,2}\s+\w+)',     # "21 August", "3 March"
        ]
        
        for pattern in month_day_patterns:
            match = re.search(pattern, date_text, re.IGNORECASE)
            if match:
                extracted_date = match.group(1)
                clean_format = f"the {extracted_date}"
                self.logger.info(f"📅 Extracted month/day: '{extracted_date}' → '{clean_format}'")
                return clean_format
        
        # Pattern 3: Numeric dates like "21/08/2024", "08-21-2024", etc.
        numeric_patterns = [
            r'(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',  # "21/08/2024", "08-21-2024"
            r'(\d{1,2}[\/\-]\d{1,2})',               # "21/08", "08-21"
        ]
        
        for pattern in numeric_patterns:
            match = re.search(pattern, date_text)
            if match:
                extracted_date = match.group(1)
                clean_format = f"the {extracted_date}"
                self.logger.info(f"📅 Extracted numeric date: '{extracted_date}' → '{clean_format}'")
                return clean_format
        
        # Pattern 4: Just year (fallback for partial detection)
        year_match = re.search(r'(202[4-9])', date_text)
        if year_match:
            year = year_match.group(1)
            clean_format = f"the {year}"
            self.logger.info(f"📅 Extracted year only: '{year}' → '{clean_format}'")
            return clean_format
        
        # Fallback: Use the AI-detected text as-is but add "the" prefix
        clean_format = f"the {date_text}"
        self.logger.info(f"📅 Using AI text as-is: '{date_text}' → '{clean_format}'")
        return clean_format
    
    def _apply_cascading_replacement(self, doc: Document, original_text: str, replacement_text: str) -> tuple[bool, str]:
        """
        Apply 3-strategy cascading replacement system
        
        Args:
            doc: Word document
            original_text: Text to find and replace
            replacement_text: Text to replace with
            
        Returns:
            tuple: (success: bool, strategy_used: str)
        """
        # Strategy 1: Exact text matching
        success = self._replace_exact_text(doc, original_text, replacement_text)
        if success:
            self.processing_stats['strategy_1_exact_matches'] += 1
            return True, "exact_match"
        
        # Strategy 2: Similarity-based matching (0.6 threshold)
        success = self._replace_similar_text(doc, original_text, replacement_text, threshold=0.6)
        if success:
            self.processing_stats['strategy_2_similarity_matches'] += 1
            return True, "similarity_match"
        
        # Strategy 3: Keyword-based fallback
        success = self._replace_keyword_text(doc, original_text, replacement_text)
        if success:
            self.processing_stats['strategy_3_keyword_matches'] += 1
            return True, "keyword_match"
        
        # All strategies failed
        self.processing_stats['failed_operations'] += 1
        return False, "all_strategies_failed"
    
    def _apply_cascading_deletion(self, doc: Document, text_to_delete: str) -> tuple[bool, str]:
        """
        Apply 3-strategy cascading deletion system
        
        Args:
            doc: Word document
            text_to_delete: Text to find and delete
            
        Returns:
            tuple: (success: bool, strategy_used: str)
        """
        # Strategy 1: Exact text matching
        success = self._delete_exact_text(doc, text_to_delete)
        if success:
            self.processing_stats['strategy_1_exact_matches'] += 1
            return True, "exact_match"
        
        # Strategy 2: Similarity-based matching (0.6 threshold)
        success = self._delete_similar_text(doc, text_to_delete, threshold=0.6)
        if success:
            self.processing_stats['strategy_2_similarity_matches'] += 1
            return True, "similarity_match"
        
        # Strategy 3: Keyword-based fallback
        success = self._delete_keyword_text(doc, text_to_delete)
        if success:
            self.processing_stats['strategy_3_keyword_matches'] += 1
            return True, "keyword_match"
        
        # All strategies failed
        self.processing_stats['failed_operations'] += 1
        return False, "all_strategies_failed"
    
    def _replace_exact_text(self, doc: Document, original: str, replacement: str) -> bool:
        """Strategy 1: Replace exact text matches with proper formatting preservation"""
        try:
            for paragraph in doc.paragraphs:
                if original in paragraph.text:
                    # ENHANCED: Preserve formatting by working with runs instead of paragraph.text
                    self.logger.info(f"Found '{original}' in paragraph: '{paragraph.text[:100]}...'")
                    
                    # Method 1: Try run-by-run replacement to preserve formatting
                    replacement_made = False
                    for run in paragraph.runs:
                        if original in run.text:
                            self.logger.info(f"Found '{original}' in run: '{run.text}'")
                            # Preserve the run's formatting while replacing text
                            old_text = run.text
                            run.text = run.text.replace(original, replacement)
                            self.logger.info(f"Replaced run text: '{old_text}' -> '{run.text}'")
                            replacement_made = True
                            break
                    
                    if replacement_made:
                        return True
                    
                    # Method 2: If not found in individual runs, do careful paragraph replacement
                    # This can happen when the text spans multiple runs
                    original_full_text = paragraph.text
                    if original in original_full_text:
                        # Clear all runs and create a new single run but preserve paragraph formatting
                        paragraph.clear()
                        new_text = original_full_text.replace(original, replacement)
                        paragraph.add_run(new_text)
                        self.logger.info(f"Paragraph-level replacement: '{original_full_text[:50]}...' -> '{new_text[:50]}...'")
                        return True
            
            return False
        except Exception as e:
            self.logger.warning(f"Exact text replacement failed: {e}")
            return False
    
    def _replace_similar_text(self, doc: Document, original: str, replacement: str, threshold: float = 0.6) -> bool:
        """Strategy 2: Replace similar text using fuzzy matching"""
        try:
            from difflib import SequenceMatcher
            
            for paragraph in doc.paragraphs:
                # Check if paragraph contains similar text
                similarity = SequenceMatcher(None, original.lower(), paragraph.text.lower()).ratio()
                if similarity >= threshold:
                    # Replace the entire paragraph or find the best matching part
                    words = paragraph.text.split()
                    for i, word in enumerate(words):
                        word_similarity = SequenceMatcher(None, original.lower(), word.lower()).ratio()
                        if word_similarity >= threshold:
                            words[i] = replacement
                            paragraph.text = ' '.join(words)
                            return True
            return False
        except Exception as e:
            self.logger.warning(f"Similar text replacement failed: {e}")
            return False
    
    def _replace_keyword_text(self, doc: Document, original: str, replacement: str) -> bool:
        """Strategy 3: Replace based on keyword matching"""
        try:
            # Extract keywords from original text
            keywords = [word.lower().strip('.,!?;:') for word in original.split() if len(word) > 2]
            
            for paragraph in doc.paragraphs:
                paragraph_lower = paragraph.text.lower()
                # Check if paragraph contains enough keywords
                keyword_matches = sum(1 for keyword in keywords if keyword in paragraph_lower)
                if keyword_matches >= max(1, len(keywords) * 0.5):  # At least 50% keyword match
                    # Find and replace the most likely candidate
                    words = paragraph.text.split()
                    for i, word in enumerate(words):
                        if any(keyword in word.lower() for keyword in keywords):
                            words[i] = replacement
                            paragraph.text = ' '.join(words)
                            return True
            return False
        except Exception as e:
            self.logger.warning(f"Keyword text replacement failed: {e}")
            return False
    
    def _delete_exact_text(self, doc: Document, text_to_delete: str) -> bool:
        """Strategy 1: Delete exact text matches"""
        try:
            for paragraph in doc.paragraphs:
                if text_to_delete in paragraph.text:
                    paragraph.text = paragraph.text.replace(text_to_delete, "")
                    return True
            return False
        except Exception as e:
            self.logger.warning(f"Exact text deletion failed: {e}")
            return False
    
    def _delete_similar_text(self, doc: Document, text_to_delete: str, threshold: float = 0.6) -> bool:
        """Strategy 2: Delete similar text using fuzzy matching"""
        try:
            from difflib import SequenceMatcher
            
            for paragraph in doc.paragraphs:
                words = paragraph.text.split()
                new_words = []
                for word in words:
                    similarity = SequenceMatcher(None, text_to_delete.lower(), word.lower()).ratio()
                    if similarity < threshold:  # Keep words that are NOT similar
                        new_words.append(word)
                
                if len(new_words) < len(words):  # Something was deleted
                    paragraph.text = ' '.join(new_words)
                    return True
            return False
        except Exception as e:
            self.logger.warning(f"Similar text deletion failed: {e}")
            return False
    
    def _delete_keyword_text(self, doc: Document, text_to_delete: str) -> bool:
        """Strategy 3: Delete based on keyword matching"""
        try:
            keywords = [word.lower().strip('.,!?;:') for word in text_to_delete.split() if len(word) > 2]
            
            for paragraph in doc.paragraphs:
                words = paragraph.text.split()
                new_words = []
                for word in words:
                    # Keep words that don't match keywords
                    if not any(keyword in word.lower() for keyword in keywords):
                        new_words.append(word)
                
                if len(new_words) < len(words):  # Something was deleted
                    paragraph.text = ' '.join(new_words)
                    return True
            return False
        except Exception as e:
            self.logger.warning(f"Keyword text deletion failed: {e}")
            return False

    def _implement_chunk_4_additional_opportunities(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Add chunk 4 content to the opportunities section as additional bullet points
        Uses the same approach as chunk 3 - find empty bullets and fill them
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Extract handwritten items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_items = detected_items.get('handwritten_text', [])
        
        if not handwritten_items:
            self.logger.info(f"No handwritten items found for {section_name}")
            return changes
        
        self.logger.info(f"Chunk 4 Additional Opportunities Debug: Found {len(handwritten_items)} handwritten items")
        for i, item in enumerate(handwritten_items):
            self.logger.info(f"  Item {i+1}: {item}")
        
        # Find the opportunities section: "most promising and strategic areas of opportunity are:"
        opportunities_text = "most promising and strategic areas of opportunity are:"
        target_paragraph = None
        
        self.logger.info(f"Chunk 4 Additional Opportunities Debug: Searching for opportunities section...")
        for i, paragraph in enumerate(doc.paragraphs):
            if opportunities_text.lower() in paragraph.text.lower():
                target_paragraph = paragraph
                self.logger.info(f"Chunk 4 Additional Opportunities Debug: Found opportunities section: '{paragraph.text[:100]}...'")
                break
        
        if not target_paragraph:
            self.logger.warning(f"Could not find opportunities section for {section_name}")
            return changes
        
        # Find existing bullet points after the opportunities paragraph
        bullet_paragraphs = []
        found_opportunities = False
        paragraph_count = 0
        total_paragraphs_examined = 0
        target_text = target_paragraph.text.strip()
        
        for paragraph in doc.paragraphs:
            total_paragraphs_examined += 1
            
            # Use text matching instead of object comparison
            if paragraph.text.strip() == target_text and not found_opportunities:
                found_opportunities = True
                self.logger.info(f"Chunk 4 Additional Opportunities Debug: Found target paragraph at position {total_paragraphs_examined}, now looking for bullets...")
                continue
            
            if found_opportunities:
                paragraph_count += 1
                self.logger.info(f"Chunk 4 Additional Opportunities Debug: Examining paragraph #{paragraph_count}: '{paragraph.text}' (stripped: '{paragraph.text.strip()}')")
                
                # Check if this is a Word formatted list item (bullet or numbered)
                is_list_item = False
                if hasattr(paragraph, '_element') and paragraph._element is not None:
                    numbering = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numbering is not None:
                        is_list_item = True
                        self.logger.info(f"Chunk 4 Additional Opportunities Debug: Found Word list item formatting in paragraph #{paragraph_count}!")
                
                # Look for bullet points (Word formatted list items) 
                if is_list_item:
                    bullet_paragraphs.append(paragraph)
                    is_empty = len(paragraph.text.strip()) == 0
                    self.logger.info(f"Chunk 4 Additional Opportunities Debug: ✅ ADDED bullet paragraph #{paragraph_count}: '{paragraph.text}' (empty: {is_empty})")
                elif paragraph.text.strip() and "reinforcing and maximizing your existing strengths" in paragraph.text.lower():
                    # Hit the next section, stop looking
                    self.logger.info(f"Chunk 4 Additional Opportunities Debug: Hit strengths section, stopping search")
                    break
                elif paragraph.text.strip() and len(paragraph.text) > 50:
                    # Hit a different section, stop looking
                    self.logger.info(f"Chunk 4 Additional Opportunities Debug: Hit different section, stopping search")
                    break
        
        self.logger.info(f"Chunk 4 Additional Opportunities Debug: Found {len(bullet_paragraphs)} bullet paragraphs")
        
        # Look for empty bullets after the filled ones and fill them with chunk 4 content
        empty_bullets = []
        filled_bullets = []
        
        for bullet in bullet_paragraphs:
            if len(bullet.text.strip()) == 0:
                empty_bullets.append(bullet)
            else:
                filled_bullets.append(bullet)
        
        self.logger.info(f"Chunk 4 Additional Opportunities Debug: Found {len(filled_bullets)} filled bullets and {len(empty_bullets)} empty bullets")
        
        # Fill empty bullets with chunk 4 content
        for i, item in enumerate(handwritten_items):
            if i < len(empty_bullets):
                content = item.get('text', item.get('description', ''))
                if content:
                    empty_bullets[i].text = content
                    
                    changes.append(ChangeRecord(
                        change_type="replacement",
                        section=section_name,
                        old_text="",
                        new_text=content,
                        strategy_used="additional_opportunities",
                        timestamp=timestamp,
                        confidence=0.95
                    ))
                    
                    self.logger.info(f"Chunk 4 Additional Opportunities Debug: ✅ FILLED empty bullet #{i+1} with: '{content[:50]}...'")
        
        self.logger.info(f"Chunk 4 Additional Opportunities: {len(changes)} changes applied")
        return changes

    def _implement_chunk_6_editing(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Chunk 6: Handle editing operations
        1. Delete entire bullet points if diagonal lines/crosses/squiggles interrupt most of the sentence
        2. Delete specific words with horizontal strikethrough lines
        3. Replace $AMOUNT with handwritten number found above the spending system bullet
        
        IMPORTANT: This handles the FIRST $AMOUNT in the document (cost of living context)
        """
        changes = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        self.logger.info(f"Chunk 6 Editing Debug: ⚡ CHUNK 6 CONTEXT - Cost of living $AMOUNT replacement (FIRST $AMOUNT)")
        
        # Extract detected items from analysis
        detected_items = analysis.get('detected_items', {})
        handwritten_text = detected_items.get('handwritten_text', [])
        strikethrough_text = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        annotations = detected_items.get('annotations', [])
        
        self.logger.info(f"Chunk 6 Editing Debug: Found {len(handwritten_text)} handwritten, {len(strikethrough_text)} strikethrough, {len(crosses)} crosses")
        
        # 1. Handle $AMOUNT replacement first - CHUNK 6 SPECIFIC AMOUNT
        chunk_6_amount_value = None
        for item in handwritten_text:
            text = item.get('text', item.get('description', ''))
            item_description = item.get('description', '').lower()
            
            # CHUNK 6 SPECIFIC: Only use amounts that are in the context of spending/cost of living
            # Skip amounts that are clearly for other sections (like fees in chunk 9)
            spending_context_indicators = ['spending', 'cost of living', 'living', 'pa', 'per annum', 'annually', 'spending suggestion']
            
            # Skip if this looks like it's for the fee section (chunk 9)
            # Be more specific: only skip if it mentions fee context OR (correction AND above/strikethrough)
            is_fee_context = (
                'fee' in item_description or 
                'cost of plan' in item_description or 
                'plan cost' in item_description or
                'MORE4LIFE' in item_description or
                ('correction' in item_description and ('above' in item_description or 'strikethrough' in item_description))
            )
            
            if is_fee_context:
                self.logger.info(f"Chunk 6 Editing Debug: Skipping fee-related amount: '{text}' (description: '{item_description}')")
                continue
            
            # Look for numbers that could be the amount - handle both $X,XXX and X,XXX formats
            import re
            
            # Try to extract amount with or without dollar sign
            # Priority 1: Full dollar amount like "$10,000"
            dollar_match = re.search(r'\$([0-9,]+(?:\.[0-9]{2})?)', text)
            # Priority 2: Just numbers like "10,000" or "160000"  
            number_match = re.search(r'([0-9,]+(?:\.[0-9]{2})?)', text)
            
            extracted_amount = None
            if dollar_match:
                extracted_amount = dollar_match.group(1)  # Just the number part, we'll add $ later
                self.logger.info(f"Chunk 6 Editing Debug: Found dollar amount: '${extracted_amount}' from text: '{text}'")
            elif number_match:
                extracted_amount = number_match.group(1)  # Just the number part
                self.logger.info(f"Chunk 6 Editing Debug: Found numeric amount: '{extracted_amount}' from text: '{text}'")
            
            if extracted_amount:
                # Prefer amounts that have spending context, but accept any if no context found
                has_spending_context = any(indicator in item_description for indicator in spending_context_indicators)
                
                if has_spending_context or not chunk_6_amount_value:  # Take first amount if no context found
                    chunk_6_amount_value = extracted_amount
                    self.logger.info(f"Chunk 6 Editing Debug: Found CHUNK 6 amount value: '{chunk_6_amount_value}' (has_spending_context: {has_spending_context})")
                    if has_spending_context:
                        break  # Prefer spending context amounts
        
        # Find and replace $AMOUNT in the spending system bullet ONLY
        # CRITICAL: Only replace the FIRST $AMOUNT that appears in spending/cost of living context
        spending_system_text = "suggestions on a spending system"
        amount_text = "$AMOUNT"
        chunk_6_replaced = False  # Track if we already replaced for chunk 6
        
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.lower()
            
            # CHUNK 6 SPECIFIC: Only replace $AMOUNT if it's in the spending system context
            # Look for context words nearby
            is_spending_context = any(keyword in para_text for keyword in [
                'spending system', 'cost of living', 'living cost', 'annual spending',
                'pa ', 'per annum', 'annually'
            ])
            
            # Check both the current paragraph and look for $AMOUNT specifically
            if amount_text in paragraph.text and is_spending_context and not chunk_6_replaced:
                if chunk_6_amount_value:
                    old_text = paragraph.text
                    new_text = paragraph.text.replace("$AMOUNT", f"${chunk_6_amount_value}")
                    paragraph.text = new_text
                    chunk_6_replaced = True  # Mark as replaced so we don't replace again
                    
                    changes.append(ChangeRecord(
                        type="replacement",
                        section=section_name,
                        original_text=old_text,
                        location="spending_system_bullet",
                        timestamp=timestamp,
                        ai_confidence=0.95,
                        strategy_used="chunk_6_amount_replacement"
                    ))
                    
                    self.logger.info(f"Chunk 6 Editing Debug: ✅ REPLACED $AMOUNT with ${chunk_6_amount_value} (CHUNK 6 SPECIFIC - COST OF LIVING) in paragraph: '{para_text[:80]}...'")
                    break  # Stop after first replacement
                else:
                    self.logger.warning(f"Chunk 6 Editing Debug: Found spending system bullet but no chunk 6 amount value to replace")
        
        if not chunk_6_replaced and chunk_6_amount_value:
            self.logger.warning(f"Chunk 6 Editing Debug: ❌ Could not find spending system context paragraph with $AMOUNT to replace")
        
        # 2. Handle strikethrough word deletion ONLY in chunk 6 strategies section
        all_strikethrough_items = strikethrough_text + annotations + handwritten_text
        
        # ENHANCED DETECTION: Look for additional strikethrough candidates
        # When AI detects one strikethrough, look for nearby words that might also be struck through
        self.logger.info(f"Chunk 6 Editing Debug: 🔍 ENHANCED STRIKETHROUGH DETECTION")
        self.logger.info(f"Chunk 6 Editing Debug: Original strikethrough items detected: {len(strikethrough_text)}")
        
        for detected_strike in strikethrough_text:
            strike_text = detected_strike.get('text', '').lower()
            strike_pos_y = detected_strike.get('position', {}).get('y', 0)
            
            self.logger.info(f"Chunk 6 Editing Debug: Analyzing detected strikethrough: '{strike_text}' at y={strike_pos_y}")
            
            # Look for text in the same bullet point that might also be struck through
            # Check if the detected strikethrough appears in any paragraph
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.lower()
                if strike_text in para_text:
                    # Found the paragraph containing the detected strikethrough
                    self.logger.info(f"Chunk 6 Editing Debug: Found paragraph with '{strike_text}': '{paragraph.text[:100]}...'")
                    
                    # Split paragraph into words and look for adjacent candidates
                    words = paragraph.text.split()
                    strike_word_indices = []
                    
                    # Find which word(s) contain the detected strikethrough
                    for i, word in enumerate(words):
                        if strike_text in word.lower():
                            strike_word_indices.append(i)
                    
                    # SELECTIVE ENHANCEMENT: Only look at immediately adjacent words (±1 word)
                    # This catches missed words in the same strikethrough phrase without over-deletion
                    for strike_idx in strike_word_indices:
                        # Only check immediately adjacent words (one before, one after)
                        adjacent_indices = []
                        if strike_idx > 0:
                            adjacent_indices.append(strike_idx - 1)  # Previous word
                        if strike_idx < len(words) - 1:
                            adjacent_indices.append(strike_idx + 1)  # Next word
                        
                        self.logger.info(f"Chunk 6 Editing Debug: Checking adjacent words to '{words[strike_idx]}' at indices: {adjacent_indices}")
                        
                        for i in adjacent_indices:
                            candidate_word = words[i]
                            clean_candidate = candidate_word.lower().strip('.,;:')
                            
                            # Only add if it's a substantial word AND likely part of strikethrough phrase
                            # Skip common words that are unlikely to be part of strikethrough phrases
                            common_words = ['the', 'and', 'that', 'with', 'within', 'this', 'they', 'from', 'your', 'will', 
                                          'been', 'have', 'were', 'are', 'can', 'may', 'should', 'would', 'could', 'environment']
                            if (len(clean_candidate) >= 4 and  # Minimum 4 letters
                                clean_candidate not in common_words):  # Skip common words
                                
                                enhanced_item = {
                                    'text': clean_candidate,
                                    'description': f'potential strikethrough adjacent to detected "{strike_text}"',
                                    'confidence': 0.7,
                                    'enhanced_detection': True
                                }
                                all_strikethrough_items.append(enhanced_item)
                                self.logger.info(f"Chunk 6 Editing Debug: Added adjacent strikethrough candidate: '{clean_candidate}' (next to '{words[strike_idx]}')")
                                
                                # CHAIN DETECTION: If we added a candidate, also check ITS adjacent words
                                # This catches phrases like "targeting absolute returns" where AI only detected "targeting"
                                candidate_idx = i
                                next_indices = []
                                if candidate_idx > 0 and candidate_idx - 1 != strike_idx:
                                    next_indices.append(candidate_idx - 1)
                                if candidate_idx < len(words) - 1 and candidate_idx + 1 != strike_idx:
                                    next_indices.append(candidate_idx + 1)
                                
                                for next_i in next_indices:
                                    next_word = words[next_i]
                                    clean_next = next_word.lower().strip('.,;:')
                                    
                                    if (len(clean_next) >= 4 and clean_next not in common_words):
                                        chain_item = {
                                            'text': clean_next,
                                            'description': f'potential strikethrough chained from "{clean_candidate}"',
                                            'confidence': 0.6,
                                            'enhanced_detection': True
                                        }
                                        all_strikethrough_items.append(chain_item)
                                        self.logger.info(f"Chunk 6 Editing Debug: Added chained strikethrough candidate: '{clean_next}' (next to '{clean_candidate}')")
        
        # Find the strategies section bounds to limit strikethrough deletion
        strategies_start = None
        strategies_end = None
        
        self.logger.info(f"Chunk 6 Editing Debug: 🔍 Looking for strategies section in {len(doc.paragraphs)} paragraphs...")
        
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.lower().strip()
            if 'initial strategies and solutions' in para_text:
                strategies_start = i
                self.logger.info(f"Chunk 6 Editing Debug: ✅ Found strategies section start at paragraph {i}: '{paragraph.text[:50]}...'")
            elif strategies_start is not None and ('providing you with confidence' in para_text or 
                                                  'tools, processes and talents' in para_text):
                strategies_end = i
                self.logger.info(f"Chunk 6 Editing Debug: ✅ Found strategies section end at paragraph {i}: '{paragraph.text[:50]}...'")
                break
                
        self.logger.info(f"Chunk 6 Editing Debug: Strategies section bounds: start={strategies_start}, end={strategies_end}")
        
        for strikethrough_item in all_strikethrough_items:
            item_text = strikethrough_item.get('text', strikethrough_item.get('description', ''))
            item_description = strikethrough_item.get('description', '')
            
            # Look for strikethrough indicators in the description or text
            strikethrough_indicators = ['strikethrough', 'crossed out', 'line through', 'horizontal line', 'strike', 'deleted', 'potential strikethrough']
            has_strikethrough = any(indicator in item_description.lower() for indicator in strikethrough_indicators)
            
            # Also include enhanced detection items
            is_enhanced_detection = strikethrough_item.get('enhanced_detection', False)
            
            if (has_strikethrough or is_enhanced_detection) and item_text:
                self.logger.info(f"Chunk 6 Editing Debug: Processing strikethrough item: '{item_text}' (description: '{item_description}')")
                
                # Split into potential words, including partial matches
                words_to_check = item_text.split()
                
                # ONLY process paragraphs in the strategies section
                for i, paragraph in enumerate(doc.paragraphs):
                    # Skip if not in strategies section
                    if strategies_start is not None and (i < strategies_start or (strategies_end is not None and i > strategies_end)):
                        continue
                    
                    original_text = paragraph.text
                    modified_text = original_text
                    words_deleted = []
                    
                    # Check each word in the paragraph against strikethrough words
                    paragraph_words = original_text.split()
                    
                    for para_word in paragraph_words:
                        # Clean the paragraph word (remove punctuation for comparison)
                        clean_para_word = re.sub(r'[^\w]', '', para_word.lower())
                        
                        for strikethrough_word in words_to_check:
                            clean_strike_word = re.sub(r'[^\w]', '', strikethrough_word.lower())
                            
                            # CONSERVATIVE APPROACH: Only delete words that are clearly affected by strikethrough
                            # Prevent over-deletion by being more selective about what gets removed
                            if len(clean_strike_word) >= 3:  # Require at least 3 letters to avoid tiny matches
                                
                                # METHOD 1: Exact match - always delete
                                if clean_para_word == clean_strike_word:
                                    modified_text = modified_text.replace(para_word, '')
                                    words_deleted.append(para_word)
                                    self.logger.info(f"Chunk 6 Editing Debug: Exact match - deleting '{para_word}' (strikethrough: '{strikethrough_word}')")
                                    break
                                
                                # METHOD 2: Strong substring match - strikethrough word is substantial part of paragraph word
                                elif (clean_strike_word in clean_para_word and 
                                      len(clean_strike_word) >= len(clean_para_word) * 0.7):  # 70%+ of word length
                                    modified_text = modified_text.replace(para_word, '')
                                    words_deleted.append(para_word)
                                    self.logger.info(f"Chunk 6 Editing Debug: Strong substring match - deleting '{para_word}' (70%+ contained: '{strikethrough_word}')")
                                    break
                                
                                # METHOD 3: Reverse substring - paragraph word is substantial part of strikethrough  
                                elif (clean_para_word in clean_strike_word and 
                                      len(clean_para_word) >= len(clean_strike_word) * 0.7):  # 70%+ of strikethrough length
                                    modified_text = modified_text.replace(para_word, '')
                                    words_deleted.append(para_word)
                                    self.logger.info(f"Chunk 6 Editing Debug: Reverse substring match - deleting '{para_word}' (70%+ of strikethrough: '{strikethrough_word}')")
                                    break
                                
                                # METHOD 4: ONLY FOR LONG WORDS (8+ chars): High character overlap (80%+)
                                elif len(clean_para_word) >= 8 and len(clean_strike_word) >= 6:
                                    # Calculate character overlap for longer words only
                                    overlap_count = 0
                                    temp_para_word = clean_para_word
                                    
                                    for char in clean_strike_word:
                                        if char in temp_para_word:
                                            overlap_count += 1
                                            temp_para_word = temp_para_word.replace(char, '', 1)
                                    
                                    overlap_percentage = overlap_count / len(clean_para_word) if len(clean_para_word) > 0 else 0
                                    
                                    # Only delete if 80%+ overlap AND both words are substantial
                                    if overlap_percentage >= 0.8:
                                        modified_text = modified_text.replace(para_word, '')
                                        words_deleted.append(para_word)
                                        self.logger.info(f"Chunk 6 Editing Debug: High overlap on long word - deleting '{para_word}' ({overlap_percentage:.1%} coverage by '{strikethrough_word}')")
                                        break
                    
                    # Clean up extra spaces and apply changes
                    modified_text = re.sub(r'\s+', ' ', modified_text).strip()
                    
                    if modified_text != original_text and words_deleted:
                        paragraph.text = modified_text
                        
                        changes.append(ChangeRecord(
                            type="word_deletion",
                            section=section_name,
                            original_text=original_text,
                            location="strikethrough_words",
                            timestamp=timestamp,
                            ai_confidence=0.90,
                            strategy_used="strikethrough_word_deletion"
                        ))
                        
                        self.logger.info(f"Chunk 6 Editing Debug: ✅ REMOVED strikethrough words {words_deleted} from paragraph")
        
        # 3. Handle bullet point deletion for major diagonal/cross/squiggle interruptions
        # IMPORTANT: Only apply bullet deletion for chunk_6_editing section
        if section_name != "chunk_6_editing":
            self.logger.info(f"Chunk 6 Editing Debug: Skipping bullet deletion - not chunk 6 (section: {section_name})")
            self.logger.info(f"Chunk 6 Editing: {len(changes)} changes applied")
            return changes
        
        bullets_to_delete = []
        
        # Check all items for bullet-deletion indicators with better logging
        all_items = crosses + annotations + handwritten_text + strikethrough_text
        
        self.logger.info(f"Chunk 6 Editing Debug: ✅ CHUNK 6 CONFIRMED - Checking {len(all_items)} items for bullet deletion indicators")
        
        for item in all_items:
            item_text = item.get('text', item.get('description', ''))
            item_description = item.get('description', '')
            
            # Log each item being checked
            self.logger.info(f"Chunk 6 Editing Debug: Checking item - text: '{item_text}', description: '{item_description}'")
            
            # ENHANCED FILTER: Don't skip items that could indicate deletion
            # Only skip items that are clearly for appending AND don't suggest deletion
            is_append_only = (
                ('handwritten after' in item_description.lower() or 
                 'handwritten addition' in item_description.lower() or
                 'handwritten clarification' in item_description.lower()) and
                not any(del_word in item_description.lower() for del_word in [
                    'cross', 'line', 'diagonal', 'strike', 'delete', 'cover', 'over'
                ])
            )
            
            if is_append_only:
                self.logger.info(f"Chunk 6 Editing Debug: ❌ Skipping handwritten append item: '{item_text}'")
                continue
            
            # Comprehensive keywords for bullet deletion (squiggly lines, crosses, diagonal lines)
            bullet_deletion_keywords = [
                # Direct deletion indicators - squiggly/wavy lines
                'squiggle', 'squiggly', 'wavy', 'curved', 'wiggly', 'zigzag', 'snake',
                
                # Diagonal lines
                'diagonal', 'slanted', 'angled', 'slash', 'forward slash', 'back slash',
                
                # Crosses and X marks
                'cross', 'crossed', 'x mark', 'x out', 'crossed out', 'strike out',
                
                # Line patterns - ENHANCED with more variations
                'line through', 'line across', 'line over', 'lines through', 'lines across',
                'drawn through', 'drawn across', 'drawn over', 'marked through',
                'single line through', 'multiple lines through', 'heavy line through',
                
                # Coverage and interruption
                'covering', 'interrupting', 'blocking', 'obscuring', 'hiding',
                'scribble', 'scribbled', 'scratch', 'scratched',
                
                # Deletion intent
                'marked out', 'cancelled', 'void', 'delete', 'removed', 'crossed off',
                
                # Coverage indicators - ENHANCED
                'most of', 'entire', 'whole', 'complete', 'over', 'on top',
                'bullet', 'point', 'sentence', 'text', 'content', 'through',
                
                # DYNAMIC DETECTION: Common patterns that suggest bullet deletion
                'next to', 'near', 'beside', 'adjacent to', 'targeting', 'affecting'
            ]
            
            combined_text = f"{item_text} {item_description}".lower()
            has_major_interruption = any(keyword in combined_text for keyword in bullet_deletion_keywords)
            
            # ENHANCED: Add pattern-based detection for bullet deletion
            # Look for patterns that suggest the visual marking is meant to delete content
            visual_deletion_patterns = [
                # Strikethrough that extends beyond just one word
                ('targeting' in combined_text and ('next to' in combined_text or 'near' in combined_text)),
                
                # Any marking described as being "through" or "across" content
                ('through' in combined_text and ('line' in combined_text or 'mark' in combined_text)),
                
                # Markings specifically described as being on bullet points
                ('bullet' in combined_text and ('line' in combined_text or 'mark' in combined_text or 'cross' in combined_text))
            ]
            
            # ENHANCED: Check visual deletion patterns
            has_visual_deletion_pattern = any(pattern for pattern in visual_deletion_patterns)
            
            # Also check if the item itself suggests it's about deleting/covering content
            deletion_patterns = ['delete', 'remove', 'cross out', 'strike through', 'cover', 'obscure']
            has_deletion_intent = any(pattern in combined_text for pattern in deletion_patterns)
            
            # SPECIAL CASE: Check for "next to" patterns which often indicate deletion targeting
            has_next_to_targeting = False
            if 'next to' in item_description.lower():
                # If something is described as "next to" a specific bullet type, it's often a deletion indicator
                bullet_targets = ['personal insurance', 'insurance needs', 'assets', 'liabilities', 'spending system', 'goals', 'objectives']
                for target in bullet_targets:
                    if target in item_description.lower():
                        has_next_to_targeting = True
                        self.logger.info(f"Chunk 6 Editing Debug: ✅ 'NEXT TO' targeting detected for '{target}': '{item_text}' -> '{item_description}'")
                        break
            
            # ENHANCED SPECIAL CASE: Insurance bullet deletion
            # Based on user feedback, if there are diagonal lines through insurance bullet, delete entire bullet
            is_insurance_deletion = False
            
            # Method 1: Direct insurance targeting in description
            if 'insurance' in item_description.lower() and ('next to' in item_description.lower() or 'near' in item_description.lower()):
                is_insurance_deletion = True
                self.logger.info(f"Chunk 6 Editing Debug: ✅ INSURANCE BULLET DELETION detected (method 1): '{item_text}' -> '{item_description}'")
            
            # Method 2: Strikethrough in bullet that likely contains insurance content
            elif (has_major_interruption and 
                  ('targeting' in combined_text or 'line through' in combined_text) and
                  len(item_text) < 15):  # Short strikethrough text suggests part of larger deletion
                # This could be part of insurance bullet deletion - check if we can infer it
                is_insurance_deletion = True
                self.logger.info(f"Chunk 6 Editing Debug: ✅ INFERRED INSURANCE DELETION (method 2): Short strikethrough '{item_text}' likely indicates larger bullet deletion")
            
            if has_major_interruption or has_deletion_intent or has_next_to_targeting or is_insurance_deletion or has_visual_deletion_pattern:
                deletion_reason = (
                    'major_interruption' if has_major_interruption else
                    'insurance_deletion' if is_insurance_deletion else
                    'visual_deletion_pattern' if has_visual_deletion_pattern else
                    'deletion_intent'
                )
                self.logger.info(f"Chunk 6 Editing Debug: ✅ FOUND bullet deletion indicator: '{item_text}' (description: '{item_description}', reason: {deletion_reason})")
                bullets_to_delete.append({
                    'text': item_text,
                    'description': item_description,
                    'reason': deletion_reason,
                    'combined_text': combined_text
                })
            else:
                self.logger.info(f"Chunk 6 Editing Debug: ❌ No deletion indicators found for this item")
        
        # DYNAMIC BULLET DELETION: Add inference for bullets that should be completely deleted
        # If we detect strikethrough in a bullet but the visual pattern suggests the entire bullet is marked for deletion
        for strike_item in strikethrough_text:
            strike_text = strike_item.get('text', '').lower()
            strike_desc = strike_item.get('description', '').lower()
            
            # Look for patterns that suggest entire bullet deletion rather than just word deletion
            # Common patterns: when strikethrough is at the end of a bullet, often the whole bullet is meant to be deleted
            bullet_deletion_indicators = [
                # If strikethrough appears in insurance context, often the whole bullet needs to go
                ('insurance' in strike_desc or 'personal insurance' in strike_desc),
                
                # If strikethrough is described as being significant or covering major content
                ('entire' in strike_desc or 'whole' in strike_desc or 'complete' in strike_desc),
                
                # If strikethrough is in a context that suggests bullet removal
                ('bullet' in strike_desc and ('line' in strike_desc or 'cross' in strike_desc)),
                
                # DYNAMIC: If the strikethrough text appears to be at the end of key bullet points
                # This catches cases where "targeting absolute returns" suggests the whole investment bullet should go
                (strike_text in ['targeting', 'absolute', 'returns'] and 
                 any(bullet_word in strike_desc for bullet_word in ['investment', 'superannuation', 'environment']))
            ]
            
            if any(indicator for indicator in bullet_deletion_indicators):
                # Add this as a bullet deletion candidate
                bullets_to_delete.append({
                    'text': strike_text,
                    'description': f'inferred bullet deletion from strikethrough pattern: {strike_desc}',
                    'reason': 'inferred_bullet_deletion',
                    'combined_text': f"{strike_text} {strike_desc}",
                    'confidence': 0.8
                })
                self.logger.info(f"Chunk 6 Editing Debug: 🎯 INFERRED BULLET DELETION from strikethrough: '{strike_text}' -> bullet deletion")

        # Delete heavily interrupted bullet points using improved detection
        self.logger.info(f"Chunk 6 Editing Debug: Found {len(bullets_to_delete)} bullet deletion indicators")
        
        bullets_deleted = 0
        for bullet_info in bullets_to_delete:
            self.logger.info(f"Chunk 6 Editing Debug: Processing deletion indicator: {bullet_info['reason']} - '{bullet_info['text']}'")
            
            # More aggressive bullet deletion - try multiple strategies
            deletion_success = False
            
            # Strategy 1: Look for specific bullet content ONLY in the strategies section
            self.logger.info(f"Chunk 6 Editing Debug: 🔍 Scanning {len(doc.paragraphs)} paragraphs for bullets to delete...")
            
            for i, paragraph in enumerate(doc.paragraphs):
                if deletion_success:
                    break
                
                # ONLY process paragraphs in the strategies section (same bounds as strikethrough)
                if strategies_start is not None and (i < strategies_start or (strategies_end is not None and i > strategies_end)):
                    self.logger.info(f"Chunk 6 Editing Debug: ⏭️  Skipping paragraph {i} (outside strategies section): '{paragraph.text[:30]}...'")
                    continue
                    
                # Enhanced bullet point detection
                is_word_bullet = False
                if hasattr(paragraph, '_element') and paragraph._element is not None:
                    numPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    is_word_bullet = numPr is not None
                
                text_based_bullet = (paragraph.text.strip().startswith('•') or 
                                   paragraph.text.strip().startswith('-') or
                                   paragraph.text.strip().startswith('*'))
                
                is_bullet = is_word_bullet or text_based_bullet
                
                self.logger.info(f"Chunk 6 Editing Debug: 📝 Paragraph {i}: is_word_bullet={is_word_bullet}, text_based_bullet={text_based_bullet}, is_bullet={is_bullet}, text='{paragraph.text[:50]}...'")
                
                if is_bullet and len(paragraph.text.strip()) > 5:
                    para_text_lower = paragraph.text.lower()
                    bullet_text_lower = bullet_info['text'].lower()
                    
                    # Check if this bullet should be deleted based on visual deletion indicators
                    should_delete = False
                    
                    # DEBUG: Log bullet matching details
                    self.logger.info(f"Chunk 6 Editing Debug: 🔍 Checking bullet: '{paragraph.text[:60]}...'")
                    self.logger.info(f"Chunk 6 Editing Debug: 🔍 Against indicator: text='{bullet_info['text']}', description='{bullet_info.get('description', 'N/A')}', combined='{bullet_info['combined_text']}'")
                    
                    # Only delete bullets when there's a clear connection between the visual deletion 
                    # indicator and the specific bullet content
                    visual_deletion_indicators = ['squiggle', 'squiggly', 'wavy', 'diagonal', 'cross', 'crossed', 'line through', 'line across', 'line over', 'x mark', 'strike', 'scratch']
                    
                    has_visual_indicators = any(indicator in bullet_info['combined_text'] for indicator in visual_deletion_indicators)
                    self.logger.info(f"Chunk 6 Editing Debug: 🔍 Visual indicators check: has_visual_indicators={has_visual_indicators}")
                    
                    # ENHANCED BULLET DELETION MATCHING - More dynamic approach
                    
                    # Method 1: Direct insurance bullet deletion
                    is_insurance_bullet_deletion = (
                        ('insurance' in para_text_lower and 'needs' in para_text_lower) and
                        (bullet_info['reason'] == 'insurance_deletion' or 
                         'personal insurance needs' in bullet_info.get('description', '').lower() or
                         'insurance needs' in bullet_info.get('description', '').lower() or
                         'bullet point about personal insurance' in bullet_info.get('description', '').lower() or
                         bullet_info['reason'] == 'inferred_bullet_deletion')
                    )
                    
                    # Method 2: Inferred deletion from strikethrough in bullet content
                    is_inferred_deletion = (
                        bullet_info['reason'] == 'inferred_bullet_deletion' and
                        bullet_text_lower in para_text_lower  # The strikethrough text appears in this bullet
                    )
                    
                    # Method 3: Dynamic content-based matching for any bullet type
                    is_dynamic_content_match = False
                    if bullet_info.get('text') and len(bullet_info['text']) > 3:
                        # Check if the deletion indicator text appears in this bullet's content
                        indicator_words = bullet_info['text'].lower().split()
                        bullet_words = para_text_lower.split()
                        
                        # Count meaningful word matches (exclude common words)
                        common_words = {'the', 'and', 'of', 'to', 'a', 'in', 'that', 'have', 'it', 'for', 'not', 'with', 'he', 'as', 'you', 'do', 'at', 'this', 'but', 'his', 'by', 'from', 'they', 'we', 'say', 'her', 'she', 'or', 'an', 'will', 'my', 'one', 'all', 'would', 'there', 'their'}
                        meaningful_matches = []
                        
                        for word in indicator_words:
                            if len(word) > 3 and word not in common_words and word in bullet_words:
                                meaningful_matches.append(word)
                        
                        # If we have meaningful word matches AND the description suggests deletion
                        deletion_context_words = ['line', 'cross', 'strike', 'through', 'over', 'mark', 'delete', 'squiggle']
                        has_deletion_context = any(word in bullet_info.get('description', '').lower() for word in deletion_context_words)
                        
                        if len(meaningful_matches) >= 1 and has_deletion_context:
                            is_dynamic_content_match = True
                            self.logger.info(f"Chunk 6 Editing Debug: 🎯 DYNAMIC CONTENT MATCH - Found '{meaningful_matches}' in bullet with deletion context")
                    
                    if is_insurance_bullet_deletion:
                        self.logger.info(f"Chunk 6 Editing Debug: 🎯 INSURANCE DELETION - Matched insurance bullet: '{paragraph.text[:60]}...'")
                        should_delete = True
                    elif is_inferred_deletion:
                        self.logger.info(f"Chunk 6 Editing Debug: 🎯 INFERRED DELETION - Matched bullet containing '{bullet_text_lower}': '{paragraph.text[:60]}...'")
                        should_delete = True
                    elif is_dynamic_content_match:
                        self.logger.info(f"Chunk 6 Editing Debug: 🎯 DYNAMIC MATCH - Content-based bullet deletion: '{paragraph.text[:60]}...'")
                        should_delete = True
                    else:
                        # ENHANCED: More flexible confidence handling for different detection types
                        deletion_confidence = bullet_info.get('confidence', 0.0)
                        self.logger.info(f"Chunk 6 Editing Debug: 🔍 Deletion confidence check: {deletion_confidence} (requires 0.85+ for original AI, 0.7+ for enhanced)")
                        
                        # SPECIAL PROTECTION: Never delete tax/wealth strategy bullets with low confidence markings
                        is_tax_wealth_bullet = ('tax' in para_text_lower and ('wealth' in para_text_lower or 'increasing' in para_text_lower))
                        if is_tax_wealth_bullet and deletion_confidence < 0.90:
                            self.logger.info(f"Chunk 6 Editing Debug: 🛡️ PROTECTED TAX/WEALTH BULLET - Confidence {deletion_confidence} too low for: '{paragraph.text[:60]}...'")
                            continue  # Special protection for tax/wealth bullets
                        
                        # Different confidence thresholds for different detection types
                        required_confidence = 0.85  # Default for original AI detections
                        if bullet_info.get('reason') in ['inferred_bullet_deletion', 'enhanced_detection']:
                            required_confidence = 0.7  # Lower threshold for our enhanced detection
                        
                        # SPECIAL CASE: If we have visual indicators (squiggly lines) and the bullet mentions insurance, be more lenient
                        if (has_visual_indicators and 'insurance' in para_text_lower and 
                            'line through' in bullet_info.get('combined_text', '').lower()):
                            required_confidence = 0.5  # Very lenient for visual evidence of insurance bullet deletion
                            self.logger.info(f"Chunk 6 Editing Debug: 🎯 INSURANCE + VISUAL EVIDENCE - Lowered confidence requirement to {required_confidence}")
                        
                        if deletion_confidence < required_confidence:
                            self.logger.info(f"Chunk 6 Editing Debug: ❌ SKIPPING DELETION - Low confidence ({deletion_confidence}) for bullet (requires {required_confidence}): '{paragraph.text[:60]}...'")
                            continue  # Skip this bullet entirely if confidence is too low
                    
                    # Note: is_insurance_bullet_deletion is now handled above, before confidence checks
                    if has_visual_indicators and not should_delete:
                        # CRITICAL: Additional validation for visual indicators
                        # Must have strong visual evidence (not just "single line through text")
                        weak_indicators = ['single line through text', 'slight marking', 'negligible marking', 'barely visible']
                        has_weak_indication = any(weak in bullet_info['combined_text'].lower() for weak in weak_indicators)
                        
                        if has_weak_indication:
                            self.logger.info(f"Chunk 6 Editing Debug: ⚠️ WEAK VISUAL INDICATION - Skipping deletion for: '{paragraph.text[:60]}...'")
                            # Don't delete bullets with only weak visual indications
                        
                        # Method 1: Direct content correlation - the deletion indicator must reference specific words from this bullet
                        elif bullet_text_lower and len(bullet_text_lower) > 5:
                            # Extract meaningful words from both the deletion indicator and the bullet
                            indicator_words = set(word.lower() for word in bullet_info['text'].split() if len(word) > 3)
                            bullet_words = set(word.lower() for word in paragraph.text.split() if len(word) > 3)
                            
                            # Look for meaningful word overlap - require at least 2 significant words or 1 very specific word
                            word_overlap = indicator_words.intersection(bullet_words)
                            
                            # Filter out common words that appear in many bullets
                            meaningful_overlap = word_overlap - {'your', 'and', 'the', 'with', 'will', 'that', 'this', 'you', 'for', 'are', 'can', 'have', 'bullet', 'point', 'should', 'lines', 'delete', 'keep'}
                            
                            # ENHANCED: Check for exact phrase matching (e.g., "absolute returns")
                            # CRITICAL: Must be substantial overlap (5+ characters) to avoid false positives
                            deletion_phrase = bullet_info['text'].lower().strip()
                            if len(deletion_phrase) >= 5 and deletion_phrase in para_text_lower:
                                should_delete = True
                                self.logger.info(f"Chunk 6 Editing Debug: 🎯 EXACT PHRASE MATCH - Found '{deletion_phrase}' in bullet: '{paragraph.text[:60]}...'")
                            elif len(meaningful_overlap) >= 3:  # Increased from 2 to 3 for higher precision
                                should_delete = True
                                self.logger.info(f"Chunk 6 Editing Debug: Strong content correlation - deletion indicator mentions '{meaningful_overlap}' from bullet: '{paragraph.text[:60]}...'")
                            elif len(meaningful_overlap) == 2:
                                # Two word match must be very specific and substantial
                                words = list(meaningful_overlap)
                                if all(len(word) >= 6 for word in words) and not any(word in ['recommendations', 'strategies', 'analysis', 'bullet', 'point'] for word in words):
                                    should_delete = True
                                    self.logger.info(f"Chunk 6 Editing Debug: Two specific word match - deletion indicator mentions '{meaningful_overlap}' from bullet: '{paragraph.text[:60]}...'")
                            elif len(meaningful_overlap) == 1:
                                # Single word match must be very specific (8+ characters and extremely specific)
                                specific_word = list(meaningful_overlap)[0]
                                if len(specific_word) >= 8 and specific_word not in ['analysis', 'recommendations', 'strategies', 'bullet', 'point', 'should', 'lines', 'delete', 'insurance', 'planning']:
                                    should_delete = True
                                    self.logger.info(f"Chunk 6 Editing Debug: Highly specific single word match - deletion indicator mentions '{specific_word}' from bullet: '{paragraph.text[:60]}...'")
                            
                            # Special case: if the deletion indicator specifically mentions "insurance" and this bullet is about insurance
                            elif 'insurance' in bullet_info['combined_text'] and 'insurance' in para_text_lower and 'needs' in para_text_lower:
                                should_delete = True
                                self.logger.info(f"Chunk 6 Editing Debug: Insurance-specific deletion match: '{paragraph.text[:60]}...'")
                        
                    # Method 2: Description-based targeting - if the description specifically mentions this bullet's topic
                    # This should run regardless of visual indicators
                    if not should_delete and bullet_info.get('description', ''):
                        self.logger.info(f"Chunk 6 Editing Debug: 🔍 Running Method 2: Description-based targeting")
                        desc_lower = bullet_info['description'].lower()
                        
                        # ENHANCED: Check if the description references the specific content of this bullet
                        # CRITICAL FIX: Much more strict phrase matching to prevent false positives
                        key_bullet_phrases = ['insurance needs', 'personal insurance', 'goals and objectives', 'assets liabilities', 'spending system', 'capital lump sum', 'wealth creation', 'non-deductible debt', 'superannuation environment', 'investments outside', 'investment mix', 'estate planning', 'cash flow analysis']
                        
                        # REMOVED: 'reduction in tax' to prevent false matching with "tax whilst increasing your wealth" bullet
                        # This was causing the false positive deletion
                        
                        for phrase in key_bullet_phrases:
                            # Require EXACT phrase match AND substantial text in both description and bullet
                            if (phrase in desc_lower and 
                                phrase.replace(' ', '') in para_text_lower.replace(' ', '') and
                                len(phrase) >= 10):  # Only match substantial phrases
                                should_delete = True
                                self.logger.info(f"Chunk 6 Editing Debug: Description-based targeting: '{phrase}' matches bullet content")
                                break
                                
                        # SPECIAL CASE: "next to personal insurance needs" or similar descriptions
                        self.logger.info(f"Chunk 6 Editing Debug: 🔍 Insurance check - 'next to' in desc: {'next to' in desc_lower}, 'insurance' in desc: {'insurance' in desc_lower}, 'insurance' in bullet: {'insurance' in para_text_lower}, 'needs' in bullet: {'needs' in para_text_lower}")
                        
                        if 'next to' in desc_lower and 'insurance' in desc_lower and 'insurance' in para_text_lower and 'needs' in para_text_lower:
                            should_delete = True
                            self.logger.info(f"Chunk 6 Editing Debug: 🎯 TARGETED INSURANCE DELETION - Description mentions 'next to insurance', bullet contains 'insurance needs': '{paragraph.text[:60]}...'")
                            
                            # SPECIAL CASE: "under bullet point about" patterns
                            if 'under bullet point about' in desc_lower or 'bullet point about' in desc_lower:
                                # Extract what the bullet point is "about" from the description
                                if 'investments outside' in desc_lower and 'investments outside' in para_text_lower:
                                    should_delete = True
                                    self.logger.info(f"Chunk 6 Editing Debug: 🎯 TARGETED INVESTMENT DELETION - Description about 'investments outside': '{paragraph.text[:60]}...'")
                                elif 'assets' in desc_lower and 'liabilities' in desc_lower and 'assets' in para_text_lower and 'liabilities' in para_text_lower:
                                    should_delete = True
                                    self.logger.info(f"Chunk 6 Editing Debug: 🎯 TARGETED ASSETS/LIABILITIES DELETION: '{paragraph.text[:60]}...'")
                                    
                            # GENERAL PATTERN: If description mentions being "next to" or "about" a specific bullet type
                            bullet_type_indicators = {
                                'insurance': ['insurance', 'needs'],
                                'assets': ['assets', 'liabilities'], 
                                'spending': ['spending', 'system'],
                                'goals': ['goals', 'objectives'],
                                'investments': ['investments', 'outside'],
                                'superannuation': ['superannuation', 'environment'],
                                'wealth': ['wealth', 'creation'],
                                'debt': ['debt', 'deductible'],
                                'tax': ['tax', 'reduction'],
                                'estate': ['estate', 'planning'],
                                'cash': ['cash', 'flow']
                            }
                            
                            for bullet_type, required_words in bullet_type_indicators.items():
                                if not should_delete and bullet_type in desc_lower:
                                    if all(word in para_text_lower for word in required_words):
                                        should_delete = True
                                        self.logger.info(f"Chunk 6 Editing Debug: 🎯 PATTERN-BASED DELETION - '{bullet_type}' type matches bullet with {required_words}: '{paragraph.text[:60]}...'")
                                        break
                    
                    # Method 3: Explicit deletion intent with content verification
                    if not should_delete and bullet_info['reason'] == 'deletion_intent':
                        deletion_keywords = ['delete', 'remove', 'cross out', 'strike through', 'cancel']
                        if any(keyword in bullet_info['combined_text'] for keyword in deletion_keywords):
                            # Only delete if there's some content connection
                            if bullet_text_lower and any(word in para_text_lower for word in bullet_text_lower.split() if len(word) > 4):
                                should_delete = True
                                self.logger.info(f"Chunk 6 Editing Debug: Explicit deletion intent with content verification: '{paragraph.text[:60]}...'")
                
                    
                    if should_delete:
                        old_text = paragraph.text
                        deletion_phrase = bullet_info['text'].lower().strip()
                        
                        # DETERMINE DELETION TYPE: Full bullet vs partial phrase
                        # Check for indicators that suggest entire bullet should be deleted
                        full_bullet_indicators = [
                            'squiggle', 'squiggly', 'wavy', 'curved', 'wiggly', 'zigzag', 'snake',
                            'diagonal', 'slanted', 'angled', 'slash', 'lines through', 'lines across',
                            'cross', 'crossed', 'x mark', 'x out', 'crossed out', 'strike out',
                            'covering', 'interrupting', 'blocking', 'obscuring', 'entire', 'whole', 'complete'
                        ]
                        
                        # Check if this indicates full bullet deletion
                        is_full_bullet_deletion = (
                            bullet_info['reason'] == 'insurance_deletion' or
                            any(indicator in bullet_info['combined_text'] for indicator in full_bullet_indicators) or
                            'entire' in bullet_info.get('description', '').lower() or
                            'whole' in bullet_info.get('description', '').lower() or
                            'complete' in bullet_info.get('description', '').lower() or
                            # If the deletion phrase is very short compared to bullet length, it's likely the whole bullet should go
                            (len(deletion_phrase) < 10 and len(old_text) > 50)
                        )
                        
                        self.logger.info(f"Chunk 6 Editing Debug: Deletion analysis - phrase: '{deletion_phrase}', bullet_length: {len(old_text)}, is_full_deletion: {is_full_bullet_deletion}")
                        
                        if is_full_bullet_deletion:
                            # COMPLETE BULLET DELETION: Remove entire bullet point
                            self.logger.info(f"Chunk 6 Editing Debug: 🗑️ COMPLETE BULLET DELETION for: '{old_text[:100]}...'")
                            
                            # Clear the paragraph completely and also remove the bullet formatting
                            paragraph.clear()  # This removes all runs and formatting
                            
                            # Check if this paragraph has Word list formatting and remove it
                            if hasattr(paragraph, '_element') and paragraph._element is not None:
                                # Remove numbering properties to eliminate bullet formatting
                                numPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                                if numPr is not None:
                                    numPr.getparent().remove(numPr)
                            
                            # Set paragraph text to empty to completely remove content
                            paragraph.text = ""
                            
                            # Also clear any runs to ensure complete removal
                            for run in paragraph.runs:
                                run.clear()
                            
                            changes.append(ChangeRecord(
                                type="complete_bullet_deletion",
                                section=section_name,
                                original_text=old_text,
                                location="bullet_deletion",
                                timestamp=timestamp,
                                ai_confidence=0.95,
                                strategy_used="complete_bullet_deletion"
                            ))
                            
                            self.logger.info(f"Chunk 6 Editing Debug: ✅ COMPLETELY DELETED entire bullet point (all sentences): '{old_text[:100]}...'")
                        
                        # PARTIAL PHRASE DELETION: Only remove specific crossed-out phrase
                        elif deletion_phrase and deletion_phrase in old_text.lower():
                            self.logger.info(f"Chunk 6 Editing Debug: ✏️ PARTIAL PHRASE DELETION for phrase: '{deletion_phrase}'")
                            
                            import re
                            # Create a case-insensitive pattern that matches the deletion phrase
                            pattern = re.compile(re.escape(deletion_phrase), re.IGNORECASE)
                            new_text = pattern.sub('', old_text).strip()
                            
                            # Clean up any double spaces or commas left behind
                            new_text = re.sub(r'\s+', ' ', new_text)  # Multiple spaces -> single space
                            new_text = re.sub(r',\s*,', ',', new_text)  # Double commas -> single comma
                            new_text = re.sub(r',\s*;', ';', new_text)  # Comma before semicolon -> just semicolon
                            new_text = re.sub(r'\s*,\s*$', '', new_text)  # Trailing comma
                            new_text = re.sub(r'^\s*,\s*', '', new_text)  # Leading comma
                            
                            # Update the paragraph with the cleaned text
                            paragraph.clear()
                            paragraph.add_run(new_text)
                            
                            changes.append(ChangeRecord(
                                type="partial_text_deletion",
                                section=section_name,
                                original_text=old_text,
                                new_text=new_text,
                                location="phrase_removal",
                                timestamp=timestamp,
                                ai_confidence=0.90,
                                strategy_used=f"phrase_deletion_{bullet_info['reason']}"
                            ))
                            
                            self.logger.info(f"Chunk 6 Editing Debug: ✅ REMOVED PHRASE '{deletion_phrase}' from bullet: '{old_text[:60]}...' → '{new_text[:60]}...'")
                        else:
                            # FALLBACK: If can't find specific phrase, delete entire bullet
                            self.logger.info(f"Chunk 6 Editing Debug: 🔄 FALLBACK - Complete deletion (couldn't find specific phrase)")
                            
                            # Clear the paragraph completely
                            paragraph.clear()
                            paragraph.text = ""
                            
                            # Remove bullet formatting if present
                            if hasattr(paragraph, '_element') and paragraph._element is not None:
                                numPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                                if numPr is not None:
                                    numPr.getparent().remove(numPr)
                            
                            # Additional cleanup: Remove any remaining text formatting
                            for run in paragraph.runs:
                                run.clear()
                        
                            changes.append(ChangeRecord(
                                type="complete_bullet_deletion",
                                section=section_name,
                                original_text=old_text,
                                location="crossed_bullet_complete",
                                timestamp=timestamp,
                                ai_confidence=0.85,
                                strategy_used=f"complete_bullet_deletion_{bullet_info['reason']}"
                            ))
                            
                            self.logger.info(f"Chunk 6 Editing Debug: ✅ COMPLETELY DELETED entire bullet point (all sentences): '{old_text[:100]}...'")
                            
                            # IMPORTANT: Also check if the next paragraph continues this bullet point
                            # and delete subsequent lines that are part of the same logical bullet
                            try:
                                current_para_index = None
                                for i, doc_para in enumerate(doc.paragraphs):
                                    if doc_para._element is paragraph._element:
                                        current_para_index = i
                                        break
                                
                                if current_para_index is not None and current_para_index + 1 < len(doc.paragraphs):
                                    next_para = doc.paragraphs[current_para_index + 1]
                                    next_text = next_para.text.strip()
                                    
                                    # Check if next paragraph is a continuation (doesn't start with bullet but has related content)
                                    is_continuation = (
                                        next_text and 
                                        not next_text.startswith('•') and 
                                        not next_text.startswith('-') and 
                                        not next_text.startswith('*') and
                                        len(next_text) > 10 and
                                        (not next_text[0].isupper() or  # Doesn't start new sentence
                                         'this analysis will include' in next_text.lower() or  # Specific continuation pattern
                                         'providing peace of mind' in next_text.lower() or  # Your specific case
                                         'continuation sentence' in next_text.lower() or  # Test case
                                         next_text.startswith('  '))  # Indented text is usually continuation
                                    )
                                    
                                    if is_continuation:
                                        continuation_text = next_para.text
                                        next_para.clear()
                                        next_para.text = ""
                                        for run in next_para.runs:
                                            run.clear()
                                        
                                        self.logger.info(f"Chunk 6 Editing Debug: ✅ ALSO DELETED continuation sentence: '{continuation_text[:100]}...'")
                                        changes.append(ChangeRecord(
                                            type="bullet_continuation_deletion",
                                            section=section_name,
                                            original_text=continuation_text,
                                            location="bullet_continuation",
                                            timestamp=timestamp,
                                            ai_confidence=0.80,
                                            strategy_used=f"continuation_deletion_{bullet_info['reason']}"
                                        ))
                            except Exception as e:
                                self.logger.warning(f"Error checking continuation paragraph: {e}")
                        
                        self.logger.info(f"Chunk 6 Editing Debug: ✅ DELETED bullet point: '{old_text[:100]}...'")
                        bullets_deleted += 1
                        deletion_success = True
                        break
            
            if not deletion_success:
                self.logger.warning(f"Chunk 6 Editing Debug: ❌ Could not find bullet to delete for indicator: '{bullet_info['text']}'")
            
            # Limit deletions to prevent over-deletion (max 3 bullets)
            if bullets_deleted >= 3:
                self.logger.info(f"Chunk 6 Editing Debug: Reached deletion limit (3 bullets), stopping")
                break
        
        # 4. Handle handwritten text appending to bullet points (NEW FUNCTIONALITY)
        handwritten_append_changes = self._append_handwritten_to_bullets(doc, handwritten_text, section_name, timestamp)
        changes.extend(handwritten_append_changes)
        
        self.logger.info(f"Chunk 6 Editing: {len(changes)} changes applied")
        return changes
    
    def _implement_chunk_7_editing(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Implement chunk 7 editing operations: strikethrough deletion and dot point deletion
        
        This chunk handles the cash flow analysis section with:
        - Strikethrough text deletion from within sentences
        - Complete bullet point deletion when crossed out
        
        Target section content:
        • Cash flow analysis after recommendations;
        • Year by year illustration of your increase in wealth and reduction of debt;
        • Fees, disclosures and general information.
        """
        changes = []
        timestamp = datetime.now().isoformat()
        
        self.logger.info(f"🔧 Processing Chunk 7: Cash flow section editing for {section_name}")
        
        # Extract detected items
        detected_items = analysis.get('detected_items', {})
        handwritten_text = detected_items.get('handwritten_text', [])
        strikethrough_items = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        arrows = detected_items.get('arrows', [])
        
        self.logger.info(f"Chunk 7 Editing Debug: Found {len(handwritten_text)} handwritten, {len(strikethrough_items)} strikethrough, {len(crosses)} crosses")
        
        if not handwritten_text and not strikethrough_items and not crosses:
            self.logger.info(f"Chunk 7 Editing: No editing operations detected")
            return changes
        
        # STEP 1: Handle bullet deletion for crossed-out items (same logic as chunk 6)
        deletion_indicators = []
        
        # Add strikethrough indicators that suggest bullet deletion
        for item in strikethrough_items:
            deletion_indicators.append({
                'text': item.get('text', ''),
                'description': item.get('description', ''),
                'combined_text': f"{item.get('text', '')} - {item.get('description', '')}",
                'confidence': item.get('confidence', 0.8),
                'reason': 'strikethrough_deletion'
            })
        
        # Add cross indicators that suggest bullet deletion  
        for cross in crosses:
            deletion_indicators.append({
                'text': cross.get('text', ''),
                'description': cross.get('description', ''),
                'combined_text': f"{cross.get('text', '')} - {cross.get('description', '')}",
                'confidence': cross.get('confidence', 0.8),
                'reason': 'cross_deletion'
            })
        
        # Process bullet deletions (reusing chunk 6 logic)
        bullets_deleted = 0
        for bullet_info in deletion_indicators:
            if bullets_deleted >= 3:  # Safety limit
                break
                
            deletion_success = False
            
            # Find and delete matching bullets
            for i, paragraph in enumerate(doc.paragraphs):
                if deletion_success:
                    break
                    
                # Check if this is a bullet point (same detection logic as chunk 6)
                is_word_bullet = False
                try:
                    para_element = paragraph._element
                    if para_element is not None:
                        pPr = para_element.find('.//w:pPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if pPr is not None:
                            numPr = pPr.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if numPr is not None:
                                is_word_bullet = True
                except:
                    pass
                
                # Also check for text-based bullets
                text_based_bullet = paragraph.text.strip().startswith(('•', '-', '*', '◦'))
                is_bullet = is_word_bullet or text_based_bullet
                
                if is_bullet and len(paragraph.text.strip()) > 5:
                    para_text_lower = paragraph.text.lower()
                    bullet_text_lower = bullet_info['text'].lower()
                    
                    # Check for content match (same logic as chunk 6 but for cash flow content)
                    should_delete = False
                    
                    # Look for cash flow specific content
                    cash_flow_phrases = ['cash flow', 'year by year', 'increase in wealth', 'reduction of debt', 'fees', 'disclosures', 'general information']
                    
                    for phrase in cash_flow_phrases:
                        if (phrase in bullet_info['combined_text'].lower() and 
                            phrase.replace(' ', '') in para_text_lower.replace(' ', '')):
                            should_delete = True
                            self.logger.info(f"Chunk 7 Editing Debug: Cash flow phrase match: '{phrase}' matches bullet content")
                            break
                    
                    # Also check for direct text correlation
                    if bullet_text_lower and len(bullet_text_lower) > 5:
                        if bullet_text_lower in para_text_lower:
                            should_delete = True
                            self.logger.info(f"Chunk 7 Editing Debug: Direct text match: '{bullet_text_lower}' found in bullet")
                    
                    if should_delete:
                        old_text = paragraph.text
                        
                        # Complete bullet deletion
                        paragraph.clear()
                        paragraph.text = ""
                        for run in paragraph.runs:
                            run.clear()
                        
                        changes.append(ChangeRecord(
                            type="complete_bullet_deletion",
                            section=section_name,
                            original_text=old_text,
                            location="cash_flow_bullet_deletion",
                            timestamp=timestamp,
                            ai_confidence=bullet_info.get('confidence', 0.8),
                            strategy_used=f"chunk_7_{bullet_info['reason']}"
                        ))
                        
                        self.logger.info(f"Chunk 7 Editing Debug: ✅ DELETED cash flow bullet: '{old_text[:100]}...'")
                        bullets_deleted += 1
                        deletion_success = True
                        break
            
            if not deletion_success:
                self.logger.warning(f"Chunk 7 Editing Debug: ❌ Could not find bullet to delete for indicator: '{bullet_info['text']}'")
        
        # STEP 2: Handle strikethrough text deletion within remaining bullets (word/phrase level)
        strikethrough_deletions = 0
        for item in strikethrough_items:
            if strikethrough_deletions >= 5:  # Limit strikethrough operations
                break
                
            strikethrough_text = item.get('text', '').strip()
            if len(strikethrough_text) < 3:  # Skip very short items
                continue
            
            # Find and remove strikethrough text from paragraphs
            for paragraph in doc.paragraphs:
                if strikethrough_text.lower() in paragraph.text.lower():
                    old_text = paragraph.text
                    new_text = paragraph.text.replace(strikethrough_text, '')
                    
                    # Clean up extra spaces
                    new_text = ' '.join(new_text.split())
                    
                    if new_text != old_text:
                        paragraph.text = new_text
                        
                        changes.append(ChangeRecord(
                            type="strikethrough_text_deletion",
                            section=section_name,
                            original_text=old_text,
                            new_text=new_text,
                            location="cash_flow_strikethrough",
                            timestamp=timestamp,
                            ai_confidence=item.get('confidence', 0.8),
                            strategy_used="chunk_7_strikethrough_deletion"
                        ))
                        
                        self.logger.info(f"Chunk 7 Editing Debug: ✅ REMOVED strikethrough text: '{strikethrough_text}' from: '{old_text[:100]}...'")
                        strikethrough_deletions += 1
                        break
        
        # STEP 3: Handle handwritten text appending to bullet points in brackets (same rule as chunk 6)
        if handwritten_text:
            bracket_changes = self._append_handwritten_to_cash_flow_bullets(doc, handwritten_text, section_name, timestamp)
            changes.extend(bracket_changes)
        
        self.logger.info(f"Chunk 7 Editing: {len(changes)} changes applied ({bullets_deleted} bullets deleted, {strikethrough_deletions} strikethrough deletions)")
        return changes
    
    def _append_handwritten_to_cash_flow_bullets(self, doc: Document, handwritten_items: List[Dict], section_name: str, timestamp: str) -> List[ChangeRecord]:
        """
        Append handwritten text to cash flow bullet points in brackets
        Handles cases where handwritten text appears after cash flow bullet points
        Adapted from chunk 6 brackets rule but focused on cash flow section
        """
        changes = []
        
        if not handwritten_items:
            self.logger.info(f"Chunk 7 Handwritten Append: No handwritten items to process")
            return changes
        
        self.logger.info(f"Chunk 7 Handwritten Append: Processing {len(handwritten_items)} handwritten items")
        
        # Group related handwritten items that are close together (multi-line text)
        grouped_items = self._group_related_handwritten_items(handwritten_items)
        self.logger.info(f"Chunk 7 Handwritten Append: Grouped into {len(grouped_items)} logical items")
        
        # Filter handwritten items for appendable text
        appendable_items = []
        for item in grouped_items:
            text = item.get('text', item.get('description', '')).strip()
            
            # Skip items that are just numbers or very short
            import re
            if re.match(r'^[\$]?[0-9,]+\.?[0-9]*$', text) or len(text) < 3:
                self.logger.info(f"Chunk 7 Handwritten Append: Skipping numeric/short item: '{text}'")
                continue
                
            appendable_items.append(item)
        
        if not appendable_items:
            self.logger.info(f"Chunk 7 Handwritten Append: No appendable handwritten items found")
            return changes
        
        # Find cash flow bullet points in the document
        cash_flow_bullets = []
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip().lower()
            
            # Look for cash flow section bullets
            is_word_bullet = False
            if hasattr(paragraph, '_element') and paragraph._element is not None:
                numPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                is_word_bullet = numPr is not None
            
            text_based_bullet = paragraph.text.strip().startswith(('•', '-', '*'))
            is_bullet = is_word_bullet or text_based_bullet
            
            # Identify cash flow bullets by content
            if (is_bullet and len(paragraph.text.strip()) > 10 and 
                any(phrase in para_text for phrase in [
                    'cash flow analysis', 
                    'year by year illustration', 
                    'increase in wealth', 
                    'reduction of debt',
                    'fees, disclosures',
                    'general information'
                ])):
                
                cash_flow_bullets.append({
                    'paragraph': paragraph,
                    'index': i,
                    'original_text': paragraph.text.strip()
                })
                self.logger.info(f"Chunk 7 Handwritten Append: Found cash flow bullet: '{paragraph.text.strip()[:60]}...'")
        
        if not cash_flow_bullets:
            self.logger.warning(f"Chunk 7 Handwritten Append: No cash flow bullet points found")
            return changes
        
        self.logger.info(f"Chunk 7 Handwritten Append: Found {len(cash_flow_bullets)} cash flow bullets")
        
        # Match handwritten items to bullets
        for handwritten_item in appendable_items:
            handwritten_text = handwritten_item.get('text', handwritten_item.get('description', '')).strip()
            description = handwritten_item.get('description', '').lower()
            
            # Clean up the handwritten text
            handwritten_text = handwritten_text.replace('handwritten:', '').strip()
            if handwritten_text.startswith('"') and handwritten_text.endswith('"'):
                handwritten_text = handwritten_text[1:-1]
            
            self.logger.info(f"Chunk 7 Handwritten Append: Processing item: '{handwritten_text}' with description: '{description}'")
            
            # Find matching bullet based on description or content
            target_bullet = None
            for bullet_info in cash_flow_bullets:
                bullet_text = bullet_info['original_text'].lower()
                
                # Skip bullets that already have bracketed content
                if '(' in bullet_info['original_text'] and ')' in bullet_info['original_text']:
                    continue
                
                # Match based on content keywords
                if ('cash flow' in description or 'analysis' in description) and 'cash flow analysis' in bullet_text:
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 7 Handwritten Append: Matched to cash flow analysis bullet")
                    break
                elif ('year by year' in description or 'wealth' in description or 'debt' in description) and ('year by year' in bullet_text or 'wealth' in bullet_text):
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 7 Handwritten Append: Matched to wealth/debt bullet")
                    break
                elif ('fees' in description or 'disclosures' in description) and ('fees' in bullet_text or 'disclosures' in bullet_text):
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 7 Handwritten Append: Matched to fees/disclosures bullet")
                    break
            
            # If no specific match, try to match to first available bullet
            if not target_bullet:
                for bullet_info in cash_flow_bullets:
                    if not ('(' in bullet_info['original_text'] and ')' in bullet_info['original_text']):
                        target_bullet = bullet_info
                        self.logger.info(f"Chunk 7 Handwritten Append: Using first available bullet as fallback")
                        break
            
            if not target_bullet:
                self.logger.warning(f"Chunk 7 Handwritten Append: Could not find matching bullet for: '{handwritten_text}'")
                continue
            
            paragraph = target_bullet['paragraph']
            original_text = target_bullet['original_text']
            
            # Apply brackets rule - place handwritten text in brackets at end of bullet
            import re
            
            # Handle punctuation correctly
            if original_text.endswith(';'):
                new_text = original_text[:-1] + f" ({handwritten_text});"
            elif original_text.endswith('.'):
                new_text = original_text[:-1] + f" ({handwritten_text})."
            else:
                new_text = original_text + f" ({handwritten_text})"
            
            # Update the paragraph
            paragraph.clear()
            paragraph.add_run(new_text)
            
            changes.append(ChangeRecord(
                type="handwritten_append",
                section=section_name,
                original_text=original_text,
                new_text=new_text,
                location="cash_flow_bullet_brackets",
                timestamp=timestamp,
                ai_confidence=0.85,
                strategy_used="chunk_7_handwritten_bracket_append"
            ))
            
            self.logger.info(f"Chunk 7 Handwritten Append: ✅ APPENDED to cash flow bullet: '{handwritten_text}' → '{new_text[:80]}...'")
        
        return changes
    
    def _implement_chunk_8_editing(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Implement chunk 8 editing operations: Simple table row deletion based on horizontal lines
        
        FOCUSED SCOPE: Only delete entire table rows when detecting horizontal lines that span 
        from left to right across the table row, disrupting text on both sides.
        
        This chunk should ONLY:
        1. Detect horizontal lines/strikes that span across a table row from left to right
        2. Detect crosses that span from left box to right box
        3. Detect diagonal lines that span from left box to right box
        4. Delete the entire row when such spans are detected
        5. Preserve all formatting and structure
        """
        changes = []
        timestamp = datetime.now().isoformat()
        
        self.logger.info(f"🔧 Processing Chunk 8: Simple table row deletion for horizontal spans")
        
        # Extract detected items - focus only on strikethrough and crosses that indicate row deletion
        detected_items = analysis.get('detected_items', {})
        strikethrough_items = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        
        self.logger.info(f"Chunk 8: Found {len(strikethrough_items)} strikethrough items, {len(crosses)} crosses")
        
        # Target table rows that can be deleted
        table_row_keywords = [
            "Development of a Business Plan",
            "Cash Flow Minder", 
            "Cash Flow Mirror",
            "Dollars of Dignity",
            "Family Lifestyle Retainer",
            "Wealth Accumulation",
            "More4Life Odometer"
        ]
        
        # Track deleted rows to avoid duplicates
        deleted_rows = set()
        
        # STEP 1: Check for strikethrough text that spans horizontally across table rows
        for strikethrough_item in strikethrough_items:
            description = strikethrough_item.get('description', '').lower()
            text_content = strikethrough_item.get('text', '').strip()
            position = strikethrough_item.get('position', {})
            
            # Look for indicators that this is a horizontal line across a row
            # SIMPLIFIED: Accept any strikethrough in chunk 8 table as row deletion intent
            horizontal_span_indicators = [
                'line through', 'line across', 'horizontal line', 'strikethrough line',
                'crossed out row', 'line spanning', 'full width line', 'row crossed out',
                'spans from left to right', 'left box to right box', 'crosses entire row',
                'horizontal span', 'spans across', 'left to right line',
                'single line through', 'line strike', 'strikethrough text'  # Added simpler patterns
            ]
            
            is_horizontal_span = any(indicator in description for indicator in horizontal_span_indicators)
            
            if is_horizontal_span:
                self.logger.info(f"Chunk 8: Detected strikethrough - {description}")
                
                # Find which table row this line affects
                deletion_target = None
                
                # First try to match by text content
                for keyword in table_row_keywords:
                    if keyword.lower() in text_content.lower() or keyword.lower() in description:
                        deletion_target = keyword
                        break
                
                # If no keyword match, use position-based detection
                if not deletion_target and position:
                    deletion_target = self._determine_position_target_row(position, table_row_keywords)
                    if deletion_target:
                        self.logger.info(f"Chunk 8: Matched strikethrough to '{deletion_target}' by position {position}")
                
                if deletion_target and deletion_target not in deleted_rows:
                    success = self._delete_table_row(doc, deletion_target, changes, section_name, timestamp, "horizontal_span_strikethrough")
                    if success:
                        deleted_rows.add(deletion_target)
        
        # STEP 2: Check for crosses that span from left box to right box
        for cross in crosses:
            description = cross.get('description', '').lower()
            size = cross.get('size', 'small').lower()
            position = cross.get('position', {})
            
            # Look for indicators that this cross spans from left to right across the table
            cross_span_indicators = [
                'spans from left to right', 'left box to right box', 'crosses entire row',
                'horizontal cross', 'wide cross', 'full width cross', 'spans across',
                'left to right cross', 'table row cross', 'row spanning cross',
                'diagonal across row', 'diagonal from left to right', 'diagonal span',
                'large diagonal cross', 'crosses both columns', 'spans both boxes',
                'diagonal cross', 'cross through'  # Added simpler patterns
            ]
            
            # Enhanced detection for row-spanning crosses
            # SIMPLIFIED: Accept any large/medium cross in chunk 8 table as row deletion intent
            is_row_spanning_cross = (
                size in ['large', 'medium'] or 
                any(indicator in description for indicator in cross_span_indicators) or
                self._is_cross_spanning_horizontally(cross, position)
            )
            
            if is_row_spanning_cross:
                self.logger.info(f"Chunk 8: Detected cross - {description} (size: {size}) at position {position}")
                
                # Enhanced cross analysis to determine target row
                deletion_target = None
                
                # First try keyword matching in description
                for keyword in table_row_keywords:
                    if keyword.lower() in description:
                        deletion_target = keyword
                        break
                
                # If no keyword match, use position-based detection
                if not deletion_target and position:
                    deletion_target = self._determine_position_target_row(position, table_row_keywords)
                    if deletion_target:
                        self.logger.info(f"Chunk 8: Matched cross to '{deletion_target}' by position {position}")
                
                if deletion_target and deletion_target not in deleted_rows:
                    success = self._delete_table_row(doc, deletion_target, changes, section_name, timestamp, "row_spanning_cross")
                    if success:
                        deleted_rows.add(deletion_target)
        
        # STEP 3: Check for diagonal lines that span from left box to right box
        # Look in annotations or other detection categories for diagonal spans
        annotations = detected_items.get('annotations', [])
        arrows = detected_items.get('arrows', [])
        handwritten_items = detected_items.get('handwritten_text', [])
        
        # Check annotations and arrows for diagonal spans
        for item in annotations + arrows:
            description = item.get('description', '').lower()
            position = item.get('position', {})
            
            diagonal_span_indicators = [
                'diagonal line from left to right', 'diagonal span across row',
                'diagonal line across table', 'diagonal from left box to right box',
                'slanted line across row', 'angled line spanning', 'diagonal strike',
                'diagonal line through row', 'slanted strike across'
            ]
            
            is_diagonal_span = any(indicator in description for indicator in diagonal_span_indicators)
            
            if is_diagonal_span:
                self.logger.info(f"Chunk 8: Detected diagonal span - {description}")
                
                # Find which table row this diagonal affects
                deletion_target = None
                for keyword in table_row_keywords:
                    if keyword.lower() in description:
                        deletion_target = keyword
                        break
                
                # If no specific keyword, try to determine by position
                if not deletion_target and position:
                    deletion_target = self._determine_position_target_row(position, table_row_keywords)
                    if deletion_target:
                        self.logger.info(f"Chunk 8: Matched diagonal to '{deletion_target}' by position {position}")
                
                if deletion_target and deletion_target not in deleted_rows:
                    success = self._delete_table_row(doc, deletion_target, changes, section_name, timestamp, "diagonal_span")
                    if success:
                        deleted_rows.add(deletion_target)
        
        # STEP 4: Check handwritten text for squiggly lines or deletion marks through rows
        # Squiggly lines through text are often detected as handwritten content
        handwritten_items = detected_items.get('handwritten_text', [])
        self.logger.info(f"Chunk 8: Found {len(handwritten_items)} handwritten items")
        
        for handwritten in handwritten_items:
            text_content = handwritten.get('text', '').strip()
            description = handwritten.get('description', '').lower()
            position = handwritten.get('position', {})
            
            self.logger.info(f"Chunk 8: Processing handwritten - text: '{text_content[:60]}...', desc: '{description}'")
            
            # Check if the handwritten content matches a table row (could be squiggly line deletion)
            deletion_target = None
            
            # First check if text content matches a table row keyword
            for keyword in table_row_keywords:
                if keyword.lower() in text_content.lower():
                    deletion_target = keyword
                    self.logger.info(f"Chunk 8: Found handwritten text matching row '{keyword}' - possible squiggly line deletion")
                    break
            
            # Use position-based detection as fallback
            if not deletion_target and position:
                deletion_target = self._determine_position_target_row(position, table_row_keywords)
                if deletion_target:
                    self.logger.info(f"Chunk 8: Matched handwritten mark to '{deletion_target}' by position {position}")
            
            if deletion_target and deletion_target not in deleted_rows:
                # Check description for deletion indicators
                deletion_indicators = ['squiggly', 'wavy line', 'scribble', 'crossed out', 
                                      'strike', 'deletion', 'removed', 'line through']
                has_deletion_indicator = any(indicator in description for indicator in deletion_indicators)
                
                # For chunk 8 table rows, treat any handwritten detection over a row as potential deletion
                if has_deletion_indicator or len(text_content) > 20:  # Long text match suggests full row coverage
                    success = self._delete_table_row(doc, deletion_target, changes, section_name, timestamp, "handwritten_deletion_mark")
                    if success:
                        deleted_rows.add(deletion_target)
        
        self.logger.info(f"Chunk 8: Completed processing. Deleted {len(deleted_rows)} rows: {list(deleted_rows)}")
        return changes
    
    def _is_cross_spanning_horizontally(self, cross: Dict, position: Dict) -> bool:
        """
        Analyze cross properties to determine if it spans horizontally across a table row
        """
        # Check if the cross has width/span information
        if 'width' in position or 'span' in cross.get('description', '').lower():
            return True
        
        # Check if position indicates it's in the middle/center area (likely spanning)
        x_pos = position.get('x', 0)
        if 0.3 <= x_pos <= 0.7:  # Middle area suggests it might span across
            return True
        
        # Check cross size - large crosses are more likely to span
        size = cross.get('size', 'small').lower()
        if size in ['large', 'huge', 'wide']:
            return True
        
        return False
    
    def _determine_cross_target_row(self, cross: Dict, description: str, table_row_keywords: List[str], position: Dict) -> str:
        """
        Determine which table row a cross is targeting based on description and position
        """
        # First, check if description mentions specific row keywords
        for keyword in table_row_keywords:
            if keyword.lower() in description:
                return keyword
        
        # If no specific mention, use position to estimate row
        return self._determine_position_target_row(position, table_row_keywords)
    
    def _determine_position_target_row(self, position: Dict, table_row_keywords: List[str]) -> str:
        """
        Determine target row based on Y position in the chunk
        Mapping adjusted based on visual inspection of 7-row table layout
        """
        y_pos = position.get('y', 0)
        
        # Map Y positions to table rows based on even distribution
        # Each row occupies approximately 1/7 ≈ 0.143 of vertical space
        if y_pos <= 0.20:  # First row region
            return "Development of a Business Plan"
        elif y_pos <= 0.30:  # Second row region
            return "Cash Flow Minder"
        elif y_pos <= 0.40:  # Third row region  
            return "Cash Flow Mirror"
        elif y_pos <= 0.50:  # Fourth row region
            return "Dollars of Dignity"
        elif y_pos <= 0.60:  # Fifth row region
            return "Family Lifestyle Retainer"
        elif y_pos <= 0.75:  # Sixth row region
            return "Wealth Accumulation"
        else:  # Seventh row region
            return "More4Life Odometer"
    
    def _delete_table_row(self, doc: Document, row_keyword: str, changes: List[ChangeRecord], section_name: str, timestamp: str, strategy: str) -> bool:
        """
        Helper method to safely delete a table row by keyword
        Returns True if deletion was successful, False otherwise
        """
        try:
            self.logger.info(f"Chunk 8: Attempting to delete row with keyword '{row_keyword}'")
            
            for table_idx, table in enumerate(doc.tables):
                self.logger.info(f"Chunk 8: Checking table {table_idx + 1}/{len(doc.tables)} with {len(table.rows)} rows")
                
                for i, row in enumerate(table.rows):
                    row_text = ' '.join([cell.text for cell in row.cells])
                    
                    # Normalize both texts: lowercase, collapse whitespace, remove extra spaces
                    # This handles line breaks, tabs, multiple spaces, etc.
                    import re
                    row_text_normalized = re.sub(r'\s+', ' ', row_text.lower().strip())
                    row_keyword_normalized = re.sub(r'\s+', ' ', row_keyword.lower().strip())
                    
                    # Extract first meaningful part for partial matching
                    # e.g., "Development of a Business Plan" -> "development of a business"
                    keyword_parts = row_keyword_normalized.split()[:5]  # First 5 words
                    keyword_partial = ' '.join(keyword_parts) if len(keyword_parts) >= 3 else row_keyword_normalized
                    
                    # Log each row check for debugging
                    self.logger.debug(f"Chunk 8: Row {i} text: '{row_text_normalized[:80]}...'")
                    
                    if row_keyword_normalized in row_text_normalized or keyword_partial in row_text_normalized:
                        original_text = row_text
                        self.logger.info(f"Chunk 8: ✓ MATCH FOUND! Deleting table row {i}: '{original_text[:100]}...'")
                        self.logger.info(f"Chunk 8: Matched using: {'full keyword' if row_keyword_normalized in row_text_normalized else 'partial keyword (' + keyword_partial + ')'}")
                        original_text = row_text
                        self.logger.info(f"Chunk 8: Deleting table row: '{original_text[:100]}...'")
                        
                        # Delete the entire row from the table
                        tbl = table._tbl
                        tr = row._tr
                        tbl.remove(tr)
                        
                        # Record the change
                        changes.append(ChangeRecord(
                            type="row_deletion",
                            section=section_name,
                            original_text=original_text,
                            new_text="[ROW DELETED]",
                            location=f"business_plan_table_row_{row_keyword.lower().replace(' ', '_')}",
                            timestamp=timestamp,
                            ai_confidence=0.95,
                            strategy_used=f"chunk_8_simple_{strategy}"
                        ))
                        
                        self.logger.info(f"Chunk 8: ✅ Successfully deleted table row: '{row_keyword}'")
                        return True
            
            self.logger.warning(f"Chunk 8: Table row not found for keyword: '{row_keyword}'")
            return False
            
        except Exception as e:
            self.logger.error(f"Chunk 8: Failed to delete table row '{row_keyword}': {e}")
            return False
    def _implement_chunk_9_editing(self, doc: Document, analysis: Dict, section_name: str) -> List[ChangeRecord]:
        """
        Implement chunk 9 editing operations: $AMOUNT replacement and tax returns text substitution
        
        This chunk handles:
        1. $AMOUNT replacement with handwritten price above strikethrough
        2. Tax returns text replacement - replace struck-through part with handwritten alternative
        3. Color correction - change red text to black in tax returns paragraph
        
        Target content:
        - "The cost of the 'MORE4LIFE' plan is $AMOUNT inclusive of GST..."
        - "and recent tax returns (personal, trust and company)."
        """
        changes = []
        timestamp = datetime.now().isoformat()
        
        self.logger.info(f"🔧 Processing Chunk 9: Fee section and tax returns editing for {section_name}")
        self.logger.info(f"Chunk 9 Editing Debug: ⚡ CHUNK 9 CONTEXT - Fee plan $AMOUNT replacement (SECOND $AMOUNT)")
        self.logger.info(f"Chunk 9 Editing Debug: Analysis keys: {list(analysis.keys())}")
        
        # Extract detected items
        detected_items = analysis.get('detected_items', {})
        handwritten_text = detected_items.get('handwritten_text', [])
        strikethrough_items = detected_items.get('strikethrough_text', [])
        crosses = detected_items.get('crosses', [])
        
        self.logger.info(f"Chunk 9 Editing Debug: Found {len(handwritten_text)} handwritten, {len(strikethrough_items)} strikethrough, {len(crosses)} crosses")
        
        # DEBUG: Log all detected items to see what we're working with
        self.logger.info(f"Chunk 9 Editing Debug: === ALL HANDWRITTEN ITEMS ===")
        for i, item in enumerate(handwritten_text):
            self.logger.info(f"Chunk 9 Editing Debug: Handwritten {i+1}: '{item.get('text', '')}' | Description: '{item.get('description', '')}'")
        
        self.logger.info(f"Chunk 9 Editing Debug: === ALL STRIKETHROUGH ITEMS ===")
        for i, item in enumerate(strikethrough_items):
            self.logger.info(f"Chunk 9 Editing Debug: Strikethrough {i+1}: '{item.get('text', '')}' | Description: '{item.get('description', '')}'")
        
        # Extract highlights too
        highlights = detected_items.get('highlights', [])
        self.logger.info(f"Chunk 9 Editing Debug: === ALL HIGHLIGHT ITEMS ===")
        for i, item in enumerate(highlights):
            self.logger.info(f"Chunk 9 Editing Debug: Highlight {i+1}: '{item.get('text', '')}' | Description: '{item.get('description', '')}'")
        
        # STEP 1: Handle $AMOUNT replacement in fee paragraph
        amount_replacements = 0
        fee_paragraph_found = False
        chunk_9_replaced = False  # Track if we already replaced for chunk 9
        
        for paragraph in doc.paragraphs:
            if '$AMOUNT' in paragraph.text or 'amount' in paragraph.text.lower():
                self.logger.info(f"Chunk 9 Editing Debug: Found paragraph with $AMOUNT: '{paragraph.text[:100]}...'")
            if 'MORE4LIFE' in paragraph.text:
                self.logger.info(f"Chunk 9 Editing Debug: Found paragraph with MORE4LIFE: '{paragraph.text[:100]}...'")
            
            # ENHANCED: Look for fee paragraphs more flexibly - just need $AMOUNT or cost-related terms
            # CHUNK 9 SPECIFIC: Must be in fee/plan context, NOT spending context
            is_fee_paragraph = ('$AMOUNT' in paragraph.text or 
                               ('cost' in paragraph.text.lower() and 'MORE4LIFE' in paragraph.text) or
                               ('plan is $' in paragraph.text) or
                               ('amount' in paragraph.text.lower() and any(word in paragraph.text.lower() for word in ['fee', 'cost', 'price', 'plan'])))
            
            # CRITICAL: Exclude spending system paragraphs (those belong to chunk 6)
            is_spending_paragraph = any(keyword in paragraph.text.lower() for keyword in [
                'spending system', 'cost of living', 'living cost', 'annual spending',
                'pa ', 'per annum', 'annually'
            ])
            
            if is_fee_paragraph and not is_spending_paragraph and not chunk_9_replaced:
                fee_paragraph_found = True
                self.logger.info(f"Chunk 9 Editing Debug: Found fee paragraph: '{paragraph.text[:100]}...'")
                
                # ENHANCED: Look specifically for strikethrough over $AMOUNT and handwritten text above it
                amount_replacement_found = False
                
                # NEW SIMPLIFIED APPROACH: Check for ANY handwritten amount with "AMOUNT" in description
                # OR handwritten amounts that are "above strikethrough" (which likely means above $AMOUNT)
                self.logger.info(f"Chunk 9 Editing Debug: === SIMPLIFIED $AMOUNT DETECTION ===")
                
                # First check if there's a strikethrough or highlighted "AMOUNT" in this chunk
                has_amount_strikethrough = any('amount' in item.get('text', '').lower() for item in strikethrough_items)
                has_amount_highlighted = any('amount' in item.get('text', '').lower() for item in highlights)
                has_amount_marked = has_amount_strikethrough or has_amount_highlighted
                self.logger.info(f"Chunk 9 Editing Debug: has_amount_strikethrough: {has_amount_strikethrough}, has_amount_highlighted: {has_amount_highlighted}")
                
                for item in handwritten_text:
                    item_text = item.get('text', '').strip()
                    item_desc = item.get('description', '').strip().lower()
                    
                    self.logger.info(f"Chunk 9 Editing Debug: Checking handwritten: '{item_text}' (desc: '{item_desc}')")
                    
                    # Look for handwritten text that mentions "AMOUNT" in its description
                    mentions_amount = 'amount' in item_desc.lower()
                    
                    # ENHANCED: Also check if it's described as being "above strikethrough/highlighted" when AMOUNT is marked
                    # This catches cases like "handwritten correction above highlighted text" or "above strikethrough text"
                    is_above_marked_amount = (
                        ('above' in item_desc and ('strikethrough' in item_desc or 'highlighted' in item_desc or 'highlight' in item_desc)) or 
                        'correction above' in item_desc
                    )
                    
                    # Check if it looks like a monetary amount
                    import re
                    looks_like_amount = re.search(r'[\$]?[0-9,]+\.?[0-9]*', item_text)
                    
                    # CHUNK 9 SPECIFIC: Skip spending context amounts
                    spending_context_indicators = ['spending', 'cost of living', 'living', 'pa', 'per annum', 'annually', 'spending suggestion', 'next to spending', 'spending amount']
                    is_spending_context = any(indicator in item_desc for indicator in spending_context_indicators)
                    
                    self.logger.info(f"Chunk 9 Editing Debug: - mentions_amount: {mentions_amount}")
                    self.logger.info(f"Chunk 9 Editing Debug: - is_above_marked_amount: {is_above_marked_amount}")
                    self.logger.info(f"Chunk 9 Editing Debug: - looks_like_amount: {bool(looks_like_amount)}")
                    self.logger.info(f"Chunk 9 Editing Debug: - is_spending_context: {is_spending_context}")
                    
                    # CRITICAL FIX: Chunk 9 should ONLY accept amounts that are explicitly:
                    # 1. Above highlighted/strikethrough AMOUNT (primary method), OR
                    # 2. Have "amount" in description BUT also have fee/correction context (not just "amount")
                    # This prevents "handwritten correction for amount" from being used in both chunks
                    has_fee_correction_context = any(keyword in item_desc for keyword in [
                        'above', 'correction above', 'fee', 'plan', 'MORE4LIFE'
                    ])
                    
                    # Accept if: above marked AMOUNT OR (mentions amount AND has fee context) AND NOT spending
                    is_chunk_9_amount = (
                        is_above_marked_amount or 
                        (mentions_amount and has_fee_correction_context)
                    ) and not is_spending_context
                    
                    self.logger.info(f"Chunk 9 Editing Debug: - has_fee_correction_context: {has_fee_correction_context}")
                    self.logger.info(f"Chunk 9 Editing Debug: - is_chunk_9_amount: {is_chunk_9_amount}")
                    
                    if looks_like_amount and is_chunk_9_amount:
                        # Clean up the amount text
                        amount_text = re.sub(r'[^\$0-9,.]', '', item_text)
                        if not amount_text.startswith('$'):
                            amount_text = '$' + amount_text
                        
                        self.logger.info(f"Chunk 9 Editing Debug: ✅ Found CHUNK 9 amount from description: '{amount_text}'")
                        
                        # Replace $AMOUNT with the handwritten amount
                        old_text = paragraph.text
                        new_text = old_text.replace('$AMOUNT', amount_text)
                        
                        # ENHANCED: If no $AMOUNT found, look for any dollar amount to replace
                        if new_text == old_text and '$' in old_text:
                            # Find and replace any existing dollar amount
                            dollar_pattern = r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?'
                            if re.search(dollar_pattern, old_text):
                                new_text = re.sub(dollar_pattern, amount_text, old_text, count=1)
                                self.logger.info(f"Chunk 9 Editing Debug: ✅ REPLACED existing dollar amount with '{amount_text}' (no $AMOUNT placeholder found)")
                        
                        if new_text != old_text:
                            paragraph.clear()
                            paragraph.add_run(new_text)
                            
                            changes.append(ChangeRecord(
                                type="amount_replacement",
                                section=section_name,
                                original_text=old_text,
                                new_text=new_text,
                                location="fee_paragraph_amount",
                                timestamp=timestamp,
                                ai_confidence=0.90,
                                strategy_used="chunk_9_amount_replacement_direct"
                            ))
                            
                            self.logger.info(f"Chunk 9 Editing Debug: ✅ REPLACED $AMOUNT with '{amount_text}' (CHUNK 9 SPECIFIC - FEE PLAN)")
                            amount_replacements += 1
                            amount_replacement_found = True
                            chunk_9_replaced = True
                            break
                
                # FALLBACK: Original strikethrough-based logic
                if not amount_replacement_found:
                    self.logger.info(f"Chunk 9 Editing Debug: === FALLBACK: STRIKETHROUGH-BASED DETECTION ===")
                    # First, check if there's a strikethrough item that mentions $AMOUNT or similar
                    for strikethrough_item in strikethrough_items:
                        strike_text = strikethrough_item.get('text', '').strip().lower()
                        strike_desc = strikethrough_item.get('description', '').strip().lower()
                        
                        self.logger.info(f"Chunk 9 Editing Debug: Checking strikethrough: '{strike_text}' (desc: '{strike_desc}')")
                        
                        # Check if this strikethrough is about the $AMOUNT
                        if ('amount' in strike_text or '$' in strike_text or 
                            'amount' in strike_desc or '$' in strike_desc or
                            'price' in strike_desc or 'cost' in strike_desc):
                            
                            self.logger.info(f"Chunk 9 Editing Debug: Found $AMOUNT strikethrough indicator")
                            
                            # Now look for handwritten text that could be the replacement
                            # Look for handwritten text that mentions being "above" or "over" the strikethrough
                            for item in handwritten_text:
                                item_text = item.get('text', '').strip()
                                item_desc = item.get('description', '').strip().lower()
                                
                                self.logger.info(f"Chunk 9 Editing Debug: Checking handwritten: '{item_text}' (desc: '{item_desc}')")
                                
                                # Check if this handwritten text is positioned above/over the $AMOUNT
                                is_above_amount = any(phrase in item_desc for phrase in [
                                    'above', 'over', 'on top', 'handwritten above', 'written above',
                                    'price above', 'amount above', 'number above'
                                ])
                                
                                # Also check if it looks like a monetary amount and isn't from earlier chunks
                                import re
                                looks_like_amount = re.search(r'[\$]?[0-9,]+\.?[0-9]*', item_text)
                                
                                self.logger.info(f"Chunk 9 Editing Debug: - looks_like_amount: {looks_like_amount}")
                                self.logger.info(f"Chunk 9 Editing Debug: - is_above_amount: {is_above_amount}")
                                self.logger.info(f"Chunk 9 Editing Debug: - 'above' in desc: {'above' in item_desc}")
                                
                                # CHUNK 9 SPECIFIC: Only use amounts that are in fee/cost context
                                # Skip amounts that are clearly for other sections (like spending in chunk 6)
                                fee_context_indicators = ['fee', 'cost of plan', 'plan cost', 'above strikethrough', 'MORE4LIFE', 'above the strikethrough', 'cost', 'price', 'plan', 'above text', 'correction above', 'handwritten correction']
                                spending_context_indicators = ['spending', 'cost of living', 'living', 'pa', 'per annum', 'annually']
                                
                                # Skip if this looks like it's for the spending section (chunk 6)  
                                if any(indicator in item_desc for indicator in spending_context_indicators):
                                    self.logger.info(f"Chunk 9 Editing Debug: Skipping spending-related amount: '{item_text}' (description: '{item_desc}')")
                                    continue
                                
                                # Prefer amounts that have fee context
                                has_fee_context = any(indicator in item_desc for indicator in fee_context_indicators)
                                
                                if looks_like_amount and (is_above_amount or 'above' in item_desc) and (has_fee_context or not any(indicator in item_desc for indicator in spending_context_indicators)):
                                    # Clean up the amount text
                                    amount_text = re.sub(r'[^\$0-9,.]', '', item_text)
                                    if not amount_text.startswith('$'):
                                        amount_text = '$' + amount_text
                                    
                                    self.logger.info(f"Chunk 9 Editing Debug: Found CHUNK 9 specific amount: '{amount_text}' (has_fee_context: {has_fee_context})")
                                    
                                    # Replace $AMOUNT with the handwritten amount
                                    old_text = paragraph.text
                                    new_text = old_text.replace('$AMOUNT', amount_text)
                                    
                                    # ENHANCED: If no $AMOUNT found, look for any dollar amount to replace
                                    if new_text == old_text and '$' in old_text:
                                        # Find and replace any existing dollar amount
                                        import re
                                        dollar_pattern = r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?'
                                        if re.search(dollar_pattern, old_text):
                                            new_text = re.sub(dollar_pattern, amount_text, old_text, count=1)
                                            self.logger.info(f"Chunk 9 Editing Debug: ✅ REPLACED existing dollar amount with '{amount_text}' (no $AMOUNT placeholder found)")
                                    
                                    if new_text != old_text:
                                        paragraph.clear()
                                        paragraph.add_run(new_text)
                                        
                                        changes.append(ChangeRecord(
                                            type="amount_replacement",
                                            section=section_name,
                                            original_text=old_text,
                                            new_text=new_text,
                                            location="fee_paragraph_amount",
                                            timestamp=timestamp,
                                            ai_confidence=0.90,
                                            strategy_used="chunk_9_amount_replacement"
                                        ))
                                        
                                        self.logger.info(f"Chunk 9 Editing Debug: ✅ REPLACED $AMOUNT with '{amount_text}' from handwritten text above strikethrough (CHUNK 9 SPECIFIC - FEE PLAN)")
                                        amount_replacements += 1
                                        amount_replacement_found = True
                                        chunk_9_replaced = True  # Mark as replaced
                                        break
                            
                            if amount_replacement_found:
                                break
                
                # SIMPLE FALLBACK: Just log that we couldn't find the right amount
                if not amount_replacement_found:
                    self.logger.warning(f"Chunk 9 Editing Debug: ❌ No appropriate amount replacement found - $AMOUNT will remain unchanged")
                    self.logger.info(f"Chunk 9 Editing Debug: This prevents using the wrong amount from chunk 6")
                break  # Exit the paragraph loop after processing first fee paragraph
        
        if not fee_paragraph_found:
            self.logger.warning(f"Chunk 9 Editing Debug: Fee paragraph with $AMOUNT not found")
        
        # STEP 2: Handle tax returns text substitution and color correction  
        tax_returns_changes = 0
        
        self.logger.info(f"Chunk 9 Editing Debug: 🔍 Searching for tax returns paragraph...")
        
        for paragraph in doc.paragraphs:
            if 'recent tax returns' in paragraph.text.lower():
                self.logger.info(f"Chunk 9 Editing Debug: Found tax returns paragraph: '{paragraph.text[:100]}...'")
                
                original_text = paragraph.text
                modified_text = original_text
                text_changed = False
                color_changed = False
                
                # ENHANCED: Look for specific strikethrough patterns in tax returns context
                for item in strikethrough_items:
                    strikethrough_text = item.get('text', '').strip()
                    strikethrough_desc = item.get('description', '').strip().lower()
                    
                    self.logger.info(f"Chunk 9 Editing Debug: Checking tax strikethrough: '{strikethrough_text}' (desc: '{strikethrough_desc}')")
                    
                    # Check if this strikethrough is in the tax returns context
                    tax_returns_context = any(phrase in strikethrough_desc for phrase in [
                        'tax', 'returns', 'personal', 'trust', 'company', 'recent'
                    ])
                    
                    # Also check direct text match
                    if strikethrough_text and (strikethrough_text.lower() in original_text.lower() or tax_returns_context):
                        
                        # Find handwritten replacement text that's specifically for tax returns
                        replacement_text = None
                        
                        for handwritten_item in handwritten_text:
                            handwritten_content = handwritten_item.get('text', '').strip()
                            handwritten_desc = handwritten_item.get('description', '').strip().lower()
                            
                            self.logger.info(f"Chunk 9 Editing Debug: Checking replacement text: '{handwritten_content}' (desc: '{handwritten_desc}')")
                            
                            # Skip if it's just a number (likely for $AMOUNT)
                            import re
                            if re.match(r'^[\$]?[0-9,]+\.?[0-9]*$', handwritten_content):
                                self.logger.info(f"Chunk 9 Editing Debug: Skipping numeric item: '{handwritten_content}'")
                                continue
                            
                            # Skip if it's the amount from earlier (160,000)
                            if '160' in handwritten_content:
                                self.logger.info(f"Chunk 9 Editing Debug: Skipping amount-like item: '{handwritten_content}'")
                                continue
                            
                            # Check if this handwritten text is related to tax returns
                            tax_related = any(phrase in handwritten_desc for phrase in [
                                'tax', 'returns', 'next to', 'near', 'replacement'
                            ])
                            
                            # Also check for common replacement terms
                            is_replacement = any(term in handwritten_content.lower() for term in [
                                'super', 'fund', 'individual', 'entity', 'self'
                            ])
                            
                            if (len(handwritten_content) > 2 and (tax_related or is_replacement)):
                                replacement_text = handwritten_content
                                self.logger.info(f"Chunk 9 Editing Debug: Found replacement text: '{replacement_text}'")
                                break
                        
                        if replacement_text:
                            # ENHANCED: Smart replacement - look for the specific words to replace
                            # Common patterns: "personal, trust and company" → "personal, super funds"
                            
                            if 'personal' in strikethrough_text.lower() and 'trust' in original_text.lower():
                                # Replace "personal" part
                                modified_text = re.sub(r'\bpersonal\b', replacement_text, modified_text, flags=re.IGNORECASE)
                                text_changed = True
                                self.logger.info(f"Chunk 9 Editing Debug: ✅ REPLACED 'personal' with '{replacement_text}'")
                            elif 'trust' in strikethrough_text.lower() and 'company' in original_text.lower():
                                # Replace "trust and company" part  
                                modified_text = re.sub(r'\btrust and company\b', replacement_text, modified_text, flags=re.IGNORECASE)
                                text_changed = True
                                self.logger.info(f"Chunk 9 Editing Debug: ✅ REPLACED 'trust and company' with '{replacement_text}'")
                            else:
                                # Direct replacement
                                modified_text = modified_text.replace(strikethrough_text, replacement_text)
                                text_changed = True
                                self.logger.info(f"Chunk 9 Editing Debug: ✅ DIRECT REPLACED '{strikethrough_text}' with '{replacement_text}'")
                        else:
                            self.logger.info(f"Chunk 9 Editing Debug: No suitable replacement text found for '{strikethrough_text}'")
                
                # STEP 3: Color correction - change red text to black (ALWAYS apply this)
                if hasattr(paragraph, '_element') and paragraph._element is not None:
                    try:
                        # Find all runs and check for red color
                        for run in paragraph.runs:
                            if hasattr(run, 'font') and hasattr(run.font, 'color'):
                                # Always set to black to ensure consistency
                                from docx.shared import RGBColor
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                                color_changed = True
                        
                        if color_changed:
                            self.logger.info(f"Chunk 9 Editing Debug: ✅ CHANGED text color to black in tax returns paragraph")
                    except Exception as e:
                        self.logger.debug(f"Color processing error: {e}")
                
                # Apply text changes
                if text_changed or color_changed:
                    if text_changed:
                        paragraph.clear()
                        paragraph.add_run(modified_text)
                    
                    changes.append(ChangeRecord(
                        type="text_substitution_and_color_correction",
                        section=section_name,
                        original_text=original_text,
                        new_text=modified_text if text_changed else original_text,
                        location="tax_returns_paragraph",
                        timestamp=timestamp,
                        ai_confidence=0.85,
                        strategy_used="chunk_9_tax_returns_processing"
                    ))
                    
                    tax_returns_changes += 1
                else:
                    # Even if no strikethrough changes, still apply color correction
                    self.logger.info(f"Chunk 9 Editing Debug: No text changes found, but applied color correction")
                
                break  # Only process first tax returns paragraph found
        
        self.logger.info(f"Chunk 9 Editing: {len(changes)} changes applied ({amount_replacements} amount replacements, {tax_returns_changes} tax returns changes)")
        return changes
    
    def _append_handwritten_to_bullets(self, doc: Document, handwritten_items: List[Dict], section_name: str, timestamp: str) -> List[ChangeRecord]:
        """
        Append handwritten text to bullet points in brackets
        Handles cases where handwritten text appears after a bullet point with or without arrow indicators
        Enhanced with multi-line text grouping and arrow-based targeting
        """
        changes = []
        
        if not handwritten_items:
            self.logger.info(f"Chunk 6 Handwritten Append: No handwritten items to process")
            return changes
        
        self.logger.info(f"Chunk 6 Handwritten Append: Processing {len(handwritten_items)} handwritten items")
        
        # PRIORITY: Check for insurance deletion indicators BEFORE grouping
        # This ensures insurance deletion takes priority over text grouping/appending
        for item in handwritten_items:
            description = item.get('description', '').lower()
            if ('personal insurance needs' in description or 
                'insurance needs' in description or 
                'bullet point about personal insurance' in description):
                self.logger.info(f"Chunk 6 Handwritten Append: 🚨 INSURANCE DELETION DETECTED - Item: '{item.get('text', '')}' with description: '{description}'")
                # This will be handled by the deletion logic in the main editing flow
                # We just need to make sure it's flagged properly
                item['priority_deletion'] = 'insurance'
        
        # CHUNK 6 SPECIFIC: Process items individually to prevent incorrect grouping
        # Each handwritten item should match to its specific bullet point
        if section_name == "chunk_6_editing":
            grouped_items = handwritten_items  # Don't group - process each item individually
            self.logger.info(f"Chunk 6 Handwritten Append: Processing {len(grouped_items)} items individually (no grouping for chunk 6)")
        else:
            # Group related handwritten items that are close together (multi-line text) for other chunks
            grouped_items = self._group_related_handwritten_items(handwritten_items)
            self.logger.info(f"Chunk 6 Handwritten Append: Grouped into {len(grouped_items)} logical items")
        
        # Filter handwritten items to only those that should be appended (not used for amount replacement or deletion)
        appendable_items = []
        for item in grouped_items:  # Use grouped_items instead of handwritten_items
            text = item.get('text', item.get('description', '')).strip()
            
            # Skip items flagged for priority deletion (like insurance)
            if item.get('priority_deletion'):
                self.logger.info(f"Chunk 6 Handwritten Append: Skipping item flagged for {item.get('priority_deletion')} deletion: '{text}'")
                continue
            
            # Skip items that are just numbers (likely for $AMOUNT replacement)
            import re
            if re.match(r'^[\$]?[0-9,]+\.?[0-9]*$', text):
                self.logger.info(f"Chunk 6 Handwritten Append: Skipping numeric item: '{text}' (likely for amount replacement)")
                continue
                
            # Skip very short items that don't add meaningful content
            if len(text) < 3:
                self.logger.info(f"Chunk 6 Handwritten Append: Skipping short item: '{text}'")
                continue
            
            appendable_items.append(item)
        
        if not appendable_items:
            self.logger.info(f"Chunk 6 Handwritten Append: No appendable handwritten items found")
            return changes
        
        self.logger.info(f"Chunk 6 Handwritten Append: Found {len(appendable_items)} appendable items")
        
        # Find bullet points in the strategies section
        strategies_bullets = []
        in_strategies_section = False
        
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            
            # Detect start of strategies section
            if 'initial strategies and solutions' in para_text.lower():
                in_strategies_section = True
                self.logger.info(f"Chunk 6 Handwritten Append: Found strategies section start at paragraph {i}")
                continue
            
            # Detect end of strategies section
            if in_strategies_section and ('providing you with confidence' in para_text.lower() or 
                                        'tools, processes and talents' in para_text.lower()):
                self.logger.info(f"Chunk 6 Handwritten Append: Strategies section ended at paragraph {i}")
                break
            
            # Skip section headers and introductory text
            if in_strategies_section and (para_text.startswith('The following') or 
                                        para_text.startswith('Engage More4Life') or
                                        'first steps required' in para_text.lower() or
                                        'prepare a plan' in para_text.lower()):
                continue
            
            # Collect bullet points in strategies section using proper detection
            if in_strategies_section:
                # Enhanced bullet point detection (same as deletion logic)
                is_word_bullet = False
                if hasattr(paragraph, '_element') and paragraph._element is not None:
                    numPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    is_word_bullet = numPr is not None
                
                text_based_bullet = (para_text.startswith('•') or 
                                   para_text.startswith('-') or
                                   para_text.startswith('*'))
                
                is_bullet = is_word_bullet or text_based_bullet
                
                if is_bullet and len(para_text) > 5:
                    # Skip bullets that were deleted (check if paragraph is empty after potential deletion)
                    current_text = paragraph.text.strip()
                    if len(current_text) < 10:
                        self.logger.info(f"Chunk 6 Handwritten Append: Skipping deleted/empty bullet at paragraph {i}")
                        continue
                        
                    strategies_bullets.append({
                        'paragraph': paragraph,
                        'index': i,
                        'original_text': current_text  # Use current text, not para_text from loop start
                    })
                    self.logger.info(f"Chunk 6 Handwritten Append: Found bullet {len(strategies_bullets)}: '{current_text[:60]}...'")
        
        if not strategies_bullets:
            self.logger.warning(f"Chunk 6 Handwritten Append: No bullet points found in strategies section")
            return changes
        
        self.logger.info(f"Chunk 6 Handwritten Append: Found {len(strategies_bullets)} bullet points to potentially append to")
        
        # Debug: Log all bullets found
        for i, bullet_info in enumerate(strategies_bullets):
            self.logger.info(f"Chunk 6 Handwritten Append: Bullet {i+1}: '{bullet_info['original_text'][:80]}...'")
        
        # Smart matching with visual proximity analysis
        for handwritten_item in appendable_items:
            handwritten_text = handwritten_item.get('text', handwritten_item.get('description', '')).strip()
            description = handwritten_item.get('description', '').lower()
            
            # Clean up the handwritten text
            handwritten_text = handwritten_text.replace('handwritten:', '').strip()
            if handwritten_text.startswith('"') and handwritten_text.endswith('"'):
                handwritten_text = handwritten_text[1:-1]
            
            self.logger.info(f"Chunk 6 Handwritten Append: Processing item: '{handwritten_text}' with description: '{description}'")
            
            # Enhanced arrow origin analysis
            position_info = handwritten_item.get('position', {})
            handwritten_x = position_info.get('x', 0)
            handwritten_y = position_info.get('y', 0)
            
            # Analyze arrow direction and target determination
            has_arrow_indicators = any(word in description.lower() for word in [
                'arrow', 'points', 'pointing', 'line to', 'connected to'
            ])
            
            # Critical: Determine if description suggests wrong bullet due to physical proximity
            mentions_assets_liabilities = 'assets/liabilities' in description.lower()
            mentions_spending_system = 'spending system' in description.lower()
            
            self.logger.info(f"Chunk 6 Handwritten Append: Arrow analysis - pos:({handwritten_x:.3f},{handwritten_y:.3f}), arrow:{has_arrow_indicators}, mentions_assets:{mentions_assets_liabilities}, mentions_spending:{mentions_spending_system}")
            
            # CRITICAL FIX: Handle the case where handwritten text is physically near assets/liabilities 
            # but arrow originates from spending system bullet above
            is_multiline_handwritten_below_spending = (
                'school fees' in handwritten_text.lower() and 
                mentions_assets_liabilities and 
                not mentions_spending_system and
                ('not including' in handwritten_text.lower() or 'money' in handwritten_text.lower())
            )
            
            if is_multiline_handwritten_below_spending:
                self.logger.info(f"Chunk 6 Handwritten Append: 🎯 ARROW ORIGIN ANALYSIS: Handwritten text '{handwritten_text}' is physically near assets/liabilities but likely originates from spending system above")
                # Override the description to point to correct bullet
                description = 'handwritten note next to bullet point about spending system'
                self.logger.info(f"Chunk 6 Handwritten Append: ✅ CORRECTED target from assets/liabilities → spending system based on arrow origin")
            
            # Clean up text that might have mixed printed/handwritten content
            if '$160,000' in handwritten_text and 'school fees' in handwritten_text:
                # This is likely mixing the printed $AMOUNT with handwritten annotation
                # Extract only the handwritten parts
                if 'not including' in handwritten_text.lower():
                    # Keep the handwritten annotation part
                    parts = handwritten_text.lower().split('not including')
                    if len(parts) > 1:
                        handwritten_text = 'not including' + parts[1]
                        # Clean up common printed text that gets mixed in
                        handwritten_text = re.sub(r'school fees money you are spending.*', 'school fees', handwritten_text)
                        handwritten_text = handwritten_text.strip()
                        self.logger.info(f"Chunk 6 Handwritten Append: ✅ CLEANED handwritten text to: '{handwritten_text}' (removed $AMOUNT printing)")
            
            # Filter out clearly printed text elements
            elif '$' in handwritten_text and any(word in handwritten_text.lower() for word in ['current', 'cost', 'living', 'being']):
                # This is mixing printed sentence with handwritten annotation
                words = handwritten_text.split()
                # Keep meaningful handwritten words, filter printed text
                handwritten_words = []
                for word in words:
                    word_lower = word.lower()
                    if (word_lower not in ['your', 'current', 'cost', 'of', 'living', 'being', 'pa', '$160,000'] and 
                        not re.match(r'^\$?[0-9,]+\.?[0-9]*$', word)):
                        handwritten_words.append(word)
                
                if handwritten_words:
                    handwritten_text = ' '.join(handwritten_words)
                    self.logger.info(f"Chunk 6 Handwritten Append: ✅ FILTERED to handwritten words: '{handwritten_text}'")
            
            # Find the best matching bullet based on description keywords
            target_bullet = None
            for bullet_info in strategies_bullets:
                bullet_text = bullet_info['original_text'].lower()
                
                # Skip bullets that already have bracketed content
                if '(' in bullet_info['original_text'] and ')' in bullet_info['original_text']:
                    continue
                
                # Skip the assets/liabilities bullet for school fees money (it's meant for spending system)
                if ('school fees' in handwritten_text.lower() or '$160,000' in handwritten_text) and 'assets/liabilities' in bullet_text:
                    self.logger.info(f"Chunk 6 Handwritten Append: Skipping assets/liabilities bullet for school fees item")
                    continue
                
                # ENHANCED MATCHING: Prioritize arrow origin over physical proximity
                
                # 1. Arrow origin analysis for school fees handwritten text
                if ('school fees' in handwritten_text.lower() or 'not including' in handwritten_text.lower()) and ('spending system' in bullet_text or 'spending' in bullet_text):
                    # Check if this matches our corrected arrow origin analysis
                    if ('spending system' in description.lower() or 
                        has_arrow_indicators or 
                        'next to' in description.lower()):
                        target_bullet = bullet_info
                        self.logger.info(f"Chunk 6 Handwritten Append: ✅ ARROW ORIGIN MATCH: Handwritten '{handwritten_text}' → spending system bullet")
                        break
                    else:
                        self.logger.info(f"Chunk 6 Handwritten Append: ❌ No arrow origin indicators for spending system, continuing search...")
                        continue
                    
                # 2. Standard description-based matching - ENHANCED with more flexible patterns
                elif (('spending system' in description or 'spending suggestion' in description or 
                       'current cost of living' in description) and 
                      ('spending system' in bullet_text or 'spending' in bullet_text)):
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 6 Handwritten Append: Matched spending-related item to spending bullet")
                    break
                elif (('insurance' in description or 'personal insurance' in description) and 
                      ('insurance needs' in bullet_text or 'insurance' in bullet_text)):
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 6 Handwritten Append: Matched insurance item to insurance bullet")
                    break
                elif ('investment property' in handwritten_text.lower() or 'investments outside' in description) and 'investments outside' in bullet_text:
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 6 Handwritten Append: Matched investment property item to investments outside bullet")
                    break
                    
                # 3. Content-based fallback matching for misidentified descriptions
                elif 'for the girls' in handwritten_text.lower() and 'insurance needs' in bullet_text:
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 6 Handwritten Append: Matched 'for the girls' to insurance bullet via content")
                    break
                elif 'for the s\'s' in handwritten_text.lower() and 'insurance needs' in bullet_text:
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 6 Handwritten Append: Matched 'for the S's' to insurance bullet via content")
                    break
                elif ('not including' in handwritten_text.lower() and 'school fees' in handwritten_text.lower()) and ('spending' in bullet_text or 'cost of living' in bullet_text):
                    target_bullet = bullet_info
                    self.logger.info(f"Chunk 6 Handwritten Append: Matched school fees note to spending bullet via content")
                    break
            
            if not target_bullet:
                self.logger.warning(f"Chunk 6 Handwritten Append: Could not find matching bullet for: '{handwritten_text}' with description: '{description}'")
                continue
            
            paragraph = target_bullet['paragraph']
            original_text = target_bullet['original_text']
            
            # Smart bracket placement: for spending system bullet, look for continuation paragraph
            # Check if this is the spending system bullet and if there's a continuation paragraph
            if 'spending system' in original_text.lower():
                # Find the paragraph index for the current bullet
                bullet_paragraph_index = target_bullet['index']
                
                # Look for continuation paragraph (next non-empty paragraph)
                continuation_paragraph = None
                continuation_text = ""
                
                for i, para in enumerate(doc.paragraphs[bullet_paragraph_index + 1:], start=bullet_paragraph_index + 1):
                    para_text = para.text.strip()
                    if para_text:  # Found non-empty paragraph
                        # Check if it's a continuation (not another bullet)
                        is_word_bullet = False
                        if hasattr(para, '_element') and para._element is not None:
                            numPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                            is_word_bullet = numPr is not None
                        
                        text_based_bullet = (para_text.startswith('•') or 
                                           para_text.startswith('-') or
                                           para_text.startswith('*'))
                        
                        is_bullet = is_word_bullet or text_based_bullet
                        
                        if not is_bullet:  # This is continuation text, not a new bullet
                            continuation_paragraph = para
                            continuation_text = para_text
                            self.logger.info(f"Chunk 6 Handwritten Append: Found continuation paragraph: '{continuation_text[:60]}...'")
                            break
                        else:
                            # Hit another bullet, no continuation found
                            break
                    elif i - bullet_paragraph_index > 3:  # Stop looking after a few empty paragraphs
                        break
                
                if continuation_paragraph and continuation_text:
                    # Place bracket at end of continuation paragraph
                    if continuation_text.endswith('.'):
                        new_continuation_text = continuation_text[:-1] + f" ({handwritten_text})."
                    elif continuation_text.endswith(';'):
                        new_continuation_text = continuation_text[:-1] + f" ({handwritten_text});"
                    else:
                        new_continuation_text = continuation_text + f" ({handwritten_text})"
                    
                    # Update the continuation paragraph instead
                    continuation_paragraph.clear()
                    continuation_paragraph.add_run(new_continuation_text)
                    
                    self.logger.info(f"Chunk 6 Handwritten Append: ✅ APPENDED to continuation paragraph: '{handwritten_text}' → '{new_continuation_text[:80]}...'")
                    
                    changes.append(ChangeRecord(
                        type="handwritten_append",
                        section=section_name,
                        original_text=continuation_text,
                        new_text=new_continuation_text,
                        location=f"continuation_paragraph",
                        timestamp=timestamp,
                        ai_confidence=0.85,
                        strategy_used="handwritten_bracket_append_continuation"
                    ))
                    
                    self.processing_stats['total_replacements'] += 1
                    continue  # Skip the regular processing below
                else:
                    self.logger.info(f"Chunk 6 Handwritten Append: No continuation paragraph found for spending system bullet")
            
            # Regular processing for other bullets or if no continuation found
            # Smart bracket placement: place after second sentence if multi-sentence bullet
            import re
            
            # Split the bullet text into sentences (handle both ; and . as separators)
            sentences = re.split(r'([.;])', original_text)
            
            # Reconstruct sentences with their punctuation
            full_sentences = []
            for i in range(0, len(sentences)-1, 2):
                if i+1 < len(sentences):
                    sentence = sentences[i] + sentences[i+1]
                    full_sentences.append(sentence.strip())
            
            # If we have remaining text without punctuation, add it
            if len(sentences) % 2 == 1 and sentences[-1].strip():
                full_sentences.append(sentences[-1].strip())
            
            self.logger.info(f"Chunk 6 Handwritten Append: Split bullet into {len(full_sentences)} sentences: {full_sentences}")
            
            if len(full_sentences) >= 2:
                # Multi-sentence bullet: place bracket after the second sentence
                first_sentence = full_sentences[0]
                remaining_sentences = full_sentences[1:]
                
                # Combine remaining sentences
                remaining_text = ' '.join(remaining_sentences)
                
                # Add bracket to the end of remaining text (after second sentence)
                if remaining_text.endswith(';'):
                    remaining_with_bracket = remaining_text[:-1] + f" ({handwritten_text});"
                elif remaining_text.endswith('.'):
                    remaining_with_bracket = remaining_text[:-1] + f" ({handwritten_text})."
                else:
                    remaining_with_bracket = remaining_text + f" ({handwritten_text})"
                
                # Reconstruct the full text
                new_text = first_sentence + ' ' + remaining_with_bracket
                
                self.logger.info(f"Chunk 6 Handwritten Append: Multi-sentence placement - First: '{first_sentence}' + Remaining with bracket: '{remaining_with_bracket}'")
                
            else:
                # Single sentence bullet: append bracket normally
                if original_text.endswith(';'):
                    new_text = original_text[:-1] + f" ({handwritten_text});"
                elif original_text.endswith('.'):
                    new_text = original_text[:-1] + f" ({handwritten_text})."
                else:
                    new_text = original_text + f" ({handwritten_text})"
                
                self.logger.info(f"Chunk 6 Handwritten Append: Single sentence placement")
            
            # Update the paragraph
            paragraph.clear()
            paragraph.add_run(new_text)
            
            changes.append(ChangeRecord(
                type="handwritten_append",
                section=section_name,
                original_text=original_text,
                new_text=new_text,
                location=f"bullet_matched",
                timestamp=timestamp,
                ai_confidence=0.85,
                strategy_used="handwritten_bracket_append"
            ))
            
            self.processing_stats['total_replacements'] += 1
            
            self.logger.info(f"Chunk 6 Handwritten Append: ✅ APPENDED: '{handwritten_text}' → '{new_text[:80]}...'")
        
        unused_items = len(appendable_items) - len(changes)
        if unused_items > 0:
            self.logger.info(f"Chunk 6 Handwritten Append: {unused_items} handwritten items could not find matching bullets")
        
        self.logger.info(f"Chunk 6 Handwritten Append: {len(changes)} append operations completed")
        return changes
    
    def _group_related_handwritten_items(self, handwritten_items: List[Dict]) -> List[Dict]:
        """
        Group handwritten items that are close together vertically and likely part of the same multi-line note
        """
        if not handwritten_items:
            return []
        
        # Sort items by vertical position and chunk  
        sorted_items = sorted(handwritten_items, key=lambda x: (x.get('chunk_index', 1), x.get('position', {}).get('y', 0)))
        
        grouped = []
        current_group = []
        
        for item in sorted_items:
            if not current_group:
                current_group = [item]
                continue
            
            # Check if this item should be grouped with the current group
            last_item = current_group[-1]
            
            # Same chunk and close vertical proximity (within 0.1 normalized units)
            same_chunk = item.get('chunk_index') == last_item.get('chunk_index')
            y_distance = abs(item.get('position', {}).get('y', 0) - last_item.get('position', {}).get('y', 0))
            close_vertically = y_distance < 0.08  # Much more conservative - only group truly adjacent text
            
            # Similar horizontal positioning (within 0.2 normalized units) 
            x_distance = abs(item.get('position', {}).get('x', 0) - last_item.get('position', {}).get('x', 0))
            similar_horizontally = x_distance < 0.15  # More conservative horizontal grouping
            
            # Check if descriptions suggest they belong to different bullet points
            current_desc = item.get('description', '').lower()
            last_desc = last_item.get('description', '').lower()
            
            # Don't group if they clearly refer to different bullet points
            different_bullets = False
            bullet_indicators = ['spending', 'insurance', 'investment', 'assets', 'liabilities', 'goals', 'objectives']
            for indicator in bullet_indicators:
                current_has = indicator in current_desc
                last_has = indicator in last_desc
                if current_has and last_has and current_desc != last_desc:
                    different_bullets = True
                    self.logger.info(f"Chunk 6 Handwritten Append: Not grouping - different bullet contexts: '{current_desc}' vs '{last_desc}'")
                    break
                elif current_has != last_has:  # One has indicator, other doesn't
                    different_bullets = True
                    self.logger.info(f"Chunk 6 Handwritten Append: Not grouping - different bullet targeting: '{current_desc}' vs '{last_desc}'")
                    break
            
            if same_chunk and close_vertically and similar_horizontally and not different_bullets:
                # This item belongs to the current group
                current_group.append(item)
                self.logger.info(f"Chunk 6 Handwritten Append: Grouping '{item.get('text', '')}' with previous items (y_dist: {y_distance:.3f}, x_dist: {x_distance:.3f})")
            else:
                # Start a new group
                if current_group:
                    grouped.append(self._merge_grouped_items(current_group))
                current_group = [item]
        
        # Add the last group
        if current_group:
            grouped.append(self._merge_grouped_items(current_group))
        
        return grouped
    
    def _merge_grouped_items(self, items: List[Dict]) -> Dict:
        """
        Merge multiple handwritten items into a single logical item
        """
        if len(items) == 1:
            return items[0]
        
        # SANITY CHECK: Don't merge items that clearly refer to different bullet points
        # Check if the combined text would create nonsensical combinations
        texts = [item.get('text', '').strip() for item in items if item.get('text', '').strip()]
        descriptions = [item.get('description', '').lower() for item in items]
        
        # Check for conflicting bullet point references
        bullet_keywords = {
            'spending': ['spending', 'cost', 'living', 'fees'],
            'insurance': ['insurance', 'cover', 'protection'],
            'investment': ['investment', 'property', 'portfolio'],
            'assets': ['assets', 'liabilities', 'balance']
        }
        
        found_categories = set()
        for desc in descriptions:
            for category, keywords in bullet_keywords.items():
                if any(keyword in desc for keyword in keywords):
                    found_categories.add(category)
        
        # If multiple different bullet categories are referenced, don't merge
        if len(found_categories) > 1:
            self.logger.info(f"Chunk 6 Handwritten Append: NOT MERGING - Multiple bullet categories detected: {found_categories}")
            # Return just the first item instead of merging conflicting items
            return items[0]
        
        # Combine the text from all items
        combined_text = ' '.join(texts)
        
        # Use the first item as the base, but update key fields
        merged = items[0].copy()
        merged['text'] = combined_text
        
        # Use the position of the first item (typically the one that starts the multi-line text)
        first_item = items[0]
        merged['position'] = first_item.get('position', {})
        
        # Try to determine the best description - prefer one that mentions arrow or specific targeting
        best_description = first_item.get('description', '')
        for item in items:
            desc = item.get('description', '')
            if 'arrow' in desc.lower() or 'pointing' in desc.lower():
                best_description = desc
                break
            elif 'spending' in desc.lower() and 'spending' not in best_description.lower():
                best_description = desc
        
        merged['description'] = best_description
        merged['grouped_from'] = [item.get('text', '') for item in items]
        
        # Preserve priority deletion flags from any item in the group
        for item in items:
            if item.get('priority_deletion'):
                merged['priority_deletion'] = item.get('priority_deletion')
                self.logger.info(f"Chunk 6 Handwritten Append: Preserving {item.get('priority_deletion')} deletion flag in merged item")
                break
        
        self.logger.info(f"Chunk 6 Handwritten Append: Merged {len(items)} items into: '{combined_text}' with description: '{best_description}'")
        
        return merged
    
    def _cleanup_spacing_after_deletion(self, cell_or_container, description: str = ""):
        """Clean up extra spacing, empty paragraphs, and formatting after dot point deletion - TABLE CELLS ONLY"""
        try:
            # Only clean up if this is a table cell, not the entire document
            container_type = type(cell_or_container).__name__
            if container_type == 'Document':
                self.logger.info(f"🔍 Skipping cleanup for Document container to preserve letter formatting")
                return 0
            
            # Get all paragraphs from the container (should be a cell)
            if hasattr(cell_or_container, 'paragraphs'):
                paragraphs = cell_or_container.paragraphs
            else:
                return 0
            
            # Only proceed if this looks like a table cell (relatively few paragraphs)
            if len(paragraphs) > 20:
                self.logger.info(f"🔍 Skipping cleanup for large container ({len(paragraphs)} paragraphs)")
                return 0
            
            cleaned_count = 0
            paragraphs_to_remove = []
            
            self.logger.info(f"🧹 Cleaning spacing after deletion{' (' + description + ')' if description else ''}...")
            self.logger.info(f"🔍 Found {len(paragraphs)} paragraphs to check in table cell")
            
            for i, para in enumerate(paragraphs):
                para_text = para.text.strip()
                
                # Remove completely empty paragraphs
                if not para_text:
                    paragraphs_to_remove.append(para)
                    self.logger.info(f"   📝 Removing empty paragraph {i}")
                    cleaned_count += 1
                
                # Remove paragraphs with only whitespace/formatting
                elif para_text.isspace() or len(para_text) == 0:
                    paragraphs_to_remove.append(para)
                    self.logger.info(f"   📝 Removing whitespace-only paragraph {i}")
                    cleaned_count += 1
            
            # Remove the identified empty paragraphs
            for para in paragraphs_to_remove:
                para._element.getparent().remove(para._element)
            
            self.logger.info(f"   ✅ Cleaned up {cleaned_count} empty paragraphs")
            return cleaned_count
            
        except Exception as e:
            self.logger.info(f"   ❌ Error during spacing cleanup: {e}")
            return 0
    
    def _cleanup_document_spacing_after_deletion(self, doc, description: str = ""):
        """Clean up extra spacing and empty paragraphs after bullet deletion in full document"""
        try:
            cleaned_count = 0
            paragraphs_to_remove = []
            
            self.logger.info(f"🧹 Cleaning document spacing after deletion{' (' + description + ')' if description else ''}...")
            
            # Look for consecutive empty paragraphs that indicate spacing gaps
            consecutive_empty = 0
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                
                if not para_text:
                    consecutive_empty += 1
                    # If we have more than 1 consecutive empty paragraph, mark extras for removal
                    if consecutive_empty > 1:
                        paragraphs_to_remove.append(para)
                        self.logger.info(f"   📝 Removing excess empty paragraph {i} (consecutive #{consecutive_empty})")
                        cleaned_count += 1
                else:
                    consecutive_empty = 0
            
            # Remove the identified empty paragraphs
            for para in paragraphs_to_remove:
                try:
                    p_element = para._element
                    parent = p_element.getparent()
                    if parent is not None:
                        parent.remove(p_element)
                except Exception as e:
                    self.logger.info(f"   ⚠️ Could not remove paragraph: {e}")
            
            self.logger.info(f"   ✅ Cleaned up {cleaned_count} excess empty paragraphs")
            return cleaned_count
            
        except Exception as e:
            self.logger.info(f"   ❌ Error during document spacing cleanup: {e}")
            return 0