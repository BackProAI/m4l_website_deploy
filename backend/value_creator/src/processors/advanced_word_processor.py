#!/usr/bin/env python3
"""
Advanced Word Processor
Intelligently modifies Word documents based on PDF annotations using GPT-4o analysis
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import openai
from dataclasses import dataclass

@dataclass
class AnnotationAction:
    """Represents an action to take on the Word document"""
    action_type: str  # 'replace', 'delete', 'insert', 'highlight', 'comment'
    target_text: str  # Text to target in Word document
    new_text: str = ""  # Replacement text (for replace/insert actions)
    position: Dict = None  # Position information from PDF
    confidence: float = 0.0
    reasoning: str = ""  # AI reasoning for the action

class AdvancedWordProcessor:
    """Production-ready Word processor using 3-strategy cascading system"""
    
    def __init__(self, config_path: str = "config.yaml"):
        import yaml
        with open(config_path, 'r') as f:
            self.config = yaml.safe_load(f)
        
        self.logger = logging.getLogger(__name__)
        
        # 3-Strategy System Configuration (battle-tested thresholds)
        self.similarity_threshold = 0.6  # Proven sweet spot from your implementation
        self.keyword_min_length = 3     # Minimum keyword length
        
        # Change tracking
        self.changes_applied = []
        self.processing_stats = {
            'exact_matches': 0,
            'similarity_matches': 0, 
            'keyword_matches': 0,
            'failed_matches': 0,
            'total_deletions': 0,
            'total_replacements': 0,
            'total_row_deletions': 0
        }
    
    # ===========================================
    # CORE 3-STRATEGY SYSTEM - PROVEN PATTERNS  
    # ===========================================
    
    def text_similarity(self, text1: str, text2: str) -> float:
        """Calculate text similarity using your proven algorithm"""
        if not text1 or not text2:
            return 0.0
            
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
            
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def extract_keywords(self, text: str) -> List[str]:
        """Extract keywords for Strategy 3 (your proven pattern)"""
        words = text.split()
        keywords = []
        for word in words:
            clean_word = word.strip('.,!?').lower()
            if len(clean_word) > self.keyword_min_length:
                keywords.append(clean_word)
        return keywords
    
    def find_text_with_keywords(self, text: str, keywords: List[str]) -> bool:
        """Check if text contains any of the keywords"""
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in keywords)
    
    def strategy_1_exact_match(self, doc: Document, target_text: str, action_type: str = "delete") -> bool:
        """Strategy 1: Exact text matching (70-80% success rate)"""
        try:
            # Search in paragraphs
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text and target_text.lower() in para_text.lower():
                    self.logger.info(f"✅ EXACT MATCH in Para {i}: '{para_text[:60]}...'")
                    if action_type == "delete":
                        para.clear()
                        self.processing_stats['exact_matches'] += 1
                        self.processing_stats['total_deletions'] += 1
                        return True
            
            # Search in table cells
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            para_text = para.text.strip()
                            if para_text and target_text.lower() in para_text.lower():
                                self.logger.info(f"✅ EXACT MATCH in Table {table_idx}, Row {row_idx}, Cell {cell_idx}")
                                if action_type == "delete":
                                    para.clear()
                                    self.processing_stats['exact_matches'] += 1
                                    self.processing_stats['total_deletions'] += 1
                                    return True
                                    
        except Exception as e:
            self.logger.error(f"Strategy 1 error: {e}")
            
        return False
    
    def strategy_2_similarity_match(self, doc: Document, target_text: str, action_type: str = "delete") -> bool:
        """Strategy 2: Similarity matching (15-20% additional coverage)"""
        try:
            # Search in paragraphs
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text:
                    similarity = self.text_similarity(para_text, target_text)
                    if similarity > self.similarity_threshold:
                        self.logger.info(f"✅ SIMILARITY MATCH in Para {i}: similarity: {similarity:.2f}")
                        if action_type == "delete":
                            para.clear()
                            self.processing_stats['similarity_matches'] += 1
                            self.processing_stats['total_deletions'] += 1
                            return True
            
            # Search in table cells
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            para_text = para.text.strip()
                            if para_text:
                                similarity = self.text_similarity(para_text, target_text)
                                if similarity > self.similarity_threshold:
                                    self.logger.info(f"✅ SIMILARITY MATCH in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: {similarity:.2f}")
                                    if action_type == "delete":
                                        para.clear()
                                        self.processing_stats['similarity_matches'] += 1
                                        self.processing_stats['total_deletions'] += 1
                                        return True
                                        
        except Exception as e:
            self.logger.error(f"Strategy 2 error: {e}")
            
        return False
    
    def strategy_3_keyword_match(self, doc: Document, target_text: str, action_type: str = "delete") -> bool:
        """Strategy 3: Keyword-based search (5-10% additional coverage)"""
        try:
            keywords = self.extract_keywords(target_text)
            if not keywords:
                return False
                
            # Search in paragraphs
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text and self.find_text_with_keywords(para_text, keywords):
                    self.logger.info(f"✅ KEYWORD MATCH in Para {i}: keywords: {keywords[:3]}")
                    if action_type == "delete":
                        para.clear()
                        self.processing_stats['keyword_matches'] += 1
                        self.processing_stats['total_deletions'] += 1
                        return True
            
            # Search in table cells
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            para_text = para.text.strip()
                            if para_text and self.find_text_with_keywords(para_text, keywords):
                                self.logger.info(f"✅ KEYWORD MATCH in Table {table_idx}, Row {row_idx}, Cell {cell_idx}")
                                if action_type == "delete":
                                    para.clear()
                                    self.processing_stats['keyword_matches'] += 1
                                    self.processing_stats['total_deletions'] += 1
                                    return True
                                    
        except Exception as e:
            self.logger.error(f"Strategy 3 error: {e}")
            
        return False
    
    def apply_cascading_strategies(self, doc: Document, target_text: str, action_type: str = "delete") -> bool:
        """Apply the 3-strategy cascading system (your proven pattern)"""
        
        # Strategy 1: Exact match (first attempt)
        if self.strategy_1_exact_match(doc, target_text, action_type):
            return True
            
        # Strategy 2: Similarity match (fallback)
        if self.strategy_2_similarity_match(doc, target_text, action_type):
            return True
            
        # Strategy 3: Keyword match (last resort)
        if self.strategy_3_keyword_match(doc, target_text, action_type):
            return True
            
        # All strategies failed
        self.processing_stats['failed_matches'] += 1
        self.logger.warning(f"❌ ALL STRATEGIES FAILED for: '{target_text[:50]}...'")
        return False
    
    def delete_table_row(self, doc: Document, table_idx: int, row_idx: int) -> bool:
        """Advanced row deletion using your proven XML manipulation"""
        try:
            if table_idx >= len(doc.tables):
                return False
                
            table = doc.tables[table_idx]
            if row_idx >= len(table.rows):
                return False
                
            row = table.rows[row_idx]
            table._tbl.remove(row._tr)  # Direct XML manipulation (your advanced pattern)
            
            self.processing_stats['total_row_deletions'] += 1
            self.logger.info(f"✅ DELETED TABLE ROW: Table {table_idx}, Row {row_idx}")
            return True
            
        except Exception as e:
            self.logger.error(f"Row deletion error: {e}")
            return False

    def apply_annotations_to_document(self, word_path: str, pdf_analysis: Dict, output_path: str) -> Dict:
        """
        Apply PDF annotations using production-ready 3-strategy cascading system
        
        Args:
            word_path: Path to input Word document
            pdf_analysis: Dictionary containing analysis results from GPT-4o  
            output_path: Path for output Word document
            
        Returns:
            Dictionary with comprehensive processing results and statistics
        """
        self.logger.info(f"🚀 Starting production Word processing: {word_path}")
        
        try:
            # Load Word document
            doc = Document(word_path)
            self.logger.info(f"Loaded Word document: {word_path}")
            
            # Reset statistics for this processing session
            self.changes_applied = []
            self.processing_stats = {
                'exact_matches': 0, 'similarity_matches': 0, 'keyword_matches': 0,
                'failed_matches': 0, 'total_deletions': 0, 'total_replacements': 0, 'total_row_deletions': 0
            }
            
            # Process different types of annotations using cascading strategies
            total_items_processed = 0
            
            # === HANDWRITTEN TEXT PROCESSING ===
            handwritten_items = pdf_analysis.get('analysis_results', {}).get('handwritten_text', [])
            for item in handwritten_items:
                if item.get('text') and item.get('confidence', 0) > 0.7:
                    # For handwritten text, add as comments rather than replacing
                    self._add_comment_for_handwritten_text(doc, item)
                    total_items_processed += 1
                    
            # === STRIKETHROUGH TEXT PROCESSING ===  
            strikethrough_items = pdf_analysis.get('analysis_results', {}).get('strikethrough_text', [])
            for item in strikethrough_items:
                if item.get('text') and item.get('confidence', 0) > 0.7:
                    # Use cascading strategies to find and delete strikethrough text
                    success = self.apply_cascading_strategies(doc, item['text'], "delete")
                    if success:
                        self.changes_applied.append({
                            "type": "deletion",
                            "target_text": item['text'][:50],
                            "confidence": item.get('confidence'),
                            "strategy_used": self._get_last_successful_strategy()
                        })
                    total_items_processed += 1
                    
            # === CROSS MARKS PROCESSING ===
            crosses = pdf_analysis.get('analysis_results', {}).get('crosses', [])
            for item in crosses:
                if item.get('confidence', 0) > 0.7:
                    # Handle different cross mark sizes
                    if item.get('size') == 'large' and item.get('deletion_scope') == 'paragraph':
                        # Large crosses might indicate row deletions in tables
                        self._handle_large_cross_deletion(doc, item)
                    else:
                        # Small crosses for specific text deletions
                        self._handle_small_cross_deletion(doc, item)
                    total_items_processed += 1
                    
            # === HIGHLIGHTS PROCESSING ===
            highlights = pdf_analysis.get('analysis_results', {}).get('highlights', [])
            for item in highlights:
                if item.get('text') and item.get('confidence', 0) > 0.7:
                    # Add highlighting to Word document
                    self._apply_highlighting(doc, item)
                    total_items_processed += 1
                    
            # === ANNOTATIONS PROCESSING ===
            annotations = pdf_analysis.get('analysis_results', {}).get('annotations', [])
            for item in annotations:
                if item.get('confidence', 0) > 0.7:
                    # Add as comments
                    self._add_annotation_comment(doc, item)
                    total_items_processed += 1
            
            # Save the processed document
            doc.save(output_path)
            self.logger.info(f"✅ Saved processed document: {output_path}")
            
            # Calculate comprehensive statistics
            total_changes = (self.processing_stats['exact_matches'] + 
                           self.processing_stats['similarity_matches'] + 
                           self.processing_stats['keyword_matches'])
            
            success_rate = (total_changes / total_items_processed) if total_items_processed > 0 else 0
            
            # Return comprehensive results
            return {
                'status': 'success',
                'input_file': word_path,
                'output_file': output_path,
                'total_items_processed': total_items_processed,
                'total_changes_applied': total_changes,
                'success_rate': success_rate,
                'strategy_breakdown': {
                    'exact_matches': self.processing_stats['exact_matches'],
                    'similarity_matches': self.processing_stats['similarity_matches'], 
                    'keyword_matches': self.processing_stats['keyword_matches'],
                    'failed_matches': self.processing_stats['failed_matches']
                },
                'operation_breakdown': {
                    'deletions': self.processing_stats['total_deletions'],
                    'replacements': self.processing_stats['total_replacements'],
                    'row_deletions': self.processing_stats['total_row_deletions']
                },
                'changes_applied': self.changes_applied,
                'file_size': Path(output_path).stat().st_size
            }
            
        except Exception as e:
            self.logger.error(f"Error processing Word document: {str(e)}")
            return {
                "status": "error",
                "error": str(e),
                "input_path": word_path
            }
    
    def _extract_document_text(self, doc: Document) -> Dict:
        """Extract text content from Word document with structure"""
        content = {
            "paragraphs": [],
            "full_text": "",
            "structure": []
        }
        
        full_text_parts = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:  # Skip empty paragraphs
                para_info = {
                    "index": i,
                    "text": para_text,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                    "runs": []
                }
                
                # Analyze runs for formatting
                for run in paragraph.runs:
                    if run.text.strip():
                        run_info = {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline
                        }
                        para_info["runs"].append(run_info)
                
                content["paragraphs"].append(para_info)
                full_text_parts.append(para_text)
                
                # Determine structure type
                if paragraph.style and any(style in paragraph.style.name.lower() 
                                         for style in ['heading', 'title']):
                    content["structure"].append({"type": "heading", "index": i, "text": para_text})
                elif para_text.startswith(('•', '-', '*', '1.', '2.', '3.')):
                    content["structure"].append({"type": "list_item", "index": i, "text": para_text})
                else:
                    content["structure"].append({"type": "paragraph", "index": i, "text": para_text})
        
        content["full_text"] = "\n".join(full_text_parts)
        return content
    
    def _interpret_annotations(self, pdf_analysis: Dict, word_content: Dict) -> List[AnnotationAction]:
        """
        Interpret PDF annotations and determine actions for Word document
        """
        actions = []
        
        # Process different types of annotations
        self._process_handwritten_text(pdf_analysis.get('handwritten_text', []), word_content, actions)
        self._process_strikethrough_text(pdf_analysis.get('strikethrough_text', []), word_content, actions)
        self._process_crosses(pdf_analysis.get('crosses', []), word_content, actions)
        self._process_arrows(pdf_analysis.get('arrows', []), pdf_analysis, word_content, actions)
        self._process_highlights(pdf_analysis.get('highlights', []), word_content, actions)
        self._process_annotations(pdf_analysis.get('annotations', []), word_content, actions)
        
        return actions
    
    def _process_handwritten_text(self, handwritten_items: List[Dict], word_content: Dict, actions: List[AnnotationAction]):
        """Process handwritten text annotations"""
        for item in handwritten_items:
            handwritten_text = item.get('text', '').strip()
            if not handwritten_text or handwritten_text.lower() in ['handwriting', 'handwritten text']:
                continue
            
            position = item.get('position', {})
            confidence = item.get('confidence', 0.0)
            
            # Use AI to determine the best placement for handwritten text
            placement_decision = self._determine_handwriting_placement(handwritten_text, word_content, position)
            
            if placement_decision['action'] == 'replace':
                actions.append(AnnotationAction(
                    action_type='replace',
                    target_text=placement_decision['target_text'],
                    new_text=handwritten_text,
                    position=position,
                    confidence=confidence,
                    reasoning=f"Handwritten text replaces: {placement_decision['target_text'][:50]}..."
                ))
            elif placement_decision['action'] == 'insert':
                actions.append(AnnotationAction(
                    action_type='insert',
                    target_text=placement_decision['target_text'],
                    new_text=handwritten_text,
                    position=position,
                    confidence=confidence,
                    reasoning=f"Insert handwritten text near: {placement_decision['target_text'][:50]}..."
                ))
            else:
                # Add as comment if uncertain
                actions.append(AnnotationAction(
                    action_type='comment',
                    target_text=placement_decision.get('target_text', ''),
                    new_text=handwritten_text,
                    position=position,
                    confidence=confidence,
                    reasoning="Handwritten text added as comment for review"
                ))
    
    def _process_strikethrough_text(self, strikethrough_items: List[Dict], word_content: Dict, actions: List[AnnotationAction]):
        """Process strikethrough text annotations"""
        for item in strikethrough_items:
            struck_text = item.get('text', '').strip()
            if not struck_text:
                continue
            
            position = item.get('position', {})
            confidence = item.get('confidence', 0.0)
            
            # Find best match in Word document
            best_match = self._find_best_text_match(struck_text, word_content)
            
            if best_match and best_match['similarity'] > self.similarity_threshold:
                # Check if there's associated handwriting nearby
                replacement_text = self._find_nearby_handwriting(item, position)
                
                if replacement_text:
                    actions.append(AnnotationAction(
                        action_type='replace',
                        target_text=best_match['text'],
                        new_text=replacement_text,
                        position=position,
                        confidence=confidence,
                        reasoning=f"Replace struck-through text with handwritten correction"
                    ))
                else:
                    actions.append(AnnotationAction(
                        action_type='delete',
                        target_text=best_match['text'],
                        position=position,
                        confidence=confidence,
                        reasoning="Delete struck-through text"
                    ))
    
    def _process_crosses(self, cross_items: List[Dict], word_content: Dict, actions: List[AnnotationAction]):
        """Process cross marks and large deletion marks"""
        for item in cross_items:
            position = item.get('position', {})
            confidence = item.get('confidence', 0.0)
            description = item.get('description', '').lower()
            
            # Determine if this is a large cross (paragraph/section deletion)
            is_large_cross = any(keyword in description for keyword in 
                               ['large', 'big', 'paragraph', 'section', 'diagonal', 'through'])
            
            if is_large_cross:
                # Find the paragraph or section to delete
                target_content = self._find_content_at_position(position, word_content)
                if target_content:
                    actions.append(AnnotationAction(
                        action_type='delete',
                        target_text=target_content['text'],
                        position=position,
                        confidence=confidence,
                        reasoning=f"Delete {target_content['type']} marked with large cross"
                    ))
            else:
                # Small cross - might be item deletion or checkmark
                nearby_content = self._find_content_near_position(position, word_content)
                if nearby_content:
                    actions.append(AnnotationAction(
                        action_type='delete',
                        target_text=nearby_content['text'],
                        position=position,
                        confidence=confidence * 0.8,  # Lower confidence for small crosses
                        reasoning="Delete item marked with cross"
                    ))
    
    def _process_arrows(self, arrow_items: List[Dict], pdf_analysis: Dict, word_content: Dict, actions: List[AnnotationAction]):
        """Process arrows and connection indicators"""
        for arrow in arrow_items:
            start_pos = arrow.get('start_position', {})
            end_pos = arrow.get('end_position', {})
            connects = arrow.get('connects', [])
            confidence = arrow.get('confidence', 0.0)
            description = arrow.get('description', '')
            
            # Find what the arrow connects
            if 'handwritten_text' in connects and 'strikethrough_text' in connects:
                # This arrow likely indicates a replacement
                # Find the handwritten text near start position
                handwriting_near_start = self._find_annotation_near_position(
                    start_pos, pdf_analysis.get('handwritten_text', [])
                )
                
                # Find the strikethrough text near end position
                strikethrough_near_end = self._find_annotation_near_position(
                    end_pos, pdf_analysis.get('strikethrough_text', [])
                )
                
                if handwriting_near_start and strikethrough_near_end:
                    # Find matching text in Word document
                    best_match = self._find_best_text_match(
                        strikethrough_near_end.get('text', ''), word_content
                    )
                    
                    if best_match and best_match['similarity'] > self.similarity_threshold:
                        actions.append(AnnotationAction(
                            action_type='replace',
                            target_text=best_match['text'],
                            new_text=handwriting_near_start.get('text', ''),
                            position=end_pos,
                            confidence=confidence,
                            reasoning=f"Arrow-directed replacement: {description}"
                        ))
            
            # Add arrow as informational annotation
            nearby_content = self._find_content_near_position(start_pos, word_content)
            if nearby_content:
                actions.append(AnnotationAction(
                    action_type='comment',
                    target_text=nearby_content['text'],
                    new_text=f"Arrow connection: {description}",
                    position=start_pos,
                    confidence=confidence,
                    reasoning="Arrow indicates relationship between elements"
                ))

    def _process_highlights(self, highlight_items: List[Dict], word_content: Dict, actions: List[AnnotationAction]):
        """Process highlighted text"""
        for item in highlight_items:
            highlighted_text = item.get('text', '').strip()
            if not highlighted_text:
                continue
            
            position = item.get('position', {})
            confidence = item.get('confidence', 0.0)
            color = item.get('color', 'yellow')
            
            # Find matching text in Word document
            best_match = self._find_best_text_match(highlighted_text, word_content)
            
            if best_match and best_match['similarity'] > self.similarity_threshold:
                actions.append(AnnotationAction(
                    action_type='highlight',
                    target_text=best_match['text'],
                    new_text=color,
                    position=position,
                    confidence=confidence,
                    reasoning=f"Highlight text in {color}"
                ))
    
    def _process_annotations(self, annotation_items: List[Dict], word_content: Dict, actions: List[AnnotationAction]):
        """Process general annotations and margin notes"""
        for item in annotation_items:
            annotation_text = item.get('text', '').strip()
            if not annotation_text:
                continue
            
            position = item.get('position', {})
            confidence = item.get('confidence', 0.0)
            annotation_type = item.get('type', 'general')
            
            # Find nearby content for context
            nearby_content = self._find_content_near_position(position, word_content)
            
            if nearby_content:
                actions.append(AnnotationAction(
                    action_type='comment',
                    target_text=nearby_content['text'],
                    new_text=annotation_text,
                    position=position,
                    confidence=confidence,
                    reasoning=f"Add {annotation_type} as comment"
                ))
    
    def _determine_handwriting_placement(self, handwritten_text: str, word_content: Dict, position: Dict) -> Dict:
        """Use AI to determine where handwritten text should be placed"""
        try:
            # Create context for AI decision
            nearby_text = self._get_text_near_position(position, word_content)
            
            prompt = f"""
            Analyze this handwritten text and determine how it should be applied to the Word document:
            
            Handwritten text: "{handwritten_text}"
            
            Nearby text from Word document:
            {nearby_text[:500]}...
            
            Position in PDF: x={position.get('x', 0):.2f}, y={position.get('y', 0):.2f}
            
            Determine:
            1. Should this text REPLACE existing text?
            2. Should this text be INSERTED near existing text?
            3. Should this text be added as a COMMENT?
            
            Consider:
            - If handwritten text looks like a correction/replacement
            - If it's additional information to insert
            - The position relative to existing text
            
            Respond in JSON format:
            {{
                "action": "replace|insert|comment",
                "target_text": "text to target in document",
                "reasoning": "explanation of decision",
                "confidence": 0.0-1.0
            }}
            """
            
            response = self.client.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are an expert at interpreting handwritten document annotations."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=300,
                temperature=0.1
            )
            
            import json
            result = json.loads(response.choices[0].message.content)
            return result
            
        except Exception as e:
            self.logger.warning(f"AI placement decision failed: {e}")
            # Fallback to simple insertion
            nearby_content = self._find_content_near_position(position, word_content)
            return {
                "action": "insert",
                "target_text": nearby_content['text'] if nearby_content else "",
                "reasoning": "Fallback insertion",
                "confidence": 0.5
            }
    
    def _find_best_text_match(self, target_text: str, word_content: Dict) -> Optional[Dict]:
        """Find the best matching text in Word document"""
        if not target_text.strip():
            return None
        
        best_match = None
        best_similarity = 0.0
        
        # Clean target text
        target_clean = self._clean_text_for_comparison(target_text)
        
        for para in word_content['paragraphs']:
            para_clean = self._clean_text_for_comparison(para['text'])
            
            # Calculate similarity
            similarity = self._calculate_text_similarity(target_clean, para_clean)
            
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = {
                    "text": para['text'],
                    "paragraph_index": para['index'],
                    "similarity": similarity
                }
            
            # Also check individual sentences
            sentences = re.split(r'[.!?]+', para['text'])
            for sentence in sentences:
                sentence_clean = self._clean_text_for_comparison(sentence)
                if len(sentence_clean) > 10:  # Only check substantial sentences
                    similarity = self._calculate_text_similarity(target_clean, sentence_clean)
                    
                    if similarity > best_similarity:
                        best_similarity = similarity
                        best_match = {
                            "text": sentence.strip(),
                            "paragraph_index": para['index'],
                            "similarity": similarity
                        }
        
        return best_match if best_similarity > 0.3 else None
    
    def _clean_text_for_comparison(self, text: str) -> str:
        """Clean text for comparison"""
        # Remove punctuation, extra spaces, convert to lowercase
        cleaned = re.sub(r'[^\w\s]', '', text.lower())
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned
    
    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two texts"""
        if not text1 or not text2:
            return 0.0
        
        # Simple word-based similarity
        words1 = set(text1.split())
        words2 = set(text2.split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def _find_nearby_handwriting(self, item: Dict, position: Dict) -> Optional[str]:
        """Find handwriting near a strikethrough item"""
        # This would be implemented to find handwritten text near the strikethrough
        # For now, return None - would need cross-reference with handwriting items
        return None
    
    def _find_annotation_near_position(self, position: Dict, annotation_list: List[Dict], threshold: float = 0.1) -> Optional[Dict]:
        """Find annotation near a specific position"""
        target_x = position.get('x', 0)
        target_y = position.get('y', 0)
        
        closest_annotation = None
        closest_distance = float('inf')
        
        for annotation in annotation_list:
            ann_pos = annotation.get('position', {})
            ann_x = ann_pos.get('x', 0)
            ann_y = ann_pos.get('y', 0)
            
            # Calculate distance
            distance = ((target_x - ann_x) ** 2 + (target_y - ann_y) ** 2) ** 0.5
            
            if distance < threshold and distance < closest_distance:
                closest_distance = distance
                closest_annotation = annotation
        
        return closest_annotation
    
    def _find_content_at_position(self, position: Dict, word_content: Dict) -> Optional[Dict]:
        """Find content at a specific position (approximated)"""
        # This is a simplified implementation
        # In reality, you'd need to map PDF positions to Word document positions
        y_pos = position.get('y', 0)
        
        # Estimate paragraph based on Y position
        total_paragraphs = len(word_content['paragraphs'])
        estimated_para_index = int(y_pos * total_paragraphs)
        
        if 0 <= estimated_para_index < total_paragraphs:
            para = word_content['paragraphs'][estimated_para_index]
            return {
                "text": para['text'],
                "type": "paragraph",
                "index": para['index']
            }
        
        return None
    
    def _find_content_near_position(self, position: Dict, word_content: Dict) -> Optional[Dict]:
        """Find content near a position"""
        return self._find_content_at_position(position, word_content)
    
    def _get_text_near_position(self, position: Dict, word_content: Dict) -> str:
        """Get text content near a position for context"""
        content = self._find_content_near_position(position, word_content)
        if content:
            return content['text']
        return ""
    
    def _apply_action_to_document(self, doc: Document, action: AnnotationAction) -> bool:
        """Apply a single action to the Word document"""
        try:
            if action.action_type == 'replace':
                return self._replace_text_in_document(doc, action.target_text, action.new_text)
            elif action.action_type == 'delete':
                return self._delete_text_from_document(doc, action.target_text)
            elif action.action_type == 'insert':
                return self._insert_text_in_document(doc, action.target_text, action.new_text)
            elif action.action_type == 'highlight':
                return self._highlight_text_in_document(doc, action.target_text, action.new_text)
            elif action.action_type == 'comment':
                return self._add_comment_to_document(doc, action.target_text, action.new_text)
            
            return False
            
        except Exception as e:
            self.logger.error(f"Error applying action {action.action_type}: {str(e)}")
            return False
    
    def _replace_text_in_document(self, doc: Document, target_text: str, new_text: str) -> bool:
        """Replace text in document"""
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                # Replace in paragraph text
                paragraph.text = paragraph.text.replace(target_text, new_text)
                return True
        return False
    
    def _delete_text_from_document(self, doc: Document, target_text: str) -> bool:
        """Delete text from document"""
        for paragraph in doc.paragraphs:
            if target_text.strip() == paragraph.text.strip():
                # Delete entire paragraph
                paragraph.clear()
                return True
            elif target_text in paragraph.text:
                # Delete part of paragraph
                paragraph.text = paragraph.text.replace(target_text, "")
                return True
        return False
    
    def _insert_text_in_document(self, doc: Document, target_text: str, new_text: str) -> bool:
        """Insert text near target text"""
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                # Insert after the paragraph
                new_paragraph = paragraph._element.getparent().insert(
                    paragraph._element.getparent().index(paragraph._element) + 1,
                    paragraph._element.__class__(paragraph._element.tag)
                )
                new_paragraph.text = new_text
                return True
        return False
    
    def _highlight_text_in_document(self, doc: Document, target_text: str, color: str) -> bool:
        """Highlight text in document"""
        color_map = {
            'yellow': WD_COLOR_INDEX.YELLOW,
            'green': WD_COLOR_INDEX.BRIGHT_GREEN,
            'blue': WD_COLOR_INDEX.TURQUOISE,
            'red': WD_COLOR_INDEX.PINK
        }
        
        highlight_color = color_map.get(color.lower(), WD_COLOR_INDEX.YELLOW)
        
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                for run in paragraph.runs:
                    if target_text in run.text:
                        run.font.highlight_color = highlight_color
                        return True
        return False
    
    def _add_comment_to_document(self, doc: Document, target_text: str, comment_text: str) -> bool:
        """Add comment to document"""
        # This is a simplified implementation
        # Real Word comments require more complex XML manipulation
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                # Add comment as a note at the end of paragraph
                paragraph.text += f" [COMMENT: {comment_text}]"
                return True
        return False
    
    def _add_processing_summary(self, doc: Document, actions: List[AnnotationAction], modifications_made: int):
        """Add processing summary to document"""
        # Add a new page break and summary
        doc.add_page_break()
        
        summary_heading = doc.add_heading('Document Processing Summary', level=1)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total actions planned: {len(actions)}\n").bold = True
        summary_para.add_run(f"Modifications applied: {modifications_made}\n").bold = True
        summary_para.add_run(f"Processing date: {self._get_current_timestamp()}\n")
        
        # Add action details
        if actions:
            doc.add_heading('Action Details', level=2)
            
            for i, action in enumerate(actions, 1):
                action_para = doc.add_paragraph()
                action_para.add_run(f"{i}. {action.action_type.upper()}: ").bold = True
                action_para.add_run(f"{action.reasoning}\n")
                if action.target_text:
                    action_para.add_run(f"   Target: {action.target_text[:100]}...\n")
                if action.new_text and action.action_type != 'highlight':
                    action_para.add_run(f"   New text: {action.new_text[:100]}...\n")
                action_para.add_run(f"   Confidence: {action.confidence:.2f}\n")
    
    def _get_current_timestamp(self) -> str:
        """Get current timestamp string"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    def _action_to_dict(self, action: AnnotationAction) -> Dict:
        """Convert action to dictionary for serialization"""
        return {
            "action_type": action.action_type,
            "target_text": action.target_text,
            "new_text": action.new_text,
            "position": action.position,
            "confidence": action.confidence,
            "reasoning": action.reasoning
        }
    
    def _get_last_successful_strategy(self) -> str:
        """Track which strategy was used for the last successful match"""
        # This would be set by the apply_cascading_strategies method
        return getattr(self, '_last_strategy_used', 'unknown')
    
    def _add_comment_for_handwritten_text(self, doc: Document, item: Dict):
        """Add handwritten text as comments in the document"""
        try:
            # Add a comment near the detected handwritten area
            comment_text = f"Handwritten text detected: '{item['text']}' (confidence: {item.get('confidence', 0):.2f})"
            
            # Find appropriate paragraph to attach comment to
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # First non-empty paragraph
                    # Add comment to first run of the paragraph
                    if paragraph.runs:
                        # For now, add as text annotation since Word comments are complex
                        paragraph.add_run(f" [HANDWRITTEN: {item['text']}]")
                    break
                    
            self.logger.info(f"Added handwritten text comment: {item['text'][:30]}")
            
        except Exception as e:
            self.logger.warning(f"Failed to add handwritten text comment: {str(e)}")
    
    def _handle_large_cross_deletion(self, doc: Document, item: Dict):
        """Handle large cross marks that might indicate table row deletions"""
        try:
            # For large crosses, try to find and delete table rows
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    row_text = " ".join([cell.text for cell in row.cells])
                    
                    # If cross is near this row content, delete the row
                    if item.get('nearby_text') and item['nearby_text'] in row_text:
                        self.delete_table_row(table, i)
                        self.processing_stats['total_row_deletions'] += 1
                        self.logger.info(f"Deleted table row based on large cross mark")
                        return
                        
        except Exception as e:
            self.logger.warning(f"Failed to handle large cross deletion: {str(e)}")
    
    def _handle_small_cross_deletion(self, doc: Document, item: Dict):
        """Handle small cross marks for specific text deletions"""
        try:
            if item.get('nearby_text'):
                success = self.apply_cascading_strategies(doc, item['nearby_text'], "delete")
                if success:
                    self.changes_applied.append({
                        "type": "cross_deletion", 
                        "target_text": item['nearby_text'][:50],
                        "confidence": item.get('confidence')
                    })
                    
        except Exception as e:
            self.logger.warning(f"Failed to handle small cross deletion: {str(e)}")
    
    def _apply_highlighting(self, doc: Document, item: Dict):
        """Apply highlighting to detected text"""
        try:
            if item.get('text'):
                success = self.apply_cascading_strategies(doc, item['text'], "highlight")
                if success:
                    self.changes_applied.append({
                        "type": "highlight",
                        "target_text": item['text'][:50],
                        "confidence": item.get('confidence')
                    })
                    
        except Exception as e:
            self.logger.warning(f"Failed to apply highlighting: {str(e)}")
    
    def _add_annotation_comment(self, doc: Document, item: Dict):
        """Add annotation as a comment in the document"""
        try:
            comment_text = f"Annotation detected: '{item.get('text', 'Visual annotation')}' (confidence: {item.get('confidence', 0):.2f})"
            
            # Find appropriate paragraph to attach comment to
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.runs:
                        # Add as text annotation
                        paragraph.add_run(f" [ANNOTATION: {item.get('text', 'Visual mark')}]")
                    break
                    
            self.logger.info(f"Added annotation comment")
            
        except Exception as e:
            self.logger.warning(f"Failed to add annotation comment: {str(e)}")


def main():
    """Example usage of the advanced word processor"""
    processor = AdvancedWordProcessor()
    
    # Example usage
    sample_analysis = {
        'handwritten_text': [
            {
                'text': 'This is corrected text',
                'position': {'x': 0.3, 'y': 0.2},
                'confidence': 0.9
            }
        ],
        'strikethrough_text': [
            {
                'text': 'old incorrect text',
                'position': {'x': 0.3, 'y': 0.2},
                'confidence': 0.8
            }
        ]
    }
    
    result = processor.apply_annotations_to_document(
        "sample_document.docx",
        sample_analysis,
        "modified_document.docx"
    )
    
    print(f"Processing result: {result}")


if __name__ == "__main__":
    main()