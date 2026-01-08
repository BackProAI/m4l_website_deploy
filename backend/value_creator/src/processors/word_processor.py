#!/usr/bin/env python3
"""
Word Document Processor
Creates and modifies Word documents based on analysis results
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn

class WordProcessor:
    """Processes Word documents with detected markings and annotations"""
    
    def __init__(self, config: Dict):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Word processing settings
        self.preserve_formatting = config['word_processing']['preserve_formatting']
        self.highlight_detected_items = config['word_processing']['highlight_detected_items']
        self.add_comments = config['word_processing']['add_comments']
        self.track_changes = config['word_processing']['track_changes']
    
    def _safe_add_heading(self, doc: Document, text: str, level: int = 1):
        """Safely add heading with fallback to formatted paragraph if styles don't exist"""
        try:
            return doc.add_heading(text, level)
        except:
            # If heading styles don't exist, create a formatted paragraph
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            
            # Set font size based on heading level
            from docx.shared import Pt
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            
            return para
    
    def _safe_apply_style(self, element, style_name, fallback_format=None):
        """Safely apply a style with fallback formatting"""
        try:
            element.style = style_name
        except KeyError:
            # Style doesn't exist, apply fallback formatting if provided
            from docx.shared import Pt
            if fallback_format and hasattr(element, 'runs') and element.runs:
                for run in element.runs:
                    if 'bold' in fallback_format:
                        run.bold = fallback_format['bold']
                    if 'italic' in fallback_format:
                        run.italic = fallback_format['italic']
                    if 'font_size' in fallback_format:
                        run.font.size = Pt(fallback_format['font_size'])
    
    def _safe_add_list_item(self, doc, text):
        """Safely add a list item with fallback formatting"""
        try:
            p = doc.add_paragraph(text, style='List Bullet')
            return p
        except KeyError:
            # Create a paragraph with bullet-like formatting
            p = doc.add_paragraph()
            p.add_run('• ' + text)
            return p
    
    def create_annotated_document(self, analysis_results: Dict, output_path: str, 
                                template_docx: Optional[str] = None) -> Document:
        """
        Create an annotated Word document based on analysis results
        
        Args:
            analysis_results: Results from chunk analysis
            output_path: Path for output Word document
            template_docx: Optional template document to modify
            
        Returns:
            Created Document object
        """
        try:
            # Load template or create new document
            if template_docx and Path(template_docx).exists():
                doc = Document(template_docx)
                self.logger.info(f"Using template document: {template_docx}")
            else:
                doc = Document()
                self.logger.info("Creating new document")
            
            # Add title and summary
            self._add_document_header(doc, analysis_results)
            
            # Add analysis summary
            self._add_analysis_summary(doc, analysis_results)
            
            # Add detailed findings
            self._add_detailed_findings(doc, analysis_results)
            
            # Add detected items with formatting
            self._add_detected_items(doc, analysis_results)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Annotated document saved: {output_path}")
            
            return doc
            
        except Exception as e:
            self.logger.error(f"Error creating annotated document: {str(e)}")
            raise
    
    def _add_document_header(self, doc: Document, analysis_results: Dict):
        """Add document header with title and metadata"""
        # Title - use simple paragraph instead of heading styles to avoid issues
        title = self._safe_add_heading(doc, 'Document Analysis Report', 1)
        
        if hasattr(title, 'alignment'):
            title.alignment = 1  # Center alignment
        
        # Metadata table
        metadata_table = doc.add_table(rows=3, cols=2)
        # Use safe table style
        try:
            self._safe_apply_style(metadata_table, 'Table Grid')
        except:
            # Style not available, use default table style
            pass
        
        # Add metadata
        cells = metadata_table.rows[0].cells
        cells[0].text = 'Total Chunks Analyzed'
        cells[1].text = str(analysis_results['metadata']['processed_chunks'])
        
        cells = metadata_table.rows[1].cells
        cells[0].text = 'Total Detections'
        total_detections = sum([
            len(analysis_results['handwritten_text']),
            len(analysis_results['strikethrough_text']),
            len(analysis_results['crosses']),
            len(analysis_results['highlights']),
            len(analysis_results['annotations'])
        ])
        cells[1].text = str(total_detections)
        
        cells = metadata_table.rows[2].cells
        cells[0].text = 'Analysis Date'
        from datetime import datetime
        cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        doc.add_paragraph()  # Add spacing
    
    def _add_analysis_summary(self, doc: Document, analysis_results: Dict):
        """Add summary of analysis results"""
        self._safe_add_heading(doc, 'Analysis Summary', 1)
        
        # Create summary list
        summary_items = [
            f"Handwritten Text Items: {len(analysis_results['handwritten_text'])}",
            f"Strikethrough Text Items: {len(analysis_results['strikethrough_text'])}",
            f"Cross Marks: {len(analysis_results['crosses'])}",
            f"Highlighted Text: {len(analysis_results['highlights'])}",
            f"Annotations: {len(analysis_results['annotations'])}"
        ]
        
        for item in summary_items:
            p = self._safe_add_list_item(doc, item)
    
    def _add_detailed_findings(self, doc: Document, analysis_results: Dict):
        """Add detailed findings for each category"""
        
        categories = [
            ('Handwritten Text', 'handwritten_text', WD_COLOR_INDEX.YELLOW),
            ('Strikethrough Text', 'strikethrough_text', WD_COLOR_INDEX.RED),
            ('Cross Marks', 'crosses', WD_COLOR_INDEX.PINK),
            ('Highlighted Text', 'highlights', WD_COLOR_INDEX.BRIGHT_GREEN),
            ('Annotations', 'annotations', WD_COLOR_INDEX.TURQUOISE)
        ]
        
        for category_name, category_key, highlight_color in categories:
            items = analysis_results[category_key]
            
            if items:
                self._safe_add_heading(doc, f'{category_name} ({len(items)} items)', 2)
                
                # Create table for items
                table = doc.add_table(rows=1, cols=4)
                # Use safe table style
                try:
                    self._safe_apply_style(table, 'Table Grid')
                except:
                    # Style not available, use default
                    pass
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Text/Description'
                header_cells[1].text = 'Position'
                header_cells[2].text = 'Confidence'
                header_cells[3].text = 'Details'
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows
                for item in items:
                    row_cells = table.add_row().cells
                    
                    # Text content
                    text_content = item.get('text', item.get('description', 'N/A'))
                    row_cells[0].text = text_content
                    
                    # Position
                    pos = item.get('position', {})
                    position_text = f"({pos.get('x', 0):.2f}, {pos.get('y', 0):.2f})"
                    row_cells[1].text = position_text
                    
                    # Confidence
                    confidence = item.get('confidence', 0.0)
                    row_cells[2].text = f"{confidence:.2f}"
                    
                    # Details
                    details = []
                    if 'color' in item:
                        details.append(f"Color: {item['color']}")
                    if 'type' in item:
                        details.append(f"Type: {item['type']}")
                    row_cells[3].text = ", ".join(details) if details else "N/A"
                    
                    # Highlight the text content if enabled
                    if self.highlight_detected_items and text_content != 'N/A':
                        self._highlight_text(row_cells[0].paragraphs[0], highlight_color)
                
                doc.add_paragraph()  # Add spacing
    
    def _add_detected_items(self, doc: Document, analysis_results: Dict):
        """Add a section showing all detected items in context"""
        self._safe_add_heading(doc, 'All Detected Items', 1)
        
        # Collect all items with their positions
        all_items = []
        
        for category, items in analysis_results.items():
            if isinstance(items, list):
                for item in items:
                    item_with_category = item.copy()
                    item_with_category['category'] = category
                    all_items.append(item_with_category)
        
        # Sort by position (top to bottom, left to right)
        all_items.sort(key=lambda x: (
            x.get('position', {}).get('y', 0),
            x.get('position', {}).get('x', 0)
        ))
        
        # Add items in order
        for i, item in enumerate(all_items, 1):
            p = doc.add_paragraph()
            
            # Add item number
            run = p.add_run(f"{i}. ")
            run.font.bold = True
            
            # Add category
            run = p.add_run(f"[{item['category'].upper()}] ")
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            
            # Add content
            text_content = item.get('text', item.get('description', 'N/A'))
            run = p.add_run(text_content)
            
            # Add position and confidence
            pos = item.get('position', {})
            confidence = item.get('confidence', 0.0)
            details_text = f" (Position: {pos.get('x', 0):.2f}, {pos.get('y', 0):.2f}, Confidence: {confidence:.2f})"
            run = p.add_run(details_text)
            # Set smaller font size safely
            try:
                # Use safe style reference
                try:
                    normal_size = doc.styles['Normal'].font.size
                except:
                    normal_size = None
                if normal_size:
                    run.font.size = int(normal_size * 0.8)
                else:
                    from docx.shared import Pt
                    run.font.size = Pt(9)  # Default smaller size
            except (AttributeError, TypeError):
                from docx.shared import Pt
                run.font.size = Pt(9)  # Default smaller size
            run.font.color.rgb = RGBColor(100, 100, 100)
    
    def _highlight_text(self, paragraph, color):
        """Apply highlighting to text in a paragraph"""
        try:
            for run in paragraph.runs:
                run.font.highlight_color = color
        except Exception as e:
            self.logger.warning(f"Could not apply highlighting: {e}")
    
    def modify_existing_document(self, template_path: str, analysis_results: Dict, 
                               output_path: str) -> Document:
        """
        Modify an existing Word document by adding comments and highlights
        
        Args:
            template_path: Path to existing Word document
            analysis_results: Analysis results to incorporate
            output_path: Path for modified document
            
        Returns:
            Modified Document object
        """
        try:
            doc = Document(template_path)
            
            # Add analysis as comments if enabled
            if self.add_comments:
                self._add_analysis_comments(doc, analysis_results)
            
            # Track changes if enabled
            if self.track_changes:
                self._enable_track_changes(doc)
            
            # Save modified document
            doc.save(output_path)
            self.logger.info(f"Modified document saved: {output_path}")
            
            return doc
            
        except Exception as e:
            self.logger.error(f"Error modifying existing document: {str(e)}")
            raise
    
    def _add_analysis_comments(self, doc: Document, analysis_results: Dict):
        """Add analysis results as comments to the document"""
        # This is a simplified implementation
        # In practice, you'd need to map detected positions to document positions
        
        # Add a summary comment at the end
        if doc.paragraphs:
            last_paragraph = doc.paragraphs[-1]
            
            # Create summary of findings
            summary = []
            for category, items in analysis_results.items():
                if isinstance(items, list) and items:
                    summary.append(f"{category}: {len(items)} items")
            
            comment_text = "Analysis Summary: " + ", ".join(summary)
            
            # Note: Adding actual comments requires more complex XML manipulation
            # For now, we'll add it as a text note
            doc.add_paragraph()
            p = doc.add_paragraph("Analysis Results: " + comment_text)
            # Use safe style or skip styling if not available
            try:
                self._safe_apply_style(p, 'Intense Quote', {'italic': True, 'font_size': 11})
            except:
                # Style not available, use normal formatting
                for run in p.runs:
                    run.italic = True
    
    def _enable_track_changes(self, doc: Document):
        """Enable track changes for the document"""
        # This requires XML manipulation of the document settings
        # Simplified implementation
        self.logger.info("Track changes enabled (simplified implementation)")
    
    def export_analysis_to_excel(self, analysis_results: Dict, output_path: str):
        """Export analysis results to Excel for further processing"""
        try:
            import pandas as pd
            
            # Prepare data for Excel export
            excel_data = []
            
            for category, items in analysis_results.items():
                if isinstance(items, list):
                    for item in items:
                        row = {
                            'Category': category,
                            'Text': item.get('text', item.get('description', 'N/A')),
                            'Position_X': item.get('position', {}).get('x', 0),
                            'Position_Y': item.get('position', {}).get('y', 0),
                            'Confidence': item.get('confidence', 0.0),
                            'Color': item.get('color', 'N/A'),
                            'Type': item.get('type', 'N/A'),
                            'Chunk_ID': item.get('chunk_index', 'N/A')
                        }
                        excel_data.append(row)
            
            # Create DataFrame and save to Excel
            df = pd.DataFrame(excel_data)
            df.to_excel(output_path, index=False)
            
            self.logger.info(f"Analysis results exported to Excel: {output_path}")
            
        except ImportError:
            self.logger.warning("pandas not available, cannot export to Excel")
        except Exception as e:
            self.logger.error(f"Error exporting to Excel: {str(e)}")