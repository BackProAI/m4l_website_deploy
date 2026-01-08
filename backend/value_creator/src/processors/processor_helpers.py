#!/usr/bin/env python3
"""
Helper methods for Unified Section Processor
Advanced Word document manipulation methods and utilities
"""

from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Inches
import logging

class ProcessorHelpers:
    """Helper methods for the unified section processor"""
    
    def __init__(self, processor):
        self.processor = processor
        self.logger = logging.getLogger(__name__)
    
    # ===== ADDITIONAL HELPER METHODS =====
    
    def _add_text_to_bullet_point(self, doc: Document, bullet_text: str, handwritten_text: str) -> Tuple[bool, str]:
        """Add handwritten text to existing bullet points"""
        # Strategy 1: Find exact bullet match
        for paragraph in doc.paragraphs:
            if bullet_text.lower() in paragraph.text.lower():
                # Add handwritten text to the end of the bullet
                paragraph.add_run(f" {handwritten_text}")
                return True, "exact"
        
        # Strategy 2: Find similar bullet
        for paragraph in doc.paragraphs:
            if self.processor._text_similarity(paragraph.text, bullet_text) > 0.6:
                paragraph.add_run(f" {handwritten_text}")
                return True, "similarity"
        
        return False, "failed"
    
    def _is_section_table(self, table, section_name: str) -> bool:
        """Identify if table belongs to specific section (customize based on your document structure)"""
        # This is a placeholder - you'd implement logic to identify tables by content/position
        # For example, check first row content, table position, etc.
        
        if not table.rows:
            return False
        
        # Example: Check if first row contains section-specific keywords
        first_row_text = ""
        if table.rows:
            first_row_text = " ".join([cell.text.lower() for cell in table.rows[0].cells])
        
        section_keywords = {
            "section_1_4": ["items discussed", "action taken", "left box", "right box"],
            "section_1_2": ["goals", "achieved", "bullet"],
            "section_1_3": ["portfolio", "selection"]
        }
        
        keywords = section_keywords.get(section_name, [])
        return any(keyword in first_row_text for keyword in keywords)
    
    def _identify_rows_for_deletion(self, table, analysis: Dict) -> List[int]:
        """Identify which table rows should be deleted based on analysis"""
        rows_to_delete = []
        
        # Check if analysis indicates specific rows for deletion
        row_deletion_info = analysis.get('row_deletion_rule', {})
        
        if row_deletion_info.get('delete_entire_row', False):
            # If analysis specifies row numbers
            if 'row_numbers' in row_deletion_info:
                rows_to_delete = row_deletion_info['row_numbers']
            else:
                # Find rows where both left and right boxes have deletion marks
                left_box = analysis.get('left_box_analysis', {})
                right_box = analysis.get('right_box_analysis', {})
                
                left_all_marked = left_box.get('all_sentences_have_deletion_marks', False)
                right_all_marked = right_box.get('all_sentences_have_deletion_marks', False)
                
                if left_all_marked and right_all_marked:
                    # Delete all content rows (skip header if exists)
                    for i in range(1, len(table.rows)):  # Skip first row (header)
                        rows_to_delete.append(i)
        
        return rows_to_delete
    
    def _delete_table_row(self, table, row_index: int) -> bool:
        """Delete a specific table row using direct XML manipulation"""
        try:
            if 0 <= row_index < len(table.rows):
                row = table.rows[row_index]
                table._tbl.remove(row._tr)  # Direct XML manipulation
                return True
        except Exception as e:
            self.logger.error(f"Failed to delete table row {row_index}: {str(e)}")
        
        return False
    
    def _add_processing_summary_to_document(self, doc: Document):
        """Add comprehensive processing summary to document"""
        try:
            # Add page break
            doc.add_page_break()
            
            # Add heading
            heading = doc.add_heading('Document Processing Summary', level=1)
            
            # Add processing statistics
            stats_para = doc.add_paragraph()
            stats_para.add_run("Processing Statistics:\n").bold = True
            
            total_changes = (self.processor.processing_stats['exact_matches'] + 
                           self.processor.processing_stats['similarity_matches'] + 
                           self.processor.processing_stats['keyword_matches'])
            
            stats_para.add_run(f"• Total sections processed: {self.processor.processing_stats['sections_processed']}\n")
            stats_para.add_run(f"• Total changes applied: {total_changes}\n")
            stats_para.add_run(f"• Exact matches: {self.processor.processing_stats['exact_matches']}\n")
            stats_para.add_run(f"• Similarity matches: {self.processor.processing_stats['similarity_matches']}\n")
            stats_para.add_run(f"• Keyword matches: {self.processor.processing_stats['keyword_matches']}\n")
            stats_para.add_run(f"• Failed matches: {self.processor.processing_stats['failed_matches']}\n")
            
            # Add operation breakdown
            ops_para = doc.add_paragraph()
            ops_para.add_run("Operations Performed:\n").bold = True
            ops_para.add_run(f"• Deletions: {self.processor.processing_stats['total_deletions']}\n")
            ops_para.add_run(f"• Replacements: {self.processor.processing_stats['total_replacements']}\n")
            ops_para.add_run(f"• Row deletions: {self.processor.processing_stats['total_row_deletions']}\n")
            
            # Add timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_para.add_run(f"Processing completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            
            # Add detailed changes if available
            if self.processor.changes_applied:
                changes_heading = doc.add_heading('Detailed Changes', level=2)
                
                for i, change in enumerate(self.processor.changes_applied[:20], 1):  # Limit to first 20
                    change_para = doc.add_paragraph()
                    change_para.add_run(f"{i}. {change.type.upper()}: ").bold = True
                    change_para.add_run(f"Section {change.section}\n")
                    change_para.add_run(f"   Original: {change.original_text[:100]}...\n")
                    if change.new_text:
                        change_para.add_run(f"   New: {change.new_text[:100]}...\n")
                    change_para.add_run(f"   Strategy: {change.strategy_used}\n")
                
                if len(self.processor.changes_applied) > 20:
                    doc.add_paragraph(f"... and {len(self.processor.changes_applied) - 20} more changes")
            
            # Add errors if any
            if self.processor.section_errors:
                errors_heading = doc.add_heading('Processing Errors', level=2)
                for error in self.processor.section_errors:
                    error_para = doc.add_paragraph()
                    error_para.add_run(f"• {error}")
            
        except Exception as e:
            self.logger.error(f"Failed to add processing summary: {str(e)}")
    
    def get_section_ai_prompts(self) -> Dict[str, str]:
        """Get AI prompts for each section type (customize for your specific needs)"""
        return {
            "section_1_4": """
            SECTION LAYOUT: This section has TWO BOXES side by side (LEFT BOX and RIGHT BOX)
            
            CRITICAL MARK DETECTION RULES:
            1. DIAGONAL LINES/CROSSES: Delete affected sentence(s) completely
            2. HORIZONTAL STRIKETHROUGHS: Delete affected sentence(s) completely  
            3. HORIZONTAL ARROWS: Replace text (find original text and replacement text)
            4. ROW DELETION RULE: If BOTH boxes have marks through ALL sentences → delete entire row
            
            Return JSON:
            {
              "left_box_analysis": {
                "sentences_to_delete": [...],
                "sentences_to_replace": [...],
                "all_sentences_have_deletion_marks": true/false
              },
              "right_box_analysis": {...},
              "row_deletion_rule": {
                "delete_entire_row": true/false,
                "deletion_reason": "..."
              }
            }
            """,
            
            "section_1_2": """
            This section contains bullet points with potential handwritten additions.
            
            DETECTION RULES:
            1. Look for handwritten text NEAR existing bullet points
            2. Ignore checkmarks - only extract actual text
            3. Associate handwritten text with nearby bullet points
            
            Return JSON:
            {
              "handwritten_text": [
                {
                  "text": "handwritten content",
                  "nearby_bullet_text": "existing bullet point text",
                  "confidence": 0.85
                }
              ]
            }
            """,
            
            "section_1_3": """
            Portfolio selection section with circles and crosses.
            
            DETECTION RULES:
            1. CIRCLES around text = KEEP this item
            2. X MARKS or CROSSES = DELETE this item
            3. Diagonal lines through bullet points = DELETE bullet
            
            Return JSON:
            {
              "circled_items": [...],
              "crossed_items": [...],
              "marked_bullet_points": [...]
            }
            """
        }


def main():
    """Example usage of the unified section processor"""
    from .unified_section_processor import UnifiedSectionProcessor
    processor = UnifiedSectionProcessor()
    
    # Example section analyses (would come from your AI analysis)
    sample_analyses = {
        "section_1_4": {
            "left_box_analysis": {
                "sentences_to_delete": [
                    {
                        "sentence_text": "Look at maximising your superannuation",
                        "mark_type": "diagonal_line",
                        "confidence": 0.9
                    }
                ],
                "sentences_to_replace": [],
                "all_sentences_have_deletion_marks": False
            },
            "right_box_analysis": {
                "sentences_to_delete": [],
                "sentences_to_replace": [
                    {
                        "original_text": "age pension entitlements",
                        "replacement_text": "retirement planning",
                        "mark_type": "horizontal_arrow",
                        "confidence": 0.85
                    }
                ],
                "all_sentences_have_deletion_marks": False
            },
            "row_deletion_rule": {
                "delete_entire_row": False,
                "deletion_reason": None
            }
        }
    }
    
    # Process all sections
    result = processor.process_all_sections(
        section_analyses=sample_analyses,
        base_document_path="data/input/template.docx",
        output_path="data/output/processed_document.docx"
    )
    
    print(f"Processing result: {result}")


if __name__ == "__main__":
    main()